"""docx2typed tool smoke suite.

Invokes every CLI command and every MCP tool against a small fixture,
asserting each either succeeds or fails with a stable error code. Exit 0 only
when every tool behaves.

Usage:
    python -m scripts.tool_smoke [--workdir DIR]
"""
from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent


def _cli(*args: object, timeout: int = 600) -> tuple[int, str]:
    result = subprocess.run(
        ["python", "-m", "scripts", *(str(a) for a in args)],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
        timeout=timeout,
    )
    return result.returncode, (result.stdout or result.stderr).strip()


MCP_DRIVER_LOOP = r"""
import json, sys, importlib
server = importlib.import_module("scripts.mcp_server")
for line in sys.stdin:
    request = json.loads(line)
    try:
        out = getattr(server, request["tool"])(**request["args"])
        print("OK " + json.dumps(out, ensure_ascii=True)[:400], flush=True)
    except Exception as exc:
        print("ERR " + str(exc)[:200], flush=True)
"""


class Mcp:
    """Persistent stdio session against the MCP server module, so that
    workdir_open state survives across tool calls."""

    def __init__(self) -> None:
        self.proc = subprocess.Popen(
            [sys.executable, "-c", MCP_DRIVER_LOOP],
            cwd=REPO_ROOT,
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            text=True,
        )
        assert self.proc.stdin is not None and self.proc.stdout is not None

    def __call__(self, tool: str, **kwargs) -> str:
        self.proc.stdin.write(json.dumps({"tool": tool, "args": kwargs}) + "\n")
        self.proc.stdin.flush()
        line = self.proc.stdout.readline()
        return line.strip() if line else "ERR driver closed"

    def close(self) -> None:
        self.proc.terminate()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--workdir", default=None)
    args = parser.parse_args(argv)
    scratch = Path(args.workdir) if args.workdir else Path(__file__).parent.parent / ".smoke"
    if scratch.exists():
        shutil.rmtree(scratch)
    scratch.mkdir(parents=True)

    # fixture: small docx with revision, table, comment, header
    fixture = scratch / "fixture.docx"
    _make_fixture(fixture)
    workdir = scratch / "wd"
    results: list[tuple[str, str, str]] = []

    def check(name: str, ok: bool, detail: str = "") -> None:
        results.append((name, "PASS" if ok else "FAIL", detail))
        print(f"{'PASS' if ok else 'FAIL':4} {name}" + (f"  {detail[:90]}" if not ok and detail else ""))

    # ---- CLI commands ----
    rc, out = _cli("extract", fixture, "-o", workdir)
    check("cli extract", rc == 0, out)
    rc, out = _cli("validate", workdir)
    check("cli validate", rc == 0, out)
    rc, out = _cli("view", workdir)
    check("cli view", rc == 0, out)
    rc, out = _cli("view", workdir, "--mode", "style")
    check("cli view style", rc == 0, out)
    rc, out = _cli("build", workdir, "-o", scratch / "built.docx")
    check("cli build", rc == 0, out)
    rc, out = _cli("verify", workdir, scratch / "built.docx")
    check("cli verify", rc == 0, out)
    rc, out = _cli("normalize", "--help")
    check("cli normalize help", rc == 0, out)
    rc, out = _cli("audit", "--help")
    check("cli audit help", rc == 0, out)
    rc, out = _cli("edit", "status", workdir)
    check("cli edit status", rc == 0, out)
    rc, out = _cli("edit", "refresh", workdir)
    check("cli edit refresh", rc == 0, out)

    # decide: settle all revisions onto a new baseline (comments survive)
    out_all = scratch / "all.docx"
    wd_all = scratch / "all-wd"
    rc, out = _cli("decide", "accept-all", "--workdir", workdir, "--output", out_all, "--workdir-out", wd_all)
    check("cli decide accept-all", rc == 0, out)
    if rc == 0:
        inv_path = wd_all / "revisions.json"
        if inv_path.exists():
            inv = json.loads(inv_path.read_text(encoding="utf-8"))
            editable = [r for r in inv["revisions"] if r.get("editable")]
            if editable:
                key = editable[0]["revision_key"]
                fp = key.split("|")[3]
                rc, out = _cli("decide", "accept", key, "--workdir", wd_all, "--fingerprint", fp)
                check("cli decide accept", rc == 0, out)
        # edit sync on the settled baseline
        edit_path = wd_all / "edit.md"
        if edit_path.exists():
            text = edit_path.read_text(encoding="utf-8")
            marker = text.find("-->", text.find("<!--@p id="))
            if marker >= 0:
                tail = text.find("\n", marker) + 1
                end = text.find("\n", tail)
                if end < 0:
                    end = len(text)
                text = text[:tail] + text[tail:end] + "改" + text[end:]
                edit_path.write_text(text, encoding="utf-8")
        rc, out = _cli("edit", "sync", wd_all)
        check("cli edit sync", rc == 0, out)
        # comment-delete
        fmt = json.loads((wd_all / "format.json").read_text(encoding="utf-8"))
        cid = next((r.get("part_entry_id") for r in fmt["paragraphs"] if r.get("part_key") == "comments"), None)
        if cid:
            rc, out = _cli("decide", "comment-delete", cid, "--workdir", wd_all)
            check("cli decide comment-delete", rc == 0, out)
        # table op
        out_tbl = scratch / "tbl.docx"
        wd_tbl = scratch / "tbl-wd"
        rc, out = _cli("decide", "table-insert-row", "T0", "--workdir", wd_all, "--output", out_tbl, "--workdir-out", wd_tbl, "--args", "0")
        check("cli table-insert-row", rc == 0, out)

    # error paths: bad input
    rc, out = _cli("extract", scratch / "missing.docx", "-o", scratch / "x")
    check("cli extract missing -> error code 1", rc == 1, out)
    rc, out = _cli("decide", "accept", "bogus|key", "--workdir", workdir, "--fingerprint", "x" * 12)
    check("cli decide bad key -> error code 1", rc == 1, out)

    # ---- MCP tools ----
    mcp = Mcp()
    r = mcp("workdir_open", workdir=str(wd_all))
    check("mcp workdir_open", r.startswith("OK"), r)
    r = mcp("workdir_status")
    check("mcp workdir_status", r.startswith("OK"), r)
    r = mcp("list_paragraphs")
    check("mcp list_paragraphs", r.startswith("OK"), r)
    r = mcp("get_paragraph", paragraph_id="P0")
    check("mcp get_paragraph", r.startswith("OK"), r)
    r = mcp("replace_text", paragraph_id="P0", old="前", new="前改")
    check("mcp replace_text", r.startswith("OK"), r)
    r = mcp("batch_edit", paragraph_id="P0", edits=[{"region": 0, "old": "前改", "new": "前改2"}])
    check("mcp batch_edit", r.startswith("OK"), r)
    r = mcp("diff_preview")
    check("mcp diff_preview", r.startswith("OK"), r)
    r = mcp("commit_sync")
    check("mcp commit_sync", r.startswith("OK"), r)
    r = mcp("build_docx", output=str(scratch / "mcp-built.docx"))
    check("mcp build_docx", r.startswith("OK"), r)
    r = mcp("verify_output", output=str(scratch / "mcp-built.docx"))
    check("mcp verify_output", r.startswith("OK"), r)
    r = mcp("insert_paragraph", after_id="P0", text="新段")
    check("mcp insert_paragraph", r.startswith("OK"), r)
    r = mcp("delete_paragraph", paragraph_id="P0")
    check("mcp delete_paragraph", r.startswith("OK"), r)
    r = mcp("revert")
    check("mcp revert", r.startswith("OK"), r)
    # comment tools
    r = mcp("delete_comment", comment_id="99999")
    check("mcp delete_comment missing -> stable error", r.startswith("ERR") and "comment-not-found" in r, r)
    # table tools (with a table-bearing fixture workdir)
    tbl_fixture = scratch / "tbl-fixture.docx"
    _make_table_fixture(tbl_fixture)
    tbl_wd = scratch / "tbl-wd2"
    _cli("extract", tbl_fixture, "-o", tbl_wd)
    mcp("workdir_open", workdir=str(tbl_wd))
    r = mcp("table_insert_row", table_ref="T0", after=0, output=str(scratch / "mcp-tbl.docx"), workdir_out=str(scratch / "mcp-tbl-wd"))
    check("mcp table_insert_row", r.startswith("OK"), r)
    r = mcp("table_delete_col", table_ref="T0", col=0, output=str(scratch / "mcp-tbl2.docx"), workdir_out=str(scratch / "mcp-tbl2-wd"))
    check("mcp table_delete_col", r.startswith("OK"), r)
    r = mcp("table_merge_cells", table_ref="T0", row=0, col=0, span=2, output=str(scratch / "mcp-tbl3.docx"), workdir_out=str(scratch / "mcp-tbl3-wd"), discard_content=True)
    check("mcp table_merge_cells", r.startswith("OK"), r)
    r = mcp("table_merge_cells", table_ref="T0", row=1, col=0, span=2, output=str(scratch / "mcp-tbl3b.docx"), workdir_out=str(scratch / "mcp-tbl3b-wd"))
    check("mcp table_merge_cells guard (content)", not r.startswith("OK") and "merge-would-discard-content" in r, r)
    r = mcp("table_split_cells", table_ref="T0", row=0, col=0, span=2, output=str(scratch / "mcp-tbl4.docx"), workdir_out=str(scratch / "mcp-tbl4-wd"))
    check("mcp table_split_cells", r.startswith("OK"), r)

    mcp.close()
    failed = sum(1 for _, status, _ in results if status == "FAIL")
    print(f"\n{len(results) - failed}/{len(results)} tool checks passed")
    return 1 if failed else 0


def _make_fixture(path: Path) -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前文")
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "tester")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "插入"
    run.append(text)
    ins.append(run)
    paragraph._p.append(ins)
    paragraph.add_run("后文")
    commented = paragraph.add_run("批注目标")
    document.add_comment(commented, text="评论", author="tester")
    document.sections[0].header.paragraphs[0].text = "页眉"
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    document.save(str(path))


def _make_table_fixture(path: Path) -> None:
    from docx import Document

    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    document.save(str(path))


if __name__ == "__main__":
    sys.exit(main())

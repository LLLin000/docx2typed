"""Generate the deterministic release-corpus fixtures (corpus/release/*.docx).

Every fixture is synthetic, committed, and byte-stable across runs so the
release task suite is reproducible on any machine (no dependence on the
gitignored real corpus). Paragraph ids are deterministic by construction;
task definitions in capabilities/tasks/*.json reference them.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parent.parent


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
V_NS = "urn:schemas-microsoft-com:vml"
FOOTNOTES_REL = f"{R_NS}/footnotes"
ENDNOTES_REL = f"{R_NS}/endnotes"
FOOTNOTES_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
ENDNOTES_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"


def _set_east_asia(run, font: str) -> None:
    run.font.name = font
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def _plain(output: Path) -> None:
    document = Document()
    document.add_paragraph("本发明涉及生物医用材料技术领域。剂量为 20 mg。")  # P0
    document.add_paragraph("实施例1采用 250 mg 剂量。")  # P1 (avoid "25 mg" collision)
    document.add_paragraph("The quick brown fox.")  # P2
    document.add_paragraph("结束段落。")  # P3
    document.add_paragraph("ABC denotes the control group.")  # P4
    document.add_paragraph("重复句子内容 重复句子内容。")  # P5 (ambiguous target)
    document.save(output)


def _styled(output: Path) -> None:
    document = Document()
    p0 = document.add_paragraph()
    p0.add_run("前缀")
    r = p0.add_run("中文区域内容")
    _set_east_asia(r, "宋体")
    p0.add_run("后缀")
    p1 = document.add_paragraph()
    p1.add_run("Lead ")
    r = p1.add_run("EnglishRegion")
    r.font.name = "Times New Roman"
    p1.add_run(" tail")
    p2 = document.add_paragraph()  # two adjacent regions: cross-region target
    r = p2.add_run("第一区域")
    _set_east_asia(r, "宋体")
    r = p2.add_run("第二区域")
    r.font.name = "Times New Roman"
    p3 = document.add_paragraph()  # leading/trailing formatting spaces
    p3.add_run(" 前导空格文字  ")
    p4 = document.add_paragraph()  # CJK + English regions for batch edit
    r = p4.add_run("智能响应")
    _set_east_asia(r, "宋体")
    p4.add_run("材料与 ")
    r = p4.add_run("ABC")
    r.font.name = "Times New Roman"
    p4.add_run(" 对照组相比。")
    document.save(output)


def _add_hyperlink(document: Document, paragraph, text: str, target: str) -> str:
    relationship_id = document.part.relate_to(target, R_NS + "/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return relationship_id


def _anchors(output: Path) -> None:
    document = Document()
    p0 = document.add_paragraph()
    p0.add_run("超链接：")
    _add_hyperlink(document, p0, "点击此处", "https://example.com/docx2typed")
    p1 = document.add_paragraph()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "0")
    start.set(qn("w:name"), "bm1")
    p1._p.append(start)
    p1.add_run("关键内容")
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "0")
    p1._p.append(end)
    comment_run = p1.add_run("批注区域")
    document.add_comment(comment_run, text="锚点附近批注", author="tester")
    document.save(output)


def _table(output: Path) -> None:
    document = Document()
    document.add_paragraph("表前")
    t0 = document.add_table(rows=3, cols=3)
    for row, values in enumerate([["A", "B", "C"], ["a1", "PVA", "a3"], ["b1", "b2", "b3"]]):
        for col, value in enumerate(values):
            t0.cell(row, col).text = value
    nested = t0.cell(1, 1).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "内层"
    t1 = document.add_table(rows=3, cols=3)
    for row, values in enumerate([["X", "Y", "Z"], ["x1", "x2", "x3"], ["y1", "PVA", "y3"]]):
        for col, value in enumerate(values):
            t1.cell(row, col).text = value
    document.add_paragraph("表后")
    document.save(output)


def _boxes(output: Path) -> None:
    from lxml import etree

    document = Document()
    document.add_paragraph("正文前")
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    pict = etree.SubElement(run._r, f"{{{W_NS}}}pict", nsmap={"v": V_NS})
    shape = etree.SubElement(pict, f"{{{V_NS}}}shape")
    shape.set("style", "width:100pt;height:50pt")
    textbox = etree.SubElement(shape, f"{{{V_NS}}}textbox")
    txbx = etree.SubElement(textbox, f"{{{W_NS}}}txbxContent")
    tp = etree.SubElement(txbx, f"{{{W_NS}}}p")
    tr = etree.SubElement(tp, f"{{{W_NS}}}r")
    tt = etree.SubElement(tr, f"{{{W_NS}}}t")
    tt.text = "框内文字"
    document.add_paragraph("正文后")
    document.save(output)


def _add_reference_part(
    files: dict[str, bytes],
    *,
    marker: str,
    reference_tag: str,
    part_name: str,
    relationship_type: str,
    content_type: str,
    content: str,
) -> None:
    document_xml = files["word/document.xml"].decode("utf-8")
    text_position = document_xml.find(f">{marker}</w:t>")
    if text_position < 0:
        raise ValueError(f"missing marker: {marker}")
    paragraph_end = document_xml.find("</w:p>", text_position)
    reference = f'<w:r><w:{reference_tag} w:id="1"/></w:r>'
    document_xml = document_xml[:paragraph_end] + reference + document_xml[paragraph_end:]
    files["word/document.xml"] = document_xml.encode("utf-8")
    files[part_name] = content.encode("utf-8")
    relationships = files["word/_rels/document.xml.rels"].decode("utf-8")
    ids = [int(value) for value in re.findall(r'Id="rId(\d+)"', relationships)]
    relationship_id = f"rId{max(ids, default=0) + 1}"
    relationship = (
        f'<Relationship Id="{relationship_id}" Type="{relationship_type}" '
        f'Target="{part_name.removeprefix("word/")}"/>'
    )
    relationships = relationships.replace("</Relationships>", relationship + "</Relationships>")
    files["word/_rels/document.xml.rels"] = relationships.encode("utf-8")
    content_types = files["[Content_Types].xml"].decode("utf-8")
    override = f'<Override PartName="/{part_name}" ContentType="{content_type}"/>'
    content_types = content_types.replace("</Types>", override + "</Types>")
    files["[Content_Types].xml"] = content_types.encode("utf-8")


def _parts(output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Draft v1"
    section.footer.paragraphs[0].text = "Page"
    document.add_paragraph("正文段落")
    document.add_paragraph("FOOTNOTE-SLOT")
    document.add_paragraph("ENDNOTE-SLOT")
    document.save(output)
    with zipfile.ZipFile(output) as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    _add_reference_part(
        files,
        marker="FOOTNOTE-SLOT",
        reference_tag="footnoteReference",
        part_name="word/footnotes.xml",
        relationship_type=FOOTNOTES_REL,
        content_type=FOOTNOTES_CONTENT_TYPE,
        content=(
            f'<w:footnotes xmlns:w="{W_NS}">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/>'
            "</w:r></w:p></w:footnote>"
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
            "<w:continuationSeparator/></w:r></w:p></w:footnote>"
            '<w:footnote w:id="1"><w:p><w:r><w:t>脚注内容</w:t>'
            "</w:r></w:p></w:footnote></w:footnotes>"
        ),
    )
    _add_reference_part(
        files,
        marker="ENDNOTE-SLOT",
        reference_tag="endnoteReference",
        part_name="word/endnotes.xml",
        relationship_type=ENDNOTES_REL,
        content_type=ENDNOTES_CONTENT_TYPE,
        content=(
            f'<w:endnotes xmlns:w="{W_NS}">'
            '<w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/>'
            "</w:r></w:p></w:endnote>"
            '<w:endnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
            "<w:continuationSeparator/></w:r></w:p></w:endnote>"
            '<w:endnote w:id="1"><w:p><w:r><w:t>尾注内容</w:t>'
            "</w:r></w:p></w:endnote></w:endnotes>"
        ),
    )
    with tempfile.NamedTemporaryFile(prefix="parts-fixture-", suffix=".docx", dir=output.parent, delete=False) as temp:
        temp_path = Path(temp.name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for name, data in files.items():
                archive.writestr(name, data)
        os.replace(temp_path, output)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _revision_run(text: str, *, deleted: bool = False) -> ET.Element:
    run = ET.Element(f"{{{W_NS}}}r")
    t = ET.SubElement(run, f"{{{W_NS}}}{'delText' if deleted else 't'}")
    t.text = text
    return run


def _revision(document: Document, paragraph, tag: str, w_id: str, text: str, *, deleted: bool = False) -> None:
    rev = OxmlElement(tag)
    rev.set(qn("w:id"), w_id)
    rev.set(qn("w:author"), "审稿人")
    rev.set(qn("w:date"), "2026-08-06T10:12:00Z")
    run = OxmlElement("w:r")
    t = OxmlElement("w:delText" if deleted else "w:t")
    t.text = text
    run.append(t)
    rev.append(run)
    paragraph._p.append(rev)


def _revisions(output: Path) -> None:
    document = Document()
    p0 = document.add_paragraph("修订前文")
    _revision(document, p0, "w:ins", "1", "已插入内容")  # 1
    p0.add_run("修订后文")
    p1 = document.add_paragraph("保留")
    _revision(document, p1, "w:del", "2", "旧文本", deleted=True)  # 2
    # paragraph-mark revision: self-closing ins inside pPr>rPr (standard shape)
    p2 = document.add_paragraph("段落标记段")
    ppr = p2._p.get_or_add_pPr()
    mark_rpr = OxmlElement("w:rPr")
    mark = OxmlElement("w:ins")
    mark.set(qn("w:id"), "3")
    mark.set(qn("w:author"), "审稿人")
    mark.set(qn("w:date"), "2026-08-06T10:12:00Z")
    mark_rpr.append(mark)
    ppr.append(mark_rpr)
    # field containing an insertion revision (opaque interior)
    p3 = document.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), ' PAGE ')
    fld_ins = OxmlElement("w:ins")
    fld_ins.set(qn("w:id"), "4")
    fld_ins.set(qn("w:author"), "审稿人")
    fld_ins.set(qn("w:date"), "2026-08-06T10:12:00Z")
    fld_run = OxmlElement("w:r")
    fld_t = OxmlElement("w:t")
    fld_t.text = "字段内插入"
    fld_run.append(fld_t)
    fld_ins.append(fld_run)
    fld.append(fld_ins)
    p3._p.append(fld)
    # math containing an insertion revision (opaque interior)
    p4 = document.add_paragraph()
    math = OxmlElement("m:oMath")
    math_ins = OxmlElement("w:ins")
    math_ins.set(qn("w:id"), "5")
    math_ins.set(qn("w:author"), "审稿人")
    math_ins.set(qn("w:date"), "2026-08-06T10:12:00Z")
    m_run = OxmlElement("w:r")
    m_t = OxmlElement("w:t")
    m_t.text = "公式内插入"
    m_run.append(m_t)
    math_ins.append(m_run)
    math.append(math_ins)
    p4._p.append(math)
    p5 = document.add_paragraph()
    _revision(document, p5, "w:ins", "6", "修订五")  # 6
    p6 = document.add_paragraph()
    _revision(document, p6, "w:del", "7", "修订六", deleted=True)  # 7
    p7 = document.add_paragraph()
    _revision(document, p7, "w:ins", "8", "修订七")  # 8
    document.save(output)
    with zipfile.ZipFile(output) as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    settings = files["word/settings.xml"]
    settings = re.sub(rb"(<w:settings[^>]*>)", rb"\1<w:trackChanges/>", settings, count=1)
    files["word/settings.xml"] = settings
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)


def _comments(output: Path) -> None:
    document = Document()
    for i, (target, text) in enumerate(
        [("关键一", "批注一内容"), ("关键二", "批注二内容"), ("关键三", "批注三内容")],
        start=1,
    ):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(target)
        document.add_comment(run, text=text, author="审稿人")
    document.save(output)


def _norm(output: Path) -> None:
    document = Document()
    document.add_paragraph("水H₂O 温度25°C Ca²⁺ 与 H₂SO₄ 反应。")
    document.add_paragraph("平方公式 a² + b² = c²。")
    document.save(output)


def _large(output: Path) -> None:
    document = Document()
    for index in range(3000):
        if index == 1500:
            document.add_paragraph("大文档锚点词")
        else:
            document.add_paragraph(f"段落 {index} 内容。")
    table = document.add_table(rows=10, cols=10)
    for row in range(10):
        for col in range(10):
            table.cell(row, col).text = f"L{row}{col}"
    document.save(output)


def _review(output: Path) -> None:
    """Comments + pending revisions + trackChanges off: the ambiguous-signal
    document used by the comment-preservation acceptance case."""
    document = Document()
    for target, text in [("关键一", "批注一内容"), ("关键二", "批注二内容"), ("关键三", "批注三内容")]:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(target)
        document.add_comment(run, text=text, author="审稿人")
    _revision(document, document.paragraphs[0], "w:ins", "21", "修订甲")
    _revision(document, document.paragraphs[1], "w:del", "22", "修订乙", deleted=True)
    document.save(output)


W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
W16_NS = "http://schemas.microsoft.com/office/word/2018/wordml"
W16DU_NS = "http://schemas.microsoft.com/office/word/2020/wordml/word16du"


def _move_revision(paragraph, tag: str, w_id: str, text: str) -> None:
    rev = OxmlElement(tag)
    rev.set(qn("w:id"), w_id)
    rev.set(qn("w:author"), "审稿人")
    rev.set(qn("w:date"), "2026-08-06T10:12:00Z")
    run = OxmlElement("w:r")
    t = OxmlElement("w:delText" if tag == "w:moveFrom" else "w:t")
    t.text = text
    run.append(t)
    rev.append(run)
    paragraph._p.append(rev)


def _move_conflict(output: Path) -> None:
    """Run/paragraph/move/conflict revision dialect: move_from/move_to
    containers plus overlapping (conflict-shaped) revision pairs."""
    document = Document()
    p0 = document.add_paragraph("移动前文")
    _move_revision(p0, "w:moveFrom", "31", "移动前")
    _move_revision(p0, "w:moveTo", "32", "移动后")
    p0.add_run("移动后文")
    p1 = document.add_paragraph()
    outer = OxmlElement("w:ins")
    outer.set(qn("w:id"), "33")
    outer.set(qn("w:author"), "审稿人甲")
    outer.set(qn("w:date"), "2026-08-06T10:12:00Z")
    inner = OxmlElement("w:ins")
    inner.set(qn("w:id"), "34")
    inner.set(qn("w:author"), "审稿人乙")
    inner.set(qn("w:date"), "2026-08-07T09:00:00Z")
    inner_run = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = "冲突文本"
    inner_run.append(inner_t)
    inner.append(inner_run)
    outer.append(inner)
    outer_run = OxmlElement("w:r")
    outer_t = OxmlElement("w:t")
    outer_t.text = "冲突外"
    outer_run.append(outer_t)
    outer.append(outer_run)
    p1._p.append(outer)
    p2 = document.add_paragraph("替换对：")
    _revision(document, p2, "w:del", "35", "旧文本", deleted=True)
    _revision(document, p2, "w:ins", "36", "新文本")
    document.save(output)
    with zipfile.ZipFile(output) as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    settings = files["word/settings.xml"]
    settings = re.sub(rb"(<w:settings[^>]*>)", rb"\1<w:trackChanges/>", settings, count=1)
    files["word/settings.xml"] = settings
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)


COMMENTS_EXTENSIBLE_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    f'<w15:commentsEx xmlns:w15="{W15_NS}">'
    '<w15:commentEx w15:paraId="1A2B3C4D" w15:done="0" w15:commentId="0"/>'
    "</w15:commentsEx>"
)
PEOPLE_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    f'<w15:people xmlns:w15="{W15_NS}">'
    '<w15:person w15:author="审稿人">'
    '<w15:presenceInfo w15:providerId="A0000000-0000-0000-0000-000000000000" '
    'w15:userId="reviewer-1"/>'
    "</w15:person>"
    "</w15:people>"
)
COMMENTS_EXTENSIBLE_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtensible+xml"
)
PEOPLE_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.people+xml"
)
COMMENTS_EXTENSIBLE_REL = "http://schemas.microsoft.com/office/2006/relationships/commentsExtensible"
PEOPLE_REL = "http://schemas.microsoft.com/office/2006/relationships/people"


def _modern_comments(output: Path) -> None:
    """Minimal public modern-comments fixture synthesized from the pinned
    MS-DOCX w15 spec (commentsExtensible + people package graph).  The
    licensed M365-generated private reference remains pending-not-run."""
    document = Document()
    p0 = document.add_paragraph()
    run = p0.add_run("现代批注目标")
    document.add_comment(run, text="现代批注内容", author="审稿人")
    document.save(output)
    with zipfile.ZipFile(output) as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    doc_xml = files["word/document.xml"].decode("utf-8")
    doc_xml = doc_xml.replace(
        '<w:document ',
        f'<w:document xmlns:w15="{W15_NS}" ',
        1,
    )
    doc_xml = doc_xml.replace(
        '<w:commentRangeStart w:id="0"/>',
        '<w:commentRangeStart w:id="0" w15:paraId="1A2B3C4D"/>',
    )
    doc_xml = doc_xml.replace(
        '<w:commentReference w:id="0"/>',
        '<w:commentReference w:id="0" w15:paraId="1A2B3C4D"/>',
    )
    files["word/document.xml"] = doc_xml.encode("utf-8")
    files["word/commentsExtensible.xml"] = COMMENTS_EXTENSIBLE_XML.encode("utf-8")
    files["word/people.xml"] = PEOPLE_XML.encode("utf-8")
    relationships = files["word/_rels/document.xml.rels"].decode("utf-8")
    ids = [int(value) for value in re.findall(r'Id="rId(\d+)"', relationships)]
    next_id = max(ids, default=0) + 1
    for part, rel_type, target in (
        ("commentsExtensible", COMMENTS_EXTENSIBLE_REL, "commentsExtensible.xml"),
        ("people", PEOPLE_REL, "people.xml"),
    ):
        relationship = (
            f'<Relationship Id="rId{next_id}" Type="{rel_type}" Target="{target}"/>'
        )
        relationships = relationships.replace("</Relationships>", relationship + "</Relationships>")
        next_id += 1
    files["word/_rels/document.xml.rels"] = relationships.encode("utf-8")
    content_types = files["[Content_Types].xml"].decode("utf-8")
    for part, content_type in (
        ("commentsExtensible.xml", COMMENTS_EXTENSIBLE_CONTENT_TYPE),
        ("people.xml", PEOPLE_CONTENT_TYPE),
    ):
        override = f'<Override PartName="/word/{part}" ContentType="{content_type}"/>'
        content_types = content_types.replace("</Types>", override + "</Types>")
    files["[Content_Types].xml"] = content_types.encode("utf-8")
    with tempfile.NamedTemporaryFile(prefix="modern-comments-", suffix=".docx", dir=output.parent, delete=False) as temp:
        temp_path = Path(temp.name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for name, data in files.items():
                archive.writestr(name, data)
        os.replace(temp_path, output)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _dialect_ns(output: Path) -> None:
    """Namespace-dialect probe: declares w14/w15/w16/w16du and pins one
    attribute from each (w14:paraId, w15:paraId, w16:docId,
    w16du:conflictMode) on a normal document."""
    document = Document()
    document.add_paragraph("命名空间方言探测段落")
    document.save(output)
    with zipfile.ZipFile(output) as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    doc_xml = files["word/document.xml"].decode("utf-8")
    doc_xml = doc_xml.replace(
        "<w:body>",
        (
            f'<w:body xmlns:w14="{W14_NS}" xmlns:w15="{W15_NS}" '
            f'xmlns:w16="{W16_NS}" xmlns:w16du="{W16DU_NS}">'
        ),
        1,
    )
    doc_xml = doc_xml.replace(
        "<w:pPr>",
        '<w:pPr w14:paraId="3F2E1D0C" w15:paraId="4E5F6071">',
        1,
    )
    files["word/document.xml"] = doc_xml.encode("utf-8")
    settings = files["word/settings.xml"].decode("utf-8")
    settings = settings.replace(
        "<w:settings ",
        (
            f'<w:settings xmlns:w14="{W14_NS}" xmlns:w15="{W15_NS}" '
            f'xmlns:w16="{W16_NS}" xmlns:w16du="{W16DU_NS}" '
        ),
        1,
    )
    settings = settings.replace(
        "</w:settings>",
        (
            '<w16:docId w16:val="{D1A2B3C4-0000-0000-0000-000000000000}"/>'
            '<w16du:conflictMode w16du:val="1"/>'
            "</w:settings>"
        ),
        1,
    )
    files["word/settings.xml"] = settings.encode("utf-8")
    with tempfile.NamedTemporaryFile(prefix="dialect-ns-", suffix=".docx", dir=output.parent, delete=False) as temp:
        temp_path = Path(temp.name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for name, data in files.items():
                archive.writestr(name, data)
        os.replace(temp_path, output)
    finally:
        if temp_path.exists():
            temp_path.unlink()


BUILDERS = {
    "plain.docx": _plain,
    "styled.docx": _styled,
    "anchors.docx": _anchors,
    "table.docx": _table,
    "boxes.docx": _boxes,
    "parts.docx": _parts,
    "revisions.docx": _revisions,
    "comments.docx": _comments,
    "review.docx": _review,
    "norm.docx": _norm,
    "large.docx": _large,
    "modern-comments.docx": _modern_comments,
    "dialect-ns.docx": _dialect_ns,
    "move-conflict.docx": _move_conflict,
}


_FIXED_TIME = (2026, 8, 8, 0, 0, 0)


def _canonicalize(path: Path) -> None:
    """Rewrite a generated docx with fixed zip timestamps and fixed core
    properties so regeneration is byte-stable across platforms and days."""
    import io

    with zipfile.ZipFile(path) as archive:
        entries = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    core = entries.get("docProps/core.xml", b"")
    if core:
        core = re.sub(
            rb"<dcterms:created[^>]*>.*?</dcterms:created>",
            f'<dcterms:created xsi:type="dcterms:W3CDTF">{_FIXED_TIME[0]:04d}-{_FIXED_TIME[1]:02d}-{_FIXED_TIME[2]:02d}T00:00:00Z</dcterms:created>'.encode(),
            core,
        )
        core = re.sub(
            rb"<dcterms:modified[^>]*>.*?</dcterms:modified>",
            f'<dcterms:modified xsi:type="dcterms:W3CDTF">{_FIXED_TIME[0]:04d}-{_FIXED_TIME[1]:02d}-{_FIXED_TIME[2]:02d}T00:00:00Z</dcterms:modified>'.encode(),
            core,
        )
        entries["docProps/core.xml"] = core
    for name in list(entries):
        if name.startswith("word/"):
            entries[name] = re.sub(
                rb'w:date="[^"]*"',
                f'w:date="{_FIXED_TIME[0]:04d}-{_FIXED_TIME[1]:02d}-{_FIXED_TIME[2]:02d}T00:00:00Z"'.encode(),
                entries[name],
            )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in sorted(entries.items()):
            info = zipfile.ZipInfo(name, date_time=_FIXED_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(info, data)


def generate(outdir: str | Path) -> Path:
    out = Path(outdir).resolve()
    out.mkdir(parents=True, exist_ok=True)
    from scripts.create_complex_fixture import create_fixture

    for name, builder in BUILDERS.items():
        target = out / name
        builder(target)
        _canonicalize(target)
    complex_target = out / "complex.docx"
    create_fixture(complex_target)
    _canonicalize(complex_target)
    return out


# ---------------------------------------------------------------------------
# Seeded-corruption calibration corpus (corpus/calibration)
# ---------------------------------------------------------------------------
# Every calibration fixture is deterministic (fixed zip timestamps, fixed
# bytes).  The just-inside / just-over pairs target the S profile limits from
# qualification/resource_profiles.json; L/X pairs are generated at runtime by
# scripts/resource_limits.py (documented in corpus/manifest.json coverage).

CALIBRATION_INSIDE_S = {
    "uncompressed_xml_bytes": 100 * 1024 * 1024,  # S: 100 MiB (inclusive)
    "nesting_depth": 512,  # S: 512 levels (inclusive)
    "text_node_bytes": 1024 * 1024,  # S: 1 MiB leaf limit (inclusive)
    "zip_parts": 2000,  # S: 2000 parts (inclusive)
}


def _cal_base_files() -> dict[str, bytes]:
    """Minimal deterministic package skeleton for calibration fixtures."""
    return {
        "[Content_Types].xml": (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Default Extension="bin" ContentType="application/vnd.openxmlformats-officedocument.oleObject"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>"
        ).encode("utf-8"),
        "_rels/.rels": (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/>'
            "</Relationships>"
        ).encode("utf-8"),
        "word/_rels/document.xml.rels": (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
        ).encode("utf-8"),
        "word/document.xml": (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            f'<w:document xmlns:w="{W_NS}"><w:body>'
            "<w:p><w:r><w:t>calibration</w:t></w:r></w:p>"
            "</w:body></w:document>"
        ).encode("utf-8"),
        "docProps/core.xml": (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
            'xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" '
            'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
            '<dcterms:created xsi:type="dcterms:W3CDTF">2026-08-08T00:00:00Z</dcterms:created>'
            '<dcterms:modified xsi:type="dcterms:W3CDTF">2026-08-08T00:00:00Z</dcterms:modified>'
            "</cp:coreProperties>"
        ).encode("utf-8"),
        "docProps/app.xml": (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties">'
            "<Application>docx2typed-calibration</Application>"
            "</Properties>"
        ).encode("utf-8"),
    }


def _cal_write(output: Path, files: dict[str, bytes]) -> None:
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in sorted(files.items()):
            info = zipfile.ZipInfo(name, date_time=_FIXED_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(info, data)


def _cal_invalid_zip(output: Path) -> None:
    """A zip that starts like one but is truncated/garbled."""
    raw = _cal_base_files()["word/document.xml"] * 4
    output.write_bytes(b"PK\x03\x04" + raw[: len(raw) // 2] + b"\x00\xff garbage")


def _cal_non_zip_bytes(output: Path) -> None:
    output.write_bytes(b"not a zip file " * 64)


def _cal_malformed_document_xml(output: Path) -> None:
    files = _cal_base_files()
    doc = files["word/document.xml"]
    files["word/document.xml"] = doc[: max(1, len(doc) // 2)]
    _cal_write(output, files)


def _cal_missing_document_xml(output: Path) -> None:
    files = _cal_base_files()
    del files["word/document.xml"]
    del files["word/_rels/document.xml.rels"]
    _cal_write(output, files)


def _cal_high_compression(output: Path, *, over: bool) -> None:
    """Uncompressed XML at/over the S 100 MiB limit, split into parts under
    the 1 MiB leaf limit; the over variant also carries duplicate entry
    names (high-compression + duplicate probe)."""
    files = _cal_base_files()
    target = CALIBRATION_INSIDE_S["uncompressed_xml_bytes"] + (1 if over else 0)
    header = f'<w:p xmlns:w="{W_NS}"><w:r><w:t xml:space="preserve">'
    footer = "</w:t></w:r></w:p>"
    overhead = len(header) + len(footer)
    base_doc = len(files["word/document.xml"])
    chunk_target = (1024 * 1024 * 9) // 10  # 90% of the S leaf limit
    payload_total = max(0, target - base_doc - overhead)
    count = max(1, (payload_total + chunk_target - 1) // chunk_target)
    payload_total = max(0, target - base_doc - count * overhead)
    base_payload = payload_total // count
    sizes = [base_payload] * count
    sizes[-1] += payload_total - base_payload * count
    for index, size in enumerate(sizes):
        files[f"word/t{index:05d}.xml"] = (
            header + (" " * size) + footer
        ).encode("utf-8")
    if over:
        # Duplicate entry: same member name written twice (intentional; the
        # harness gate rejects duplicate entries).  Suppress the zipfile
        # UserWarning so regeneration stays quiet and deterministic.
        import warnings

        dup = files["word/t00000.xml"]
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
                for name, data in sorted(files.items()):
                    info = zipfile.ZipInfo(name, date_time=_FIXED_TIME)
                    info.compress_type = zipfile.ZIP_DEFLATED
                    archive.writestr(info, data)
                    if name == "word/t00000.xml":
                        archive.writestr(info, dup)  # second identical member
        return
    _cal_write(output, files)


def _cal_deep_nesting(output: Path, *, over: bool) -> None:
    """Element nesting depth at/over the S limit (512 levels); the
    document/body/p wrappers add 3 levels and the r/t leaves 2 more."""
    files = _cal_base_files()
    depth = CALIBRATION_INSIDE_S["nesting_depth"] - 5 + (1 if over else 0)
    inner = "<w:r><w:t>depth</w:t></w:r>"
    for _ in range(depth):
        inner = f"<w:ins>{inner}</w:ins>"
    files["word/document.xml"] = (
        f'<w:document xmlns:w="{W_NS}"><w:body><w:p>{inner}</w:p></w:body></w:document>'
    ).encode("utf-8")
    _cal_write(output, files)


def _cal_oversized_text(output: Path, *, over: bool) -> None:
    """Single text node at/over the S 1 MiB leaf limit."""
    files = _cal_base_files()
    limit = CALIBRATION_INSIDE_S["text_node_bytes"]
    size = limit + 1 if over else limit - 512
    files["word/document.xml"] = (
        f'<w:document xmlns:w="{W_NS}"><w:body><w:p><w:r>'
        f"<w:t xml:space=\"preserve\">{'x' * size}</w:t>"
        "</w:r></w:p></w:body></w:document>"
    ).encode("utf-8")
    _cal_write(output, files)


def _cal_many_parts(output: Path, *, over: bool) -> None:
    """Total ZIP part count at/over the S limit (2000), including base
    parts."""
    files = _cal_base_files()
    target = CALIBRATION_INSIDE_S["zip_parts"] + (1 if over else 0)
    base = len(files)
    index = 0
    while len(files) < target:
        files[f"word/media/t{index:05d}.bin"] = f"tiny-part-{index:05d}".encode("utf-8")
        index += 1
    _cal_write(output, files)


def _cal_relationship_cycle(output: Path) -> None:
    """Relationship cycle (word/document.xml -> word/a.xml -> back) plus a
    relationship targeting a missing part."""
    files = _cal_base_files()
    files["word/a.xml"] = (
        f'<w:document xmlns:w="{W_NS}"><w:body><w:p><w:r><w:t>a</w:t></w:r></w:p></w:body></w:document>'
    ).encode("utf-8")
    files["word/b.xml"] = (
        f'<w:document xmlns:w="{W_NS}"><w:body><w:p><w:r><w:t>b</w:t></w:r></w:p></w:body></w:document>'
    ).encode("utf-8")
    files["word/_rels/document.xml.rels"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" '
        'Target="a.xml"/>'
        '<Relationship Id="rId2" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" '
        'Target="missing.xml"/>'
        "</Relationships>"
    ).encode("utf-8")
    files["word/_rels/a.xml.rels"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" '
        'Target="../word/document.xml"/>'
        "</Relationships>"
    ).encode("utf-8")
    _cal_write(output, files)


CALIBRATION_BUILDERS = {
    "cal-invalid-zip.docx": (lambda out: _cal_invalid_zip(out)),
    "cal-non-zip-bytes.docx": (lambda out: _cal_non_zip_bytes(out)),
    "cal-malformed-document-xml.docx": (lambda out: _cal_malformed_document_xml(out)),
    "cal-missing-document-xml.docx": (lambda out: _cal_missing_document_xml(out)),
    "cal-high-compression-inside-S.docx": (lambda out: _cal_high_compression(out, over=False)),
    "cal-high-compression-over-S.docx": (lambda out: _cal_high_compression(out, over=True)),
    "cal-deep-nesting-inside-S.docx": (lambda out: _cal_deep_nesting(out, over=False)),
    "cal-deep-nesting-over-S.docx": (lambda out: _cal_deep_nesting(out, over=True)),
    "cal-oversized-text-inside-S.docx": (lambda out: _cal_oversized_text(out, over=False)),
    "cal-oversized-text-over-S.docx": (lambda out: _cal_oversized_text(out, over=True)),
    "cal-many-parts-inside-S.docx": (lambda out: _cal_many_parts(out, over=False)),
    "cal-many-parts-over-S.docx": (lambda out: _cal_many_parts(out, over=True)),
    "cal-relationship-cycle.docx": (lambda out: _cal_relationship_cycle(out)),
}


def generate_calibration(outdir: str | Path) -> Path:
    out = Path(outdir).resolve()
    out.mkdir(parents=True, exist_ok=True)
    for name, builder in CALIBRATION_BUILDERS.items():
        builder(out / name)
    return out


def _model_manifest_path(release_dir: str | Path) -> Path:
    """The typed-model manifest lives next to the release fixtures it
    describes, so regeneration into a scratch outdir never touches the
    committed corpus (with the default ``corpus/release`` outdir this is
    the committed file)."""
    return Path(release_dir).resolve() / "model-manifest.json"


def _model_hash(docx_path: Path, work_root: Path) -> str:
    """Hash of the extracted typed model (typed.md + paragraph records +
    styles) — rsid attributes and environment noise are normalized away."""
    import subprocess

    workdir = work_root / docx_path.stem
    subprocess.run(
        ["python", "-m", "scripts", "extract", str(docx_path), "-o", str(workdir)],
        cwd=REPO_ROOT, check=True, capture_output=True,
    )
    hasher = hashlib.sha256()
    for name in ("typed.md", "styles.json"):
        hasher.update((workdir / name).read_bytes())
    fmt = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    records = json.dumps(fmt.get("paragraphs", []), ensure_ascii=False, sort_keys=True)
    hasher.update(records.encode("utf-8"))
    return hasher.hexdigest()[:16]


def check_models(release_dir: str | Path, work_root: str | Path, *, write: bool = False) -> int:
    """Compare the extracted typed models of the fixtures against the
    committed manifest (or write it with --write)."""
    release = Path(release_dir).resolve()
    work = Path(work_root).resolve()
    work.mkdir(parents=True, exist_ok=True)
    manifest_path = _model_manifest_path(release)
    current = {
        path.name: _model_hash(path, work)
        for path in sorted(release.glob("*.docx"))
    }
    if write:
        manifest_path.write_text(
            json.dumps({"schema": "docx2typed-fixture-model-1", "fixtures": current}, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        print(f"model manifest written: {manifest_path}")
        return 0
    committed = json.loads(manifest_path.read_text(encoding="utf-8"))["fixtures"]
    mismatches = {name: (committed.get(name), current[name]) for name in current if committed.get(name) != current[name]}
    if mismatches:
        for name, (expected, got) in mismatches.items():
            print(f"MISMATCH {name}: expected {expected} got {got}")
        return 1
    print(f"model manifest OK ({len(current)} fixtures)")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate corpus/release fixtures")
    parser.add_argument("--outdir", default="corpus/release")
    parser.add_argument("--calibration-dir", default="corpus/calibration")
    parser.add_argument("--with-calibration", action="store_true", help="also regenerate corpus/calibration fixtures")
    parser.add_argument("--check-models", action="store_true", help="verify regenerated fixtures' typed models against the committed manifest")
    parser.add_argument("--write-models", action="store_true", help="write the typed-model manifest for the current fixtures")
    parser.add_argument("--write-manifest", action="store_true", help="write corpus/manifest.json (schema docx2typed-fixture-manifest-2)")
    parser.add_argument("--manifest-path", default="corpus/manifest.json", help="corpus manifest output path (default: corpus/manifest.json)")
    parser.add_argument("--work", default="/tmp/fixture-models")
    args = parser.parse_args()
    if args.check_models or args.write_models:
        return check_models(args.outdir, args.work, write=args.write_models)
    out = generate(args.outdir)
    print(f"generated {len(BUILDERS) + 1} fixtures in {out}")
    if args.with_calibration:
        cal = generate_calibration(args.calibration_dir)
        print(f"generated {len(CALIBRATION_BUILDERS)} calibration fixtures in {cal}")
    if args.write_manifest:
        manifest_path = write_manifest(
            args.outdir,
            args.calibration_dir,
            args.with_calibration or None,
            manifest_path=args.manifest_path,
        )
        print(f"manifest written: {manifest_path}")
    return 0


# ---------------------------------------------------------------------------
# Fixture corpus manifest (corpus/manifest.json, docx2typed-fixture-manifest-2)
# ---------------------------------------------------------------------------

FIXTURE_DIALECT_FEATURES: dict[str, list[str]] = {
    "plain.docx": ["iso-classic", "body-paragraphs", "cjk", "plain-runs"],
    "styled.docx": ["iso-classic", "cjk-fonts-rtl", "style-spans", "east-asia-fonts", "leading-trailing-spaces"],
    "anchors.docx": ["iso-classic", "hyperlink", "bookmarks", "classic-comments"],
    "table.docx": ["iso-classic", "nested-tables", "table-cells"],
    "boxes.docx": ["iso-classic", "text-boxes", "vml-drawing"],
    "parts.docx": ["iso-classic", "headers-footers", "footnotes-endnotes"],
    "revisions.docx": ["iso-classic", "run-revisions", "paragraph-revisions", "opaque-fields-math-drawings", "track-changes-on"],
    "comments.docx": ["iso-classic", "classic-comments"],
    "review.docx": ["iso-classic", "classic-comments", "run-revisions", "ambiguous-review-state"],
    "norm.docx": ["iso-classic", "cjk-fonts-rtl", "unicode-vertical"],
    "large.docx": ["iso-classic", "large-packages", "pathological-packages"],
    "complex.docx": ["iso-classic", "sdt-content-controls", "opaque-fields-math-drawings", "sections", "images", "custom-xml", "tables"],
    "modern-comments.docx": ["modern-comments", "w14-w15-w16-w16du", "commentsExtensible", "people"],
    "dialect-ns.docx": ["w14-w15-w16-w16du"],
    "move-conflict.docx": ["iso-classic", "run-revisions", "move-revisions", "conflict-revisions"],
}

CALIBRATION_CORRUPTION: dict[str, str] = {
    "cal-invalid-zip.docx": "invalid-zip",
    "cal-non-zip-bytes.docx": "non-zip-bytes",
    "cal-malformed-document-xml.docx": "malformed-document-xml",
    "cal-missing-document-xml.docx": "missing-document-xml",
    "cal-high-compression-inside-S.docx": "high-compression-duplicate",
    "cal-high-compression-over-S.docx": "high-compression-duplicate",
    "cal-deep-nesting-inside-S.docx": "deep-xml-nesting",
    "cal-deep-nesting-over-S.docx": "deep-xml-nesting",
    "cal-oversized-text-inside-S.docx": "oversized-text-node",
    "cal-oversized-text-over-S.docx": "oversized-text-node",
    "cal-many-parts-inside-S.docx": "many-tiny-parts",
    "cal-many-parts-over-S.docx": "many-tiny-parts",
    "cal-relationship-cycle.docx": "relationship-cycle",
}

FULL_ELIGIBLE = [
    "cli.extract",
    "cli.build",
    "cli.validate",
    "cli.verify",
    "mcp.*",
    "office.open",
    "office.render",
    "office.save",
    "office.reopen",
    "office.retention",
]
CALIBRATION_ELIGIBLE = ["resource.qualification", "office.open"]


def _toolchain() -> dict[str, str]:
    """Frozen generator-toolchain lock (issue #42).

    Static by design: runtime-probed python/lxml versions differ across
    platforms and would break byte-identical regeneration. The lock names
    the generator family and the determinism contract; exact probe versions
    are recorded in evidence, not in the reproducible manifest."""
    return {
        "generator": "scripts/release_fixtures.py",
        "toolchain": "python-docx 1.2 + lxml, deterministic builders",
        "zip": "ZIP_DEFLATED, fixed 2026-08-08 timestamps",
        "reproducibility": "byte-identical across platforms and days",
    }


def _manifest_entry(
    rel_path: str,
    *,
    root: str | Path,
    tier: str,
    features: list[str],
    eligible: list[str],
    model_hash: str | None,
    provenance_note: str,
    corruption_type: str | None = None,
) -> dict[str, object]:
    from scripts.qualify_adapters import file_sha256

    path = (Path(root).resolve() / rel_path)
    signature: dict[str, object] = {"package_sha256": file_sha256(path)}
    if model_hash:
        signature["typed_model_sha256_prefix"] = model_hash
    dialect: dict[str, object] = {"features": features}
    if corruption_type:
        dialect["corruption_type"] = corruption_type
    stable = hashlib.sha256(rel_path.encode("utf-8")).hexdigest()
    prefix = "cal" if tier == "calibration" else "fx"
    anonymous_id = f"{prefix}-{int(stable[:8], 16) % 1000000:06d}"
    return {
        "anonymous_id": anonymous_id,
        "tier": tier,
        "path": rel_path,
        "sha256": file_sha256(path),
        "size_bytes": path.stat().st_size,
        "provenance": {
            "kind": "synthetic" if tier != "private" else "private-pending",
            "generator": provenance_note,
            "license": "CC0-1.0" if tier == "public" else ("unlicensed (deterministic corruption bytes)" if tier == "calibration" else "private-license-pending"),
            "consent": "n/a synthetic" if tier != "private" else "written owner authorization pending",
            "origin_url": None,
        },
        "toolchain": _toolchain(),
        "dialect": dialect,
        "expected_signatures": signature,
        "eligible_cells": eligible,
        "privacy": {
            "class": f"{tier}-synthetic" if tier != "private" else "private-real",
            "owner": "LLLin000",
            "retention": "release-eol-plus-2y" if tier != "private" else "per-authorization",
        },
        "mutation_lineage": [
            "baseline: deterministic synthetic generation by scripts/release_fixtures.py",
        ],
    }


def write_manifest(
    release_dir: str | Path,
    calibration_dir: str | Path,
    with_calibration: bool | None = None,
    manifest_path: str | Path = "corpus/manifest.json",
    corpus_root: str | Path | None = None,
) -> Path:
    """Regenerate the corpus manifest from the release fixtures.

    Deterministic: anonymous ids, hashes and sizes derive from the fixture
    bytes; only the generated date varies (excluded from the semantic hash
    pin by the qualification runner's canonicalization).

    ``corpus_root`` is where the ``release/`` and ``calibration/`` fixtures
    live (default: the committed ``corpus/`` dir).  Regeneration into a
    scratch outdir passes the scratch dir so the manifest hashes the
    regenerated bytes — a drifted committed fixture then shows up as a
    byte-identity mismatch instead of being re-hashed away.
    """
    from scripts.fixture_manifest import build_manifest

    release = Path(release_dir).resolve()
    calibration = Path(calibration_dir).resolve()
    corpus_root = Path(corpus_root).resolve() if corpus_root is not None else (REPO_ROOT / "corpus").resolve()
    model_manifest = json.loads(_model_manifest_path(release).read_text(encoding="utf-8"))
    model_hashes = model_manifest.get("fixtures", {})
    if with_calibration is None:
        with_calibration = any(calibration.glob("cal-*.docx"))
    fixtures: list[dict[str, object]] = []
    for name, features in sorted(FIXTURE_DIALECT_FEATURES.items()):
        if not (release / name).is_file():
            continue
        fixtures.append(
            _manifest_entry(
                f"release/{name}",
                root=corpus_root,
                tier="public",
                features=features,
                eligible=FULL_ELIGIBLE,
                model_hash=model_hashes.get(name),
                provenance_note=f"scripts/release_fixtures.py::{name.replace('.docx', '').replace('-', '_')}",
            )
        )
    if with_calibration:
        for name, corruption in sorted(CALIBRATION_CORRUPTION.items()):
            if not (calibration / name).is_file():
                continue
            fixtures.append(
                _manifest_entry(
                    f"calibration/{name}",
                    root=corpus_root,
                    tier="calibration",
                    features=[corruption],
                    eligible=CALIBRATION_ELIGIBLE,
                    model_hash=None,
                    provenance_note=f"scripts/release_fixtures.py::calibration::{corruption}",
                    corruption_type=corruption,
                )
            )
    dialect_inventory: dict[str, list[str]] = {}
    for name, features in sorted(FIXTURE_DIALECT_FEATURES.items()):
        if not (release / name).is_file():
            continue
        for feature in features:
            dialect_inventory.setdefault(feature, []).append(f"release/{name}")
    calibration_inventory: dict[str, list[str]] = {}
    if with_calibration:
        for name, corruption in sorted(CALIBRATION_CORRUPTION.items()):
            if (calibration / name).is_file():
                calibration_inventory.setdefault(corruption, []).append(f"calibration/{name}")
    coverage = {
        "dialect_inventory": dialect_inventory,
        "calibration_inventory": calibration_inventory,
        "notes": (
            "Just-inside/just-over pairs are committed at the S profile; "
            "L/X pairs are generated deterministically at runtime by "
            "scripts/resource_limits.py and enforced by the resource-profiles "
            "qualification check."
        ),
    }
    private_pending = [
        {
            "anonymous_id": "pvt-00001",
            "description": "licensed M365-generated modern-comments reference (calibrates package graph + roundtrip signature)",
            "status": "pending-not-run",
            "reason": "licensed M365 host/fixture not supplied on this runner",
        }
    ]
    manifest = build_manifest(
        fixtures,
        coverage=coverage,
        generated="2026-08-13",
        private_pending=private_pending,
    )
    target = Path(manifest_path).resolve()
    target.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return target


def _manifest_diff(relative: str, committed: Path, regenerated: Path) -> str:
    """Compact unified diff between the committed and regenerated files."""
    import difflib

    a = committed.read_text(encoding="utf-8").splitlines()
    b = regenerated.read_text(encoding="utf-8").splitlines()
    lines = list(
        difflib.unified_diff(
            a,
            b,
            fromfile=f"committed {relative}",
            tofile=f"regenerated {relative}",
            lineterm="",
            n=1,
        )
    )
    head = "\n".join(lines[:24])
    if len(lines) > 24:
        head += f"\n… {len(lines) - 24} more diff lines"
    return f"{relative} differs:\n{head}"


def verify_corpus(root: str | Path, work_root: str | Path, scratch: str | Path) -> dict[str, object]:
    """Regenerate release fixtures, calibration fixtures, the typed-model
    manifest, and the corpus manifest into ``scratch`` and assert the two
    manifests are byte-identical to the committed ones.

    Runs the same regeneration pipeline as the release_fixtures entrypoint
    (generate -> generate_calibration -> check_models -> write_manifest).
    A mismatch is corpus drift (stale committed fixtures/manifest or a
    nondeterministic generator) and fails the gate with the diff — never a
    pass.
    """
    root = Path(root).resolve()
    scratch = Path(scratch).resolve()
    release = scratch / "release"
    calibration = scratch / "calibration"
    work = Path(work_root).resolve() / "models"
    generate(release)
    generate_calibration(calibration)
    check_models(release, work, write=True)
    manifest_path = write_manifest(
        release,
        calibration,
        with_calibration=True,
        manifest_path=scratch / "manifest.json",
        corpus_root=scratch,
    )
    pairs = [
        ("corpus/manifest.json", manifest_path),
        ("corpus/release/model-manifest.json", release / "model-manifest.json"),
    ]
    diffs: list[str] = []
    for relative, regenerated in pairs:
        committed = root / relative
        if not committed.is_file():
            diffs.append(f"{relative}: committed file missing")
        elif committed.read_bytes() != regenerated.read_bytes():
            diffs.append(_manifest_diff(relative, committed, regenerated))
    ok = not diffs
    detail = (
        "corpus manifest and model manifest byte-identical after regeneration"
        if ok
        else "regeneration differs from committed: " + "; ".join(d.splitlines()[0] for d in diffs)
    )
    return {"ok": ok, "diffs": diffs, "detail": detail}


if __name__ == "__main__":
    raise SystemExit(main())

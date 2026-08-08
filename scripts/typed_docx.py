"""DOCX extraction, typed workdirs, byte patching, and independent verify."""
from __future__ import annotations

from dataclasses import dataclass, field
import argparse
from difflib import SequenceMatcher
import hashlib
import json
import os
from pathlib import Path
import re
import shutil
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from typing import Any, Iterable

try:
    from .typed_core import (
        NS_R,
        NS_W,
        AnchorNode,
        InlineNode,
        OpaqueNode,
        Paragraph,
        RangeNode,
        RevisionNode,
        StyleRegistry,
        TextNode,
        TypedDocument,
        TypedError,
        choose_base_style,
        canonical_xml,
        content_signature,
        contains_opaque,
        element_end_xml,
        element_start_xml,
        etree_xml,
        local_name,
        merge_adjacent_text,
        parse_typed,
        qname,
        serialize_typed,
        skeleton,
        style_id_for_rpr,
        visible_text,
        visible_text_original,
        w,
        xml_escape,
    )
    from .xml_walker import TagCursor, find_element_range, find_open_tag_end, iter_tags
except ImportError:
    from typed_core import (
        NS_R,
        NS_W,
        AnchorNode,
        InlineNode,
        OpaqueNode,
        Paragraph,
        RangeNode,
        RevisionNode,
        StyleRegistry,
        TextNode,
        TypedDocument,
        TypedError,
        choose_base_style,
        canonical_xml,
        content_signature,
        contains_opaque,
        element_end_xml,
        element_start_xml,
        etree_xml,
        local_name,
        merge_adjacent_text,
        parse_typed,
        qname,
        serialize_typed,
        skeleton,
        style_id_for_rpr,
        visible_text,
        visible_text_original,
        w,
        xml_escape,
    )
    from xml_walker import TagCursor, find_element_range, find_open_tag_end, iter_tags


class ValidationError(TypedError):
    """A workdir cannot be safely built."""


@dataclass
class ParagraphSlice:
    index: int
    start: int
    end: int
    raw: bytes
    container_path: tuple[str, int, ...] = ()
    table_index: int = -1


@dataclass
class TableSlice:
    table_index: int
    start: int
    end: int
    cells: list[ParagraphSlice]
    body_level: bool = False


@dataclass
class SdtSlice:
    sdt_index: int
    start: int
    end: int
    cells: list[ParagraphSlice]


@dataclass
class BoxSlice:
    box_index: int
    start: int
    end: int
    cells: list[ParagraphSlice]
    parent_paragraph: int = -1  # body paragraph slice index containing the box
    ordinal_in_parent: int = 0  # box ordinal within that paragraph


@dataclass
class DocumentSlices:
    xml: bytes
    body_start: int
    body_end: int
    paragraphs: list[ParagraphSlice]
    tables: list[TableSlice] = field(default_factory=list)
    boxes: list[BoxSlice] = field(default_factory=list)
    sdts: list[SdtSlice] = field(default_factory=list)


@dataclass
class ParsedDocx:
    document: TypedDocument
    styles: StyleRegistry
    tokens: dict[str, dict[str, Any]]
    slices: DocumentSlices
    token_table: Any = None


@dataclass
class ValidatedWorkdir:
    path: Path
    format_data: dict[str, Any]
    styles: StyleRegistry
    typed: TypedDocument
    baseline: TypedDocument
    baseline_tokens: dict[str, dict[str, Any]]
    template_path: Path
    template_xml: bytes
    template_slices: DocumentSlices
    live_paragraphs: list[Paragraph]
    baseline_by_id: dict[str, Paragraph]
    warnings: list[str]


_P_OPEN_RE = re.compile(r"^<(?:[A-Za-z_][\w.-]*:)?p(?:\s[^>]*?)?/?>")
_P_ATTR_RE = re.compile(
    r'\s+(?P<name>[A-Za-z_][\w.-]*(?::[A-Za-z_][\w.-]*)?)\s*=\s*(?P<value>"[^"]*"|\'[^\']*\')'
)
_EPHEMERAL_P_ATTRS = {
    "paraId",
    "textId",
    "rsidP",
    "rsidR",
    "rsidRDefault",
    "rsidDel",
    "rsidRPr",
}




def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def sha256_file(path: Path) -> str:
    return sha256_bytes(path.read_bytes())
def json_bytes(data: Any) -> bytes:
    return (json.dumps(data, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode("utf-8")


def zip_manifest(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {name: sha256_bytes(archive.read(name)) for name in sorted(archive.namelist())}


def locate_document_xml(xml: bytes) -> DocumentSlices:
    """Locate ``w:body`` direct paragraphs plus table cell paragraphs without
    rewriting XML.

    Body paragraphs keep their flat index; cell paragraphs (inside
    ``w:tbl > w:tr > w:tc``, including nested tables) are collected per table
    with a container path ``("tbl", t, "tr", r, "tc", c, "p", p)`` where
    ``t`` is the global table ordinal and ``r``/``c``/``p`` are ordinals
    within their parent.
    """
    stack: list[tuple[str, int, int]] = []  # (name, start, ordinal_in_parent)
    ordinals: dict[tuple[int, int, str], int] = {}  # (parent_depth, parent_start, name) -> next ordinal
    body_depth: int | None = None
    body_start = -1
    body_end = -1
    paragraphs: list[ParagraphSlice] = []
    tables: list[TableSlice] = []
    tables_by_index: dict[int, TableSlice] = {}
    table_ordinal = -1
    boxes: list[BoxSlice] = []
    boxes_by_index: dict[int, BoxSlice] = {}
    box_ordinal = -1
    sdts: list[SdtSlice] = []
    sdts_by_index: dict[int, SdtSlice] = {}
    sdt_ordinal = -1
    open_sdts: list[tuple[int, int]] = []  # (sdt_index, start)
    paragraph_starts: list[tuple[int, int, tuple[str, int, ...], int]] = []
    open_tables: list[tuple[int, int]] = []  # (table_index, start)
    open_boxes: list[tuple[int, int]] = []  # (box_index, start)

    for tag in iter_tags(xml):
        closing, self_closing = tag.closing, tag.self_closing
        name = tag.name
        if closing:
            if not stack or stack[-1][0] != name:
                raise ValidationError(f"malformed document XML nesting near {tag.raw_name}")
            depth = len(stack)
            if (
                name == "p"
                and body_depth is not None
                and depth == body_depth + 2
            ):
                start = stack[-1][1]
                paragraph_starts.append((start, tag.end, (), -1))
            elif (
                name == "p"
                and body_depth is not None
                and depth > body_depth + 1
                and (open_tables or open_boxes or open_sdts)
            ):
                start = stack[-1][1]
                if open_boxes and _is_box_paragraph_stack(stack, body_depth):
                    path = _box_path(stack, body_depth, open_boxes[-1][0])
                    paragraph_starts.append((start, tag.end, path, -1))
                elif open_sdts and _is_sdt_paragraph_stack(stack, body_depth):
                    path = _sdt_path(stack, body_depth, open_sdts[-1][0])
                    paragraph_starts.append((start, tag.end, path, -1))
                elif open_tables and _is_cell_paragraph_stack(stack, body_depth):
                    t_index = open_tables[-1][0]
                    path = _cell_path(stack, body_depth)
                    paragraph_starts.append((start, tag.end, path, t_index))
            if name == "tbl" and open_tables and open_tables[-1][1] == stack[-1][1]:
                tables[open_tables[-1][0]].end = tag.end
                open_tables.pop()
            if name == "txbxContent" and open_boxes and open_boxes[-1][1] == stack[-1][1]:
                boxes[open_boxes[-1][0]].end = tag.end
                open_boxes.pop()
            if name == "sdtContent" and open_sdts and open_sdts[-1][1] == stack[-1][1]:
                sdts[open_sdts[-1][0]].end = tag.end
                open_sdts.pop()
            if name == "body" and body_depth is not None and depth == body_depth + 1:
                body_end = tag.start
            stack.pop()
            continue
        depth = len(stack)
        if name == "body" and body_depth is None:
            body_depth = depth
            body_start = tag.end
        parent_key = (depth - 1, stack[-1][1] if stack else -1, name)
        ordinal = ordinals.get(parent_key, 0)
        ordinals[parent_key] = ordinal + 1
        if body_depth is not None:
            if name == "p" and depth == body_depth + 1:
                if self_closing:
                    paragraph_starts.append((tag.start, tag.end, (), -1))
                else:
                    stack.append((name, tag.start, ordinal))
                continue
            if name == "sdtContent" and depth > body_depth + 1 and _is_sdt_content_stack(stack, body_depth):
                sdt_ordinal += 1
                sdt = SdtSlice(sdt_ordinal, tag.start, -1, [])
                sdts.append(sdt)
                sdts_by_index[sdt_ordinal] = sdt
                open_sdts.append((sdt_ordinal, tag.start))
            if name == "txbxContent" and depth > body_depth + 1 and _in_body_paragraph(stack, body_depth):
                box_ordinal += 1
                parent = _containing_body_paragraph(stack, body_depth)
                ordinal_in_parent = sum(1 for existing in boxes if existing.parent_paragraph == parent)
                box = BoxSlice(box_ordinal, tag.start, -1, [], parent, ordinal_in_parent)
                boxes.append(box)
                boxes_by_index[box_ordinal] = box
                open_boxes.append((box_ordinal, tag.start))
            if name == "tbl" and depth == body_depth + 1:
                table_ordinal += 1
                table = TableSlice(table_ordinal, tag.start, -1, [], body_level=True)
                tables.append(table)
                tables_by_index[table_ordinal] = table
                open_tables.append((table_ordinal, tag.start))
                stack.append((name, tag.start, ordinal))
                continue
            if name == "tbl" and depth > body_depth + 1 and stack and stack[-1][0] == "tc":
                # nested table inside a cell
                table_ordinal += 1
                table = TableSlice(table_ordinal, tag.start, -1, [])
                tables.append(table)
                tables_by_index[table_ordinal] = table
                open_tables.append((table_ordinal, tag.start))
        if not self_closing:
            stack.append((name, tag.start, ordinal))
        if self_closing and name == "body":
            body_end = tag.start
    if stack:
        raise ValidationError("document XML has unclosed elements")
    if open_tables:
        raise ValidationError("document XML has unclosed tables")
    if body_start < 0 or body_end < body_start:
        raise ValidationError("document XML has no direct w:body")
    for index, (start, end, path, t_index) in enumerate(paragraph_starts):
        paragraph = ParagraphSlice(index, start, end, xml[start:end], path, t_index)
        if t_index < 0:
            if path and path[0] == "box":
                boxes_by_index[path[1]].cells.append(paragraph)
            elif path and path[0] == "sdt":
                sdts_by_index[path[1]].cells.append(paragraph)
            else:
                paragraphs.append(paragraph)
        else:
            tables_by_index[t_index].cells.append(paragraph)
    if open_boxes:
        raise ValidationError("document XML has unclosed text boxes")
    if open_sdts:
        raise ValidationError("document XML has unclosed sdt content controls")
    return DocumentSlices(xml, body_start, body_end, paragraphs, tables, boxes, sdts)


def _in_body_paragraph(stack: list[tuple[str, int, int]], body_depth: int) -> bool:
    """Whether the open elements include a direct body paragraph."""
    return any(name == "p" and depth == body_depth + 1 for depth, (name, _, _) in enumerate(stack))


def _containing_body_paragraph(stack: list[tuple[str, int, int]], body_depth: int) -> int:
    """Slice index of the direct body paragraph containing the box (body
    paragraph ordinals are contiguous 0-based within body)."""
    for depth, (name, _, ordinal) in enumerate(stack):
        if name == "p" and depth == body_depth + 1:
            return ordinal
    return -1


def _is_sdt_content_stack(stack: list[tuple[str, int, int]], body_depth: int) -> bool:
    """Whether the next element opens a body-level sdtContent: the open chain
    is body > sdt > (sdtPr | sdtEndPr)*."""
    chain = [name for name, _, _ in stack[body_depth + 1:]]
    if not chain or chain[0] != "sdt":
        return False
    return set(chain[1:]) <= {"sdtPr", "sdtEndPr"}


def _is_sdt_paragraph_stack(stack: list[tuple[str, int, int]], body_depth: int) -> bool:
    """Whether the closing p is a direct child of sdtContent at body level."""
    chain = [name for name, _, _ in stack[body_depth + 1:]]
    if not chain or chain[-1] != "p":
        return False
    return len(chain) == 3 and chain[0] == "sdt" and chain[1] == "sdtContent"


def _sdt_path(stack: list[tuple[str, int, int]], body_depth: int, sdt_index: int) -> tuple[str, int, ...]:
    p_ordinal = 0
    for name, _, ordinal in stack[body_depth + 1:]:
        if name == "p":
            p_ordinal = ordinal
    return ("sdt", sdt_index, "p", p_ordinal)


def _is_box_paragraph_stack(stack: list[tuple[str, int, int]], body_depth: int) -> bool:
    """Whether the chain above the last container is txbxContent > p (the box
    paragraph itself is a direct child of txbxContent; nested tables inside
    boxes are not extracted in v1)."""
    chain = [name for name, _, _ in stack[body_depth + 1:]]
    if not chain or chain[-1] != "p":
        return False
    # the paragraph must be a DIRECT child of txbxContent: no intermediate
    # container (like tbl/tr/tc) between it and the box
    return "txbxContent" in chain and not any(
        name in ("tbl", "tr", "tc") for name in chain
    )


def _box_path(stack: list[tuple[str, int, int]], body_depth: int, box_index: int) -> tuple[str, int, ...]:
    p_ordinal = 0
    for name, _, ordinal in stack[body_depth + 1:]:
        if name == "p":
            p_ordinal = ordinal
    return ("box", box_index, "p", p_ordinal)


def _is_cell_paragraph_stack(stack: list[tuple[str, int, int]], body_depth: int) -> bool:
    """Whether the open-element stack above body is a tbl>tr>tc>…>p chain
    (nested tables allowed)."""
    chain = [name for name, _, _ in stack[body_depth + 1:]]
    if not chain or chain[-1] != "p":
        return False
    if len(chain) < 4:
        return False
    # walk from the paragraph upward; every level must be tbl/tr/tc/p
    # in the pattern ..., tbl, tr, tc, p (repeating for nested tables)
    depth_ok = True
    for level, name in enumerate(reversed(chain[:-1])):
        if level % 3 == 0:
            depth_ok = depth_ok and name == "tc"
        elif level % 3 == 1:
            depth_ok = depth_ok and name == "tr"
        else:
            depth_ok = depth_ok and name == "tbl"
        if not depth_ok:
            return False
    return True


def _cell_path(stack: list[tuple[str, int, int]], body_depth: int) -> tuple[str, int, ...]:
    """Container path for a cell paragraph: per-parent ordinals of every
    container on the chain (tbl/tr/tc/p, nested tables included) — this is
    the navigation path, not the display id."""
    path: list[str | int] = []
    for name, _, ordinal in stack[body_depth + 1:]:
        path.append(name)
        path.append(ordinal)
    return tuple(path)


def _raw_p_parts(raw: bytes) -> tuple[str, str]:
    text = raw.decode("utf-8")
    opening = _P_OPEN_RE.match(text)
    if not opening:
        raise ValidationError("direct paragraph has no recognizable w:p opening")
    p_open = opening.group(0)
    if p_open.endswith("/>"):
        # A touched paragraph must render as a paired element, never self-closing.
        p_open = p_open[:-2] + ">"
    # pPr ranges can nest (w:pPrChange carries a w:pPr); the walker returns
    # the OUTERMOST first element (self-closing pPr is skipped, matching the
    # historical behavior of the tag-depth scan).
    rng = find_element_range(raw, "pPr")
    if rng is not None:
        return p_open, raw[rng[0]:rng[1]].decode("utf-8")
    return p_open, ""


def _attrs(element: ET.Element) -> dict[str, str]:
    return {qname(key): value for key, value in element.attrib.items()}


def _token_ids(nodes: Iterable[Any]) -> list[list[str]]:
    values: list[list[str]] = []
    for node in nodes:
        if isinstance(node, (RangeNode, RevisionNode)):
            values.append([node.token_id, node.kind])
            values.extend(_token_ids(node.children))
        elif isinstance(node, (AnchorNode, InlineNode, OpaqueNode)):
            values.append([node.token_id, node.kind])
    return values


def _assign_default_style(nodes: Iterable[Any], style_id: str) -> None:
    for node in nodes:
        if isinstance(node, TextNode) and not node.style_id:
            node.style_id = style_id
        elif isinstance(node, (RangeNode, RevisionNode)):
            _assign_default_style(node.children, style_id)


def _contains_structural(nodes: Iterable[Any]) -> bool:
    return any(not isinstance(node, TextNode) for node in nodes)


def _iter_anchor_nodes(nodes: Iterable[Any]) -> Iterable[AnchorNode]:
    for node in nodes:
        if isinstance(node, AnchorNode):
            yield node
        elif isinstance(node, (RangeNode, RevisionNode)):
            yield from _iter_anchor_nodes(node.children)


def _validate_anchor_pairs(paragraphs: Iterable[Paragraph], scope: str) -> None:
    pairs = {
        "bookmark": ("bookmark-start", "bookmark-end"),
        "comment": ("comment-start", "comment-end"),
    }
    positions: dict[str, dict[str, list[tuple[str, int]]]] = {
        kind: {"start": [], "end": []} for kind in pairs
    }
    for paragraph_index, paragraph in enumerate(paragraphs):
        for node in _iter_anchor_nodes(paragraph.nodes):
            for family, (start_kind, end_kind) in pairs.items():
                if node.kind == start_kind:
                    anchor_id = node.attrs.get("w:id") or node.attrs.get("id")
                    if not anchor_id:
                        raise ValidationError(f"{scope} {family} start has no ID")
                    positions[family]["start"].append((anchor_id, paragraph_index))
                elif node.kind == end_kind:
                    anchor_id = node.attrs.get("w:id") or node.attrs.get("id")
                    if not anchor_id:
                        raise ValidationError(f"{scope} {family} end has no ID")
                    positions[family]["end"].append((anchor_id, paragraph_index))
    for family, sides in positions.items():
        starts = sides["start"]
        ends = sides["end"]
        start_ids = [anchor_id for anchor_id, _ in starts]
        end_ids = [anchor_id for anchor_id, _ in ends]
        if sorted(start_ids) != sorted(end_ids):
            raise ValidationError(f"{scope} {family} anchors are not paired")
        if len(start_ids) != len(set(start_ids)):
            raise ValidationError(f"{scope} {family} anchor IDs are duplicated")
        for anchor_id, start_paragraph in starts:
            end_paragraph = next(index for candidate, index in ends if candidate == anchor_id)
            if start_paragraph > end_paragraph:
                raise ValidationError(f"{scope} {family} anchor range is reversed: {anchor_id}")


class _TokenTable:
    def __init__(self) -> None:
        self.records: dict[str, dict[str, Any]] = {}
        self.next_id = 0

    def add(self, kind: str, **record: Any) -> str:
        token_id = f"N{self.next_id}"
        self.next_id += 1
        self.records[token_id] = {"kind": kind, **record}
        return token_id


def _empty_rpr() -> str:
    return f'<w:rPr xmlns:w="{NS_W}"/>'


def _parse_run(element: ET.Element, styles: StyleRegistry, tokens: _TokenTable) -> list[Any]:
    rpr = next((child for child in element if local_name(child.tag) == "rPr"), None)
    rpr_xml = etree_xml(rpr) if rpr is not None else _empty_rpr()
    style_id = styles.ensure(rpr_xml)
    output: list[Any] = []
    known_inline = {
        "t", "tab", "br", "cr", "noBreakHyphen", "softHyphen", "sym",
        "commentReference", "footnoteRef", "endnoteRef", "annotationRef",
        "separator", "continuationSeparator", "lastRenderedPageBreak",
    }
    format_change = None
    if rpr is not None:
        for child in rpr:
            if local_name(child.tag) == "rPrChange":
                format_change = child
                break
    for child in element:
        name = local_name(child.tag)
        if name == "rPr":
            continue
        if name == "t" or name == "delText":
            output.append(TextNode(style_id, child.text or ""))
        elif name in known_inline:
            token_id = tokens.add(
                name if name != "cr" else "cr",
                raw=etree_xml(child),
                attrs=_attrs(child),
                style_id=style_id,
            )
            output.append(InlineNode(token_id, name if name != "cr" else "cr", style_id, _attrs(child)))
        else:
            token_id = tokens.add(
                "unsupported-run",
                raw=etree_xml(element),
                attrs={"tag": qname(child.tag)},
                style_id=style_id,
            )
            return [OpaqueNode(token_id, "unsupported-run", {"tag": qname(child.tag)})]
    if format_change is not None:
        token_id = tokens.add(
            "rpr-change", raw=etree_xml(format_change),
            attrs={"tag": "w:rPrChange"}, style_id=style_id,
        )
        output.append(InlineNode(token_id, "rpr-change", style_id, {"tag": "w:rPrChange"}))
    if not output:
        token_id = tokens.add("empty-run", raw=etree_xml(element), attrs={}, style_id=style_id)
        return [InlineNode(token_id, "empty-run", style_id, {})]
    return merge_adjacent_text(output)


def _parse_container(children: Iterable[ET.Element], styles: StyleRegistry, tokens: _TokenTable) -> list[Any]:
    output: list[Any] = []
    for element in children:
        name = local_name(element.tag)
        if name in {"pPr", "proofErr"}:
            if name == "proofErr":
                token_id = tokens.add("opaque", raw=etree_xml(element), attrs={"tag": qname(element.tag)})
                output.append(OpaqueNode(token_id, "opaque", {"tag": qname(element.tag)}))
            continue
        if name == "r":
            output.extend(_parse_run(element, styles, tokens))
            continue
        if name == "hyperlink":
            token_id = tokens.add(
                "hyperlink",
                open=element_start_xml(element),
                close=element_end_xml(element),
                attrs=_attrs(element),
            )
            output.append(RangeNode(token_id, "hyperlink", _attrs(element), _parse_container(list(element), styles, tokens)))
            continue
        if name in {"ins", "del", "moveFrom", "moveTo"}:
            kind = {
                "ins": "insert",
                "del": "delete",
                "moveFrom": "move_from",
                "moveTo": "move_to",
            }[name]
            token_id = tokens.add(
                "revision",
                open=element_start_xml(element),
                close=element_end_xml(element),
                attrs=_attrs(element),
            )
            output.append(
                RevisionNode(
                    token_id,
                    kind,
                    _attrs(element),
                    _parse_container(list(element), styles, tokens),
                )
            )
            continue
        if name in {"bookmarkStart", "bookmarkEnd", "commentRangeStart", "commentRangeEnd"}:
            kind = {
                "bookmarkStart": "bookmark-start",
                "bookmarkEnd": "bookmark-end",
                "commentRangeStart": "comment-start",
                "commentRangeEnd": "comment-end",
            }[name]
            attrs = _attrs(element)
            token_id = tokens.add(kind, raw=etree_xml(element), attrs=attrs)
            output.append(AnchorNode(token_id, kind, attrs))
            continue
        token_id = tokens.add("opaque", raw=etree_xml(element), attrs={"tag": qname(element.tag)})
        output.append(OpaqueNode(token_id, "opaque", {"tag": qname(element.tag)}))
    return merge_adjacent_text(output)


def _parse_paragraph(
    element: ET.Element,
    raw: bytes,
    index: int,
    styles: StyleRegistry,
    tokens: _TokenTable,
    *,
    paragraph_id: str | None = None,
    container_path: tuple[str, int, ...] = (),
    original_index: int | None = None,
) -> Paragraph:
    p_open, ppr_xml = _raw_p_parts(raw)
    children = _parse_container(list(element), styles, tokens)
    section_bearing = "sectPr" in ppr_xml
    mark_revision = _parse_mark_revision(ppr_xml, tokens)
    if paragraph_id is None:
        paragraph_id = f"P{index}"
    if original_index is None:
        original_index = index
    paragraph = Paragraph(
        paragraph_id=paragraph_id,
        base_style="",
        nodes=children,
        p_open=p_open,
        ppr=ppr_xml,
        raw_xml=raw.decode("utf-8"),
        section_bearing=section_bearing,
        editable=not contains_opaque(children),
        original_index=original_index,
        mark_revision=mark_revision,
    )
    if container_path:
        paragraph.container_path = container_path
    return paragraph


_MARK_TAG_RE = re.compile(
    r"<w:(ins|del)(?:\s+[^>]*)?/>"
)


def _parse_mark_revision(ppr_xml: str, tokens: _TokenTable) -> dict[str, Any] | None:
    """Paragraph-mark revision: a self-closing w:ins/w:del inside pPr/rPr.

    Only the mark itself is recorded; the surrounding rPr stays with the
    template bytes. Returns {"kind", "token_id", "attrs"} or None.
    """
    if not ppr_xml:
        return None
    match = _MARK_TAG_RE.search(ppr_xml)
    if match is None:
        return None
    kind = {"ins": "insert", "del": "delete"}[match.group(1)]
    attrs = _parse_attrs_xml(match.group(0))
    token_id = tokens.add(
        "paragraph-mark",
        raw=match.group(0),
        attrs=attrs,
    )
    return {"kind": kind, "token_id": token_id, "attrs": attrs}


def _parse_attrs_xml(tag_xml: str) -> dict[str, str]:
    """Attribute map of a self-contained XML tag.

    The fragment carries no xmlns declarations, so it is parsed inside a
    wrapper that declares every python-docx namespace prefix (plus w16du,
    which python-docx does not ship); attribute names are normalized to
    qname form.
    """
    from docx.oxml.ns import nsmap
    from .typed_core import NS_W16DU

    declarations = " ".join(
        f'xmlns:{prefix}="{uri}"'
        for prefix, uri in {**nsmap, "w16du": NS_W16DU}.items()
    )
    wrapper = f"<docx2typed-root {declarations}>{tag_xml}</docx2typed-root>"
    root = ET.fromstring(wrapper)
    return {qname(key): value for key, value in root[0].attrib.items()}


def _inject_mark_revision(ppr: str, mark: dict[str, Any], tokens: dict[str, dict[str, Any]]) -> str:
    """Inject the paragraph-mark revision XML into pPr's rPr (Word places the
    pilcrow revision on the paragraph mark run properties)."""
    record = tokens.get(mark["token_id"])
    mark_xml = str(record.get("raw", "")) if record and record.get("raw") else ""
    if not mark_xml:
        raise ValidationError(f"missing paragraph-mark XML: {mark['token_id']}")
    rpr_match = re.search(r"<w:rPr(?:\s+[^>]*)?>", ppr)
    if rpr_match:
        return ppr[: rpr_match.end()] + mark_xml + ppr[rpr_match.end():]
    sect_match = re.search(r"<w:sectPr", ppr)
    if sect_match:
        # OOXML CT_PPr order: rPr must precede sectPr.
        return ppr[: sect_match.start()] + f"<w:rPr>{mark_xml}</w:rPr>" + ppr[sect_match.start():]
    close_match = re.search(r"</w:pPr>", ppr)
    if close_match:
        return ppr[: close_match.start()] + f"<w:rPr>{mark_xml}</w:rPr>" + ppr[close_match.start():]
    if not ppr:
        return f"<w:pPr><w:rPr>{mark_xml}</w:rPr></w:pPr>"
    return ppr + f"<w:rPr>{mark_xml}</w:rPr>"


def parse_document_xml(xml: bytes, *, styles: StyleRegistry | None = None) -> ParsedDocx:
    slices = locate_document_xml(xml)
    try:
        root = ET.fromstring(xml)
    except ET.ParseError as exc:
        raise ValidationError(f"invalid document XML: {exc}") from exc
    body = next((child for child in root.iter() if local_name(child.tag) == "body"), None)
    if body is None:
        raise ValidationError("document XML has no body element")
    registry = styles or StyleRegistry()
    tokens = _TokenTable()
    paragraphs: list[Paragraph] = []
    slice_index = 0
    cell_slices: dict[tuple[str, int, ...], ParagraphSlice] = {
        cell.container_path: cell for table in slices.tables for cell in table.cells
    }
    for child in list(body):
        if local_name(child.tag) != "p":
            continue
        if slice_index >= len(slices.paragraphs):
            raise ValidationError("XML paragraph locator disagrees with parsed body")
        paragraph = _parse_paragraph(child, slices.paragraphs[slice_index].raw, slice_index, registry, tokens)
        paragraphs.append(paragraph)
        slice_index += 1
    if slice_index != len(slices.paragraphs):
        raise ValidationError("XML paragraph locator missed a direct body paragraph")
    table_paragraphs: list[Paragraph] = []
    scan_index_by_id: dict[str, int] = {}
    for index, paragraph in enumerate(paragraphs):
        scan_index_by_id[paragraph.paragraph_id] = slices.paragraphs[index].index
    for sdt in slices.sdts:
        for cell in sdt.cells:
            element = _find_sdt_paragraph(body, cell.container_path, sdt)
            if element is None:
                raise ValidationError(f"sdt paragraph locator disagrees with parsed body: {cell.container_path}")
            paragraph_id = _sdt_paragraph_id(cell.container_path)
            paragraph = _parse_paragraph(
                element,
                cell.raw,
                -1,
                registry,
                tokens,
                paragraph_id=paragraph_id,
                container_path=cell.container_path,
                original_index=-1,
            )
            scan_index_by_id[paragraph_id] = cell.index
            table_paragraphs.append(paragraph)
    for box in slices.boxes:
        for cell in box.cells:
            element = _find_box_paragraph(body, cell.container_path, box)
            if element is None:
                raise ValidationError(f"box paragraph locator disagrees with parsed body: {cell.container_path}")
            paragraph_id = _box_paragraph_id(cell.container_path)
            paragraph = _parse_paragraph(
                element,
                cell.raw,
                -1,
                registry,
                tokens,
                paragraph_id=paragraph_id,
                container_path=cell.container_path,
                original_index=-1,
            )
            scan_index_by_id[paragraph_id] = cell.index
            table_paragraphs.append(paragraph)
    for table in slices.tables:
        for cell in table.cells:
            element = _find_element_by_path(body, cell.container_path)
            if element is None:
                raise ValidationError(f"cell paragraph locator disagrees with parsed body: {cell.container_path}")
            paragraph_id = _cell_paragraph_id(cell.container_path, cell.table_index)
            paragraph = _parse_paragraph(
                element,
                cell.raw,
                -1,
                registry,
                tokens,
                paragraph_id=paragraph_id,
                container_path=cell.container_path,
                original_index=-1,
            )
            paragraph.table_index = cell.table_index
            scan_index_by_id[paragraph_id] = cell.index
            table_paragraphs.append(paragraph)
    # document order: body and cell paragraphs interleave by locator scan order
    paragraphs = sorted(
        paragraphs + table_paragraphs,
        key=lambda p: scan_index_by_id[p.paragraph_id],
    )
    _attach_freestanding_anchors(xml, paragraphs, slices, tokens)
    normal_style = registry.ensure(_empty_rpr(), label="Normal")
    previous_style = normal_style
    for paragraph in paragraphs:
        paragraph.base_style = choose_base_style(paragraph.nodes, previous_style)
        if paragraph.nodes and paragraph.base_style:
            previous_style = paragraph.base_style
    return ParsedDocx(TypedDocument({"schema": "1"}, paragraphs), registry, tokens.records, slices, tokens)


def _cell_paragraph_id(path: tuple[str, int, ...], table_index: int) -> str:
    values = dict(zip(path[::2], path[1::2]))
    return f"T{table_index}.R{values['tr']}.C{values['tc']}.P{values['p']}"


_FREESTANDING_ANCHOR_NAMES = {
    "bookmarkStart": "bookmark-start",
    "bookmarkEnd": "bookmark-end",
    "commentRangeStart": "comment-start",
    "commentRangeEnd": "comment-end",
}


def _attach_freestanding_anchors(
    xml: bytes,
    paragraphs: list[Paragraph],
    slices: DocumentSlices,
    tokens: _TokenTable,
) -> None:
    """Attach anchors that sit between paragraphs (Word places bookmark and
    comment range ends after the closing w:p or between table cells) to the
    paragraph they follow, keeping anchor pairing intact."""
    ranges: list[tuple[int, int, Paragraph]] = []
    for slice_ in slices.paragraphs:
        ranges.append((slice_.start, slice_.end, paragraphs[slice_.index] if slice_.index < len(paragraphs) else None))
    for table in slices.tables:
        for cell in table.cells:
            from .typed_core import RevisionNode  # noqa: F401
            match = next(
                (p for p in paragraphs if p.paragraph_id == _cell_paragraph_id(cell.container_path, cell.table_index)),
                None,
            )
            ranges.append((cell.start, cell.end, match))
    for box in slices.boxes:
        for cell in box.cells:
            match = next(
                (p for p in paragraphs if p.paragraph_id == _box_paragraph_id(cell.container_path)),
                None,
            )
            ranges.append((cell.start, cell.end, match))
    for sdt in slices.sdts:
        for cell in sdt.cells:
            match = next(
                (p for p in paragraphs if p.paragraph_id == _sdt_paragraph_id(cell.container_path)),
                None,
            )
            ranges.append((cell.start, cell.end, match))
    ranges.sort(key=lambda item: item[0])
    for tag in iter_tags(xml):
        name = tag.name
        kind = _FREESTANDING_ANCHOR_NAMES.get(name)
        if kind is None:
            continue
        position = tag.start
        if any(start <= position < end for start, end, _ in ranges):
            continue  # inside a paragraph: parsed with its content
        # freestanding: attach to the nearest preceding paragraph
        target: Paragraph | None = None
        for start, end, paragraph in ranges:
            if end <= position and paragraph is not None:
                target = paragraph
            else:
                break
        if target is None:
            continue
        token = tag.bytes_in(xml)
        attrs = _parse_attrs_xml(token.decode("utf-8"))
        token_id = tokens.add(kind, raw=token.decode("utf-8"), attrs=attrs)
        target.nodes.append(AnchorNode(token_id, kind, attrs))


def _sdt_paragraph_id(path: tuple[str, int, ...]) -> str:
    values = dict(zip(path[::2], path[1::2]))
    return f"S{values['sdt']}.P{values['p']}"


def _find_sdt_paragraph(body: ET.Element, path: tuple[str, int, ...], sdt: SdtSlice) -> ET.Element | None:
    values = dict(zip(path[::2], path[1::2]))
    sdts = [child for child in list(body) if local_name(child.tag) == "sdt"]
    if sdt.sdt_index >= len(sdts):
        return None
    content = next(
        (child for child in list(sdts[sdt.sdt_index]) if local_name(child.tag) == "sdtContent"),
        None,
    )
    if content is None:
        return None
    paras = [child for child in list(content) if local_name(child.tag) == "p"]
    if values["p"] >= len(paras):
        return None
    return paras[values["p"]]


def _box_paragraph_id(path: tuple[str, int, ...]) -> str:
    values = dict(zip(path[::2], path[1::2]))
    return f"B{values['box']}.P{values['p']}"


def _find_box_paragraph(body: ET.Element, path: tuple[str, int, ...], box: "BoxSlice") -> ET.Element | None:
    """Navigate to a box paragraph: the parent body paragraph, then the box's
    txbxContent (ordinal within that paragraph), then the p ordinal."""
    values = dict(zip(path[::2], path[1::2]))
    body_paras = [child for child in list(body) if local_name(child.tag) == "p"]
    if box.parent_paragraph < 0 or box.parent_paragraph >= len(body_paras):
        return None
    para = body_paras[box.parent_paragraph]
    boxes_in_para = [
        descendant for descendant in para.iter()
        if local_name(descendant.tag) == "txbxContent"
    ]
    if box.ordinal_in_parent >= len(boxes_in_para):
        return None
    content = boxes_in_para[box.ordinal_in_parent]
    paras = [child for child in list(content) if local_name(child.tag) == "p"]
    if values["p"] >= len(paras):
        return None
    return paras[values["p"]]


def _find_element_by_path(body: ET.Element, path: tuple[str, int, ...]) -> ET.Element | None:
    """Navigate to a cell paragraph by container path (locator ordinals)."""
    target: ET.Element | None = body
    kind_index = 0
    while kind_index < len(path):
        name = path[kind_index]
        ordinal = path[kind_index + 1]
        matches = [child for child in list(target) if local_name(child.tag) == name]
        if ordinal >= len(matches):
            return None
        target = matches[ordinal]
        kind_index += 2
    return target


def _format_token_ids(tokens: dict[str, dict[str, Any]]) -> dict[str, dict[str, Any]]:
    return {key: value for key, value in sorted(tokens.items())}


def _paragraph_insertion_style(paragraph: Paragraph) -> str:
    """Recorded typing context for an empty paragraph (contract: paragraph-mark
    ``w:rPr``, then the first text run's style, then the base style)."""
    if paragraph.ppr:
        try:
            root = ET.fromstring(paragraph.ppr)
        except ET.ParseError:
            root = None
        if root is not None:
            for child in root:
                if local_name(child.tag) == "rPr":
                    return style_id_for_rpr(etree_xml(child))
    for node in paragraph.nodes:
        if isinstance(node, TextNode):
            return node.style_id
        if isinstance(node, (RangeNode, RevisionNode)):
            for child in node.children:
                if isinstance(child, TextNode):
                    return child.style_id
    return paragraph.base_style


def _relative_source_path(source_path: Path, output_dir: Path) -> str:
    try:
        return os.path.relpath(source_path, output_dir)
    except ValueError:  # Windows: source and workdir on different drives
        return str(source_path)


def parse_package_document(archive: zipfile.ZipFile) -> ParsedDocx:
    """Parse document.xml plus editable parts (headers/footers/footnotes/
    endnotes) into paragraphs in extract order: header/footer parts first,
    body, then footnotes/endnotes."""
    document_xml = archive.read("word/document.xml")
    part_xmls = {
        match.group(1): archive.read(name)
        for name in archive.namelist()
        if (match := PART_KEYS_PATTERN.match(name))
    }
    parsed = parse_document_xml(document_xml)
    part_paragraphs: list[Paragraph] = []
    for part_key in sorted(
        part_xmls,
        key=lambda key: (0 if key.startswith("header") else 1 if key.startswith("footer") else 2 if key in ("footnotes", "endnotes") else 3, key),
    ):
        part_paragraphs.extend(
            parse_part_xml(
                part_xmls[part_key], part_key,
                styles=parsed.styles, tokens=parsed.token_table,
            )
        )
    parsed.document.paragraphs = (
        [p for p in part_paragraphs if p.part_key.startswith("header")]
        + parsed.document.paragraphs
        + [p for p in part_paragraphs if p.part_key.startswith("footer")]
        + [p for p in part_paragraphs if p.part_key in ("footnotes", "endnotes")]
        + [p for p in part_paragraphs if p.part_key == "comments"]
    )
    return parsed


def extract_workdir(source: str | Path, outdir: str | Path) -> Path:
    source_path = Path(source).resolve()
    output_dir = Path(outdir).resolve()
    if not source_path.exists():
        raise ValidationError(f"file not found: {source_path}")
    output_dir.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(source_path) as archive:
            parsed = parse_package_document(archive)
            document_xml = archive.read("word/document.xml")
            part_xmls = {
                match.group(1): archive.read(name)
                for name in archive.namelist()
                if (match := PART_KEYS_PATTERN.match(name))
            }
    except (zipfile.BadZipFile, KeyError) as exc:
        raise ValidationError(f"not a valid DOCX: {source_path}") from exc
    template_path = output_dir / "_template.docx"
    shutil.copy2(source_path, template_path)
    styles_path = output_dir / "styles.json"
    styles_path.write_bytes(json_bytes(parsed.styles.to_json()))
    format_data: dict[str, Any] = {
        "schema": "typed-format-1",
        "model_version": 1,
        "canonicalizer_version": 1,
        "source": source_path.name,
        "source_path": _relative_source_path(source_path, output_dir),
        "source_sha256": sha256_file(source_path),
        "template": template_path.name,
        "template_sha256": sha256_file(template_path),
        "document_xml_sha256": sha256_bytes(document_xml),
        "package_manifest": zip_manifest(template_path),
        "styles_sha256": sha256_file(styles_path),
        "source_track_enabled": source_track_enabled(source_path),
        "uses_date_utc": document_uses_date_utc(document_xml),
        "parts": {part_key: sha256_bytes(part_xmls[part_key]) for part_key in sorted(part_xmls)},
        "paragraphs": [
            {
                "id": paragraph.paragraph_id,
                "base_style": paragraph.base_style,
                "insertion_style": _paragraph_insertion_style(paragraph),
                "skeleton": skeleton(paragraph.nodes),
                "token_ids": _token_ids(paragraph.nodes),
                "section_bearing": paragraph.section_bearing,
                "editable": paragraph.editable,
                "mark_revision": paragraph.mark_revision,
                "original_index": index,
                "part_key": paragraph.part_key,
                "part_entry_id": paragraph.part_entry_id,
            }
            for index, paragraph in enumerate(parsed.document.paragraphs)
        ],
        "tokens": _format_token_ids(parsed.tokens),
    }
    format_path = output_dir / "format.json"
    format_path.write_bytes(json_bytes(format_data))
    parsed.document.meta.update(
        {
            "format": format_path.name,
            "styles": styles_path.name,
            "template": template_path.name,
            "source": source_path.name,
        }
    )
    typed_path = output_dir / "typed.md"
    typed_path.write_text(serialize_typed(parsed.document), encoding="utf-8", newline="\n")
    from .edit import generate_clean_edit  # lazy: edit.py imports this module

    generate_clean_edit(output_dir, parsed.document)
    return output_dir


def _load_workdir(path: str | Path) -> tuple[Path, dict[str, Any], StyleRegistry, TypedDocument, Path]:
    workdir = Path(path).resolve()
    if not workdir.is_dir():
        raise ValidationError(f"workdir not found: {workdir}")
    required = {"typed.md", "format.json", "styles.json", "_template.docx"}
    missing = sorted(name for name in required if not (workdir / name).exists())
    if missing:
        raise ValidationError(f"workdir missing: {', '.join(missing)}")
    try:
        format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
        styles_data = json.loads((workdir / "styles.json").read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ValidationError(f"invalid workdir JSON: {exc}") from exc
    if format_data.get("schema") != "typed-format-1" or format_data.get("model_version") != 1 or format_data.get("canonicalizer_version") != 1:
        raise ValidationError("incompatible typed workdir schema")
    styles = StyleRegistry.from_json(styles_data)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    template = workdir / str(format_data.get("template", "_template.docx"))
    if template.name != "_template.docx":
        raise ValidationError("template must be the workdir _template.docx")
    return workdir, format_data, styles, typed, template


def _validate_token_nodes(nodes: Iterable[Any], records: dict[str, Any]) -> None:
    for node in nodes:
        if isinstance(node, RevisionNode):
            record = records.get(node.token_id)
            if not record or record.get("kind") != "revision" or record.get("attrs", {}) != node.attrs:
                raise ValidationError(f"revision token changed or missing: {node.token_id}")
            _validate_token_nodes(node.children, records)
        elif isinstance(node, RangeNode):
            record = records.get(node.token_id)
            if not record or record.get("kind") != node.kind or record.get("attrs", {}) != node.attrs:
                raise ValidationError(f"range token changed or missing: {node.token_id}")
            _validate_token_nodes(node.children, records)
        elif isinstance(node, (AnchorNode, InlineNode, OpaqueNode)):
            record = records.get(node.token_id)
            if not record or record.get("kind") != node.kind:
                raise ValidationError(f"structural token changed or missing: {node.token_id}")
            if record.get("attrs", {}) != node.attrs:
                raise ValidationError(f"structural token attributes changed: {node.token_id}")
            if isinstance(node, InlineNode) and record.get("style_id", "") != node.style_id:
                raise ValidationError(f"inline token style changed: {node.token_id}")
        elif isinstance(node, TextNode):
            pass
        else:
            raise ValidationError("unknown typed AST node")


def _text_segments(nodes: Iterable[Any]) -> list[tuple[str, str]]:
    segments: list[tuple[str, str]] = []
    for node in nodes:
        if isinstance(node, TextNode):
            if node.text:
                segments.append((node.text, node.style_id))
        elif isinstance(node, RangeNode):
            segments.extend(_text_segments(node.children))
        elif isinstance(node, RevisionNode):
            if node.kind in ("insert", "move_to"):
                segments.extend(_text_segments(node.children))
    return segments


def _validate_segment_rewrite(
    old_segments: list[tuple[str, str]],
    new_segments: list[tuple[str, str]],
    paragraph_id: str,
) -> None:
    if len(old_segments) <= 1:
        return
    old_text = "".join(text for text, _ in old_segments)
    new_text = "".join(text for text, _ in new_segments)
    old_offsets: list[tuple[int, int]] = []
    new_offsets: list[tuple[int, int]] = []
    offset = 0
    for text, _ in old_segments:
        old_offsets.append((offset, offset + len(text)))
        offset += len(text)
    offset = 0
    for text, _ in new_segments:
        new_offsets.append((offset, offset + len(text)))
        offset += len(text)

    def touched(offsets: list[tuple[int, int]], start: int, end: int) -> int:
        return sum(1 for left, right in offsets if start < right and left < end)

    for tag, i1, i2, j1, j2 in SequenceMatcher(None, old_text, new_text, autojunk=False).get_opcodes():
        if tag != "equal" and (touched(old_offsets, i1, i2) > 1 or touched(new_offsets, j1, j2) > 1):
            raise ValidationError(f"cross-boundary text rewrite requires explicit style ownership: {paragraph_id}")
    for index, (new_text_node, _) in enumerate(new_segments):
        if index < len(old_segments) and new_text_node == old_segments[index][0]:
            continue
        if any(index != old_index and new_text_node == old_text_node for old_index, (old_text_node, _) in enumerate(old_segments)):
            raise ValidationError(f"cross-boundary text rewrite requires explicit style ownership: {paragraph_id}")


def _validate_cross_boundary_edit(baseline: Paragraph, current: Paragraph) -> None:
    _validate_segment_rewrite(
        _text_segments(baseline.nodes),
        _text_segments(current.nodes),
        current.paragraph_id,
    )


def _validate_styles(nodes: Iterable[Any], styles: StyleRegistry) -> None:
    for node in nodes:
        if isinstance(node, TextNode):
            styles.require(node.style_id)
        elif isinstance(node, InlineNode):
            if node.style_id:
                styles.require(node.style_id)
        elif isinstance(node, (RangeNode, RevisionNode)):
            _validate_styles(node.children, styles)


def _canonical_ppr(value: str) -> str:
    if not value:
        return ""
    # Paragraph-mark revisions are compared via content_signature; strip the
    # self-closing ins/del so pPr comparison stays structural.
    value = re.sub(r"<w:(ins|del)(?:\s+[^>]*)?/>", "", value)
    opening_end = value.find(">")
    if opening_end >= 0:
        opening = value[:opening_end]
        declarations: list[str] = []
        from docx.oxml.ns import nsmap
        from .typed_core import NS_W16DU

        for prefix, namespace in {**nsmap, "w16du": NS_W16DU}.items():
            if f"{prefix}:" in value and f"xmlns:{prefix}=" not in opening:
                declarations.append(f'xmlns:{prefix}="{namespace}"')
        if declarations:
            value = value.replace("<w:pPr", "<w:pPr " + " ".join(declarations), 1)
    return canonical_xml(value)


def _paragraph_attrs(p_open: str) -> dict[str, str]:
    if not p_open:
        return {}
    try:
        element = ET.fromstring(p_open[:-1] + "/>" if p_open.endswith(">") else p_open)
    except ET.ParseError:
        return {}
    return {qname(key): value for key, value in element.attrib.items()}


def _new_paragraph_opening(p_open: str) -> str:
    def remove_ephemeral(match: re.Match[str]) -> str:
        name = match.group("name").rsplit(":", 1)[-1]
        return "" if name in _EPHEMERAL_P_ATTRS else match.group(0)

    return _P_ATTR_RE.sub(remove_ephemeral, p_open)


def _protected_document_bytes(xml: bytes) -> bytes:
    slices = locate_document_xml(xml)
    # Table byte ranges are re-rendered from cell paragraphs, so they are not
    # protected regions; their fidelity is covered by per-cell byte checks.
    excluded = [(table.start, table.end) for table in slices.tables if table.body_level]
    excluded.extend((sdt.start, sdt.end) for sdt in slices.sdts)
    excluded.extend((paragraph.start, paragraph.end) for paragraph in slices.paragraphs)
    excluded.sort()
    merged: list[tuple[int, int]] = []
    for start, end in excluded:
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    pieces: list[bytes] = []
    cursor = 0
    for start, end in merged:
        pieces.append(xml[cursor:start])
        cursor = end
    pieces.append(xml[cursor:])
    return b"".join(pieces)


def package_guard(template: Path, output: Path, editable_parts: set[str] | None = None) -> None:
    editable_parts = editable_parts or set()
    with zipfile.ZipFile(template) as source_zip, zipfile.ZipFile(output) as output_zip:
        source_names = sorted(source_zip.namelist())
        output_names = sorted(output_zip.namelist())
        if source_names != output_names:
            raise ValidationError("DOCX package part list changed")
        for name in source_names:
            if name == "word/document.xml":
                continue
            match = PART_KEYS_PATTERN.match(name)
            if match and match.group(1) in editable_parts:
                continue
            if name in editable_parts:  # full-path keys (comments parts)
                continue
            if source_zip.read(name) != output_zip.read(name):
                raise ValidationError(f"protected DOCX part changed: {name}")
        source_xml = source_zip.read("word/document.xml")
        output_xml = output_zip.read("word/document.xml")
        if _protected_document_bytes(source_xml) != _protected_document_bytes(output_xml):
            raise ValidationError("protected document XML region changed")
        try:
            ET.fromstring(output_xml)
        except ET.ParseError as exc:
            raise ValidationError(f"built document XML is invalid: {exc}") from exc


def _render_node(
    node: Any,
    base_style: str,
    styles: StyleRegistry,
    tokens: dict[str, dict[str, Any]],
    *,
    in_delete: bool = False,
) -> str:
    if isinstance(node, TextNode):
        style = styles.require(node.style_id)
        preserve = node.text[:1].isspace() or node.text[-1:].isspace()
        space = ' xml:space="preserve"' if node.text and preserve else ""
        text_tag = "delText" if in_delete else "t"
        return f"<w:r>{style.rpr}<w:{text_tag}{space}>{xml_escape(node.text)}</w:{text_tag}></w:r>"
    if isinstance(node, InlineNode):
        record = tokens.get(node.token_id)
        if not record:
            raise ValidationError(f"missing inline token: {node.token_id}")
        if node.kind == "rpr-change":
            return ""  # injected at paragraph level (see _render_paragraph)
        if node.kind == "empty-run":
            return str(record.get("raw", ""))
        style_id = node.style_id or str(record.get("style_id", ""))
        style = styles.require(style_id) if style_id else None
        raw = str(record.get("raw", ""))
        if not raw:
            raise ValidationError(f"inline token has no XML: {node.token_id}")
        return f"<w:r>{style.rpr if style else ''}{raw}</w:r>"
    if isinstance(node, AnchorNode):
        record = tokens.get(node.token_id)
        if not record or not record.get("raw"):
            raise ValidationError(f"missing anchor XML: {node.token_id}")
        return str(record["raw"])
    if isinstance(node, OpaqueNode):
        raise ValidationError(f"opaque node cannot be synthesized: {node.token_id}")
    record = tokens.get(node.token_id)
    if not record or not record.get("open") or not record.get("close"):
        raise ValidationError(f"missing range XML: {node.token_id}")
    if isinstance(node, RevisionNode):
        inner = "".join(
            _render_node(
                child,
                base_style,
                styles,
                tokens,
                in_delete=in_delete or node.kind in ("delete", "move_from"),
            )
            for child in node.children
        )
    else:
        inner = "".join(_render_node(child, base_style, styles, tokens, in_delete=in_delete) for child in node.children)
    return f"{record['open']}{inner}{record['close']}"


def _strip_paragraph_marks(ppr: str) -> str:
    """Remove self-closing paragraph-mark revisions from pPr bytes."""
    return re.sub(r"<w:(ins|del)(?:\s+[^>]*)?/>", "", ppr)


def _collect_rpr_changes(nodes: Iterable[Any], tokens: dict[str, dict[str, Any]]) -> list[str]:
    """Raw w:rPrChange XML carried by a paragraph (format history markers)."""
    raw_parts: list[str] = []
    for node in nodes:
        if isinstance(node, InlineNode) and node.kind == "rpr-change":
            record = tokens.get(node.token_id, {})
            raw = str(record.get("raw", ""))
            if raw:
                raw_parts.append(raw)
        elif isinstance(node, (RangeNode, RevisionNode)):
            raw_parts.extend(_collect_rpr_changes(node.children, tokens))
    return raw_parts


def _render_nodes_seq(
    nodes: Iterable[Any],
    base_style: str,
    styles: StyleRegistry,
    tokens: dict[str, dict[str, Any]],
    *,
    in_delete: bool = False,
) -> str:
    """Render a node sequence, binding each rPrChange marker to the run that
    carries it (the nearest preceding run in the AST)."""
    chunks: list[str] = []
    last_run_index = -1

    def inject_history(raw: str) -> None:
        nonlocal last_run_index
        if last_run_index >= 0:
            chunk = chunks[last_run_index]
            if "</w:rPr>" in chunk:
                chunks[last_run_index] = chunk.replace("</w:rPr>", raw + "</w:rPr>", 1)
                return
            if chunk.startswith("<w:r>"):
                chunks[last_run_index] = chunk.replace("<w:r>", f"<w:r><w:rPr>{raw}</w:rPr>", 1)
                return
        chunks.append(f"<w:r><w:rPr>{raw}</w:rPr></w:r>")
        last_run_index = len(chunks) - 1

    for node in nodes:
        if isinstance(node, InlineNode) and node.kind == "rpr-change":
            raw = str(tokens.get(node.token_id, {}).get("raw", ""))
            if raw:
                inject_history(raw)
            continue
        if isinstance(node, TextNode):
            style = styles.require(node.style_id)
            preserve = node.text[:1].isspace() or node.text[-1:].isspace()
            space = ' xml:space="preserve"' if node.text and preserve else ""
            text_tag = "delText" if in_delete else "t"
            chunk = f"<w:r>{style.rpr}<w:{text_tag}{space}>{xml_escape(node.text)}</w:{text_tag}></w:r>"
            chunks.append(chunk)
            last_run_index = len(chunks) - 1
            continue
        if isinstance(node, (RangeNode, RevisionNode)):
            record = tokens.get(node.token_id)
            if not record or not record.get("open") or not record.get("close"):
                raise ValidationError(f"missing range XML: {node.token_id}")
            inner = _render_nodes_seq(
                node.children, base_style, styles, tokens,
                in_delete=in_delete or node.kind in ("delete", "move_from"),
            )
            chunk = f"{record['open']}{inner}{record['close']}"
            chunks.append(chunk)
            if chunk.startswith("<w:"):
                last_run_index = len(chunks) - 1
            continue
        chunk = _render_node(node, base_style, styles, tokens, in_delete=in_delete)
        chunks.append(chunk)
        if chunk.startswith("<w:r>"):
            last_run_index = len(chunks) - 1
    return "".join(chunks)



def _render_paragraph(paragraph: Paragraph, inherited: Paragraph, styles: StyleRegistry, tokens: dict[str, dict[str, Any]]) -> bytes:
    if paragraph.inherit:
        p_open = paragraph.p_open or inherited.p_open
        ppr = paragraph.ppr or inherited.ppr
    else:
        p_open = paragraph.p_open
        ppr = paragraph.ppr
    if not p_open:
        raise ValidationError(f"paragraph {paragraph.paragraph_id} has no template opening")
    # Template pPr bytes may already carry paragraph marks (extract keeps them
    # verbatim); render them from the AST state only, so a resolved mark
    # (mark_revision=None) disappears instead of surviving in the pPr bytes.
    if paragraph.mark_revision is not None or _MARK_TAG_RE.search(ppr) is not None:
        ppr = _strip_paragraph_marks(ppr)
    if paragraph.mark_revision:
        ppr = _inject_mark_revision(ppr, paragraph.mark_revision, tokens)
    body = _render_nodes_seq(paragraph.nodes, paragraph.base_style, styles, tokens)
    return (p_open + ppr + body + "</w:p>").encode("utf-8")


def _paragraph_placements(paragraphs: list[Paragraph], template_count: int) -> tuple[list[int | None], list[int | None]]:
    existing_order = [paragraph.original_index for paragraph in paragraphs if paragraph.original_index >= 0]
    if existing_order != sorted(existing_order):
        raise ValidationError("existing paragraph order cannot change")
    slots: list[int | None] = []
    insert_before: list[int | None] = []
    for index, paragraph in enumerate(paragraphs):
        if paragraph.original_index >= 0:
            slots.append(paragraph.original_index)
            insert_before.append(None)
            continue
        target = template_count
        for following in paragraphs[index + 1:]:
            if following.original_index >= 0:
                target = following.original_index
                break
        slots.append(None)
        insert_before.append(target)
    return slots, insert_before


def _render_sdts(
    xml: bytes,
    slices: DocumentSlices,
    sdt_paragraphs: list[Paragraph],
    baseline_by_id: dict[str, Paragraph],
    styles: StyleRegistry,
    tokens: dict[str, dict[str, Any]],
) -> dict[int, bytes]:
    """Render body-level sdt content controls: touched paragraphs render,
    untouched replay raw; the sdtPr/sdt structure bytes always come from the
    template."""
    by_id = {paragraph.paragraph_id: paragraph for paragraph in sdt_paragraphs}
    rendered: dict[int, bytes] = {}
    for sdt in slices.sdts:
        cell_slices = {cell.container_path: cell for cell in sdt.cells}
        output: list[bytes] = []
        cursor = sdt.start
        for cell in sdt.cells:
            output.append(xml[cursor:cell.start])
            paragraph = by_id.get(_sdt_paragraph_id(cell.container_path))
            baseline = baseline_by_id.get(_sdt_paragraph_id(cell.container_path))
            if paragraph is not None and baseline is not None and paragraph.nodes != baseline.nodes:
                output.append(_render_paragraph(paragraph, baseline, styles, tokens))
            else:
                output.append(cell.raw)
            cursor = cell.end
        output.append(xml[cursor:sdt.end])
        rendered[sdt.sdt_index] = b"".join(output)
    return rendered


def _render_tables(
    xml: bytes,
    slices: DocumentSlices,
    cell_paragraphs: list[Paragraph],
    baseline_by_id: dict[str, Paragraph],
    styles: StyleRegistry,
    tokens: dict[str, dict[str, Any]],
) -> dict[int, bytes]:
    """Render body-level table byte ranges from their cell paragraphs.

    Untouched cells replay raw bytes, touched cells render from the AST, and
    nested tables render recursively (bottom-up). Table structure bytes
    always come from the template range. Returns a map for body-level tables
    only; nested tables are embedded in their parent's bytes.
    """
    by_table: dict[int, list[Paragraph]] = {}
    for paragraph in cell_paragraphs:
        by_table.setdefault(paragraph.table_index, []).append(paragraph)
    renders: dict[int, bytes] = {}

    def render_table(table: TableSlice) -> bytes:
        if table.table_index in renders:
            return renders[table.table_index]
        paragraphs = by_table.get(table.table_index, [])
        paragraphs_by_path = {p.container_path: p for p in paragraphs}
        baseline_by_path = {
            p.container_path: p
            for p in (baseline_by_id[paragraph.paragraph_id] for paragraph in paragraphs)
        }
        nested = [
            candidate
            for candidate in slices.tables
            if candidate.start > table.start
            and candidate.end < table.end
            and not any(
                other.start > candidate.start and other.end < candidate.end
                for other in slices.tables
                if other is not candidate and other is not table
            )
        ]
        units: list[tuple[int, int, str, Any]] = [
            (cell.start, cell.end, "cell", cell.container_path) for cell in table.cells
        ]
        units.extend((nested_table.start, nested_table.end, "tbl", nested_table.table_index) for nested_table in nested)
        units.sort(key=lambda unit: unit[0])
        output: list[bytes] = []
        cursor = table.start
        for start, end, kind, key in units:
            output.append(xml[cursor:start])
            if kind == "tbl":
                output.append(render_table(slices.tables[key]))
            else:
                paragraph = paragraphs_by_path.get(key)
                baseline = baseline_by_path.get(key)
                if paragraph is None or baseline is None or paragraph.nodes == baseline.nodes:
                    output.append(xml[start:end])
                else:
                    output.append(_render_paragraph(paragraph, baseline, styles, tokens))
            cursor = end
        output.append(xml[cursor:table.end])
        rendered = b"".join(output)
        renders[table.table_index] = rendered
        return rendered

    for table in slices.tables:
        render_table(table)
    return {table.table_index: renders[table.table_index] for table in slices.tables if table.body_level}


def _render_boxes_in_paragraph(
    replacement: bytes,
    parent_slice: ParagraphSlice,
    boxes: list[BoxSlice],
    xml: bytes,
    cell_paragraphs: list[Paragraph],
    baseline_by_id: dict[str, Paragraph],
    styles: StyleRegistry,
    tokens: dict[str, dict[str, Any]],
) -> bytes:
    """Replace text-box byte ranges inside a rendered body paragraph:
    untouched box paragraphs replay raw bytes, touched ones render. The box
    offsets inside the replacement match the template offsets because
    untouched structure bytes (including pict/opaque runs) are replayed
    verbatim."""
    result = replacement
    for box in boxes:
        cell_slices = {cell.container_path: cell for cell in box.cells}
        paragraphs_by_path = {
            paragraph.container_path: paragraph
            for paragraph in cell_paragraphs
            if paragraph.container_path in cell_slices
        }
        output: list[bytes] = []
        cursor = box.start
        for cell in box.cells:
            output.append(xml[cursor:cell.start])
            paragraph = paragraphs_by_path.get(cell.container_path)
            baseline = baseline_by_id.get(paragraph.paragraph_id) if paragraph else None
            if paragraph is None or baseline is None or paragraph.nodes == baseline.nodes:
                output.append(cell.raw)
            else:
                output.append(_render_paragraph(paragraph, baseline, styles, tokens))
            cursor = cell.end
        output.append(xml[cursor:box.end])
        box_bytes = b"".join(output)
        rel_start = box.start - parent_slice.start
        rel_end = box.end - parent_slice.start
        result = result[:rel_start] + box_bytes + result[rel_end:]
    return result


def patch_document_xml(
    xml: bytes,
    slices: DocumentSlices,
    replacements: list[bytes],
    slots: list[int | None] | None = None,
    insert_before: list[int | None] | None = None,
    table_render: dict[int, bytes] | None = None,
    sdt_render: dict[int, bytes] | None = None,
) -> bytes:
    if slots is None:
        slots = list(range(len(replacements)))
    if len(slots) != len(replacements):
        raise ValidationError("paragraph replacement slots do not match replacements")
    if insert_before is not None and len(insert_before) != len(replacements):
        raise ValidationError("paragraph insertion positions do not match replacements")
    before: dict[int, list[bytes]] = {}
    replacement_by_slot: dict[int, bytes] = {}
    for index, (slot, replacement) in enumerate(zip(slots, replacements)):
        if slot is None:
            target = insert_before[index] if insert_before is not None else len(slices.paragraphs)
            if target is None:
                target = len(slices.paragraphs)
            if target < 0 or target > len(slices.paragraphs):
                raise ValidationError(f"invalid paragraph insertion position: {target}")
            before.setdefault(target, []).append(replacement)
        elif slot in replacement_by_slot:
            raise ValidationError(f"duplicate replacement for paragraph slot {slot}")
        else:
            replacement_by_slot[slot] = replacement
    table_render = table_render or {}
    sdt_render = sdt_render or {}
    units: list[tuple[int, int, str, int]] = []
    for index, paragraph in enumerate(slices.paragraphs):
        units.append((paragraph.start, paragraph.end, "p", index))
    for table in slices.tables:
        if table.body_level:
            units.append((table.start, table.end, "t", table.table_index))
    for sdt in slices.sdts:
        units.append((sdt.start, sdt.end, "s", sdt.sdt_index))
    units.sort(key=lambda unit: unit[0])
    output: list[bytes] = []
    cursor = 0
    for start, end, kind, key in units:
        output.append(xml[cursor:start])
        if kind == "p":
            output.extend(before.get(key, ()))
            if key in replacement_by_slot:
                output.append(replacement_by_slot[key])
        elif kind == "t":
            output.append(table_render.get(key, xml[start:end]))
        else:
            output.append(sdt_render.get(key, xml[start:end]))
        cursor = end
    output.extend(before.get(len(slices.paragraphs), ()))
    output.append(xml[cursor:])
    return b"".join(output)


PART_KEYS_PATTERN = re.compile(r"word/(header\d+|footer\d+|footnotes|endnotes|comments)\.xml$")


@dataclass
class PartSlice:
    part_key: str
    start: int
    end: int
    paragraphs: list[ParagraphSlice]
    entry_ids: list[str]  # footnote/endnote w:id per container; [] for headers


@dataclass
class PartLayout:
    """Named layout of one header/footer/footnote/endnote/comments part."""

    part_key: str
    root_start: int
    root_end: int
    paragraphs: list[ParagraphSlice]
    entry_ids: list[str]
    table_ranges: list[tuple[int, int]]
    cell_paragraphs: list[ParagraphSlice]
    entry_ranges: list[tuple[int, int, str]]  # (start, end, entry id)


def locate_part_xml(xml: bytes, part_key: str) -> PartLayout:
    """Locate content containers inside a header/footer/footnote/endnote part:
    direct paragraphs for w:hdr/w:ftr roots (plus cell paragraphs inside
    part-level tables), per-entry paragraphs for w:footnotes/w:endnotes
    roots."""
    root_name = {
        "header": "hdr",
        "footer": "ftr",
        "footnotes": "footnotes",
        "endnotes": "endnotes",
        "comments": "comments",
    }[part_key.rstrip("0123456789")]
    cursor = TagCursor(xml)
    root_start = -1
    root_end = -1
    paragraphs: list[ParagraphSlice] = []
    entry_ids: list[str] = []
    current_entry: str | None = None
    para_index = 0
    table_ranges: list[tuple[int, int]] = []
    open_tables: list[tuple[str, int]] = []  # (name, start)
    cell_paragraphs: list[ParagraphSlice] = []
    table_ordinal = -1
    tr_ordinal = 0
    tc_ordinal = 0
    cell_p_ordinal = 0
    entry_ranges: list[tuple[int, int, str]] = []  # (start, end, entry id)
    open_entries: list[tuple[str, str, int]] = []  # (name, entry_id, start)
    for tag in cursor:
        closing, self_closing = tag.closing, tag.self_closing
        name = tag.name
        if closing:
            if not cursor.stack or cursor.stack[-1][0] != name:
                raise ValidationError(f"malformed {part_key} XML nesting near {tag.raw_name}")
            depth = len(cursor.stack)
            if name == root_name and depth == 1:
                root_end = tag.start
            if name in ("footnote", "endnote", "comment") and depth == 2:
                current_entry = None
                if open_entries and open_entries[-1][2] == cursor.stack[-1][1]:
                    entry_name, entry_id, entry_start = open_entries.pop()
                    entry_ranges.append((entry_start, tag.end, entry_id))
            if name == "p" and depth == 2 and root_name in ("hdr", "ftr"):
                paragraphs.append(ParagraphSlice(para_index, cursor.stack[-1][1], tag.end, xml[cursor.stack[-1][1]:tag.end]))
                para_index += 1
            if name == "p" and depth == 3 and root_name in ("footnotes", "endnotes", "comments"):
                paragraphs.append(ParagraphSlice(para_index, cursor.stack[-1][1], tag.end, xml[cursor.stack[-1][1]:tag.end]))
                entry_ids.append(current_entry or "")
                para_index += 1
            if (
                name == "p"
                and depth >= 5
                and root_name in ("hdr", "ftr")
                and open_tables
                and _part_cell_stack(cursor.stack, open_tables)
            ):
                path = _part_cell_path(cursor.stack, open_tables, table_ordinal, tr_ordinal, tc_ordinal, cell_p_ordinal)
                cell_paragraphs.append(
                    ParagraphSlice(para_index, cursor.stack[-1][1], tag.end, xml[cursor.stack[-1][1]:tag.end], path, -1)
                )
                para_index += 1
                cell_p_ordinal += 1
            if name == "tbl" and open_tables and open_tables[-1][1] == cursor.stack[-1][1]:
                table_ranges.append((open_tables[-1][1], tag.end))
                open_tables.pop()
            cursor.pop()
            continue
        depth = len(cursor.stack)
        if name == root_name and depth == 0:
            root_start = tag.start
            if self_closing:
                root_end = tag.end
        if name in ("footnote", "endnote", "comment") and depth == 1 and root_name in ("footnotes", "endnotes", "comments"):
            current_entry = element_attr(xml, tag.start, tag.end, "w:id")
            open_entries.append((name, current_entry, tag.start))
        if name == "tbl" and depth == 1 and root_name in ("hdr", "ftr"):
            table_ordinal += 1
            tr_ordinal = 0
            tc_ordinal = 0
            open_tables.append(("tbl", tag.start))
        if name == "tr" and open_tables and depth == 2:
            tr_ordinal += 1
            tc_ordinal = 0
        if name == "tc" and open_tables and depth == 3:
            tc_ordinal += 1
            cell_p_ordinal = 0
        if not self_closing:
            cursor.stack.append((name, tag.start))
    if root_start < 0 or root_end < root_start:
        raise ValidationError(f"no {root_name} root in {part_key}")
    return PartLayout(
        part_key,
        root_start,
        root_end,
        paragraphs,
        entry_ids,
        table_ranges,
        cell_paragraphs,
        entry_ranges,
    )


def _part_cell_stack(stack: list[tuple[str, int]], open_tables: list[tuple[str, int]]) -> bool:
    """Whether the closing p is inside a part-level tbl>tr>tc chain."""
    if not open_tables:
        return False
    names = [name for name, _ in stack[-4:-1]]
    return names == ["tbl", "tr", "tc"]


def _part_cell_path(
    stack: list[tuple[str, int]],
    open_tables: list[tuple[str, int]],
    table_ordinal: int,
    tr_ordinal: int,
    tc_ordinal: int,
    p_ordinal: int,
) -> tuple[str, int, ...]:
    return ("tbl", table_ordinal, "tr", tr_ordinal - 1, "tc", tc_ordinal - 1, "p", p_ordinal)


def element_attr(xml: bytes, start: int, end: int, name: str) -> str:
    """Extract an attribute value from a tag byte range (name may be qname)."""
    fragment = xml[start:end].decode("utf-8")
    for match in _P_ATTR_RE.finditer(fragment):
        if match.group("name") == name:
            return match.group("value")[1:-1]
    return ""


def parse_part_xml(
    xml: bytes,
    part_key: str,
    *,
    styles: StyleRegistry | None = None,
    tokens: _TokenTable | None = None,
) -> list[Paragraph]:
    """Parse one header/footer/footnote/endnote part into paragraphs.

    Shares the caller's style registry and token table so part paragraphs
    and body paragraphs use one namespace.
    """
    layout = locate_part_xml(xml, part_key)
    try:
        root = ET.fromstring(xml)
    except ET.ParseError as exc:
        raise ValidationError(f"invalid {part_key} XML: {exc}") from exc
    registry = styles or StyleRegistry()
    tokens = tokens or _TokenTable()
    root_element = root
    paragraphs: list[Paragraph] = []
    for index, slice_ in enumerate(layout.paragraphs):
        element = _find_part_paragraph(root_element, part_key, index)
        if element is None:
            raise ValidationError(f"part paragraph locator disagrees with parsed body: {part_key}.P{index}")
        paragraph = _parse_paragraph(
            element,
            slice_.raw,
            -1,
            registry,
            tokens,
            paragraph_id=f"{part_key}.P{index}",
            original_index=-1,
        )
        if layout.entry_ids:
            paragraph.part_entry_id = layout.entry_ids[index]
        paragraph.part_key = part_key
        paragraphs.append(paragraph)
    for slice_ in layout.cell_paragraphs:
        element = _find_element_by_path(root_element, slice_.container_path)
        if element is None:
            raise ValidationError(f"part cell locator disagrees with parsed body: {slice_.container_path}")
        values = dict(zip(slice_.container_path[::2], slice_.container_path[1::2]))
        paragraph_id = f"{part_key}.T{values['tbl']}.R{values['tr']}.C{values['tc']}.P{values['p']}"
        paragraph = _parse_paragraph(
            element,
            slice_.raw,
            -1,
            registry,
            tokens,
            paragraph_id=paragraph_id,
            container_path=slice_.container_path,
            original_index=-1,
        )
        paragraph.part_key = part_key
        paragraphs.append(paragraph)
    return paragraphs


def _find_part_paragraph(root: ET.Element, part_key: str, index: int) -> ET.Element | None:
    """Navigate to the index-th paragraph inside the part's content containers."""
    base = part_key.rstrip("0123456789")
    paras: list[ET.Element] = []
    if base in ("header", "footer"):
        paras = [child for child in list(root) if local_name(child.tag) == "p"]
    else:
        for entry in list(root):
            if local_name(entry.tag) in ("footnote", "endnote", "comment"):
                paras.extend(child for child in list(entry) if local_name(child.tag) == "p")
    return paras[index] if index < len(paras) else None


def scan_package_revisions(path: Path) -> list[dict[str, Any]]:
    """Read-only inventory of tracked revisions across all WordprocessingML
    parts (document, headers, footers, footnotes, endnotes, comments)."""
    from .typed_core import NS_W

    revisions: list[dict[str, Any]] = []
    with zipfile.ZipFile(path) as archive:
        for name in sorted(archive.namelist()):
            if not name.endswith(".xml"):
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            for element in root.iter():
                local = local_name(element.tag)
                if local not in {"ins", "del", "moveFrom", "moveTo"}:
                    continue
                attrs = {
                    qname(key): value
                    for key, value in element.attrib.items()
                    if key != f"{{{NS_W}}}id"
                }
                text_parts: list[str] = []
                for child in element.iter():
                    if local_name(child.tag) in {"t", "delText"}:
                        text_parts.append(child.text or "")
                revisions.append(
                    {
                        "part": name,
                        "kind": {
                            "ins": "insert",
                            "del": "delete",
                            "moveFrom": "move_from",
                            "moveTo": "move_to",
                        }[local],
                        "w_id": element.attrib.get(f"{{{NS_W}}}id", ""),
                        "author": attrs.get("w:author", ""),
                        "date": attrs.get("w:date", ""),
                        "text": "".join(text_parts),
                    }
                )
    return revisions


def source_track_enabled(source_path: Path) -> bool:
    """Whether the source package has track changes enabled (settings.xml)."""
    from .typed_core import NS_W

    with zipfile.ZipFile(source_path) as archive:
        try:
            settings_xml = archive.read("word/settings.xml")
        except KeyError:
            return False
    try:
        root = ET.fromstring(settings_xml)
    except ET.ParseError:
        return False
    return any(local_name(child.tag) == "trackChanges" for child in root.iter())


def document_uses_date_utc(document_xml: bytes) -> bool:
    """Whether the source document already uses w16du:dateUtc on revisions."""
    return b"dateUtc" in document_xml


def used_revision_ids(path: Path) -> set[int]:
    """Package-wide set of w:id values used by tracked revisions."""
    from .typed_core import NS_W

    ids: set[int] = set()
    with zipfile.ZipFile(path) as archive:
        for name in sorted(archive.namelist()):
            if not name.endswith(".xml"):
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            for element in root.iter():
                local = local_name(element.tag)
                if local not in {"ins", "del", "moveFrom", "moveTo"}:
                    continue
                raw = element.attrib.get(f"{{{NS_W}}}id", "")
                if raw.isdigit():
                    ids.add(int(raw))
    return ids


def next_revision_id(path: Path, used: set[int] | None = None) -> int:
    """Lowest available non-negative w:id over the package (and ``used``)."""
    taken = used_revision_ids(path)
    if used:
        taken = taken | set(used)
    candidate = 0
    while candidate in taken:
        candidate += 1
    return candidate


def _render_parts(validated: ValidatedWorkdir, part_paragraphs: list[Paragraph]) -> dict[str, bytes]:
    """Render every editable part from its paragraphs (touched only; others
    replay the template blob). Parts without live paragraphs still render so
    fully-deleted entries (e.g. removed comments) drop out."""
    by_id = {paragraph.paragraph_id: paragraph for paragraph in part_paragraphs}
    rendered: dict[str, bytes] = {}
    with zipfile.ZipFile(validated.template_path) as archive:
        for name in archive.namelist():
            match = PART_KEYS_PATTERN.match(name)
            if not match:
                continue
            part_key = match.group(1)
            template_xml = archive.read(name)
            rendered[part_key] = render_part_xml(
                template_xml, part_key, by_id, validated.baseline_by_id,
                validated.styles, validated.format_data.get("tokens", {}),
            )
    return rendered


def render_part_xml(
    template_xml: bytes,
    part_key: str,
    paragraphs_by_id: dict[str, Paragraph],
    baseline_by_id: dict[str, Paragraph],
    styles: StyleRegistry,
    tokens: dict[str, dict[str, Any]],
) -> bytes:
    """Render one part file: untouched paragraphs replay raw bytes, touched
    ones render from the AST; part-level tables re-render their cells
    (structure bytes from the template); container structure bytes always
    come from the template."""
    layout = locate_part_xml(template_xml, part_key)

    def paragraph_bytes(key: str, start_: int, end_: int) -> bytes:
        paragraph = paragraphs_by_id.get(key)
        baseline = baseline_by_id.get(key)
        if paragraph is not None and baseline is not None and paragraph.nodes != baseline.nodes:
            return _render_paragraph(paragraph, baseline, styles, tokens)
        return template_xml[start_:end_]

    def render_table(table_start: int, table_end: int) -> bytes:
        cells = [
            cell for cell in layout.cell_paragraphs
            if cell.start > table_start and cell.end < table_end
        ]
        output: list[bytes] = []
        cursor = table_start
        for cell in sorted(cells, key=lambda cell: cell.start):
            output.append(template_xml[cursor:cell.start])
            values = dict(zip(cell.container_path[::2], cell.container_path[1::2]))
            key = f"{part_key}.T{values['tbl']}.R{values['tr']}.C{values['tc']}.P{values['p']}"
            output.append(paragraph_bytes(key, cell.start, cell.end))
            cursor = cell.end
        output.append(template_xml[cursor:table_end])
        return b"".join(output)

    entry_mode = part_key.rstrip("0123456789") in ("footnotes", "endnotes", "comments")
    output: list[bytes] = [template_xml[:layout.root_start]]
    cursor = layout.root_start
    if entry_mode:
        for entry_start, entry_end, entry_id in sorted(layout.entry_ranges, key=lambda item: item[0]):
            if entry_start < cursor:
                continue
            output.append(template_xml[cursor:entry_start])
            entry_slices = [
                s for s in layout.paragraphs if s.start >= entry_start and s.end <= entry_end
            ]
            entry_paragraphs = [
                f"{part_key}.P{s.index}" for s in entry_slices
            ]
            if not any(key in paragraphs_by_id for key in entry_paragraphs):
                # deleted entry (e.g. a removed comment): drop its bytes
                cursor = entry_end
                continue
            entry_output: list[bytes] = []
            entry_cursor = entry_start
            for s in entry_slices:
                entry_output.append(template_xml[entry_cursor:s.start])
                entry_output.append(paragraph_bytes(f"{part_key}.P{s.index}", s.start, s.end))
                entry_cursor = s.end
            entry_output.append(template_xml[entry_cursor:entry_end])
            output.append(b"".join(entry_output))
            cursor = entry_end
    else:
        for slice_ in layout.paragraphs:
            output.append(template_xml[cursor:slice_.start])
            output.append(paragraph_bytes(f"{part_key}.P{slice_.index}", slice_.start, slice_.end))
            cursor = slice_.end
    # part-level tables, in document order
    for table_start, table_end in sorted(layout.table_ranges):
        if table_start < cursor:
            continue  # nested inside an already-rendered unit (not expected in v1)
        output.append(template_xml[cursor:table_start])
        output.append(render_table(table_start, table_end))
        cursor = table_end
    output.append(template_xml[cursor:])
    return b"".join(output)



_REVISION_TAG_KINDS = {
    "ins": "insert", "del": "delete",
    "moveFrom": "move_from", "moveTo": "move_to",
}


def settle_xml_revisions(xml: bytes, action: str) -> bytes:
    """Byte-level settlement of every tracked revision in a raw XML part.

    Accept: insert/move_to unwrap (keep children), delete/move_from remove.
    Reject: insert/move_to remove, delete/move_from unwrap with w:delText
    switched to w:t. Paragraph-mark revisions (self-closing ins/del) are
    removed in both directions (the paragraph itself is never removed by
    byte settlement). Opaque interior bytes are copied verbatim; only
    revision wrapper bytes change.
    """
    remove_kinds = {
        "accept": {"delete", "move_from"},
        "reject": {"insert", "move_to"},
    }[action]
    unwrap_deltext = action == "reject"
    anchor_tags = {"bookmarkStart", "bookmarkEnd", "commentRangeStart", "commentRangeEnd"}
    stack: list[str] = []
    skip_depth = 0
    out: list[bytes] = []
    pending_anchors: list[bytes] = []
    cursor = 0
    for tag in iter_tags(xml):
        token = tag.bytes_in(xml)
        closing, self_closing = tag.closing, tag.self_closing
        name = tag.name
        if skip_depth:
            # content inside a removed container is dropped wholesale, except
            # comment/bookmark anchors which are re-anchored outside the
            # removed range (keeps anchor pairing intact)
            cursor = tag.end
            if name in anchor_tags:
                pending_anchors.append(token)
            elif closing and name in _REVISION_TAG_KINDS:
                skip_depth -= 1
                if skip_depth == 0 and pending_anchors:
                    out.extend(pending_anchors)
                    pending_anchors = []
            elif not closing and not self_closing and name in _REVISION_TAG_KINDS:
                skip_depth += 1
            continue
        out.append(xml[cursor:tag.start])
        cursor = tag.end
        if name in _REVISION_TAG_KINDS:
            kind = _REVISION_TAG_KINDS[name]
            if closing:
                if stack and stack[-1] == kind:
                    stack.pop()
                    continue  # wrapper close dropped (unwrap)
                continue
            if self_closing:
                # paragraph mark: removed in both directions
                continue
            if kind in remove_kinds:
                skip_depth += 1
                continue
            stack.append(kind)
            continue
        if closing and name == "delText" and unwrap_deltext:
            out.append(b"</w:t>")
            continue
        if not closing and not self_closing and name == "delText" and unwrap_deltext:
            out.append(re.sub(rb"<w:delText([ >])", rb"<w:t\1", token, count=1))
            continue
        out.append(token)
    if skip_depth or stack:
        raise ValidationError("malformed revision containers in settlement")
    out.append(xml[cursor:])
    return b"".join(out)


_COMMENT_PARTS = (
    "word/comments.xml",
    "word/commentsExtended.xml",
    "word/commentsIds.xml",
    "word/commentsExtensible.xml",
)


def clear_comments_from_document(xml: bytes) -> bytes:
    """Remove every comment anchor/reference from document XML (byte-level).

    Anchors are matched by local name, so alternate namespace prefixes
    behave identically to ``w:``.
    """
    out: list[bytes] = []
    cursor = 0
    for tag in iter_tags(xml):
        if tag.name in ("commentRangeStart", "commentRangeEnd", "commentReference") and tag.self_closing:
            out.append(xml[cursor:tag.start])
            cursor = tag.end
    out.append(xml[cursor:])
    return b"".join(out)


def empty_comments_part(xml: bytes) -> bytes:
    """Keep the original part root (with its namespace declarations and
    prefixes) but drop all children — an empty comments definition Word
    accepts."""
    for tag in iter_tags(xml):
        if tag.name == "comments" and not tag.closing and not tag.self_closing:
            return xml[tag.start:tag.end] + b"</w:comments>"
    return xml



_TABLE_STRUCT_NAMES = ("tbl", "tr", "tc", "tblPr", "tblGrid", "trPr", "tcPr", "gridSpan", "vMerge")


def _locate_table_elements(
    xml: bytes, table_start: int, table_end: int,
) -> tuple[list[tuple[int, int]], list[tuple[int, int]]]:
    """Row and cell byte ranges of ONE table (template offsets).

    Nested tables inside cells are excluded: only tr/tc at the table's own
    depth are collected (a nested table's rows would otherwise corrupt the
    row/cell splices of every structure operation).
    """
    cursor = TagCursor(xml, table_start, table_end)
    rows: list[tuple[int, int]] = []
    cells: list[tuple[int, int]] = []
    nested_tbls = 0
    for tag in cursor:
        name = tag.name
        if tag.closing:
            if not cursor.stack or cursor.stack[-1][0] != name:
                raise ValidationError(f"malformed table XML nesting near {tag.raw_name}")
            open_start = cursor.stack[-1][1]
            if name == "tbl" and nested_tbls:
                nested_tbls -= 1
            elif name == "tr" and not nested_tbls:
                rows.append((open_start, tag.end))
            elif name == "tc" and not nested_tbls:
                cells.append((open_start, tag.end))
            cursor.pop()
            continue
        if name == "tbl" and tag.start != table_start:
            nested_tbls += 1
        if name == "tr" or name == "tc" or not tag.self_closing:
            cursor.stack.append((name, tag.start))
    return rows, cells


def _make_empty_cell(cell_xml: bytes, *, text: str = "") -> bytes:
    """Clone a cell's structure (tcPr + one paragraph) with empty text."""
    return cell_xml


def apply_table_operation(
    xml: bytes,
    table_index: int,
    operation: str,
    *args: int,
) -> bytes:
    """Byte-level table structure operation on a body-level table.

    Operations: insert-row <after>, delete-row <row>, insert-col <after>,
    delete-col <col>, merge-cells <row> <col> <span>, split-cells <row> <col>
    <span>. Structure bytes are synthesized from the template; new cell text
    starts empty.
    """
    slices = locate_document_xml(xml)
    body_tables = [t for t in slices.tables if t.body_level]
    if table_index >= len(body_tables):
        raise ValidationError(f"table-not-found: T{table_index}")
    table = body_tables[table_index]
    rows, cells = _locate_table_elements(xml, table.start, table.end)
    if not rows:
        raise ValidationError(f"table has no rows: T{table_index}")
    if operation == "insert-row":
        after = args[0]
        if after < 0 or after >= len(rows):
            raise ValidationError(f"row index out of range: {after}")
        template_start, template_end = rows[after]
        new_row = _clone_row_with_empty_cells(xml[template_start:template_end])
        output: list[bytes] = [xml[: table.start]]
        cursor = table.start
        for row_start, row_end in rows:
            output.append(xml[cursor:row_start])
            output.append(xml[row_start:row_end])
            cursor = row_end
            if row_start == template_start:
                output.append(new_row)
        output.append(xml[cursor:])
        return b"".join(output)
    if operation == "delete-row":
        row = args[0]
        if row < 0 or row >= len(rows):
            raise ValidationError(f"row index out of range: {row}")
        output = [xml[: table.start]]
        cursor = table.start
        output.append(xml[cursor: rows[0][0]])
        cursor = rows[0][0]
        for idx, (row_start, row_end) in enumerate(rows):
            if idx == row:
                cursor = row_end
                continue
            output.append(xml[cursor:row_start])
            output.append(xml[row_start:row_end])
            cursor = row_end
        output.append(xml[cursor:])
        return b"".join(output)
    # per-row cell surgery
    row_cells: list[list[tuple[int, int]]] = []
    for row_start, row_end in rows:
        row_cells.append(
            [(cell_start, cell_end) for cell_start, cell_end in cells
             if cell_start >= row_start and cell_end <= row_end]
        )
    if operation == "insert-col":
        after = args[0]
        if after < 0:
            raise ValidationError(f"column index out of range: {after}")
        output: list[bytes] = [xml[: table.start]]
        cursor = table.start
        for (row_start, row_end), row_cells_now in zip(rows, row_cells):
            if after >= len(row_cells_now):
                raise ValidationError(f"column index out of range: {after}")
            template_start, template_end = row_cells_now[after]
            new_cell = _clone_cell_with_empty_paragraph(xml[template_start:template_end])
            output.append(xml[cursor:row_start])
            cursor = row_start
            for cell_start, cell_end in row_cells_now:
                output.append(xml[cursor:cell_start])
                output.append(xml[cell_start:cell_end])
                cursor = cell_end
                if cell_start == template_start:
                    output.append(new_cell)
            output.append(xml[cursor:row_end])
            cursor = row_end
        output.append(xml[cursor:])
        return b"".join(output)
    if operation == "delete-col":
        col = args[0]
        output: list[bytes] = [xml[: table.start]]
        cursor = table.start
        for (row_start, row_end), row_cells_now in zip(rows, row_cells):
            if col >= len(row_cells_now):
                raise ValidationError(f"column index out of range: {col}")
            output.append(xml[cursor:row_start])
            cursor = row_start
            for idx, (cell_start, cell_end) in enumerate(row_cells_now):
                if idx == col:
                    output.append(xml[cursor:cell_start])
                    cursor = cell_end
                    continue
                output.append(xml[cursor:cell_start])
                output.append(xml[cell_start:cell_end])
                cursor = cell_end
            output.append(xml[cursor:row_end])
            cursor = row_end
        output.append(xml[cursor:])
        return b"".join(output)
    if operation == "merge-cells":
        row, col, span = args
        output: list[bytes] = [xml[: table.start]]
        cursor = table.start
        for row_idx, ((row_start, row_end), row_cells_now) in enumerate(zip(rows, row_cells)):
            output.append(xml[cursor:row_start])
            cursor = row_start
            if row_idx == row:
                if col + span > len(row_cells_now):
                    raise ValidationError(f"merge span out of range: {col}+{span}")
                for idx, (cell_start, cell_end) in enumerate(row_cells_now):
                    if idx == col:
                        output.append(xml[cursor:cell_start])
                        merged = _merge_cell_bytes(xml[cell_start:cell_end], span)
                        output.append(merged)
                        cursor = row_cells_now[col + span - 1][1]
                        continue
                    if col < idx < col + span:
                        continue  # swallowed by the merge
                    output.append(xml[cursor:cell_start])
                    output.append(xml[cell_start:cell_end])
                    cursor = cell_end
                output.append(xml[cursor:row_end])
                cursor = row_end
                continue
            for cell_start, cell_end in row_cells_now:
                output.append(xml[cursor:cell_start])
                output.append(xml[cell_start:cell_end])
                cursor = cell_end
            output.append(xml[cursor:row_end])
            cursor = row_end
        output.append(xml[cursor:])
        return b"".join(output)
    if operation == "split-cells":
        row, col, span = args
        output: list[bytes] = [xml[: table.start]]
        cursor = table.start
        for row_idx, ((row_start, row_end), row_cells_now) in enumerate(zip(rows, row_cells)):
            output.append(xml[cursor:row_start])
            cursor = row_start
            if row_idx == row:
                if col >= len(row_cells_now) or span < 1:
                    raise ValidationError(f"split out of range: {col}/{span}")
                cell_start, cell_end = row_cells_now[col]
                split_parts = _split_cell_bytes(xml[cell_start:cell_end], span)
                output.append(xml[cursor:cell_start])
                output.extend(split_parts)
                cursor = cell_end
                for later_start, later_end in row_cells_now[col + 1:]:
                    output.append(xml[cursor:later_start])
                    output.append(xml[later_start:later_end])
                    cursor = later_end
                output.append(xml[cursor:row_end])
                cursor = row_end
                continue
            for cell_start, cell_end in row_cells_now:
                output.append(xml[cursor:cell_start])
                output.append(xml[cell_start:cell_end])
                cursor = cell_end
            output.append(xml[cursor:row_end])
            cursor = row_end
        output.append(xml[cursor:])
        return b"".join(output)
    raise ValidationError(f"unknown table operation: {operation}")


def _clone_cell_with_empty_paragraph(cell_xml: bytes) -> bytes:
    """Clone a cell keeping tcPr but with a single empty paragraph."""
    tc_open_end = find_open_tag_end(cell_xml, "tc")
    if tc_open_end < 0:
        return cell_xml
    tc_pr = _find_element_bytes(cell_xml, "tcPr")
    tc_pr_keep = tc_pr if tc_pr else b""
    return cell_xml[:tc_open_end] + tc_pr_keep + b"<w:p/>" + b"</w:tc>"


def _find_open_end(xml: bytes, name: str) -> int:
    return find_open_tag_end(xml, name)


def _find_element_bytes(xml: bytes, name: str) -> bytes:
    rng = find_element_range(xml, name)
    return xml[rng[0]:rng[1]] if rng is not None else b""


def _merge_cell_bytes(cell_xml: bytes, span: int) -> bytes:
    """Set gridSpan=span on a cell (or add it to tcPr)."""
    if span <= 1:
        return cell_xml
    tc_pr = _find_element_bytes(cell_xml, "tcPr")
    if tc_pr:
        if b"gridSpan" in tc_pr:
            merged = re.sub(rb'<w:gridSpan w:val="\d+"/>', f'<w:gridSpan w:val="{span}"/>'.encode(), tc_pr, count=1)
        else:
            merged = tc_pr[:-len(b"</w:tcPr>")] + f'<w:gridSpan w:val="{span}"/></w:tcPr>'.encode()
        return cell_xml.replace(tc_pr, merged, 1)
    # no tcPr: inject one before the first child
    open_end = _find_open_end(cell_xml, "tc")
    if open_end < 0:
        return cell_xml
    return (
        cell_xml[:open_end]
        + f'<w:tcPr><w:gridSpan w:val="{span}"/></w:tcPr>'.encode()
        + cell_xml[open_end:]
    )


def _split_cell_bytes(cell_xml: bytes, span: int) -> list[bytes]:
    """Reduce gridSpan to 1 and return span copies (first keeps content).

    The extra copies must not inherit the merged cell's gridSpan — each
    split cell claims exactly one grid column.
    """
    tc_pr = _find_element_bytes(cell_xml, "tcPr")
    parts: list[bytes] = []
    for index in range(span):
        if index == 0:
            if tc_pr and b"gridSpan" in tc_pr:
                single = re.sub(rb'<w:gridSpan w:val="\d+"/>', b'<w:gridSpan w:val="1"/>', tc_pr, count=1)
                parts.append(cell_xml.replace(tc_pr, single, 1))
            else:
                parts.append(cell_xml)
        else:
            clone = _clone_cell_with_empty_paragraph(cell_xml)
            if tc_pr and b"gridSpan" in tc_pr:
                clone_tc_pr = _find_element_bytes(clone, "tcPr")
                if clone_tc_pr:
                    stripped = re.sub(rb'<w:gridSpan w:val="\d+"/>', b"", clone_tc_pr, count=1)
                    clone = clone.replace(clone_tc_pr, stripped, 1)
            parts.append(clone)
    return parts


def _clone_row_with_empty_cells(row_xml: bytes) -> bytes:
    """Clone a row, preserving cell properties and clearing cell text.

    Paragraph content becomes a single empty paragraph per cell; nested
    tables inside cells are dropped (the synthesized row must be empty).
    """
    out: list[bytes] = []
    cursor = 0
    skip_depth = 0
    skip_name: str | None = None
    for tag in iter_tags(row_xml):
        token = tag.bytes_in(row_xml)
        closing, self_closing = tag.closing, tag.self_closing
        name = tag.name
        if skip_name is not None:
            cursor = tag.end
            if closing and name == skip_name:
                skip_name = None
            continue
        if not skip_depth:
            out.append(row_xml[cursor:tag.start])
        cursor = tag.end
        if name == "p" and not closing and not self_closing:
            out.append(b"<w:p/>")
            skip_depth = 1
            continue
        if name == "p" and closing and skip_depth:
            skip_depth -= 1
            continue
        if name == "tbl" and not closing and not self_closing:
            skip_name = "tbl"
            continue
        if skip_depth:
            continue
        out.append(token)
    if not skip_depth and skip_name is None:
        out.append(row_xml[cursor:])
    return b"".join(out)


def _write_patched_docx(template: Path, output: Path, document_xml: bytes, part_render: dict[str, bytes] | None = None) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    part_render = part_render or {}
    with zipfile.ZipFile(template) as source_zip, zipfile.ZipFile(output, "w") as output_zip:
        for info in source_zip.infolist():
            part_key = None
            match = PART_KEYS_PATTERN.match(info.filename)
            if match:
                part_key = match.group(1)
            if info.filename == "word/document.xml":
                data = document_xml
            elif info.filename in part_render:
                data = part_render[info.filename]  # full-path key (e.g. comments)
            elif part_key is not None and part_key in part_render:
                data = part_render[part_key]
            else:
                data = source_zip.read(info.filename)
            output_zip.writestr(info, data)



def validate_workdir(path: str | Path) -> ValidatedWorkdir:
    workdir, format_data, styles, typed, template = _load_workdir(path)
    if sha256_file(workdir / "styles.json") != format_data.get("styles_sha256"):
        raise ValidationError("styles.json changed after extract")
    if sha256_file(template) != format_data.get("template_sha256"):
        raise ValidationError("template fingerprint changed after extract")
    source_value = str(format_data.get("source_path", ""))
    if source_value:
        source_ref = Path(source_value)
        source_path = source_ref if source_ref.is_absolute() else workdir / source_ref
        if source_path.exists() and sha256_file(source_path) != format_data.get("source_sha256"):
            raise ValidationError("source fingerprint changed after extract")
    current_manifest = zip_manifest(template)
    if current_manifest != format_data.get("package_manifest"):
        raise ValidationError("template package manifest changed after extract")
    with zipfile.ZipFile(template) as archive:
        parsed = parse_package_document(archive)
        template_xml = archive.read("word/document.xml")
        part_xmls = {
            match.group(1): archive.read(name)
            for name in archive.namelist()
            if (match := PART_KEYS_PATTERN.match(name))
        }
    if sha256_bytes(template_xml) != format_data.get("document_xml_sha256"):
        raise ValidationError("template document.xml fingerprint changed after extract")
    if part_xmls and format_data.get("parts") != {
        part_key: sha256_bytes(part_xmls[part_key]) for part_key in sorted(part_xmls)
    }:
        raise ValidationError("template part fingerprints changed after extract")
    if set(parsed.styles.styles) != set(styles.styles):
        raise ValidationError("style registry does not match template styles")
    for style_id, style in parsed.styles.styles.items():
        if styles.styles[style_id].canonical != style.canonical:
            raise ValidationError(f"style registry differs from template: {style_id}")
    records = format_data.get("paragraphs", [])
    if len(records) != len(parsed.document.paragraphs):
        raise ValidationError("format paragraph baseline does not match template")
    baseline_by_id: dict[str, Paragraph] = {}
    for paragraph, record in zip(parsed.document.paragraphs, records):
        paragraph.paragraph_id = record["id"]
        if paragraph.paragraph_id in baseline_by_id:
            raise ValidationError(f"duplicate baseline paragraph ID: {paragraph.paragraph_id}")
        baseline_by_id[paragraph.paragraph_id] = paragraph
        if record.get("base_style") != paragraph.base_style:
            raise ValidationError(f"base style baseline differs for {paragraph.paragraph_id}")
        if record.get("insertion_style") != _paragraph_insertion_style(paragraph):
            raise ValidationError(f"insertion style baseline differs for {paragraph.paragraph_id}")
        if record.get("skeleton") != skeleton(paragraph.nodes):
            raise ValidationError(f"structure skeleton differs for {paragraph.paragraph_id}")
        if record.get("token_ids") != _token_ids(paragraph.nodes):
            raise ValidationError(f"structure token IDs differ for {paragraph.paragraph_id}")
        if bool(record.get("section_bearing")) != paragraph.section_bearing:
            raise ValidationError(f"section-bearing baseline differs for {paragraph.paragraph_id}")
        if (record.get("mark_revision") or None) != paragraph.mark_revision:
            raise ValidationError(f"paragraph-mark baseline differs for {paragraph.paragraph_id}")
    _validate_anchor_pairs(parsed.document.paragraphs, "template")
    expected_header = {
        "schema": "1",
        "format": "format.json",
        "styles": "styles.json",
        "template": "_template.docx",
    }
    for key, value in expected_header.items():
        if typed.meta.get(key) != value:
            raise ValidationError(f"typed header {key} must be {value}")
    baseline_ids = set(baseline_by_id)
    live_id_list = [paragraph.paragraph_id for paragraph in typed.paragraphs]
    if len(live_id_list) != len(set(live_id_list)):
        raise ValidationError("duplicate live paragraph ID")
    delete_id_list = list(typed.deletions)
    if len(delete_id_list) != len(set(delete_id_list)):
        raise ValidationError("duplicate paragraph deletion tombstone")
    live_ids = set(live_id_list)
    delete_ids = set(delete_id_list)
    if live_ids & delete_ids:
        raise ValidationError("paragraph cannot be both live and deleted")
    unknown_deletes = delete_ids - baseline_ids
    if unknown_deletes:
        unknown_text = ", ".join(sorted(unknown_deletes))
        raise ValidationError(f"unknown paragraph IDs in deletion tombstones: {unknown_text}")
    missing = baseline_ids - live_ids - delete_ids
    if missing:
        raise ValidationError("missing explicit deletion tombstone for: " + ", ".join(sorted(missing)))
    records_by_id = {record["id"]: record for record in records}
    live_paragraphs: list[Paragraph] = []
    warnings: list[str] = []
    for paragraph in typed.paragraphs:
        if paragraph.paragraph_id in baseline_by_id:
            baseline = baseline_by_id[paragraph.paragraph_id]
            record = records_by_id[paragraph.paragraph_id]
            if paragraph.inherit:
                raise ValidationError(f"existing paragraph cannot use inherit: {paragraph.paragraph_id}")
            if paragraph.base_style != baseline.base_style:
                raise ValidationError(f"base style changed: {paragraph.paragraph_id}")
            synced_segments = record.get("sync_segments")
            if synced_segments is None:
                if skeleton(paragraph.nodes) != record.get("skeleton"):
                    raise ValidationError(f"structure skeleton changed: {paragraph.paragraph_id}")
                if _token_ids(paragraph.nodes) != record.get("token_ids"):
                    raise ValidationError(f"structure token IDs changed: {paragraph.paragraph_id}")
                _validate_token_nodes(paragraph.nodes, format_data.get("tokens", {}))
                _validate_styles(paragraph.nodes, styles)
                _validate_cross_boundary_edit(baseline, paragraph)
            else:
                governed_segments = [(segment[1], segment[0]) for segment in synced_segments]
                if _text_segments(paragraph.nodes) != governed_segments:
                    if skeleton(paragraph.nodes) != record.get("sync_skeleton"):
                        raise ValidationError(f"structure skeleton changed: {paragraph.paragraph_id}")
                    _validate_token_nodes(paragraph.nodes, format_data.get("tokens", {}))
                    _validate_styles(paragraph.nodes, styles)
                    _validate_segment_rewrite(
                        governed_segments,
                        _text_segments(paragraph.nodes),
                        paragraph.paragraph_id,
                    )
                else:
                    _validate_token_nodes(paragraph.nodes, format_data.get("tokens", {}))
                    _validate_styles(paragraph.nodes, styles)
            if content_signature(paragraph) != content_signature(baseline) and contains_opaque(baseline.nodes):
                raise ValidationError(f"unsupported opaque paragraph was edited: {paragraph.paragraph_id}")
            paragraph.p_open = baseline.p_open
            paragraph.ppr = baseline.ppr
            paragraph.section_bearing = baseline.section_bearing
            paragraph.original_index = baseline.original_index
            paragraph.container_path = baseline.container_path
            paragraph.table_index = baseline.table_index
            paragraph.part_key = baseline.part_key
            paragraph.part_entry_id = baseline.part_entry_id
        else:
            if not paragraph.inherit or paragraph.inherit not in baseline_by_id:
                raise ValidationError(f"new paragraph requires existing inherit ID: {paragraph.paragraph_id}")
            inherited = baseline_by_id[paragraph.inherit]
            if inherited.container_path or inherited.part_key:
                raise ValidationError(
                    f"table-structure-immutable: new paragraphs cannot inherit from a "
                    f"table cell, text box, or header/footer/note paragraph "
                    f"({paragraph.inherit}); container structure operations are out of scope"
                )
            if inherited.section_bearing or _contains_structural(inherited.nodes):
                raise ValidationError(
                    f"new paragraph cannot inherit protected structure: {paragraph.paragraph_id}"
                )
            _assign_default_style(paragraph.nodes, inherited.base_style)
            if _contains_structural(paragraph.nodes):
                raise ValidationError(f"new paragraph cannot add structural tokens: {paragraph.paragraph_id}")
            _validate_styles(paragraph.nodes, styles)
            paragraph.base_style = inherited.base_style
            paragraph.p_open = _new_paragraph_opening(inherited.p_open)
            paragraph.ppr = inherited.ppr
            paragraph.section_bearing = False
            paragraph.original_index = -1
        live_paragraphs.append(paragraph)
    existing_order = [paragraph.original_index for paragraph in live_paragraphs if paragraph.original_index >= 0]
    if existing_order != sorted(existing_order):
        raise ValidationError("existing paragraph order cannot change")
    for deleted_id in typed.deletions:
        baseline = baseline_by_id[deleted_id]
        if deleted_id.startswith("comments."):
            continue  # comment entry removal is a deliberate decision action
        if baseline.section_bearing or _contains_structural(baseline.nodes):
            raise ValidationError(f"paragraph with protected structure cannot be deleted: {deleted_id}")
    _validate_anchor_pairs(live_paragraphs, "typed source")
    used_styles: set[str] = set()
    for paragraph in live_paragraphs:
        for node in paragraph.nodes:
            if isinstance(node, TextNode):
                used_styles.add(node.style_id)
    unused = sorted(set(styles.styles) - used_styles)
    if unused:
        warnings.append("unused styles: " + ", ".join(unused))
    return ValidatedWorkdir(
        workdir,
        format_data,
        styles,
        typed,
        parsed.document,
        format_data.get("tokens", {}),
        template,
        template_xml,
        parsed.slices,
        live_paragraphs,
        baseline_by_id,
        warnings,
    )


def build_workdir(path: str | Path, output: str | Path | None = None) -> Path:
    validated = validate_workdir(path)
    from .edit import require_clean_edit  # lazy: edit.py imports this module

    require_clean_edit(path)
    output_path = Path(output).resolve() if output else validated.path.parent / f"{validated.path.name}.docx"
    reserved_paths = {
        (validated.path / name).resolve()
        for name in {"_template.docx", "typed.md", "format.json", "styles.json"}
    }
    source_value = str(validated.format_data.get("source_path", ""))
    if source_value:
        source_ref = Path(source_value)
        source_path = source_ref if source_ref.is_absolute() else validated.path / source_ref
        reserved_paths.add(source_path.resolve())
    if output_path in reserved_paths:
        raise ValidationError(f"output path is reserved: {output_path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    body_paragraphs = [p for p in validated.live_paragraphs if not p.container_path and not p.part_key]
    replacements: list[bytes] = []
    for paragraph in body_paragraphs:
        if not paragraph.inherit:
            baseline = validated.baseline_by_id[paragraph.paragraph_id]
            if content_signature(paragraph) == content_signature(baseline):
                replacements.append(baseline.raw_xml.encode("utf-8"))
                continue
            replacements.append(_render_paragraph(paragraph, baseline, validated.styles, validated.format_data.get("tokens", {})))
        else:
            inherited = validated.baseline_by_id[paragraph.inherit]
            replacements.append(_render_paragraph(paragraph, inherited, validated.styles, validated.format_data.get("tokens", {})))
    box_paragraphs = [p for p in validated.live_paragraphs if p.container_path and p.container_path[0] == "box"]
    for index, paragraph in enumerate(body_paragraphs):
        boxes = [
            box for box in validated.template_slices.boxes
            if box.parent_paragraph == paragraph.original_index
        ]
        if boxes:
            replacements[index] = _render_boxes_in_paragraph(
                replacements[index],
                validated.template_slices.paragraphs[paragraph.original_index],
                boxes,
                validated.template_xml,
                box_paragraphs,
                validated.baseline_by_id,
                validated.styles,
                validated.format_data.get("tokens", {}),
            )
    slots, insert_before = _paragraph_placements(body_paragraphs, len(validated.template_slices.paragraphs))
    sdt_render = _render_sdts(
        validated.template_xml,
        validated.template_slices,
        [p for p in validated.live_paragraphs if p.container_path and p.container_path[0] == "sdt"],
        validated.baseline_by_id,
        validated.styles,
        validated.format_data.get("tokens", {}),
    )
    table_render = _render_tables(
        validated.template_xml,
        validated.template_slices,
        [p for p in validated.live_paragraphs if p.container_path],
        validated.baseline_by_id,
        validated.styles,
        validated.format_data.get("tokens", {}),
    )
    patched_xml = patch_document_xml(
        validated.template_xml,
        validated.template_slices,
        replacements,
        slots,
        insert_before,
        table_render,
        sdt_render,
    )
    part_render = _render_parts(
        validated,
        [p for p in validated.live_paragraphs if p.part_key],
    )
    temp_name = None
    try:
        with tempfile.NamedTemporaryFile(prefix=".typed-build-", suffix=".docx", dir=output_path.parent, delete=False) as temp:
            temp_name = temp.name
        temp_path = Path(temp_name)
        _write_patched_docx(validated.template_path, temp_path, patched_xml, part_render)
        package_guard(validated.template_path, temp_path, editable_parts=set(part_render))
        verify_workdir(validated.path, temp_path)
        os.replace(temp_path, output_path)
    finally:
        if temp_name and Path(temp_name).exists():
            Path(temp_name).unlink()
    return output_path


def _compare_output_paragraph(
    expected: Paragraph,
    actual: Paragraph,
    tokens: dict[str, dict[str, Any]],
) -> None:
    expected_ppr = expected.ppr
    if expected.mark_revision:
        expected_ppr = _inject_mark_revision(expected_ppr, expected.mark_revision, tokens)
    if _canonical_ppr(expected_ppr) != _canonical_ppr(actual.ppr):
        raise ValidationError(f"output paragraph properties differ: {expected.paragraph_id}")
    if _paragraph_attrs(expected.p_open) != _paragraph_attrs(actual.p_open):
        raise ValidationError(f"output paragraph attributes differ: {expected.paragraph_id}")
    if content_signature(expected) != content_signature(actual):
        raise ValidationError(f"output text or structure differs: {expected.paragraph_id}")
    # three-layer revision verification: final view, original view, structure.
    if visible_text(expected.nodes) != visible_text(actual.nodes):
        raise ValidationError(f"output final-view text differs: {expected.paragraph_id}")
    if visible_text_original(expected.nodes) != visible_text_original(actual.nodes):
        raise ValidationError(f"output original-view text differs: {expected.paragraph_id}")


def verify_workdir(path: str | Path, output: str | Path) -> None:
    validated = validate_workdir(path)
    from .edit import require_clean_edit  # lazy: edit.py imports this module

    require_clean_edit(path)
    output_path = Path(output).resolve()
    if not output_path.exists():
        raise ValidationError(f"output DOCX not found: {output_path}")
    package_guard(
        validated.template_path, output_path,
        editable_parts=set(validated.format_data.get("parts", {})) | set(_COMMENT_PARTS),
    )
    with zipfile.ZipFile(output_path) as archive:
        output_parsed = parse_package_document(archive)
    if len(output_parsed.document.paragraphs) != len(validated.live_paragraphs):
        raise ValidationError(
            f"output direct paragraph count differs: expected {len(validated.live_paragraphs)}, got {len(output_parsed.document.paragraphs)}"
        )
    expected: list[Paragraph] = []
    for paragraph in validated.live_paragraphs:
        expected.append(paragraph)
    body_slice_index = 0
    for index, (wanted, actual) in enumerate(zip(expected, output_parsed.document.paragraphs)):
        actual.paragraph_id = wanted.paragraph_id
        if wanted.container_path or wanted.part_key:
            # cell/box/part paragraphs are covered by their container byte
            # ranges and the paragraph comparison below; body slice indexing
            # does not apply
            _compare_output_paragraph(wanted, actual, validated.format_data.get("tokens", {}))
            continue
        if not wanted.inherit:
            baseline = validated.baseline_by_id[wanted.paragraph_id]
            if content_signature(wanted) == content_signature(baseline):
                boxes = [
                    box for box in validated.template_slices.boxes
                    if box.parent_paragraph == wanted.original_index
                ]
                expected_raw = baseline.raw_xml.encode("utf-8")
                actual_raw = output_parsed.slices.paragraphs[body_slice_index].raw
                if boxes:
                    parent_slice = validated.template_slices.paragraphs[wanted.original_index]
                    output_parent = output_parsed.slices.paragraphs[body_slice_index]
                    output_boxes = [
                        box for box in output_parsed.slices.boxes
                        if box.parent_paragraph == body_slice_index
                    ]

                    def strip_boxes(raw_bytes: bytes, ranges: list[tuple[int, int]]) -> bytes:
                        pieces: list[bytes] = []
                        cursor = 0
                        for rel_start, rel_end in ranges:
                            pieces.append(raw_bytes[cursor:rel_start])
                            cursor = rel_end
                        pieces.append(raw_bytes[cursor:])
                        return b"".join(pieces)

                    expected_ranges = [
                        (box.start - parent_slice.start, box.end - parent_slice.start)
                        for box in boxes
                    ]
                    actual_ranges = [
                        (box.start - output_parent.start, box.end - output_parent.start)
                        for box in output_boxes
                    ]
                    if strip_boxes(expected_raw, expected_ranges) != strip_boxes(actual_raw, actual_ranges):
                        raise ValidationError(f"untouched paragraph bytes differ: {wanted.paragraph_id}")
                elif actual_raw != expected_raw:
                    raise ValidationError(f"untouched paragraph bytes differ: {wanted.paragraph_id}")
        body_slice_index += 1
        _compare_output_paragraph(wanted, actual, validated.format_data.get("tokens", {}))

def extract(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(prog="docx2typed extract")
    parser.add_argument("input", help="source .docx")
    parser.add_argument("-o", "--outdir", default=".", help="typed workdir")
    args = parser.parse_args(argv)
    try:
        workdir = extract_workdir(args.input, args.outdir)
    except (OSError, zipfile.BadZipFile, TypedError) as exc:
        print(f"ERROR: {exc}")
        return 1
    print(f"workdir:  {workdir}")
    print(f"typed:    {workdir / 'typed.md'}")
    print(f"format:   {workdir / 'format.json'}")
    print(f"styles:   {workdir / 'styles.json'}")
    print(f"template: {workdir / '_template.docx'}")
    return 0


def validate(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(prog="docx2typed validate")
    parser.add_argument("workdir", help="typed workdir")
    args = parser.parse_args(argv)
    try:
        checked = validate_workdir(args.workdir)
        from .edit import require_clean_edit  # lazy: edit.py imports this module

        require_clean_edit(args.workdir)
    except (OSError, zipfile.BadZipFile, TypedError) as exc:
        print(f"ERROR: {exc}")
        return 1
    print(f"validated: {checked.path}")
    for warning in checked.warnings:
        print(f"warning: {warning}")
    return 0


def build(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(prog="docx2typed build")
    parser.add_argument("workdir", help="typed workdir")
    parser.add_argument("-o", "--output", help="output .docx")
    args = parser.parse_args(argv)
    try:
        output = build_workdir(args.workdir, args.output)
    except (OSError, zipfile.BadZipFile, TypedError) as exc:
        print(f"ERROR: {exc}")
        return 1
    print(f"built: {output}")
    return 0


def verify(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(prog="docx2typed verify")
    parser.add_argument("workdir", help="typed workdir")
    parser.add_argument("output", help="built .docx")
    args = parser.parse_args(argv)
    try:
        verify_workdir(args.workdir, args.output)
    except (OSError, zipfile.BadZipFile, TypedError) as exc:
        print(f"ERROR: {exc}")
        return 1
    print("verified: typed source, XML structure, and package integrity")
    return 0

"""Run the visible DOCX acceptance workflow and leave inspectable artifacts."""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import platform
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Callable

from docx import Document

try:
    from .create_complex_fixture import create_fixture
    from .edit import refresh_edit_projection
    from .report_html import write_report_html
    from .typed_core import TypedError
    from .typed_docx import build_workdir, extract_workdir, verify_workdir
except ImportError:
    from create_complex_fixture import create_fixture
    from edit import refresh_edit_projection
    from report_html import write_report_html
    from typed_core import TypedError
    from typed_docx import build_workdir, extract_workdir, verify_workdir


class AcceptanceFailure(RuntimeError):
    """The user-facing acceptance run did not pass."""


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _zip_manifest(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {name: _sha256(archive.read(name)) for name in sorted(archive.namelist())}


def _protected_manifest(manifest: dict[str, str]) -> dict[str, str]:
    return {name: digest for name, digest in manifest.items() if name != "word/document.xml"}


def _replace_zip_member(path: Path, member: str, replace: Callable[[bytes], bytes]) -> None:
    with zipfile.ZipFile(path) as source:
        files = {info.filename: source.read(info.filename) for info in source.infolist()}
    files[member] = replace(files[member])
    with tempfile.NamedTemporaryFile(prefix="acceptance-", suffix=".docx", dir=path.parent, delete=False) as temp:
        temp_path = Path(temp.name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for name, data in files.items():
                target.writestr(name, data)
        os.replace(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _environment() -> dict[str, Any]:
    try:
        import docx

        python_docx = getattr(docx, "__version__", "unknown")
    except Exception:
        python_docx = "unavailable"
    officecli = shutil.which("officecli")
    try:
        officecli_version = (
            subprocess.run(
                [officecli, "--version"],
                check=True,
                capture_output=True,
                text=True,
            ).stdout.strip()
            if officecli
            else None
        )
    except (OSError, subprocess.CalledProcessError):
        officecli_version = "unavailable"
    repo = Path(__file__).resolve().parents[1]
    try:
        commit = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=repo,
            check=True,
            capture_output=True,
            text=True,
        ).stdout.strip()
    except (OSError, subprocess.CalledProcessError):
        commit = "unknown"
    return {
        "python": sys.version,
        "platform": platform.platform(),
        "python_docx": python_docx,
        "engine_commit": commit,
        "officecli": officecli,
        "officecli_version": officecli_version,
        "soffice": shutil.which("soffice"),
        "pdftoppm": shutil.which("pdftoppm"),
    }


def _record(checks: list[dict[str, str]], name: str, condition: bool, evidence: str) -> None:
    status = "PASS" if condition else "FAIL"
    checks.append({"name": name, "status": status, "evidence": evidence})
    if not condition:
        raise AcceptanceFailure(f"{name}: {evidence}")


def _build_and_verify(workdir: Path, output: Path) -> Path:
    built = build_workdir(workdir, output)
    verify_workdir(workdir, built)
    return built


def _paragraph_text(path: Path, marker: str) -> str:
    document = Document(path)
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            return paragraph.text
    raise AssertionError(f"marker not found in {path.name}: {marker}")

def _render_with_officecli(documents: dict[str, Path], output_dir: Path) -> dict[str, str]:
    officecli = shutil.which("officecli")
    if not officecli:
        raise AcceptanceFailure("officecli is required for --visual")
    rendered_dir = output_dir / "officecli"
    rendered_dir.mkdir()
    rendered: dict[str, str] = {}
    for label, document in documents.items():
        screenshot = rendered_dir / f"{label}.png"
        subprocess.run(
            [
                officecli,
                "view",
                str(document),
                "screenshot",
                "--render",
                "html",
                "--grid",
                "auto",
                "-o",
                str(screenshot),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        rendered[label] = str(screenshot)
    return rendered


def run_acceptance(output: str | Path, *, visual: bool = False) -> dict[str, Any]:
    output_dir = Path(output).resolve()
    if output_dir.exists():
        raise FileExistsError(f"output already exists; choose another path: {output_dir}")
    output_dir.mkdir(parents=True)
    checks: list[dict[str, str]] = []
    artifacts: dict[str, str] = {}
    result = "FAIL"
    try:
        original = output_dir / "original.docx"
        create_fixture(original)
        artifacts["original"] = str(original)
        baseline_manifest = _zip_manifest(original)

        workdir = output_dir / "typed-workdir"
        extract_workdir(original, workdir)
        typed_path = workdir / "typed.md"
        typed_source = typed_path.read_text(encoding="utf-8")
        _record(
            checks,
            "extract exposes editable and opaque structures",
            "EDIT-ME:" in typed_source and "docx-opaque" in typed_source and "docx-inline" in typed_source,
            "typed.md contains the edit target plus structural tokens",
        )

        noop = _build_and_verify(workdir, output_dir / "00-noop.docx")
        artifacts["noop"] = str(noop)
        _record(
            checks,
            "no-op round trip",
            _protected_manifest(_zip_manifest(noop)) == _protected_manifest(baseline_manifest),
            "all non-document package parts match the original",
        )

        current = typed_source
        typed_outputs: list[Path] = []
        for index, (old, new) in enumerate(
            (
                ("EDIT-ME:", "EDIT-ONE:"),
                ("EDIT-ONE:", "EDIT-TWO:"),
                ("EDIT-TWO:", "EDIT-THREE:"),
            ),
            start=1,
        ):
            if old not in current:
                raise AcceptanceFailure(f"missing repeated-edit target: {old}")
            current = current.replace(old, new, 1)
            if index == 2:
                current = current.replace("bold", "加粗文本", 1)
            if index == 3:
                current = current.replace("line break", "人工改过的换行文本", 1)
            typed_path.write_text(current, encoding="utf-8")
            refresh_edit_projection(workdir)
            edited = _build_and_verify(workdir, output_dir / f"0{index}-typed-edit.docx")
            typed_outputs.append(edited)
            artifacts[f"typed_edit_{index}"] = str(edited)
            visible = _paragraph_text(edited, new)
            _record(
                checks,
                f"typed edit cycle {index}",
                new in visible,
                f"{edited.name} contains {new!r} after extract/build/verify",
            )
            _record(
                checks,
                f"protected package after typed edit {index}",
                _protected_manifest(_zip_manifest(edited)) == _protected_manifest(baseline_manifest),
                "non-document package parts remain unchanged",
            )

        manual = output_dir / "04-manual-edited.docx"
        manual_document = Document(typed_outputs[-1])
        manual_changed = False
        for paragraph in manual_document.paragraphs:
            for run in paragraph.runs:
                if "EDIT-THREE:" in run.text:
                    run.text = run.text.replace("EDIT-THREE:", "MANUAL-DOCX:", 1)
                    manual_changed = True
        if not manual_changed:
            raise AcceptanceFailure("could not create the simulated manual DOCX edit")
        manual_document.save(manual)
        artifacts["manual_edited"] = str(manual)

        manual_workdir = output_dir / "manual-workdir"
        extract_workdir(manual, manual_workdir)
        manual_roundtrip = _build_and_verify(manual_workdir, output_dir / "05-manual-roundtrip.docx")
        artifacts["manual_roundtrip"] = str(manual_roundtrip)
        _record(
            checks,
            "manual DOCX edit followed by re-extraction",
            "MANUAL-DOCX:" in _paragraph_text(manual_roundtrip, "MANUAL-DOCX:"),
            "a DOCX save was re-extracted, rebuilt, and independently verified",
        )

        malformed = output_dir / "malformed-workdir"
        shutil.copytree(workdir, malformed)
        malformed_typed = malformed / "typed.md"
        malformed_typed.write_text(
            malformed_typed.read_text(encoding="utf-8").replace("</span>", "", 1),
            encoding="utf-8",
        )
        try:
            build_workdir(malformed, output_dir / "malformed.docx")
        except TypedError as exc:
            _record(checks, "malformed typed source is rejected", True, str(exc))
        else:
            _record(checks, "malformed typed source is rejected", False, "build unexpectedly succeeded")

        drift = output_dir / "drift-workdir"
        shutil.copytree(workdir, drift)
        _replace_zip_member(
            drift / "_template.docx",
            "word/document.xml",
            lambda data: data.replace(b"EDIT-ME:", b"DRIFT-ME:", 1),
        )
        try:
            build_workdir(drift, output_dir / "drift.docx")
        except TypedError as exc:
            _record(checks, "template drift is rejected", True, str(exc))
        else:
            _record(checks, "template drift is rejected", False, "build unexpectedly succeeded")

        if visual:
            rendered = _render_with_officecli(
                {
                    "original": original,
                    "typed-edit-3": typed_outputs[-1],
                    "manual-roundtrip": manual_roundtrip,
                },
                output_dir,
            )
            artifacts.update({f"visual_{label}": path for label, path in rendered.items()})
            _record(
                checks,
                "officecli visual render",
                all(Path(path).exists() for path in rendered.values()),
                "screenshots were generated for original, typed edit, and manual round-trip",
            )

        result = "PASS"
    except Exception as exc:
        checks.append({"name": "acceptance runner", "status": "FAIL", "evidence": f"{type(exc).__name__}: {exc}"})
        failure = exc
    else:
        failure = None
    report = {
        "result": result,
        "environment": _environment(),
        "checks": checks,
        "artifacts": artifacts,
        "manual_workflow": "The manual DOCX step uses python-docx to simulate an ordinary save. Open manual-edited.docx and manual-roundtrip.docx in Word; use --visual for officecli screenshots.",
    }
    report_path = output_dir / "report.json"
    report_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    if failure is not None:
        raise failure
    if visual:
        html_report = write_report_html(output_dir, report)
        report["artifacts"]["html_report"] = str(html_report)
        report_path.write_text(
            json.dumps(report, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
    return report


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "-o",
        "--output",
        default="tests/fixtures/complex-docx/acceptance-output",
        help="new directory for DOCX artifacts and report",
    )
    parser.add_argument(
        "--visual",
        action="store_true",
        help="render original and edited DOCX files to officecli screenshots",
    )
    args = parser.parse_args(argv)
    try:
        report = run_acceptance(args.output, visual=args.visual)
    except Exception as exc:
        print(f"FAIL: {exc}")
        return 1
    print(f"PASS: {len(report['checks'])} acceptance checks")
    print(f"artifacts: {Path(args.output).resolve()}")
    print("open original.docx, 03-typed-edit.docx, manual-edited.docx, and manual-roundtrip.docx in Word")
    if args.visual:
        print("officecli screenshots: " + str(Path(args.output).resolve() / "officecli"))
        print("self-contained HTML report: " + str(Path(args.output).resolve() / "report.html"))
    elif not report["environment"]["officecli"]:
        print("visual render: skipped (officecli not installed)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

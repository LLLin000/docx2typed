"""Create the real Word fixture used by the complex DOCX integration test."""
from __future__ import annotations

import argparse
import base64
import io
import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
FOOTNOTES_REL = f"{R_NS}/footnotes"
ENDNOTES_REL = f"{R_NS}/endnotes"
FOOTNOTES_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
ENDNOTES_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"

# A tiny valid PNG is enough to exercise the DOCX image relationship and drawing XML.
PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_border(cell, **edges: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge, value in edges.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), value)
        element.set(qn("w:color"), "4F81BD")


def _add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def _add_hyperlink(document: Document, paragraph, text: str, target: str) -> None:
    relationship_id = document.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_bookmark(paragraph, run, bookmark_id: str, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    run._r.addprevious(start)
    run._r.addnext(end)


def _add_symbol_run(paragraph) -> None:
    run = paragraph.add_run()
    symbol = OxmlElement("w:sym")
    symbol.set(qn("w:font"), "Wingdings")
    symbol.set(qn("w:char"), "F0FC")
    run._r.append(symbol)


def _add_special_run(paragraph, tag: str, text: str | None = None, **attrs: str) -> None:
    run = paragraph.add_run()
    element = OxmlElement(f"w:{tag}")
    for key, value in attrs.items():
        element.set(qn(f"w:{key}"), value)
    if text is not None:
        element.text = text
    run._r.append(element)


def _add_opaque_paragraphs(document: Document) -> None:
    proof = document.add_paragraph("PROOFERR-SLOT")
    proof_error = OxmlElement("w:proofErr")
    proof_error.set(qn("w:type"), "spellStart")
    proof._p.append(proof_error)

    tracked = document.add_paragraph("TRACKED-CHANGE-SLOT")
    insertion = OxmlElement("w:ins")
    insertion.set(qn("w:id"), "99")
    insertion.set(qn("w:author"), "fixture")
    insertion.set(qn("w:date"), "2026-08-05T00:00:00Z")
    tracked_run = OxmlElement("w:r")
    tracked_text = OxmlElement("w:t")
    tracked_text.text = " inserted revision "
    tracked_run.append(tracked_text)
    insertion.append(tracked_run)
    tracked._p.append(insertion)

    field = document.add_paragraph("FIELD-SLOT")
    field_simple = OxmlElement("w:fldSimple")
    field_simple.set(qn("w:instr"), " PAGE ")
    field_run = OxmlElement("w:r")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_run.append(field_text)
    field_simple.append(field_run)
    field._p.append(field_simple)

    document.add_paragraph("FOOTNOTE-SLOT")
    document.add_paragraph("ENDNOTE-SLOT")
    document.add_paragraph("MATH-SLOT")
    document.add_paragraph("SDT-SLOT")


def _add_table_content(document: Document) -> None:
    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["Element", "Value", "Formatting", "Status"]
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = text
        _set_cell_shading(cell, "D9EAF7")
        _set_cell_border(cell, top="12", bottom="12", left="12", right="12")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rows = [
        ("Unicode", "² / ₂ / ±", "superscript + subscript", "preserve"),
        ("XML", "& < >", "escaped text", "round-trip"),
        ("Opaque", "field / drawing / revision", "locked", "diagnostic"),
    ]
    for row_index, row in enumerate(rows, start=1):
        for col_index, text in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = text
            _set_cell_border(cell, top="6", bottom="6", left="6", right="6")
    nested = table.cell(3, 0).add_table(rows=2, cols=2)
    nested.style = "Table Grid"
    nested.cell(0, 0).text = "nested"
    nested.cell(0, 1).text = "table"
    nested.cell(1, 0).text = "inside"
    nested.cell(1, 1).text = "cell"


def _add_math_and_structural_body_nodes(document_xml: str) -> str:
    math = (
        f'<m:oMathPara xmlns:m="{M_NS}"><m:oMath>'
        '<m:r><m:t>x</m:t></m:r><m:r><m:t>+</m:t></m:r>'
        '<m:r><m:t>y</m:t></m:r>'
        "</m:oMath></m:oMathPara>"
    )
    custom_xml = (
        f'<w:customXml xmlns:w="{W_NS}" w:element="fixture" '
        'w:uri="urn:docx2typed:fixture"><w:p><w:r><w:t>custom XML</w:t>'
        "</w:r></w:p></w:customXml>"
    )
    sdt = (
        f'<w:sdt xmlns:w="{W_NS}"><w:sdtPr><w:alias w:val="fixture content"/>'
        '</w:sdtPr><w:sdtContent><w:p><w:r><w:t>structured document tag</w:t>'
        "</w:r></w:p></w:sdtContent></w:sdt>"
    )
    insertion = math + custom_xml + sdt
    section_marker = "<w:sectPr"
    position = document_xml.rfind(section_marker)
    if position < 0:
        raise ValueError("document.xml has no section properties")
    return document_xml[:position] + insertion + document_xml[position:]


def _add_reference_part(
    files: dict[str, bytes],
    *,
    marker: str,
    reference_tag: str,
    part_name: str,
    relationship_type: str,
    content_type: str,
    content: str,
) -> None:
    document_xml = files["word/document.xml"].decode("utf-8")
    text_position = document_xml.find(f">{marker}</w:t>")
    if text_position < 0:
        raise ValueError(f"missing marker: {marker}")
    paragraph_end = document_xml.find("</w:p>", text_position)
    if paragraph_end < 0:
        raise ValueError(f"missing paragraph for marker: {marker}")
    reference = f'<w:r><w:{reference_tag} w:id="1"/></w:r>'
    document_xml = document_xml[:paragraph_end] + reference + document_xml[paragraph_end:]
    files["word/document.xml"] = document_xml.encode("utf-8")
    files[part_name] = content.encode("utf-8")

    relationships = files["word/_rels/document.xml.rels"].decode("utf-8")
    ids = [int(value) for value in re.findall(r'Id="rId(\d+)"', relationships)]
    relationship_id = f"rId{max(ids, default=0) + 1}"
    relationship = (
        f'<Relationship Id="{relationship_id}" Type="{relationship_type}" '
        f'Target="{part_name.removeprefix("word/")}"/>'
    )
    relationships = relationships.replace("</Relationships>", relationship + "</Relationships>")
    files["word/_rels/document.xml.rels"] = relationships.encode("utf-8")

    content_types = files["[Content_Types].xml"].decode("utf-8")
    override = f'<Override PartName="/{part_name}" ContentType="{content_type}"/>'
    content_types = content_types.replace("</Types>", override + "</Types>")
    files["[Content_Types].xml"] = content_types.encode("utf-8")


def _patch_package(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    files["word/document.xml"] = _add_math_and_structural_body_nodes(
        files["word/document.xml"].decode("utf-8")
    ).encode("utf-8")
    _add_reference_part(
        files,
        marker="FOOTNOTE-SLOT",
        reference_tag="footnoteReference",
        part_name="word/footnotes.xml",
        relationship_type=FOOTNOTES_REL,
        content_type=FOOTNOTES_CONTENT_TYPE,
        content=(
            f'<w:footnotes xmlns:w="{W_NS}">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/>'
            '</w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
            '<w:continuationSeparator/></w:r></w:p></w:footnote>'
            '<w:footnote w:id="1"><w:p><w:r><w:t>Fixture footnote</w:t>'
            "</w:r></w:p></w:footnote></w:footnotes>"
        ),
    )
    _add_reference_part(
        files,
        marker="ENDNOTE-SLOT",
        reference_tag="endnoteReference",
        part_name="word/endnotes.xml",
        relationship_type=ENDNOTES_REL,
        content_type=ENDNOTES_CONTENT_TYPE,
        content=(
            f'<w:endnotes xmlns:w="{W_NS}">'
            '<w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/>'
            '</w:r></w:p></w:endnote>'
            '<w:endnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
            '<w:continuationSeparator/></w:r></w:p></w:endnote>'
            '<w:endnote w:id="1"><w:p><w:r><w:t>Fixture endnote</w:t>'
            "</w:r></w:p></w:endnote></w:endnotes>"
        ),
    )
    with tempfile.NamedTemporaryFile(prefix="complex-fixture-", suffix=".docx", dir=path.parent, delete=False) as temp:
        temp_path = Path(temp.name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for name, data in files.items():
                archive.writestr(name, data)
        os.replace(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def create_fixture(output: str | Path) -> Path:
    output_path = Path(output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "docx2typed complex DOCX fixture"
    document.core_properties.subject = "Supported and opaque Word structures"
    document.core_properties.author = "docx2typed"
    document.core_properties.keywords = "DOCX, fixture, formatting, structure"

    body_style = document.styles.add_style("Fixture Body", WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = "Aptos"
    body_style.font.size = Pt(10.5)
    body_style.paragraph_format.space_after = Pt(6)
    code_style = document.styles.add_style("Fixture Code", WD_STYLE_TYPE.CHARACTER)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(31, 78, 121)

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.different_first_page_header_footer = True
    first_header = section.first_page_header.paragraphs[0]
    first_header.text = "docx2typed / COMPLEX FIXTURE"
    first_header.runs[0].bold = True
    header = section.header.paragraphs[0]
    header.text = "docx2typed / complex Word structures"
    _add_field(header, "PAGE")
    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    footer.add_run("Page ")
    _add_field(footer, "PAGE")
    footer.add_run(" of ")
    _add_field(footer, "NUMPAGES")

    document.add_paragraph("docx2typed Complex DOCX Fixture", style="Title")
    document.add_paragraph("A real Word package covering editable and preserved structures", style="Subtitle")
    document.add_paragraph("1. Formatting and inline structure", style="Heading 1")

    editable = document.add_paragraph(style="Fixture Body")
    editable.add_run("EDIT-ME: ")
    bold = editable.add_run("bold")
    bold.bold = True
    italic = editable.add_run(" italic")
    italic.italic = True
    underline = editable.add_run(" underline")
    underline.underline = True
    strike = editable.add_run(" strike")
    strike.font.strike = True
    double_strike = editable.add_run(" double-strike")
    double_strike.font.double_strike = True
    all_caps = editable.add_run(" all-caps")
    all_caps.font.all_caps = True
    small_caps = editable.add_run(" small-caps")
    small_caps.font.small_caps = True
    shadow = editable.add_run(" shadow")
    shadow.font.shadow = True
    color = editable.add_run(" color")
    color.font.color.rgb = RGBColor(192, 0, 0)
    highlight = editable.add_run(" highlight")
    highlight.font.highlight_color = WD_COLOR_INDEX.YELLOW
    superscript = editable.add_run(" superscript²")
    superscript.font.superscript = True
    subscript = editable.add_run(" subscript₂")
    subscript.font.subscript = True
    code = editable.add_run(" code")
    code.style = code_style
    code.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    escaped = editable.add_run(" XML & <angle> and  double  spaces")
    escaped.font.name = "Times New Roman"
    editable.add_run().add_tab()
    editable.add_run("tab")
    editable.add_run().add_break(WD_BREAK.LINE)
    editable.add_run("line break")
    _add_special_run(editable, "noBreakHyphen")
    _add_special_run(editable, "softHyphen")
    _add_symbol_run(editable)
    bookmarked = editable.add_run(" bookmark-target")
    _add_bookmark(editable, bookmarked, "42", "ComplexFixtureBookmark")
    comment_run = editable.add_run(" comment-target")
    document.add_comment(comment_run, text="Fixture comment with a real comments.xml part.", author="docx2typed")
    _add_hyperlink(document, editable, " external hyperlink", "https://example.com/docx2typed-fixture")

    image_paragraph = document.add_paragraph("Inline drawing/image: ", style="Fixture Body")
    image_paragraph.add_run().add_picture(io.BytesIO(PNG_1X1), width=Inches(0.35))
    image_paragraph.add_run(" image relationship preserved")

    document.add_paragraph("2. Lists, breaks, and tables", style="Heading 1")
    document.add_paragraph("Bullet item with direct formatting", style="List Bullet")
    document.add_paragraph("Numbered item with a tab\tinside", style="List Number")
    page_break = document.add_paragraph("Page-break paragraph", style="Fixture Body")
    page_break.add_run().add_break(WD_BREAK.PAGE)
    page_break.add_run("after page break")
    _add_table_content(document)

    document.add_paragraph("3. Headers, sections, and fields", style="Heading 1")
    _add_opaque_paragraphs(document)
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Inches(0.5)
    new_section.right_margin = Inches(0.5)
    new_section.header.paragraphs[0].text = "Landscape section header"
    new_section.footer.paragraphs[0].text = "Landscape footer / "
    _add_field(new_section.footer.paragraphs[0], "PAGE")
    document.add_paragraph("Second section with landscape geometry", style="Heading 1")
    second_section_body = document.add_paragraph(style="Fixture Body")
    second_section_body.add_run("The section boundary and sectPr are part of the protected baseline.")
    second_section_body.add_run(" ²").font.superscript = True

    document.save(output_path)
    _patch_package(output_path)
    return output_path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output", help="output DOCX path")
    args = parser.parse_args(argv)
    path = create_fixture(args.output)
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

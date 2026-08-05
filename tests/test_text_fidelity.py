from docx import Document

from build import build
from extract import extract
from verify import verify
from view import view_workdir


def test_literal_xml_characters_and_spaces_round_trip(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("  A & B < C > D  ")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    raw = (workdir / "typed.md").read_text(encoding="utf-8")
    assert "&amp;" in raw and "&lt;" in raw and "&gt;" in raw
    assert view_workdir(workdir, "clean", markers=False) == "  A & B < C > D  "
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    assert Document(output).paragraphs[0].text == "  A & B < C > D  "

from zipfile import ZipFile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from build import build
from extract import extract
from verify import verify


def test_edit_preserves_hyperlink_and_comment_anchors(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph("前")
    relationship_id = document.part.relate_to("https://example.com", RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "链接"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    commented_run = paragraph.add_run("批注")
    document.add_comment(commented_run, text="保留锚点", author="tester")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    typed_path = workdir / "typed.md"
    typed = typed_path.read_text(encoding="utf-8")
    assert "hyperlink" in typed and "comment-start" in typed and "comment-end" in typed
    typed_path.write_text(typed.replace("前", "后").replace("批注", "注释"), encoding="utf-8")

    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        assert "w:hyperlink" in xml
        assert "w:commentRangeStart" in xml and "w:commentRangeEnd" in xml

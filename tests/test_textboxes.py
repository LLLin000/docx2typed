"""N2: text box (txbxContent) paragraph editing — v2 container slice."""
from __future__ import annotations

import json
import zipfile
from pathlib import Path

from lxml import etree
from docx import Document

from scripts.build import build
from scripts.edit import sync_edit_projection
from scripts.extract import extract
from scripts.typed_core import parse_typed
from scripts.verify import verify

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_V = "urn:schemas-microsoft-com:vml"


def make_box_docx(path: Path, *, two_boxes: bool = False) -> None:
    document = Document()
    document.add_paragraph("正文前")

    def add_box(text: str) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        pict = etree.SubElement(run._r, f"{{{NS_W}}}pict", nsmap={"v": NS_V})
        shape = etree.SubElement(pict, f"{{{NS_V}}}shape")
        shape.set("style", "width:100pt;height:50pt")
        textbox = etree.SubElement(shape, f"{{{NS_V}}}textbox")
        txbx = etree.SubElement(textbox, f"{{{NS_W}}}txbxContent")
        tp = etree.SubElement(txbx, f"{{{NS_W}}}p")
        tr = etree.SubElement(tp, f"{{{NS_W}}}r")
        tt = etree.SubElement(tr, f"{{{NS_W}}}t")
        tt.text = text

    add_box("框内文字")
    if two_boxes:
        add_box("第二个框")
    document.add_paragraph("正文后")
    document.save(path)


def extract_box(tmp_path: Path, *, two_boxes: bool = False) -> tuple[Path, Path]:
    source = tmp_path / "box-src.docx"
    workdir = tmp_path / "box"
    make_box_docx(source, two_boxes=two_boxes)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return workdir, source


def _edit_box(workdir: Path, box_id: str, old: str, new: str) -> None:
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = f'<!--@p id="{box_id}"-->'
    start = text.index(marker)
    end = text.find("\n\n", start)
    if end < 0:
        end = len(text)
    block = text[start:end]
    assert old in block, f"'{old}' not in block: {block}"
    text = text[:start] + block.replace(old, new, 1) + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")


def test_extract_exposes_box_paragraphs(tmp_path):
    workdir, _ = extract_box(tmp_path, two_boxes=True)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    ids = [p.paragraph_id for p in typed.paragraphs]
    assert "B0.P0" in ids
    assert "B1.P0" in ids
    # the box-bearing body paragraph stays editable (pict is opaque but the
    # paragraph lock only applies to its own text)
    assert "B0.P0" in (workdir / "edit.md").read_text(encoding="utf-8")


def test_noop_roundtrip_byte_identical(tmp_path):
    workdir, source = extract_box(tmp_path, two_boxes=True)
    output = tmp_path / "noop.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(source) as archive:
        source_xml = archive.read("word/document.xml")
    with zipfile.ZipFile(output) as archive:
        output_xml = archive.read("word/document.xml")
    assert source_xml == output_xml


def test_edit_box_paragraph_via_sync(tmp_path):
    workdir, _ = extract_box(tmp_path)
    _edit_box(workdir, "B0.P0", "框内文字", "框内文字改")
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["B0.P0"]
    output = tmp_path / "edited.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "框内文字改" in xml
    assert "正文前" in xml


def test_tracked_edit_inside_box(tmp_path):
    workdir, _ = extract_box(tmp_path)
    _edit_box(workdir, "B0.P0", "框内文字", "框内文字改")
    _, _, changed = sync_edit_projection(workdir, track=True)
    assert changed == ["B0.P0"]
    output = tmp_path / "tracked.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:ins" in xml


def test_delete_box_paragraph_rejected(tmp_path):
    workdir, _ = extract_box(tmp_path)
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = '<!--@p id="B0.P0"-->'
    start = text.index(marker)
    end = text.find("\n\n", start)
    if end < 0:
        end = len(text)
    text = text[:start] + '<!--@delete id="B0.P0"-->' + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")
    try:
        sync_edit_projection(workdir)
        raise AssertionError("expected table-structure-immutable")
    except Exception as exc:
        assert "table-structure-immutable" in str(exc)


def test_new_paragraph_inheriting_box_rejected(tmp_path):
    workdir, _ = extract_box(tmp_path)
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    text = text.replace(
        '<!--@p id="B0.P0"-->',
        '<!--@new temp="N1" inherit="B0.P0"-->\n新\n\n<!--@p id="B0.P0"-->',
        1,
    )
    (workdir / "edit.md").write_text(text, encoding="utf-8")
    try:
        sync_edit_projection(workdir)
        raise AssertionError("expected table-structure-immutable")
    except Exception as exc:
        assert "table-structure-immutable" in str(exc)

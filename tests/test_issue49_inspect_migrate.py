"""Issue #49 acceptance: Python workdir manifest, inspect, and lossless
migration path.

Exercised only through public seams: the CLI entry (``scripts.main`` and the
installed ``python -m scripts``). Four criteria:

1. ``inspect`` returns one stable readiness classification, an asset table,
   reason codes, and the permitted next action without changing source
   bytes, mtimes, or lock state.
2. ``migrate SOURCE --out TARGET`` publishes only after asset closure,
   semantic-state, and observable-behavior verification; failure publishes
   no normal TARGET.
3. Template bytes, baseline-scoped identities, non-clean state,
   revisions/comments/review state, and optional-asset presence remain
   equivalent.
4. Unknown required features fail closed; explicitly classified Opaque
   attachments remain read-only and manifest-declared.
"""
from __future__ import annotations

import contextlib
import io
import json
import os
import subprocess
import sys
import uuid
from pathlib import Path

import pytest
from docx import Document

from scripts import main
from scripts.edit import classify_edit_state
from scripts.extract import extract
from scripts.inspect_migrate import WORKDIR_MANIFEST_SCHEMA, inventory_sha256
from scripts.protocol import (
    canonical_operation_input,
    derived_workdir_manifest,
    result_envelope,
    schema_bundle,
    semantic_sha256,
)

ROOT = Path(__file__).resolve().parents[1]


def _op() -> str:
    return uuid.uuid4().hex


def _make_doc(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("标题：")
    paragraph.add_run("原始文本")
    document.add_paragraph("第二段")
    document.save(path)


def _extract(source: Path, outdir: Path) -> Path:
    with contextlib.redirect_stdout(io.StringIO()):
        assert extract([str(source), "-o", str(outdir)]) == 0
    return outdir


def _dirty_edit(workdir: Path) -> None:
    path = workdir / "edit.md"
    text = path.read_text(encoding="utf-8")
    marker = text.find("-->", text.find("<!--@p id="))
    tail = text.find("\n", marker) + 1
    end = text.find("\n", tail)
    if end < 0:
        end = len(text)
    text = text[:tail] + text[tail:end] + "改" + text[end:]
    path.write_text(text, encoding="utf-8")


def _set_required_features(workdir: Path, features: list[str]) -> None:
    fmt_path = workdir / "format.json"
    fmt = json.loads(fmt_path.read_text(encoding="utf-8"))
    fmt["required_features"] = features
    fmt_path.write_text(
        json.dumps(fmt, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _snapshot(root: Path) -> dict[str, tuple[int, int]]:
    """Relative path -> (bytes, mtime_ns) for every file under root."""
    return {
        str(p.relative_to(root)): (p.stat().st_size, p.stat().st_mtime_ns)
        for p in root.rglob("*")
        if p.is_file()
    }


def _cli(*args: object) -> tuple[int, str]:
    result = subprocess.run(
        [sys.executable, "-m", "scripts", *(str(a) for a in args)],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        capture_output=True,
        timeout=600,
    )
    return result.returncode, (result.stdout or result.stderr).strip()


# --------------------------------------------------------------------------
# Criterion 1: inspect readiness, asset table, reason codes, read-only
# --------------------------------------------------------------------------

def test_inspect_ready_clean_workdir(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    _make_doc(source)
    _extract(source, workdir)

    assert main(["--json", "inspect", str(workdir)]) == 0
    result = json.loads(capsys.readouterr().out)
    assert result["schema"] == "docx2typed-result-1"
    assert result["operation"] == "inspect"
    assert result["outcome"] == "success"
    assert result["diagnostics"] == []

    data = result["data"]
    assert data["readiness"] == "ready"
    assert data["next_action"] == "migrate"
    assert data["reason_codes"] == ["ok"]
    assert data["semantic_state"]["edit"]["state"] == "clean"
    assert data["semantic_state"]["template_drift"] is False
    assert data["semantic_state"]["styles_drift"] is False

    authoritative = [a for a in data["assets"] if a["kind"] == "authoritative"]
    assert {a["path"] for a in authoritative} == {
        "typed.md", "format.json", "styles.json", "_template.docx",
    }
    for asset in authoritative:
        assert asset["presence"] == "present"
        assert asset["required"] is True
        assert asset["sha256"] and len(asset["sha256"]) == 64
        assert asset["bytes"] > 0
        assert asset["mtime_ns"] > 0
    template = next(a for a in authoritative if a["path"] == "_template.docx")
    assert template["read_only"] is True  # template baseline is immutable
    assert all(a["read_only"] is False for a in authoritative if a["path"] != "_template.docx")

    optional = [a for a in data["assets"] if a["kind"] == "optional"]
    assert {a["path"] for a in optional} >= {
        "edit.md", "edit.state.json", "revisions.json", "regions.md",
    }
    assert all(a["required"] is False for a in optional)
    assert data["source_snapshot"]["files"] == sum(
        asset["presence"] == "present" for asset in data["assets"]
    )


def test_inspect_is_read_only_no_lock_or_state_files(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    _make_doc(source)
    _extract(source, workdir)

    before = _snapshot(workdir)
    before_names = {p.name for p in workdir.iterdir()}
    assert main(["--json", "inspect", str(workdir)]) == 0
    capsys.readouterr()
    assert _snapshot(workdir) == before  # bytes and mtimes untouched
    # inspect itself adds no lock/state/ledger/evidence files
    assert {p.name for p in workdir.iterdir()} == before_names


def test_inspect_missing_path_and_file_path_fail(tmp_path, capsys):
    missing = tmp_path / "missing"
    assert main(["--json", "inspect", str(missing)]) == 1
    assert json.loads(capsys.readouterr().out)["diagnostics"][0]["code"] == "workdir-not-found"

    plain = tmp_path / "plain.txt"
    plain.write_text("x", encoding="utf-8")
    assert main(["--json", "inspect", str(plain)]) == 1
    assert json.loads(capsys.readouterr().out)["diagnostics"][0]["code"] == "workdir-invalid"

    assert main(["--json", "inspect"]) == 2  # invocation error
    assert json.loads(capsys.readouterr().out)["diagnostics"][0]["code"] == "invalid-arguments"


def test_inspect_classifies_non_workdir_as_blocked(tmp_path, capsys):
    empty = tmp_path / "empty"
    empty.mkdir()
    assert main(["--json", "inspect", str(empty)]) == 0  # classification is the result
    data = json.loads(capsys.readouterr().out)["data"]
    assert data["readiness"] == "blocked"
    assert data["next_action"] == "none"
    assert "asset-closure" in data["reason_codes"]


def test_inspect_dirty_is_classified_not_blocked(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    _make_doc(source)
    _extract(source, workdir)
    _dirty_edit(workdir)

    assert main(["--json", "inspect", str(workdir)]) == 0
    data = json.loads(capsys.readouterr().out)["data"]
    assert data["readiness"] == "ready"  # non-clean state is preserved, not flattened
    assert data["semantic_state"]["edit"]["state"] == "dirty"
    assert "non-clean-edit" in data["reason_codes"]
    assert data["next_action"] == "migrate"


def test_inspect_unknown_required_feature_blocks(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    _make_doc(source)
    _extract(source, workdir)
    _set_required_features(workdir, ["future-capability"])

    assert main(["--json", "inspect", str(workdir)]) == 0
    data = json.loads(capsys.readouterr().out)["data"]
    assert data["readiness"] == "blocked"
    assert data["next_action"] == "none"
    assert "required-feature-unsupported" in data["reason_codes"]


def test_inspect_opaque_attachments_declared_read_only(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    _make_doc(source)
    _extract(source, workdir)
    (workdir / "notes.txt").write_text("opaque note", encoding="utf-8")
    (workdir / "extra").mkdir()
    (workdir / "extra" / "blob.bin").write_bytes(b"\x00\x01opaque")

    assert main(["--json", "inspect", str(workdir)]) == 0
    data = json.loads(capsys.readouterr().out)["data"]
    assert data["readiness"] == "ready"
    assert "opaque-attachment" in data["reason_codes"]
    opaque = [a for a in data["assets"] if a["kind"] == "opaque"]
    assert {a["path"] for a in opaque} == {"notes.txt", "extra/"}
    for asset in opaque:
        assert asset["read_only"] is True
        assert asset["required"] is False
        assert asset["sha256"]
    assert data["semantic_state"]["opaque_attachment_count"] == 2


def test_engine_descriptor_lists_inspect_and_migrate(tmp_path, capsys):
    assert main(["--json", "--version"]) == 0
    descriptor = json.loads(capsys.readouterr().out)
    assert "inspect" in descriptor["commands"]["finite"]
    assert "migrate" in descriptor["commands"]["finite"]


# --------------------------------------------------------------------------
# Criterion 2: migrate publishes only after verification; failure no TARGET
# --------------------------------------------------------------------------

def test_migrate_clean_publishes_manifest_backed_target(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 0
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "success"
    assert result["data"]["operation_id"]
    assert result["data"]["workdir"]["value"] == str(target.resolve())
    assert result["evidence"][0]["kind"] == "mutation"
    assert result["evidence"][0]["operation_id"] == result["data"]["operation_id"]
    assert result["evidence"][0]["payload"]["checks"][0]["status"] in ("pass", "skipped")

    assert target.is_dir()
    manifest_path = target / "workdir.manifest.json"
    assert manifest_path.is_file()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert manifest["schema"] == WORKDIR_MANIFEST_SCHEMA
    assert manifest["manifest_version"] == 1
    assert manifest["workdir_schema"] == "typed-format-1"
    assert manifest["producer"]["operation"] == "migrate"
    assert manifest["state"]["readiness"] == "ready"
    assert manifest["state"]["edit"]["state"] == "clean"
    assert set(manifest["features"]["required"]) >= {"hybrid-fidelity", "locked-structure", "typed-mode"}

    # every source asset is enumerated with a hash; the manifest declares
    # itself as the single generated asset with a truthful null self-hash
    source_paths = {
        p.relative_to(workdir).as_posix() for p in workdir.rglob("*") if p.is_file()
    }
    non_generated = {
        a["path"] for a in manifest["assets"] if a["kind"] != "generated"
    }
    assert non_generated == source_paths
    generated = [a for a in manifest["assets"] if a["kind"] == "generated"]
    assert [a["path"] for a in generated] == ["workdir.manifest.json"]
    for asset in manifest["assets"]:
        if asset["kind"] == "generated":
            continue
        assert asset["sha256"] and len(asset["sha256"]) == 64

    # the migrated workdir is a normal typed workdir: validate passes and the
    # semantic workdir manifest is unchanged by migration
    assert main(["--json", "validate", str(target)]) == 0
    capsys.readouterr()
    assert derived_workdir_manifest(workdir) == derived_workdir_manifest(target)

    # observable behavior equivalence: source and target build byte-identical DOCX
    out1, out2 = tmp_path / "a.docx", tmp_path / "b.docx"
    assert main(["--json", "build", str(workdir), "-o", str(out1), "--operation-id", _op()]) == 0
    capsys.readouterr()
    assert main(["--json", "build", str(target), "-o", str(out2), "--operation-id", _op()]) == 0
    capsys.readouterr()
    assert out1.read_bytes() == out2.read_bytes()


def test_migrate_requires_operation_id_in_json_contract(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)

    assert main(["--json", "migrate", str(workdir), "--out", str(target)]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "operation-id-required"
    assert not target.exists()


def test_migrate_failure_publishes_no_normal_target(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    _make_doc(source)
    _extract(source, workdir)
    with open(workdir / "_template.docx", "ab") as handle:  # drift: fingerprint mismatch
        handle.write(b"DRIFT")

    assert main(["--json", "inspect", str(workdir)]) == 0
    assert "source-drift" in json.loads(capsys.readouterr().out)["data"]["reason_codes"]

    target = tmp_path / "migrated"
    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "source-drift"
    assert not target.exists()  # no normal TARGET
    assert not [p for p in tmp_path.iterdir() if ".migrate-" in p.name]  # staging cleaned


def test_migrate_unknown_required_feature_fails_closed(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    _set_required_features(workdir, ["future-capability"])

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "required-feature-unsupported"
    assert not target.exists()


def test_migrate_target_already_exists_fails(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    target.mkdir()
    (target / "stray.txt").write_text("x", encoding="utf-8")

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["diagnostics"][0]["code"] == "target-already-exists"
    assert (target / "stray.txt").read_text(encoding="utf-8") == "x"  # untouched


def test_migrate_verification_failure_fails_closed(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    typed = workdir / "typed.md"
    typed.write_text(
        typed.read_text(encoding="utf-8") + '\n<docx-opaque id="bad" kind="bad">\n',
        encoding="utf-8",
    )

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "migrate-verification-failed"
    assert not target.exists()


# --------------------------------------------------------------------------
# Criterion 3: source immutable; non-clean, review, revisions, baseline
# preserved equivalently
# --------------------------------------------------------------------------

def test_migrate_never_modifies_source(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    _dirty_edit(workdir)
    (workdir / "notes.txt").write_text("keep", encoding="utf-8")
    before = _snapshot(workdir)

    assert main(["--json", "inspect", str(workdir)]) == 0
    capsys.readouterr()
    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 0
    capsys.readouterr()

    assert _snapshot(workdir) == before  # bytes and mtimes identical
    assert {p.name for p in workdir.iterdir()} - {p.name for p in target.iterdir()} == set()
    assert sorted(p.name for p in workdir.iterdir()) == sorted(
        p.name for p in target.iterdir() if p.name != "workdir.manifest.json"
    )


def test_migrate_preserves_dirty_state_equivalently(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    _dirty_edit(workdir)

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 0
    capsys.readouterr()

    assert classify_edit_state(target)["state"] == "dirty"
    manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))
    assert manifest["state"]["edit"]["state"] == "dirty"
    assert (target / "edit.md").read_bytes() == (workdir / "edit.md").read_bytes()
    # both workdirs refuse to build until the draft is applied (equivalent behavior)
    built = tmp_path / "dirty.docx"
    assert main(["--json", "build", str(workdir), "-o", str(built), "--operation-id", _op()]) == 1
    capsys.readouterr()
    assert main(["--json", "build", str(target), "-o", str(built), "--operation-id", _op()]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "failure"  # CLI build maps non-clean state to failure
    assert not built.exists()


def test_migrate_preserves_stale_and_conflict_states(tmp_path, capsys):
    source = tmp_path / "src.docx"
    stale = tmp_path / "stale"
    conflict = tmp_path / "conflict"
    _make_doc(source)
    _extract(source, stale)
    _extract(source, conflict)

    stale_typed = stale / "typed.md"
    stale_typed.write_text(
        stale_typed.read_text(encoding="utf-8").replace("原始文本", "原始文本X"),
        encoding="utf-8",
    )
    assert classify_edit_state(stale)["state"] == "stale-clean"

    conflict_typed = conflict / "typed.md"
    conflict_typed.write_text(
        conflict_typed.read_text(encoding="utf-8").replace("原始文本", "原始文本Y"),
        encoding="utf-8",
    )
    _dirty_edit(conflict)
    assert classify_edit_state(conflict)["state"] == "conflict"

    for workdir, expected in ((stale, "stale-clean"), (conflict, "conflict")):
        target = tmp_path / f"{expected}-migrated"
        assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 0
        capsys.readouterr()
        assert classify_edit_state(target)["state"] == expected
        manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))
        assert manifest["state"]["edit"]["state"] == expected


def test_migrate_preserves_review_state_and_locks(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    review = workdir / ".review"
    (review / "snapshots").mkdir(parents=True)
    (review / "session.json").write_text('{"schema": "docx2typed-review-session-1"}', encoding="utf-8")
    (review / "history.jsonl").write_text('{"event": 1}\n', encoding="utf-8")
    (review / "snapshots" / "C0.json").write_text('{"id": "C0"}', encoding="utf-8")
    (review / "writer.lock").write_text("held", encoding="utf-8")
    (review / "inbox").mkdir()
    (review / "inbox" / "e1.json").write_text('{"id": "e1"}', encoding="utf-8")

    assert main(["--json", "inspect", str(workdir)]) == 0
    data = json.loads(capsys.readouterr().out)["data"]
    assert data["readiness"] == "ready"
    review_assets = [a for a in data["assets"] if a["role"] in ("review-session", "review-lock")]
    assert {a["path"] for a in review_assets} == {
        ".review/session.json",
        ".review/history.jsonl",
        ".review/snapshots/C0.json",
        ".review/writer.lock",
        ".review/inbox/e1.json",
    }
    lock = next(a for a in review_assets if a["path"] == ".review/writer.lock")
    assert lock["read_only"] is True

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 0
    capsys.readouterr()
    for rel in (
        ".review/session.json",
        ".review/history.jsonl",
        ".review/snapshots/C0.json",
        ".review/writer.lock",
        ".review/inbox/e1.json",
    ):
        assert (target / rel).read_bytes() == (workdir / rel).read_bytes()  # review state preserved
    manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))
    declared = {a["path"] for a in manifest["assets"] if a["role"] in ("review-session", "review-lock")}
    assert declared >= {".review/session.json", ".review/writer.lock", ".review/inbox/e1.json"}


def test_migrate_preserves_revisions_comments_decisions_and_baseline(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    # revisions + audit/decision artifacts as optional assets
    (workdir / "revisions.json").write_text('{"revisions": [{"revision_key": "k1"}]}', encoding="utf-8")
    (workdir / "decisions.json").write_text('{"revision_count": 1}', encoding="utf-8")
    before_format = (workdir / "format.json").read_bytes()
    before_template = (workdir / "_template.docx").read_bytes()

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 0
    capsys.readouterr()
    assert (target / "revisions.json").read_bytes() == (workdir / "revisions.json").read_bytes()
    assert (target / "decisions.json").read_bytes() == (workdir / "decisions.json").read_bytes()
    assert (target / "_template.docx").read_bytes() == before_template  # template bytes exact
    assert (target / "format.json").read_bytes() == before_format  # baseline identities exact

    manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))
    assert manifest["baseline"]["template_sha256"] == json.loads(before_format)["template_sha256"]
    assert manifest["baseline"]["package_manifest"] == json.loads(before_format)["package_manifest"]
    assert manifest["baseline"]["source_sha256"] == json.loads(before_format)["source_sha256"]
    assert manifest["state"]["revision_count"] == 1
    roles = {a["role"] for a in manifest["assets"]}
    assert {"revisions", "decisions"} <= roles
    # template asset is declared read-only in the manifest too
    template_asset = next(a for a in manifest["assets"] if a["path"] == "_template.docx")
    assert template_asset["read_only"] is True


# --------------------------------------------------------------------------
# Criterion 4: opaque attachments read-only + manifest-declared; unknown
# required features fail closed (covered above); idempotent replay/reuse
# --------------------------------------------------------------------------

def test_migrate_opaque_attachments_preserved_and_manifest_declared(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    (workdir / "notes.txt").write_text("opaque note", encoding="utf-8")
    (workdir / "extra").mkdir()
    (workdir / "extra" / "blob.bin").write_bytes(b"\x00\x01opaque")

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 0
    capsys.readouterr()
    assert (target / "notes.txt").read_text(encoding="utf-8") == "opaque note"
    assert (target / "extra" / "blob.bin").read_bytes() == b"\x00\x01opaque"

    manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))
    opaque = [a for a in manifest["assets"] if a["kind"] == "opaque"]
    assert {a["path"] for a in opaque} == {"notes.txt", "extra/"}
    for asset in opaque:
        assert asset["read_only"] is True
        assert asset["required"] is False
        assert asset["sha256"]

    # the published evidence declares the opaque assets too
    evidence = json.loads((target.parent / f"{target.name}.migrate.evidence.json").read_text(encoding="utf-8"))
    assert evidence["payload"]["outputs"]["target"]["opaque_assets"] == 2
    assert evidence["payload"]["outputs"]["target"]["manifest_sha256"] == semantic_sha256(manifest)


def test_migrate_operation_id_replay_and_reuse(tmp_path):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()

    rc1, out1 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc1 == 0
    assert json.loads(out1)["data"]["operation_id"] == op
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns

    # identical retry replays the original Result with no second effect
    rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc2 == 0
    assert out2 == out1
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime

    # changed canonical input (different source content) is rejected
    other = tmp_path / "other"
    _make_doc(source)
    _extract(source, other)
    (other / "typed.md").write_text(
        (other / "typed.md").read_text(encoding="utf-8").replace("原始文本", "原始文本Z"),
        encoding="utf-8",
    )
    rc3, out3 = _cli("--json", "migrate", other, "--out", target, "--operation-id", op)
    assert rc3 == 1
    reused = json.loads(out3)
    assert reused["outcome"] == "failure"
    assert reused["diagnostics"][0]["code"] == "operation-id-reused"
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime  # no second effect


def test_migrate_publishes_run_evidence_sidecar(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", op]) == 0
    result = json.loads(capsys.readouterr().out)
    evidence_path = Path(str(target) + ".migrate.evidence.json")
    assert evidence_path.is_file()
    evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
    assert evidence["schema"] == "docx2typed-run-evidence-1"
    assert evidence["operation"] == "migrate"
    assert evidence["kind"] == "mutation"
    assert evidence["operation_id"] == op
    assert evidence["payload_sha256"] == semantic_sha256(evidence["payload"])
    assert result["evidence"] == [evidence]
    assert "source" in evidence["payload"]["inputs"]
    assert "target" in evidence["payload"]["outputs"]


# --------------------------------------------------------------------------
# Criterion 5: issue #49 review-finding regressions
# --------------------------------------------------------------------------

def test_inspect_and_migrate_unreadable_source_emit_workdir_unreadable(tmp_path, capsys):
    """Finding: inventory/canonical computation sits inside JSON error
    translation, so an unreadable source yields a workdir-unreadable Result
    instead of an uncaught OSError."""
    if os.name != "nt":
        pytest.skip("exclusive-share lock is Windows-specific")
    import ctypes

    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)

    kernel32 = ctypes.windll.kernel32
    GENERIC_READ = 0x80000000
    OPEN_EXISTING = 3
    FILE_ATTRIBUTE_NORMAL = 0x80
    handle = kernel32.CreateFileW(
        str(workdir / "typed.md"),
        GENERIC_READ,
        0,  # share mode 0: every other open (including reads) fails
        None,
        OPEN_EXISTING,
        FILE_ATTRIBUTE_NORMAL,
        None,
    )
    assert handle not in (-1, ctypes.c_void_p(-1).value)
    try:
        assert main(["--json", "inspect", str(workdir)]) == 1
        result = json.loads(capsys.readouterr().out)
        assert result["outcome"] == "failure"
        assert result["diagnostics"][0]["code"] == "workdir-unreadable"

        assert main(
            ["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]
        ) == 1
        result = json.loads(capsys.readouterr().out)
        assert result["outcome"] == "failure"
        assert result["diagnostics"][0]["code"] == "workdir-unreadable"
    finally:
        kernel32.CloseHandle(handle)
    assert not target.exists()


def test_migrate_reconstructs_success_after_publish_ledger_loss(tmp_path):
    """Finding: no post-publish replay gap. After the atomic publish lands but
    the success record never reaches the ledger (crash window), a retry with
    the same operation_id reconstructs the original Result from the target
    manifest + published evidence instead of failing with
    target-already-exists, with no second effect."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()

    rc1, out1 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc1 == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns

    # Simulate the crash window: publish landed, ledger record lost.
    ledger_file = Path(str(target) + ".operation-ledger.json")
    assert ledger_file.is_file()
    ledger_file.unlink()

    rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc2 == 0
    assert out2 == out1  # exact replay of the original Result
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime  # no second effect
    assert ledger_file.is_file()  # reconstruction re-persists the record
    record = json.loads(ledger_file.read_text(encoding="utf-8"))["records"][op]
    assert record["envelope"]["outcome"] == "success"

    # A changed source reusing the operation_id is still rejected.
    other = tmp_path / "other"
    _extract(source, other)
    (other / "typed.md").write_text(
        (other / "typed.md").read_text(encoding="utf-8").replace("原始文本", "原始文本R"),
        encoding="utf-8",
    )
    rc3, out3 = _cli("--json", "migrate", other, "--out", target, "--operation-id", op)
    assert rc3 == 1
    assert json.loads(out3)["diagnostics"][0]["code"] == "operation-id-reused"


def test_migrate_pending_ledger_reruns_to_completion(tmp_path):
    """Finding: pre-publish pending ledger record. A crash before the publish
    leaves an explicitly pending record (no envelope, no target); the retry
    reruns the migration to completion and upgrades the record to success."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()

    canonical = canonical_operation_input(
        "migrate",
        {
            "source": str(workdir.resolve()),
            "out": str(target.resolve()),
            "source_inventory_sha256": inventory_sha256(workdir),
        },
    )
    ledger_file = Path(str(target) + ".operation-ledger.json")
    ledger_file.parent.mkdir(parents=True, exist_ok=True)
    ledger_file.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {
                    op: {
                        "input_sha256": canonical,
                        "envelope": None,
                        "pending": True,
                    }
                },
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )

    rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc == 0
    assert json.loads(out)["outcome"] == "success"
    assert target.is_dir()
    record = json.loads(ledger_file.read_text(encoding="utf-8"))["records"][op]
    assert record["envelope"]["outcome"] == "success"

    rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc2 == 0
    assert out2 == out  # replay, no second effect


def test_migrate_pending_ledger_reconstructs_when_publish_landed(tmp_path):
    """Finding: pending record + existing matching target (crash between
    publish and record update) reconstructs success instead of rerunning."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()

    canonical = canonical_operation_input(
        "migrate",
        {
            "source": str(workdir.resolve()),
            "out": str(target.resolve()),
            "source_inventory_sha256": inventory_sha256(workdir),
        },
    )
    ledger_file = Path(str(target) + ".operation-ledger.json")
    rc1, out1 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc1 == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns

    # Downgrade the success record to an explicitly pending pre-publish
    # record, keeping the target published.
    ledger_file.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {
                    op: {
                        "input_sha256": canonical,
                        "envelope": None,
                        "pending": True,
                    }
                },
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )

    rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc2 == 0
    assert out2 == out1  # reconstructed from manifest/evidence, exact replay
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime


def test_inspect_rejects_symlinks_anywhere_under_source(tmp_path, capsys):
    """Finding: every link under the source (top level, .review, opaque
    subtrees) blocks with a stable reason/diagnostic and is never
    dereferenced. Covers real symlinks when privileges allow and directory
    junctions (reparse points) which never require privileges."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    _make_doc(source)
    _extract(source, workdir)
    outside = tmp_path / "outside.txt"
    outside.write_text("outside", encoding="utf-8")
    (workdir / "extra").mkdir()
    (workdir / "extra" / "real.bin").write_bytes(b"real")

    links: list[str] = []
    # Directory junction: a reparse point that needs no privileges.
    try:
        subprocess.run(
            ["cmd", "/c", "mklink", "/J",
             str(workdir / "extra-link"), str(workdir / "extra")],
            check=True,
            capture_output=True,
        )
        links.append("extra-link")
    except (OSError, subprocess.CalledProcessError):
        pass
    # Real symlinks (file in .review, file and dir at top level): privileged.
    try:
        (workdir / "notes.txt").symlink_to(outside)
        (workdir / ".review").mkdir()
        (workdir / ".review" / "snapshot.json").symlink_to(outside)
        (workdir / "extra" / "blob.bin").symlink_to(outside)
        links.extend(["extra/blob.bin", ".review/snapshot.json", "notes.txt"])
    except (OSError, NotImplementedError):
        pass
    if not links:
        pytest.skip("no symlink/junction creation available on this host")

    assert main(["--json", "inspect", str(workdir)]) == 0
    data = json.loads(capsys.readouterr().out)["data"]
    assert data["readiness"] == "blocked"
    assert data["next_action"] == "none"
    assert "symlink-detected" in data["reason_codes"]
    assert data["symlinks"] == sorted(links)

    target = tmp_path / "migrated"
    assert main(
        ["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]
    ) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "symlink-detected"
    assert not target.exists()  # fail closed, nothing copied through a link
    assert (workdir / "extra" / "real.bin").read_bytes() == b"real"  # source untouched


def test_opaque_dir_digests_are_source_relative(tmp_path, capsys):
    """Finding: opaque directory digests key on source-relative paths so
    same-basename descendants under different subtrees never collide, and
    migration preserves every descendant byte-for-byte."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    (workdir / "extra").mkdir()
    (workdir / "extra" / "blob.bin").write_bytes(b"AAA")
    (workdir / "extra" / "deep").mkdir()
    (workdir / "extra" / "deep" / "blob.bin").write_bytes(b"BBB")  # same basename as extra2's
    (workdir / "extra2").mkdir()
    (workdir / "extra2" / "blob.bin").write_bytes(b"BBB")

    assert main(["--json", "inspect", str(workdir)]) == 0
    data = json.loads(capsys.readouterr().out)["data"]
    opaque = {a["path"]: a for a in data["assets"] if a["kind"] == "opaque"}
    assert set(opaque) == {"extra/", "extra2/"}
    assert opaque["extra/"]["sha256"] != opaque["extra2/"]["sha256"]  # distinct identities

    assert main(
        ["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]
    ) == 0
    capsys.readouterr()
    assert (target / "extra" / "blob.bin").read_bytes() == b"AAA"
    assert (target / "extra" / "deep" / "blob.bin").read_bytes() == b"BBB"
    assert (target / "extra2" / "blob.bin").read_bytes() == b"BBB"
    manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))
    declared = {a["path"]: a for a in manifest["assets"] if a["kind"] == "opaque"}
    assert set(declared) == {"extra/", "extra2/"}
    assert declared["extra/"]["sha256"] == opaque["extra/"]["sha256"]
    assert declared["extra/"]["sha256"] != declared["extra2/"]["sha256"]


def test_manifest_self_closure_and_schema_invariants(tmp_path, capsys):
    """Findings: the manifest truthfully declares itself as a generated
    metadata asset (null self-hash) and the tightened schema encodes the
    invariant (checks required, per-asset constraints, generated exception)."""
    jsonschema = pytest.importorskip("jsonschema")

    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)

    assert main(
        ["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]
    ) == 0
    capsys.readouterr()
    manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))

    self_entries = [
        a for a in manifest["assets"]
        if a["path"] == "workdir.manifest.json"
    ]
    assert len(self_entries) == 1
    self_entry = self_entries[0]
    assert self_entry["kind"] == "generated"
    assert self_entry["role"] == "workdir-manifest"
    assert self_entry["required"] is True
    assert self_entry["read_only"] is True
    assert self_entry["presence"] == "present"
    assert self_entry["sha256"] is None  # no pretend self-hash
    assert self_entry["bytes"] is None
    assert self_entry["mtime_ns"] is None
    assert sum(a["kind"] == "generated" for a in manifest["assets"]) == 1
    assert manifest["checks"], "manifest carries the verification checks"

    schema = schema_bundle()["schemas"]["docx2typed-workdir-manifest-1"]
    jsonschema.validate(manifest, schema)  # a real manifest validates

    # missing checks is rejected
    broken = json.loads(json.dumps(manifest))
    del broken["checks"]
    with pytest.raises(jsonschema.ValidationError):
        jsonschema.validate(broken, schema)

    # pretending a self-hash on the generated entry is rejected
    forged = json.loads(json.dumps(manifest))
    forged_self = next(a for a in forged["assets"] if a["kind"] == "generated")
    forged_self["sha256"] = "0" * 64
    with pytest.raises(jsonschema.ValidationError):
        jsonschema.validate(forged, schema)

    # a real asset must carry a real 64-hex hash
    missing_hash = json.loads(json.dumps(manifest))
    real_asset = next(a for a in missing_hash["assets"] if a["kind"] != "generated")
    real_asset["sha256"] = None
    with pytest.raises(jsonschema.ValidationError):
        jsonschema.validate(missing_hash, schema)

    # a real asset must carry a non-null byte count
    no_bytes = json.loads(json.dumps(manifest))
    real_asset = next(a for a in no_bytes["assets"] if a["kind"] != "generated")
    real_asset["bytes"] = None
    with pytest.raises(jsonschema.ValidationError):
        jsonschema.validate(no_bytes, schema)


def test_migrate_reconstruct_rejects_edited_target_asset(tmp_path):
    """Finding: reconstruction validates the actual target closure. After a
    completed publish, an edited target asset must fail closed with
    target-already-exists — reconstruction must not claim success for a
    target whose bytes no longer match the manifest declaration."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()

    rc1, out1 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc1 == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns

    # Tamper a migrated asset, then lose the ledger record like the crash
    # window does.
    (target / "typed.md").write_text(
        (target / "typed.md").read_text(encoding="utf-8") + "tampered",
        encoding="utf-8",
    )
    ledger_file = Path(str(target) + ".operation-ledger.json")
    ledger_file.unlink()

    rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc2 == 1
    result = json.loads(out2)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "target-already-exists"
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime  # no second effect


def test_migrate_reconstruct_malformed_manifest_fails_closed(tmp_path):
    """Finding: a malformed/tampered manifest never reconstructs success and
    never escapes as a traceback — the retry returns a JSON Result with
    target-already-exists."""
    for label, tamper in (
        (
            "non-json",
            lambda path, wd, op: path.write_bytes(b"{not json"),
        ),
        (
            "wrong-shape",
            lambda path, wd, op: path.write_text(
                json.dumps(
                    {
                        "schema": "docx2typed-workdir-manifest-1",
                        "manifest_version": 1,
                        "producer": {
                            "engine": "docx2typed-python",
                            "version": "0",
                            "operation": "migrate",
                            "operation_id": op,
                        },
                        "source": {
                            "identity": inventory_sha256(wd),
                            "semantic_manifest_sha256": "0" * 64,
                        },
                        "state": {"semantic_manifest_sha256": "0" * 64},
                        "checks": [],
                        "assets": "not-a-list",
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            ),
        ),
    ):
        source = tmp_path / f"src-{label}.docx"
        workdir = tmp_path / f"wd-{label}"
        target = tmp_path / f"migrated-{label}"
        _make_doc(source)
        _extract(source, workdir)
        op = _op()

        assert _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)[0] == 0
        tamper(target / "workdir.manifest.json", workdir, op)
        ledger_file = Path(str(target) + ".operation-ledger.json")
        ledger_file.unlink()

        rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
        assert rc == 1
        result = json.loads(out)  # a JSON Result, not a traceback
        assert result["outcome"] == "failure"
        assert result["diagnostics"][0]["code"] == "target-already-exists"


def test_migrate_reconstruct_rebuilds_tampered_evidence(tmp_path):
    """Finding: the evidence sidecar is replayed only when it is a valid
    exact record of this publish. A tampered payload or payload_sha256 is
    ignored and deterministically rebuilt from the validated manifest: the
    rebuilt sidecar is a valid docx2typed-run-evidence-1 record whose
    payload equals the canonical payload and whose payload_sha256 covers
    it."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()

    assert _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)[0] == 0
    manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))
    evidence_path = Path(str(target) + ".migrate.evidence.json")

    for label, tamper in (
        ("payload", lambda ev: ev["payload"]["outputs"]["target"].update(opaque_assets=999)),
        ("hash", lambda ev: ev.update(payload_sha256="0" * 64)),
    ):
        evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
        tamper(evidence)
        evidence_path.write_text(
            json.dumps(evidence, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        ledger_file = Path(str(target) + ".operation-ledger.json")
        ledger_file.unlink()

        rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
        assert rc == 0, label
        result = json.loads(out)
        assert result["outcome"] == "success", label
        rebuilt = result["evidence"][0]
        assert rebuilt["schema"] == "docx2typed-run-evidence-1", label
        assert rebuilt["operation"] == "migrate", label
        assert rebuilt["outcome"] == "success", label
        assert rebuilt["kind"] == "mutation", label
        assert rebuilt["operation_id"] == op, label
        # canonical payload rebuilt from the validated manifest, not the tamper
        assert rebuilt["payload"]["outputs"]["target"]["opaque_assets"] != 999, label
        assert rebuilt["payload"]["outputs"]["target"]["manifest_sha256"] == semantic_sha256(manifest), label
        assert rebuilt["payload"]["outputs"]["target"]["assets"] == len(manifest["assets"]), label
        assert rebuilt["payload"]["checks"] == manifest["checks"], label
        assert rebuilt["payload_sha256"] == semantic_sha256(rebuilt["payload"]), label
        # the sidecar on disk was repaired to the same valid record, so a
        # later replay finds a valid sidecar instead of rebuilding again
        assert json.loads(evidence_path.read_text(encoding="utf-8"))["payload"] == rebuilt["payload"], label


def test_migrate_source_with_preexisting_manifest_is_replaced(tmp_path, capsys):
    """Finding: a pre-existing source workdir.manifest.json is excluded from
    the generated target manifest before the single generated self-entry is
    appended. The final target carries only the newly generated manifest
    (the copied stale bytes are overwritten), closure checks stay truthful,
    and the manifest never declares two self-entries."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    stale = workdir / "workdir.manifest.json"
    stale.write_text(
        json.dumps({"schema": "stale-manifest", "note": "pre-existing"}),
        encoding="utf-8",
    )

    assert main(["--json", "inspect", str(workdir)]) == 0
    data = json.loads(capsys.readouterr().out)["data"]
    assert data["readiness"] == "ready"  # the optional manifest never blocks

    assert main(["--json", "migrate", str(workdir), "--out", str(target), "--operation-id", _op()]) == 0
    capsys.readouterr()

    manifest = json.loads((target / "workdir.manifest.json").read_text(encoding="utf-8"))
    assert manifest["schema"] == WORKDIR_MANIFEST_SCHEMA  # newly generated, not the stale copy
    self_entries = [a for a in manifest["assets"] if a["path"] == "workdir.manifest.json"]
    assert len(self_entries) == 1
    assert self_entries[0]["kind"] == "generated"
    assert sum(a["kind"] == "generated" for a in manifest["assets"]) == 1
    # every other declared asset is a real source file; the stale source
    # manifest is declared nowhere
    declared = {a["path"] for a in manifest["assets"] if a["kind"] != "generated"}
    source_paths = {
        p.relative_to(workdir).as_posix() for p in workdir.rglob("*") if p.is_file()
    }
    assert declared == source_paths - {"workdir.manifest.json"}
    # semantic identity of the migrated target equals the source (the derived
    # manifest covers the workdir assets, not the generated manifest)
    assert derived_workdir_manifest(workdir) == derived_workdir_manifest(target)


def test_migrate_malformed_ledger_record_fails_closed(tmp_path):
    """Finding: malformed/corrupt persisted ledger rows never traceback and
    never replay or rerun an unproven effect. A corrupt row for the same
    operation_id fails closed with a structured ``operation-ledger-invalid``
    Result (no mutation, no reconstruction) and stays preserved in the file;
    a fully corrupt ledger FILE (not JSON) has no operation_id-scoped row and
    is treated as no record, so reconstruction still validates the actual
    target against the source."""
    for label, corrupt in (
        ("non-dict-record", lambda records, op: records.update({op: "garbage"})),
        ("missing-input", lambda records, op: records.update({op: {"envelope": None}})),
        (
            "bad-envelope",
            lambda records, op: records.update(
                {op: {"input_sha256": "0" * 64, "envelope": {"not": "a result"}}}
            ),
        ),
    ):
        source = tmp_path / f"src-{label}.docx"
        workdir = tmp_path / f"wd-{label}"
        target = tmp_path / f"migrated-{label}"
        _make_doc(source)
        _extract(source, workdir)
        op = _op()
        ledger_file = Path(str(target) + ".operation-ledger.json")
        ledger_file.parent.mkdir(parents=True, exist_ok=True)

        def _write(records):
            ledger_file.write_text(
                json.dumps(
                    {"schema": "docx2typed-operation-ledger-1", "records": records},
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                )
                + "\n",
                encoding="utf-8",
            )

        # Corrupt row + no target: fail closed, never rerun (the effect may
        # have completed; a lost pending marker must not cause a mutation).
        records = {op: {"input_sha256": "0" * 64, "envelope": None}}
        corrupt(records, op)
        _write(records)

        rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
        assert rc == 1, label
        result = json.loads(out)
        assert result["outcome"] == "failure", label  # JSON Result, never a traceback
        assert result["diagnostics"][0]["code"] == "operation-ledger-invalid", label
        assert not target.exists(), label  # no second mutation
        # the corrupt row is preserved verbatim for inspection
        assert json.loads(ledger_file.read_text(encoding="utf-8"))["records"][op] == records[op]

        # Corrupt record + existing target: fail closed, no reconstruction,
        # no second effect. (Clear the corrupt row first so a real migrate
        # can publish the target.)
        ledger_file.unlink()
        assert _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)[0] == 0
        manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns
        records = {op: {"input_sha256": "0" * 64, "envelope": None}}
        corrupt(records, op)
        _write(records)

        rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
        assert rc2 == 1, label
        assert json.loads(out2)["diagnostics"][0]["code"] == "operation-ledger-invalid", label
        assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime, label

    # A fully corrupt ledger FILE (not even JSON) with an existing target is
    # treated as no record and reconstructs after validation.
    source = tmp_path / "src-file.docx"
    workdir = tmp_path / "wd-file"
    target = tmp_path / "migrated-file"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()
    ledger_file = Path(str(target) + ".operation-ledger.json")
    assert _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)[0] == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns
    ledger_file.write_text("{not json", encoding="utf-8")

    rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc == 0
    assert json.loads(out)["outcome"] == "success"
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime


def test_migrate_reconstruct_rejects_semantic_hash_tamper(tmp_path):
    """Finding: reconstruction recomputes the source and target semantic
    manifest identities from the actual workdirs via derived_workdir_manifest
    and compares them against the stored manifest. Altering either stored
    hash (metadata-only tamper that still shape-checks as a 64-hex value)
    fails reconstruction closed."""
    for field in ("source-semantic", "state-semantic"):
        source = tmp_path / f"src-{field}.docx"
        workdir = tmp_path / f"wd-{field}"
        target = tmp_path / f"migrated-{field}"
        _make_doc(source)
        _extract(source, workdir)
        op = _op()

        assert _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)[0] == 0
        manifest_path = target / "workdir.manifest.json"
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        if field == "source-semantic":
            manifest["source"]["semantic_manifest_sha256"] = "1" * 64
        else:
            manifest["state"]["semantic_manifest_sha256"] = "1" * 64
        manifest_path.write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        Path(str(target) + ".operation-ledger.json").unlink()  # force reconstruction

        rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
        assert rc == 1, field
        result = json.loads(out)
        assert result["outcome"] == "failure", field
        assert result["diagnostics"][0]["code"] == "target-already-exists", field


def test_migrate_prepared_envelope_byte_exact_recovery(tmp_path):
    """Finding: the exact success envelope is precomputed once the final
    manifest is known and persisted as the pending ledger record BEFORE the
    atomic publish. A retry after the publish — record still pending, external
    evidence sidecar lost/tampered — returns the byte-exact original Result,
    repairs the sidecar from the stored envelope, and upgrades the record to
    complete without changing the envelope."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()
    ledger_file = Path(str(target) + ".operation-ledger.json")
    evidence_path = Path(str(target) + ".migrate.evidence.json")

    rc1, out1 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc1 == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns

    # Simulate the crash window: the record is still pending and the
    # external evidence sidecar is corrupt. The prepared envelope alone must
    # be enough for a byte-exact replay.
    canonical = canonical_operation_input(
        "migrate",
        {
            "source": str(workdir.resolve()),
            "out": str(target.resolve()),
            "source_inventory_sha256": inventory_sha256(workdir),
        },
    )
    envelope = json.loads(out1)
    ledger_file.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {
                    op: {
                        "input_sha256": canonical,
                        "envelope": envelope,
                        "pending": True,
                    }
                },
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )
    evidence_path.write_text("{not json", encoding="utf-8")  # sidecar lost/corrupt

    rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc2 == 0
    assert out2 == out1  # byte-exact original response, not a reconstruction
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime  # no second effect
    record = json.loads(ledger_file.read_text(encoding="utf-8"))["records"][op]
    assert record.get("pending") is not True  # upgraded to complete
    assert record["envelope"] == envelope  # envelope unchanged by the upgrade
    # the sidecar was repaired from the stored exact envelope
    repaired = json.loads(evidence_path.read_text(encoding="utf-8"))
    assert repaired == envelope["evidence"][0]
    assert repaired["payload_sha256"] == semantic_sha256(repaired["payload"])


def test_migrate_prepared_envelope_pending_no_target_reruns(tmp_path):
    """Finding: a pending record carrying a prepared success envelope is
    never replayed when no publish landed — the retry reruns the migration
    safely and only then returns a success envelope."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()
    canonical = canonical_operation_input(
        "migrate",
        {
            "source": str(workdir.resolve()),
            "out": str(target.resolve()),
            "source_inventory_sha256": inventory_sha256(workdir),
        },
    )
    ledger_file = Path(str(target) + ".operation-ledger.json")
    ledger_file.parent.mkdir(parents=True, exist_ok=True)
    ledger_file.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {
                    op: {
                        "input_sha256": canonical,
                        "envelope": result_envelope(
                            "migrate",
                            "success",
                            data={
                                "operation_id": op,
                                "workdir": {
                                    "kind": "absolute",
                                    "value": str(target.resolve()),
                                },
                                "manifest": {
                                    "kind": "absolute",
                                    "value": str(
                                        (target / "workdir.manifest.json").resolve()
                                    ),
                                },
                            },
                        ),
                        "pending": True,
                    }
                },
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )
    assert not target.exists()

    rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc == 0
    result = json.loads(out)
    assert result["outcome"] == "success"  # reran, never replayed a phantom success
    assert target.is_dir()
    record = json.loads(ledger_file.read_text(encoding="utf-8"))["records"][op]
    assert record["envelope"]["outcome"] == "success"
    assert record.get("pending") is not True  # upgraded to complete


def test_migrate_envelope_less_completed_row_fails_closed(tmp_path):
    """Finding: an envelope-less row that is NOT explicitly pending is
    corrupt (the pending marker could have been lost AFTER a completed
    effect). The retry must not rerun (second mutation) nor reconstruct:
    it fails closed with operation-ledger-invalid and preserves the row."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()
    canonical = canonical_operation_input(
        "migrate",
        {
            "source": str(workdir.resolve()),
            "out": str(target.resolve()),
            "source_inventory_sha256": inventory_sha256(workdir),
        },
    )
    ledger_file = Path(str(target) + ".operation-ledger.json")
    ledger_file.parent.mkdir(parents=True, exist_ok=True)

    corrupt = {"input_sha256": canonical, "envelope": None}  # no pending marker

    # No target yet: fail closed, never rerun, never create a target.
    ledger_file.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {op: corrupt},
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )
    rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc == 1
    assert json.loads(out)["diagnostics"][0]["code"] == "operation-ledger-invalid"
    assert not target.exists()  # no second mutation
    assert json.loads(ledger_file.read_text(encoding="utf-8"))["records"][op] == corrupt

    # With a published target: fail closed instead of reconstructing.
    ledger_file.unlink()
    assert _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)[0] == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns
    ledger_file.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {op: corrupt},
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )
    rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc2 == 1
    assert json.loads(out2)["diagnostics"][0]["code"] == "operation-ledger-invalid"
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime


def test_migrate_completed_envelope_missing_required_field_fails_closed(tmp_path):
    """Finding: a completed ledger row must carry a FULL
    ``docx2typed-result-1`` envelope (schema, operation, outcome, data,
    diagnostics, evidence, engine). Removing any required field makes the
    row corrupt: the retry fails closed with operation-ledger-invalid instead
    of replaying a malformed Result or rerunning the migration. The prepared
    pending envelope is validated the same way."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()
    canonical = canonical_operation_input(
        "migrate",
        {
            "source": str(workdir.resolve()),
            "out": str(target.resolve()),
            "source_inventory_sha256": inventory_sha256(workdir),
        },
    )
    ledger_file = Path(str(target) + ".operation-ledger.json")

    rc1, out1 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc1 == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns

    def _write(row):
        ledger_file.write_text(
            json.dumps(
                {
                    "schema": "docx2typed-operation-ledger-1",
                    "records": {op: row},
                },
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
            )
            + "\n",
            encoding="utf-8",
        )

    for field in ("engine", "data", "diagnostics", "operation"):
        envelope = json.loads(out1)
        del envelope[field]
        _write({"input_sha256": canonical, "envelope": envelope})
        rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
        assert rc == 1, field
        assert json.loads(out)["diagnostics"][0]["code"] == "operation-ledger-invalid", field
        assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime, field

    # The prepared pending envelope is validated the same way: a pending row
    # missing a required envelope field is corrupt too.
    envelope = json.loads(out1)
    del envelope["evidence"]
    _write({"input_sha256": canonical, "envelope": envelope, "pending": True})
    rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc == 1
    assert json.loads(out)["diagnostics"][0]["code"] == "operation-ledger-invalid"
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime


def test_migrate_sidecar_repair_failure_keeps_pending(tmp_path):
    """Finding: recovery branch A (pending record with prepared envelope).
    When the required evidence sidecar is missing/corrupt and the repair
    write fails, the replay must NOT return recovered success and must NOT
    finalize the ledger: it emits an evidence-publish-failed Result, leaves
    the record pending so the retry can repair, and never touches the target
    a second time."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()
    ledger_file = Path(str(target) + ".operation-ledger.json")
    evidence_path = Path(str(target) + ".migrate.evidence.json")

    rc1, out1 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc1 == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns
    canonical = canonical_operation_input(
        "migrate",
        {
            "source": str(workdir.resolve()),
            "out": str(target.resolve()),
            "source_inventory_sha256": inventory_sha256(workdir),
        },
    )
    envelope = json.loads(out1)

    # Crash window: record still pending with the prepared envelope, and the
    # sidecar is both corrupt AND un-repairable (a directory blocks publish).
    ledger_file.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {
                    op: {
                        "input_sha256": canonical,
                        "envelope": envelope,
                        "pending": True,
                    }
                },
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )
    evidence_path.unlink()
    evidence_path.mkdir()  # required sidecar cannot be written

    rc2, out2 = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc2 == 1
    result = json.loads(out2)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "evidence-publish-failed"
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime  # no second effect
    record = json.loads(ledger_file.read_text(encoding="utf-8"))["records"][op]
    assert record.get("pending") is True  # left pending: the retry can repair
    assert record["envelope"] == envelope  # never finalized with a failure


def test_migrate_reconstruct_sidecar_repair_failure_fails_closed(tmp_path):
    """Finding: recovery branch B (no ledger record / pending without
    envelope). Reconstruction may only return recovered success when the
    rebuilt required evidence sidecar is durably published; a failed sidecar
    write emits an evidence-publish-failed Result with no recovered success
    and no ledger finalization."""
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)
    op = _op()
    ledger_file = Path(str(target) + ".operation-ledger.json")
    evidence_path = Path(str(target) + ".migrate.evidence.json")

    assert _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)[0] == 0
    manifest_mtime = os.stat(target / "workdir.manifest.json").st_mtime_ns

    # Crash window: ledger record lost, sidecar lost AND un-repairable.
    ledger_file.unlink()
    evidence_path.unlink()
    evidence_path.mkdir()  # required sidecar cannot be written

    rc, out = _cli("--json", "migrate", workdir, "--out", target, "--operation-id", op)
    assert rc == 1
    result = json.loads(out)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "evidence-publish-failed"
    assert os.stat(target / "workdir.manifest.json").st_mtime_ns == manifest_mtime  # no second effect
    assert not ledger_file.exists()  # no success record was fabricated


# --------------------------------------------------------------------------
# Human CLI surface preserved
# --------------------------------------------------------------------------

def test_human_inspect_and_migrate(tmp_path, capsys):
    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    target = tmp_path / "migrated"
    _make_doc(source)
    _extract(source, workdir)

    assert main(["inspect", str(workdir)]) == 0
    human = capsys.readouterr().out
    assert "readiness:   ready" in human
    assert "next action: migrate" in human
    assert "typed.md" in human

    assert main(["migrate", str(workdir), "--out", str(target)]) == 0
    human_migrate = capsys.readouterr().out
    assert f"migrated: {target}" in human_migrate
    assert "workdir.manifest.json" in human_migrate
    assert target.is_dir()

    # human migrate also refuses an existing target without touching it
    assert main(["migrate", str(workdir), "--out", str(target)]) == 1
    assert "already exists" in capsys.readouterr().out

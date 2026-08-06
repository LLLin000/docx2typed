"""R1: revision visibility and preservation."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.build import build
from scripts.edit import refresh_edit_projection, sync_edit_projection
from scripts.edit_sync import collect_document_revisions, flatten_paragraph
from scripts.extract import extract
from scripts.typed_core import RevisionNode, parse_typed, visible_text, visible_text_original
from scripts.verify import verify


def make_revision_docx(path: Path) -> None:
    """Ordinary docx with w:ins / w:del / nested revision / comment anchors."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前文")
    # insertion revision: <w:ins> containing a run with w:t
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "100")
    ins.set(qn("w:author"), "张三")
    ins.set(qn("w:date"), "2026-08-06T10:12:00Z")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "插入词"
    run.append(text)
    ins.append(run)
    paragraph._p.append(ins)
    # deletion revision: <w:del> containing a run with w:delText
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), "101")
    dele.set(qn("w:author"), "张三")
    del_run = OxmlElement("w:r")
    del_text = OxmlElement("w:delText")
    del_text.text = "旧词"
    del_run.append(del_text)
    dele.append(del_run)
    paragraph._p.append(dele)
    paragraph.add_run("后文")
    # comment anchors around the inserted text (must stay paired)
    comment_run = paragraph.add_run("批注目标")
    document.add_comment(comment_run, text="评论", author="tester")
    # nested revision: w:ins containing w:del
    nested = paragraph.add_run("")
    nested_ins = OxmlElement("w:ins")
    nested_ins.set(qn("w:id"), "102")
    nested_ins.set(qn("w:author"), "李四")
    nested_del = OxmlElement("w:del")
    nested_del.set(qn("w:id"), "103")
    nested_del.set(qn("w:author"), "李四")
    nested_run = OxmlElement("w:r")
    nested_text = OxmlElement("w:delText")
    nested_text.text = "旧内"
    nested_run.append(nested_text)
    nested_del.append(nested_run)
    nested_ins.append(nested_del)
    paragraph._p.append(nested_ins)
    document.add_paragraph("第二段")
    document.save(path)


def extract_fixture(tmp_path: Path, name: str = "rev") -> Path:
    source = tmp_path / f"{name}-src.docx"
    workdir = tmp_path / name
    make_revision_docx(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return workdir


def test_extract_parses_revision_nodes(tmp_path):
    workdir = extract_fixture(tmp_path)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    p0 = typed.paragraphs[0]
    revisions = [n for n in _walk(p0.nodes) if isinstance(n, RevisionNode)]
    assert len(revisions) == 4  # ins + del + nested ins + nested del
    by_id = {n.attrs.get("w:id"): n for n in revisions}
    assert by_id["100"].kind == "insert"
    assert by_id["101"].kind == "delete"
    assert by_id["102"].kind == "insert"  # nested container
    assert any(isinstance(c, RevisionNode) and c.attrs.get("w:id") == "103" for c in by_id["102"].children)


def _walk(nodes):
    from scripts.typed_core import RangeNode

    for node in nodes:
        yield node
        if isinstance(node, (RangeNode, RevisionNode)):
            yield from _walk(node.children)


def test_final_and_original_views(tmp_path):
    workdir = extract_fixture(tmp_path)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    p0 = typed.paragraphs[0]
    final = visible_text(p0.nodes)
    original = visible_text_original(p0.nodes)
    assert "插入词" in final and "旧词" not in final
    assert "旧词" in original and "插入词" not in original
    assert "前文" in final and "后文" in final


def test_edit_md_final_view_with_gaps(tmp_path):
    workdir = extract_fixture(tmp_path)
    edit = (workdir / "edit.md").read_text(encoding="utf-8")
    assert "插入词" in edit  # insertion visible
    assert "旧词" not in edit  # deletion hidden
    assert "revision-gap" in edit  # positions preserved
    assert edit.count("revision-gap") == 2  # two deletion gaps (101, 103)


def test_noop_roundtrip_preserves_revision_bytes(tmp_path):
    workdir = extract_fixture(tmp_path)
    output = tmp_path / "out.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(workdir / "_template.docx") as src, zipfile.ZipFile(output) as out:
        assert src.read("word/document.xml") == out.read("word/document.xml")
    with zipfile.ZipFile(output) as out:
        xml = out.read("word/document.xml").decode("utf-8")
        assert "<w:ins" in xml and "<w:delText>旧词</w:delText>" in xml
        assert "w:commentRangeStart" in xml and "w:commentRangeEnd" in xml


def test_anchor_pairs_stay_valid_on_revision_doc(tmp_path):
    workdir = extract_fixture(tmp_path)
    output = tmp_path / "out.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0


def test_revision_inventory(tmp_path):
    workdir = extract_fixture(tmp_path)
    inventory = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    assert inventory["schema"] == "typed-revisions-1"
    revisions = inventory["revisions"]
    assert any(r["kind"] == "insert" and r["text"] == "插入词" for r in revisions)
    assert any(r["kind"] == "delete" and r["text"] == "旧词" for r in revisions)
    for entry in revisions:
        assert entry["revision_key"]
        assert "part" in entry
    md = (workdir / "revisions.md").read_text(encoding="utf-8")
    assert "插入词" in md and "旧词" in md and "张三" in md


def test_regions_md_shows_gaps(tmp_path):
    workdir = extract_fixture(tmp_path)
    regions = (workdir / "regions.md").read_text(encoding="utf-8")
    assert "revision-gap" in regions


def test_flatten_keeps_gap_units(tmp_path):
    workdir = extract_fixture(tmp_path)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    units = flatten_paragraph(typed.paragraphs[0])
    gaps = [u for u in units if u.token and u.value[0] == "G"]
    assert len(gaps) == 2  # top-level deletion + deletion nested inside insert
    assert {gap.value[1] for gap in gaps} == {"delete"}
    nested = [g for g in gaps if g.range_path]
    assert len(nested) == 1  # the nested deletion sits on the insert path
    blocks = [u for u in units if u.token and u.value[0] == "IS"]
    assert len(blocks) == 2  # top-level insert + nested insert


def test_sync_edits_plain_text_around_gaps(tmp_path):
    workdir = extract_fixture(tmp_path)
    edit = (workdir / "edit.md").read_text(encoding="utf-8")
    edit = edit.replace("前文", "前文改", 1)
    (workdir / "edit.md").write_text(edit, encoding="utf-8")
    # the fixture carries pending revisions without track changes: ambiguous by
    # default, so the caller must choose --no-track for a direct edit
    _, _, changed = sync_edit_projection(workdir, track=False)
    assert changed == ["P0"]
    output = tmp_path / "out2.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as out:
        xml = out.read("word/document.xml").decode("utf-8")
        assert "前文改" in xml and "旧词" in xml  # revision delText still there
        assert "<w:ins" in xml  # revisions preserved in touched paragraph


def test_gap_mutation_rejected(tmp_path):
    workdir = extract_fixture(tmp_path)
    edit = (workdir / "edit.md").read_text(encoding="utf-8")
    # remove a revision-gap placeholder -> protected-token-mutated
    gap = re.search(r"\u27e6revision-gap[^\u27e7]*\u27e7", edit).group(0)
    (workdir / "edit.md").write_text(edit.replace(gap, "", 1), encoding="utf-8")
    try:
        sync_edit_projection(workdir)
    except Exception as exc:
        assert "protected-token-mutated" in str(exc)
    else:
        raise AssertionError("gap removal must be rejected")

from pathlib import Path

from docx import Document

from scripts.extract import extract
from scripts.view import view_workdir


def test_clean_style_and_raw_views_share_typed_source(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("普通")
    bold = paragraph.add_run("加粗")
    bold.bold = True
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    clean = view_workdir(workdir, "clean", markers=False)
    style = view_workdir(workdir, "style", markers=False)
    raw = view_workdir(workdir, "raw")

    assert clean == "普通加粗"
    assert "普通" in style and "加粗" in style and "s_" in style
    assert raw == (workdir / "typed.md").read_text(encoding="utf-8")


def test_style_view_labels_a_uniform_paragraph_style(tmp_path):
    source = tmp_path / "uniform.docx"
    workdir = tmp_path / "uniform-workdir"
    document = Document()
    run = document.add_paragraph().add_run("全段加粗")
    run.bold = True
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    style = view_workdir(workdir, "style", markers=False)
    assert "[s_" in style and "bold" in style

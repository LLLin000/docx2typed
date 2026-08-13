"""Issue #50 acceptance: immutable generations, Writer lane, durable journals,
CAS pointer, startup recovery, and external atomic publication.

Covers the four acceptance bullets through public seams:

1. Readers pin one generation; concurrent/stale writers get stable
   ``writer-busy`` / ``writer-timeout`` / ``generation-conflict``.
2. Successful mutation durability: generation assets, evidence, pointer,
   ledger, completed journal state.
3. Injected kill/ENOSPC/short-write/flush/corruption/CAS-race/lock-holder-death
   yield only complete old, complete new, or explicit ``needs-recovery``.
4. External DOCX publication and workdir commit recover without half-published
   normal state or duplicate Operation-ID effects.
"""
from __future__ import annotations

import json
import os
import subprocess
import sys
import threading
import time
import zipfile
from pathlib import Path

import pytest

from scripts import main
from scripts.protocol import (
    canonical_operation_input,
    file_sha256,
    new_operation_id,
    operation_ledger,
)
from scripts.store import (
    GENERATION_CONFLICT,
    NEEDS_RECOVERY,
    RESERVE_DEPLETED,
    STORE_DIR_NAME,
    UNSUPPORTED_BY_DESIGN,
    WRITER_BUSY,
    WRITER_TIMEOUT,
    Store,
    StoreError,
    _Kill,
    clear_faults,
    has_store,
    kill_at,
    read_root,
    set_fault,
)

ROOT = Path(__file__).resolve().parents[1]


def _make_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("第二段")
    document.save(path)


def _extract(tmp_path: Path, name: str = "wd", operation_id: str | None = None, source: Path | None = None) -> Path:
    if source is None:
        source = tmp_path / f"{name}-src.docx"
        _make_docx(source)
    workdir = tmp_path / name
    op = operation_id or new_operation_id()
    assert main(["extract", "--json", str(source), "-o", str(workdir), "--operation-id", op]) == 0
    return workdir


def _op() -> str:
    return new_operation_id()


def _simple_run(text: str = "hello\nworld\n"):
    def run(gen_dir, tx):
        (gen_dir / "typed.md").write_text(text, encoding="utf-8")
        (gen_dir / "note.txt").write_text("note\n", encoding="utf-8")
        return (
            "success",
            {"changed": ["P0"]},
            "mutation",
            {"checks": [{"name": "mutated", "status": "pass"}]},
            [],
        )

    return run


def _mutate(store: Store, *, op: str | None = None, run=None, expect: type | None = None):
    pin = store.pin()
    op = op or _op()
    canonical = canonical_operation_input("edit", {"workdir": str(store.root), "op": op})
    try:
        envelope = store.mutate(
            operation="edit",
            operation_id=op,
            canonical=canonical,
            input_sha256=pin["manifest_sha256"],
            expected_generation=pin["generation"],
            run=run or _simple_run(),
        )
        if expect is not None:
            raise AssertionError(f"expected {expect.__name__}, got success")
        return envelope
    except StoreError as exc:
        if expect is None or not isinstance(exc, expect):
            raise
        return exc


# --------------------------------------------------------------------------
# Bullet 1: writer outcomes and reader pinning
# --------------------------------------------------------------------------

def _hold_writer(workdir: Path, marker: Path, seconds: int = 30) -> subprocess.Popen:
    """Subprocess that acquires the Writer lane and holds it, signaling
    readiness by writing ``marker`` after the lock is held."""
    return subprocess.Popen(
        [
            sys.executable,
            "-c",
            (
                "import sys, time; sys.path.insert(0, %r);\n"
                "from scripts.store import Store;\n"
                "s = Store(%r);\n"
                "with s.writer():\n"
                "    import pathlib; pathlib.Path(%r).write_text('ready')\n"
                "    time.sleep(%d)\n"
            )
            % (str(ROOT), str(workdir), str(marker), seconds),
        ],
        cwd=ROOT,
    )


def _wait_marker(marker: Path, timeout: float = 30.0) -> None:
    deadline = time.monotonic() + timeout
    while not marker.exists():
        if time.monotonic() > deadline:
            raise AssertionError("lock holder never became ready")
        time.sleep(0.05)


class TestWriterOutcomes:
    def test_reader_pins_one_generation(self, tmp_path):
        """A reader that pinned generation N reads N consistently even while a
        writer commits N+1; the immutable pinned path never mixes."""
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        pinned = store.pin()
        pinned_path = pinned["path"]
        typed_before = (pinned_path / "typed.md").read_bytes()
        _mutate(store, run=_simple_run("next generation\n"))
        # The pinned generation directory is immutable: still the old content.
        assert (pinned_path / "typed.md").read_bytes() == typed_before
        # New readers pin the committed generation.
        fresh = Store.open(workdir).pin()
        assert fresh["generation"] != pinned["generation"]
        assert (fresh["path"] / "typed.md").read_text(encoding="utf-8") == "next generation\n"
        assert read_root(workdir) == fresh["path"]

    def test_writer_busy_and_timeout(self, tmp_path):
        workdir = _extract(tmp_path)
        marker = tmp_path / "holder-ready"
        holder = _hold_writer(workdir, marker)
        try:
            _wait_marker(marker)
            store = Store.open(workdir)
            # Immediate contention -> writer-busy.
            error = _mutate(store, expect=StoreError)
            assert error.code == WRITER_BUSY
            # Bounded wait that expires -> writer-timeout.
            started = time.monotonic()
            try:
                store.mutate(
                    operation="edit",
                    operation_id=_op(),
                    canonical="c",
                    input_sha256=store.pin()["manifest_sha256"],
                    expected_generation=store.pin()["generation"],
                    run=_simple_run(),
                    lock_timeout_ms=300,
                )
                raise AssertionError("expected writer-timeout")
            except StoreError as exc:
                assert exc.code == WRITER_TIMEOUT
                assert time.monotonic() - started < 10
        finally:
            holder.terminate()
            holder.wait(timeout=10)

    def test_lock_holder_death_releases_lane(self, tmp_path):
        workdir = _extract(tmp_path)
        marker = tmp_path / "holder-ready"
        holder = _hold_writer(workdir, marker, seconds=60)
        try:
            _wait_marker(marker)
            store = Store.open(workdir)
            error = _mutate(store, expect=StoreError)
            assert error.code == WRITER_BUSY
            holder.kill()  # process death: the OS advisory lock is released
            holder.wait(timeout=10)
            time.sleep(0.5)
            envelope = _mutate(store)
            assert envelope["outcome"] == "success"
        finally:
            if holder.poll() is None:
                holder.kill()
                holder.wait(timeout=10)

    def test_stale_writer_generation_conflict(self, tmp_path):
        """A writer that planned against generation N but commits after N+1
        landed receives generation-conflict, never a silent overwrite."""
        workdir = _extract(tmp_path)
        first = Store.open(workdir)
        second = Store.open(workdir)
        planned = first.pin()
        # A concurrent writer commits N+1.
        _mutate(second)
        # The stale planner now commits with expected=N -> conflict.
        canonical = canonical_operation_input("edit", {"workdir": str(workdir), "stale": True})
        with pytest.raises(StoreError) as info:
            first.mutate(
                operation="edit",
                operation_id=_op(),
                canonical=canonical,
                input_sha256=planned["manifest_sha256"],
                expected_generation=planned["generation"],
                run=_simple_run(),
            )
        assert info.value.code == GENERATION_CONFLICT
        # The committed generation is untouched.
        assert (second.pin()["path"] / "typed.md").read_text(encoding="utf-8") == "hello\nworld\n"


# --------------------------------------------------------------------------
# Bullet 2: successful mutation durability
# --------------------------------------------------------------------------

class TestDurability:
    def test_success_durability_includes_all_required_state(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        before = store.pin()
        op = _op()
        canonical = canonical_operation_input("edit", {"workdir": str(workdir), "durable": True})
        envelope = store.mutate(
            operation="edit",
            operation_id=op,
            canonical=canonical,
            input_sha256=before["manifest_sha256"],
            expected_generation=before["generation"],
            run=_simple_run(),
        )
        assert envelope["outcome"] == "success"
        after = store.pin()
        assert after["generation"] != before["generation"]
        gen_dir = after["path"]
        # Generation assets are complete and authoritative.
        assert (gen_dir / "typed.md").read_text(encoding="utf-8") == "hello\nworld\n"
        assert (gen_dir / "note.txt").is_file()
        manifest = json.loads((gen_dir / "generation.json").read_text(encoding="utf-8"))
        assert manifest["parent"] == before["generation"]
        assert manifest["operation_id"] == op
        # Evidence and ledger are durable inside the committed generation.
        assert (gen_dir / "run.evidence.json").is_file()
        evidence = json.loads((gen_dir / "run.evidence.json").read_text(encoding="utf-8"))
        assert evidence["operation_id"] == op
        ledger = json.loads((gen_dir / "operation-ledger.json").read_text(encoding="utf-8"))
        assert op in ledger["records"]
        assert ledger["records"][op]["envelope"]["outcome"] == "success"
        # Pointer selects the new generation with the manifest hash.
        pointer = json.loads((workdir / "workdir.json").read_text(encoding="utf-8"))
        assert pointer["generation"] == after["generation"]
        assert pointer["manifest_sha256"] == manifest["assets_sha256"]
        # The root is the materialized mirror of the committed generation.
        assert (workdir / "typed.md").read_text(encoding="utf-8") == "hello\nworld\n"
        # Recovery finds nothing to do and no journals linger.
        result = Store.open(workdir).recover()
        assert result["needs_recovery"] == []
        assert result["rolled_back"] == []
        tx_dir = workdir / STORE_DIR_NAME / "transactions"
        assert not tx_dir.exists() or not any(tx_dir.iterdir())

    def test_replay_same_operation_id_no_second_effect(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        op = _op()
        canonical = canonical_operation_input("edit", {"workdir": str(workdir), "replay": True})
        calls = {"count": 0}

        def run(gen_dir, tx):
            calls["count"] += 1
            return _simple_run()(gen_dir, tx)

        envelope1 = store.mutate(
            operation="edit",
            operation_id=op,
            canonical=canonical,
            input_sha256=store.pin()["manifest_sha256"],
            expected_generation=store.pin()["generation"],
            run=run,
        )
        generation1 = store.pin()["generation"]
        # Same op-id + identical canonical input: replay the original envelope.
        record = operation_ledger.lookup_persisted(op, store.ledger_dir(), directory=True)
        assert record is not None and record["input_sha256"] == canonical
        envelope2 = store.mutate(
            operation="edit",
            operation_id=op,
            canonical=canonical,
            input_sha256=store.pin()["manifest_sha256"],
            expected_generation=store.pin()["generation"],
            run=run,
        )
        assert envelope2 == envelope1
        assert calls["count"] == 1
        assert store.pin()["generation"] == generation1
        # Changed canonical input with the same op-id: rejected, no effect.
        with pytest.raises(StoreError):
            store.mutate(
                operation="edit",
                operation_id=op,
                canonical="changed-input",
                input_sha256=store.pin()["manifest_sha256"],
                expected_generation=store.pin()["generation"],
                run=run,
            )
        assert calls["count"] == 1
        assert store.pin()["generation"] == generation1

    def test_cli_operation_replay_returns_original_envelope(self, tmp_path):
        workdir = _extract(tmp_path)
        edit_md = workdir / "edit.md"
        text = edit_md.read_text(encoding="utf-8")
        marker = text.find("-->", text.find("<!--@p id="))
        tail = text.find("\n", marker) + 1
        end = text.find("\n", tail)
        text = text[:tail] + text[tail:end] + "改" + text[end:]
        edit_md.write_text(text, encoding="utf-8")
        op = _op()
        first = subprocess.run(
            [sys.executable, "-m", "scripts", "edit", "--json", "sync", str(workdir), "--operation-id", op],
            cwd=ROOT, capture_output=True, text=True, encoding="utf-8",
        )
        assert first.returncode == 0
        second = subprocess.run(
            [sys.executable, "-m", "scripts", "edit", "--json", "sync", str(workdir), "--operation-id", op],
            cwd=ROOT, capture_output=True, text=True, encoding="utf-8",
        )
        assert second.returncode == 0
        assert json.loads(first.stdout) == json.loads(second.stdout)
        # No duplicate effect: exactly one committed generation after the sync.
        generations = [
            p for p in (workdir / STORE_DIR_NAME / "generations").iterdir()
        ]
        assert len(generations) == 2  # extract birth + one sync commit


# --------------------------------------------------------------------------
# Bullet 3: injected faults yield only old / new / needs-recovery
# --------------------------------------------------------------------------

def _journal_phases(store: Store, operation_id: str) -> list[str]:
    tx_dir = store.transactions_dir / operation_id
    if not tx_dir.is_dir():
        return []
    return sorted(p.name for p in tx_dir.glob("*.json"))


def _assert_complete_old(store: Store, generation: str, content: str) -> None:
    """The pointer still selects ``generation`` and its content is intact."""
    assert store.pin()["generation"] == generation
    assert (store.pin()["path"] / "typed.md").read_text(encoding="utf-8") == content


def _old_content(store: Store) -> str:
    return (store.pin()["path"] / "typed.md").read_text(encoding="utf-8")


class TestFaultInjection:
    @pytest.fixture(autouse=True)
    def _no_faults(self):
        clear_faults()
        yield
        clear_faults()

    @pytest.mark.parametrize(
        "cut",
        [
            "journal-write-prepared",
            "journal-flush-prepared",
            "journal-rename-prepared",
            "generation-copy",
            "pointer-write",
            "pointer-flush",
            "pointer-rename",
        ],
    )
    def test_kill_before_pointer_commit_leaves_old(self, tmp_path, cut):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        old_gen = store.pin()["generation"]
        op = _op()
        canonical = canonical_operation_input("edit", {"workdir": str(workdir), cut: True})
        kill_at(cut)
        with pytest.raises(_Kill):
            store.mutate(
                operation="edit", operation_id=op, canonical=canonical,
                input_sha256=store.pin()["manifest_sha256"],
                expected_generation=old_gen, run=_simple_run(),
            )
        old_content = _old_content(store)
        clear_faults()  # the crashed process is gone; a new process recovers
        # Process death: pointer never moved -> recovery rolls back.
        recovered = Store.open(workdir).recover()
        assert recovered["needs_recovery"] == []
        _assert_complete_old(Store.open(workdir), old_gen, old_content)
        assert recovered["rolled_back"]

    @pytest.mark.parametrize(
        "cut",
        [
            "materialize",
            "ledger-write",
            "journal-write-generation-committed",
            "journal-write-completed",
        ],
    )
    def test_kill_after_pointer_commit_recovers_forward(self, tmp_path, cut):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        old_gen = store.pin()["generation"]
        op = _op()
        canonical = canonical_operation_input("edit", {"workdir": str(workdir), cut: True})
        kill_at(cut)
        with pytest.raises(_Kill):
            store.mutate(
                operation="edit", operation_id=op, canonical=canonical,
                input_sha256=store.pin()["manifest_sha256"],
                expected_generation=old_gen, run=_simple_run(),
            )
        clear_faults()
        recovered = Store.open(workdir).recover()
        assert recovered["needs_recovery"] == []
        fresh = Store.open(workdir)
        # Pointer committed -> the new generation is complete and materialized.
        assert (fresh.pin()["path"] / "typed.md").read_text(encoding="utf-8") == "hello\nworld\n"
        assert (workdir / "typed.md").read_text(encoding="utf-8") == "hello\nworld\n"
        # Ledger durable after roll-forward.
        record = operation_ledger.lookup_persisted(op, fresh.ledger_dir(), directory=True)
        assert record is not None and record["envelope"]["outcome"] == "success"

    @pytest.mark.parametrize(
        "cut",
        [
            "journal-write-intent",
            "journal-flush-intent",
            "journal-rename-intent",
        ],
    )
    def test_kill_before_prepared_recovers_old(self, tmp_path, cut):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        old_gen = store.pin()["generation"]
        op = _op()
        kill_at(cut)
        with pytest.raises(_Kill):
            store.mutate(
                operation="edit", operation_id=op, canonical="c",
                input_sha256=store.pin()["manifest_sha256"],
                expected_generation=old_gen, run=_simple_run(),
            )
        old_content = _old_content(store)
        clear_faults()
        recovered = Store.open(workdir).recover()
        assert recovered["needs_recovery"] == []
        _assert_complete_old(Store.open(workdir), old_gen, old_content)

    def test_kill_during_external_publish_rolls_back_verified_backup(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        output = tmp_path / "out.docx"
        output.write_bytes(b"prior output")
        op = _op()
        kill_at("external-publish-out.docx")

        def run(target, tx):
            staged = tx.staging("out.docx")
            staged.write_bytes(b"new output")
            tx.stage_external(output, staged, mode="replace")
            return ("success", {"output": str(output)}, "build", {"checks": []}, [])

        with pytest.raises(_Kill):
            store.mutate(
                operation="build", operation_id=op, canonical="c",
                input_sha256=store.pin()["manifest_sha256"],
                expected_generation=store.pin()["generation"],
                run=run, generation=False,
                ledger_anchor=output, ledger_directory=False,
                evidence_path=Path(str(output) + ".evidence.json"),
            )
        clear_faults()
        recovered = Store.open(workdir).recover()
        assert recovered["needs_recovery"] == []
        assert recovered["rolled_back"], recovered
        # Prior output restored from the verified backup; nothing half-published.
        assert output.read_bytes() == b"prior output"
        assert Store.open(workdir).pin()["generation"] is not None

    def test_kill_after_external_publish_recovers_forward(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        output = tmp_path / "out.docx"
        output.write_bytes(b"prior output")
        op = _op()
        kill_at("journal-write-external-published")  # publish already landed

        def run(target, tx):
            staged = tx.staging("out.docx")
            staged.write_bytes(b"new output")
            tx.stage_external(output, staged, mode="replace")
            return ("success", {"output": str(output)}, "build", {"checks": []}, [])

        with pytest.raises(_Kill):
            store.mutate(
                operation="build", operation_id=op, canonical="c",
                input_sha256=store.pin()["manifest_sha256"],
                expected_generation=store.pin()["generation"],
                run=run, generation=False,
                ledger_anchor=output, ledger_directory=False,
                evidence_path=Path(str(output) + ".evidence.json"),
            )
        clear_faults()
        assert output.read_bytes() == b"new output"
        recovered = Store.open(workdir).recover()
        assert recovered["needs_recovery"] == []
        # Roll forward: the published output is the complete new state.
        assert output.read_bytes() == b"new output"
        assert Path(str(output) + ".operation-ledger.json").is_file()
        assert Path(str(output) + ".evidence.json").is_file()

    def test_enospc_releases_reserve_and_locks_readonly(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        op = _op()
        enospc = OSError("no space left")
        enospc.errno = 28
        set_fault("journal-write-prepared", enospc)
        with pytest.raises(StoreError) as info:
            _mutate(store, op=op)
        assert info.value.code == RESERVE_DEPLETED
        clear_faults()
        # Reserve released; workdir is read-only until replenished.
        reserve = workdir / STORE_DIR_NAME / "reserve"
        assert reserve.stat().st_size < 1024 * 1024
        assert (workdir / STORE_DIR_NAME / "reserve-depleted.json").is_file()
        with pytest.raises(StoreError) as info2:
            _mutate(store, op=_op())
        assert info2.value.code == RESERVE_DEPLETED
        # Replenishing re-enables mutations.
        store.replenish_reserve()
        envelope = _mutate(store)
        assert envelope["outcome"] == "success"

    def test_short_write_corrupts_journal_to_needs_recovery(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        old_gen = store.pin()["generation"]
        op = _op()

        def truncate(temp_path):
            with open(temp_path, "r+b") as handle:
                handle.truncate(4)  # short write: a torn record lands

        set_fault("journal-flush-prepared", truncate)
        old_content = _old_content(store)
        with pytest.raises(OSError, match="short write"):
            _mutate(store, op=op)
        clear_faults()
        # The torn record was detected and rejected: recovery rolls back to
        # the complete old state; nothing half-published.
        recovered = Store.open(workdir).recover()
        assert recovered["needs_recovery"] == []
        _assert_complete_old(Store.open(workdir), old_gen, old_content)

    def test_flush_failure_leaves_no_partial_record(self, tmp_path):
        """A flush failure aborts the journal write: the record never lands,
        so recovery deterministically rolls back to the complete old state."""
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        old_gen = store.pin()["generation"]
        op = _op()
        set_fault("journal-flush-generation-committed", OSError("flush failed"))
        old_content = _old_content(store)
        with pytest.raises(OSError):
            _mutate(store, op=op)
        clear_faults()
        recovered = Store.open(workdir).recover()
        assert recovered["needs_recovery"] == []
        # The pointer was restored by the abort: complete old, never mixed.
        _assert_complete_old(Store.open(workdir), old_gen, old_content)

    def test_corrupt_pointer_yields_needs_recovery(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        old_gen = store.pin()["generation"]
        (workdir / "workdir.json").write_text("{not json", encoding="utf-8")
        with pytest.raises(StoreError) as info:
            store.mutate(
                operation="edit", operation_id=_op(), canonical="c",
                input_sha256="x" * 64, expected_generation=old_gen, run=_simple_run(),
            )
        assert info.value.code == NEEDS_RECOVERY

    def test_corrupt_generation_manifest_yields_needs_recovery(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        gen_dir = store.pin()["path"]
        (gen_dir / "generation.json").write_text("{broken", encoding="utf-8")
        with pytest.raises(StoreError) as info:
            Store.open(workdir).pin()
        assert info.value.code == NEEDS_RECOVERY

    def test_corrupt_journal_chain_yields_needs_recovery(self, tmp_path):
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        op = _op()
        canonical = canonical_operation_input("edit", {"workdir": str(workdir), "corrupt": True})
        kill_at("journal-write-prepared")
        with pytest.raises(_Kill):
            store.mutate(
                operation="edit", operation_id=op, canonical=canonical,
                input_sha256=store.pin()["manifest_sha256"],
                expected_generation=store.pin()["generation"], run=_simple_run(),
            )
        clear_faults()
        # Tamper with the intent record so the chain breaks.
        tx_dir = store.transactions_dir / op
        (tx_dir / "intent.json").write_text('{"schema":"docx2typed-transaction-journal-1","phase":"intent","tampered":true}\n', encoding="utf-8")
        recovered = Store.open(workdir).recover()
        assert recovered["needs_recovery"], recovered

    def test_cli_build_kill_recovers_no_half_published_output(self, tmp_path):
        workdir = _extract(tmp_path)
        output = tmp_path / "built.docx"
        op = _op()
        kill_at("external-publish-built.docx")
        with pytest.raises(_Kill):
            main(["build", "--json", str(workdir), "-o", str(output), "--operation-id", op])
        clear_faults()
        # The output was never published (kill before the atomic replace).
        assert not output.exists()
        Store.open(workdir).recover()
        # Retry with a fresh op-id builds normally.
        assert main(["build", "--json", str(workdir), "-o", str(output), "--operation-id", _op()]) == 0
        assert output.is_file()


# --------------------------------------------------------------------------
# Bullet 4: external DOCX publication and workdir commit idempotency
# --------------------------------------------------------------------------

class TestExternalAndIdempotency:
    def test_workdir_commit_no_duplicate_operation_effects(self, tmp_path):
        """Same Operation-ID + canonical input replays the original envelope;
        changed input is rejected; the effect never runs twice."""
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        op = _op()
        canonical = canonical_operation_input("edit", {"workdir": str(workdir), "idem": True})
        runs = {"count": 0}

        def run(gen_dir, tx):
            runs["count"] += 1
            return _simple_run()(gen_dir, tx)

        first = store.mutate(
            operation="edit", operation_id=op, canonical=canonical,
            input_sha256=store.pin()["manifest_sha256"],
            expected_generation=store.pin()["generation"], run=run,
        )
        store2 = Store.open(workdir)
        second = store2.mutate(
            operation="edit", operation_id=op, canonical=canonical,
            input_sha256=store2.pin()["manifest_sha256"],
            expected_generation=store2.pin()["generation"], run=run,
        )
        assert second == first
        assert runs["count"] == 1

    def test_build_replace_and_verify_roundtrip(self, tmp_path):
        workdir = _extract(tmp_path)
        output = tmp_path / "built.docx"
        op = _op()
        assert main(["build", "--json", str(workdir), "-o", str(output), "--operation-id", op]) == 0
        assert output.is_file()
        # Replacing the output backs it up and swaps atomically.
        op2 = _op()
        assert main(["build", "--json", str(workdir), "-o", str(output), "--operation-id", op2]) == 0
        assert output.is_file()
        assert main(["verify", "--json", str(workdir), str(output)]) == 0
        # Replay returns the original envelope without rebuilding.
        replay = subprocess.run(
            [sys.executable, "-m", "scripts", "build", "--json", str(workdir), "-o", str(output), "--operation-id", op],
            cwd=ROOT, capture_output=True, text=True, encoding="utf-8",
        )
        assert replay.returncode == 0
        assert '"operation_id":"%s"' % op in replay.stdout.replace(" ", "")

    def test_startup_recovery_runs_at_next_mutation_entry(self, tmp_path):
        """After a simulated kill, the next CLI mutation entry point performs
        startup recovery before running (no separate recover command needed)."""
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        op = _op()
        kill_at("materialize")
        with pytest.raises(_Kill):
            _mutate(store, op=op)
        clear_faults()
        # The pointer committed but the root is stale; the next entry point
        # recovers automatically.
        assert main(["validate", "--json", str(workdir)]) in (0, 1)
        result = Store.open(workdir).recover()
        assert result["needs_recovery"] == []
        assert (workdir / "typed.md").read_text(encoding="utf-8") == "hello\nworld\n"

    def test_leftover_transaction_recovers_before_next_store_mutation(self, tmp_path):
        """Issue #50 final finding: a leftover transaction followed by a
        store-backed mutation recovers (roll back/forward) instead of crashing
        on the recovery summary (mutate must index the real result)."""
        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        op = _op()
        kill_at("journal-write-prepared")  # crash before the prepared record lands
        with pytest.raises(_Kill):
            store.mutate(
                operation="edit", operation_id=op, canonical=canonical_operation_input(
                    "edit", {"workdir": str(workdir), "leftover": True}
                ),
                input_sha256=store.pin()["manifest_sha256"],
                expected_generation=store.pin()["generation"], run=_simple_run(),
            )
        clear_faults()
        # The crash left an intent-only transaction journal behind.
        assert (store.transactions_dir / op).is_dir()
        # The next store-backed mutation recovers the leftover (rolls it back)
        # and commits normally — it must not crash on the recovery summary.
        envelope = _mutate(Store.open(workdir))
        assert envelope["outcome"] == "success"
        assert not (store.transactions_dir / op).exists()
        assert Store.open(workdir).recover()["needs_recovery"] == []


class TestReviewHttpIdempotency:
    """Issue #34: every mutating review HTTP POST requires an Idempotency-Key
    with the same syntax and ledger behavior as CLI/MCP operation IDs."""

    def _serve(self, workdir: Path):
        from http.server import ThreadingHTTPServer

        from scripts.review_server import _handler_for

        server = ThreadingHTTPServer(("127.0.0.1", 0), _handler_for(workdir))
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        return server

    def _post(self, port: int, path: str, payload: dict, key: str | None):
        import urllib.error
        import urllib.request

        request = urllib.request.Request(
            f"http://127.0.0.1:{port}{path}",
            data=json.dumps(payload).encode("utf-8"),
            method="POST",
        )
        request.add_header("Content-Type", "application/json")
        if key:
            request.add_header("Idempotency-Key", key)
        try:
            with urllib.request.urlopen(request) as response:
                return response.status, json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            return exc.code, json.loads(exc.read().decode("utf-8"))

    def test_review_post_requires_and_replays_idempotency_key(self, tmp_path):
        import threading

        workdir = _extract(tmp_path)
        server = self._serve(workdir)
        port = server.server_address[1]
        try:
            # Missing key -> 400; never mutates.
            status, body = self._post(port, "/api/reviews/dispatch", {}, key=None)
            assert status == 400
            assert body.get("code") in ("idempotency-key-required", "server-error")
            # Mutating POST with a key commits a generation.
            patch = {
                "type": "patch",
                "event_id": "evt-0001",
                "origin": "human_ui",
                "author": "tester",
                "parent_snapshot": "C0",
                "paragraph_id": "P0",
                "kind": "replace",
                "before": "",
                "after": "新词",
                "target": {
                    "start_offset": 0,
                    "end_offset": 0,
                    "expected_text": "",
                    "left_context": "",
                    "right_context": "第一段",
                    "paragraph_fingerprint": "",
                    "region_fingerprint": "",
                },
            }
            status, first = self._post(port, "/api/reviews/patch", patch, key="revkey00000001")
            assert status == 200
            # Identical key + payload replays the original response.
            status, replay = self._post(port, "/api/reviews/patch", patch, key="revkey00000001")
            assert status == 200 and replay == first
            # Same key with a changed payload -> operation-id-reused.
            changed = dict(patch, event_id="evt-0002", after="别的")
            status, body = self._post(port, "/api/reviews/patch", changed, key="revkey00000001")
            assert status == 409
            assert body.get("code") == "operation-id-reused"
            # Dispatch with its own key; replay returns the same response.
            status, first_dispatch = self._post(port, "/api/reviews/dispatch", {}, key="revkey00000002")
            assert status == 200
            status, replay_dispatch = self._post(port, "/api/reviews/dispatch", {}, key="revkey00000002")
            assert status == 200 and replay_dispatch == first_dispatch
            # The review mutations were journaled as generations.
            store = Store.open(workdir)
            assert len(list(store.generations_dir.iterdir())) >= 3
            assert store.recover()["needs_recovery"] == []
        finally:
            server.shutdown()
            server.server_close()


# --------------------------------------------------------------------------
# Issue #50 findings: extract single-envelope, probe volume binding,
# CLI decide table-* create mode
# --------------------------------------------------------------------------

class TestExtractSingleEnvelope:
    def test_extract_store_init_failure_emits_one_failure_envelope(self, tmp_path, monkeypatch, capsys):
        """Findings: extract births the store BEFORE printing success; a
        failed Store.init emits exactly ONE failure envelope — never a
        success-first double emit."""
        import scripts.store as store_module

        source = tmp_path / "fail-src.docx"
        _make_docx(source)
        workdir = tmp_path / "fail-wd"
        op = _op()

        def unqualified(store_dir):
            raise store_module.UnsupportedFilesystem("simulated unqualified volume")

        monkeypatch.setattr(store_module, "_probe_or_reuse", unqualified)
        code = main(["extract", "--json", str(source), "-o", str(workdir), "--operation-id", op])
        lines = [line for line in capsys.readouterr().out.splitlines() if line.strip()]
        assert code == 1
        assert len(lines) == 1  # one failure envelope, no success-first double emit
        envelope = json.loads(lines[0])
        assert envelope["schema"] == "docx2typed-result-1"
        assert envelope["operation"] == "extract"
        assert envelope["outcome"] == "failure"
        assert envelope["data"]["operation_id"] == op
        assert envelope["diagnostics"][0]["code"] == UNSUPPORTED_BY_DESIGN
        # The extracted assets exist but the store was never born: retrying
        # with the same op-id replays the failure without a second extract
        # success envelope.
        code2 = main(["extract", "--json", str(source), "-o", str(workdir), "--operation-id", op])
        lines2 = [line for line in capsys.readouterr().out.splitlines() if line.strip()]
        assert code2 == 1
        assert len(lines2) == 1


class TestProbeVolumeBinding:
    def test_probe_cache_binds_to_volume_and_reprobes_after_move(self, tmp_path, monkeypatch):
        """Findings: the probe cache is bound to the workdir host volume
        identity; a workdir moved onto a different volume re-probes, and an
        unqualified volume rejects before any mutation."""
        import scripts.store as store_module

        workdir = _extract(tmp_path)
        store = Store.open(workdir)
        probe_path = workdir / STORE_DIR_NAME / "probe.json"
        cached = json.loads(probe_path.read_text(encoding="utf-8"))
        assert cached.get("volume_identity") is not None
        assert cached["volume_identity"] == str(os.stat(workdir / STORE_DIR_NAME).st_dev)

        # Simulate the move: the cached probe now claims a foreign volume.
        moved = dict(cached, volume_identity="foreign-volume-123")
        probe_path.write_text(json.dumps(moved, ensure_ascii=False, sort_keys=True) + "\n", encoding="utf-8")

        calls = {"n": 0}
        real_probe = store_module._probe_filesystem

        def unqualified(store_dir):
            calls["n"] += 1
            result = real_probe(store_dir)
            result["qualified"] = False
            return result

        monkeypatch.setattr(store_module, "_probe_filesystem", unqualified)
        generations_before = {p.name for p in (workdir / STORE_DIR_NAME / "generations").iterdir()}
        # Opening on the moved volume re-probes (identity mismatch) and
        # rejects before any mutation can start.
        with pytest.raises(StoreError) as info:
            Store.open(workdir)
        assert info.value.code == UNSUPPORTED_BY_DESIGN
        assert calls["n"] == 1  # re-probed, never reused the foreign cache
        generations_after = {p.name for p in (workdir / STORE_DIR_NAME / "generations").iterdir()}
        assert generations_after == generations_before  # nothing mutated

        # A qualified re-probe (same workdir, real filesystem) succeeds and
        # refreshes the cache with the current volume identity.
        monkeypatch.setattr(store_module, "_probe_filesystem", real_probe)
        reopened = Store.open(workdir)
        assert reopened.pin()["generation"] is not None
        refreshed = json.loads(probe_path.read_text(encoding="utf-8"))
        assert refreshed["volume_identity"] == cached["volume_identity"]


class TestCliDecideCreateMode:
    def test_cli_decide_table_op_create_mode_success(self, tmp_path, capsys):
        """Findings: CLI decide table-* hashes the STAGED output for
        evidence — create mode means the final path does not exist until
        publish — so the operation succeeds and the evidence records both
        the staged hash and the final path."""
        from docx import Document

        source = tmp_path / "tbl-cli.docx"
        document = Document()
        document.add_paragraph("表前")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "A2"
        table.cell(1, 0).text = "B1"
        table.cell(1, 1).text = "B2"
        document.add_paragraph("表后")
        document.save(source)
        workdir = _extract(tmp_path, name="tbl-cli-wd", source=source)
        capsys.readouterr()  # drop the extract envelope from the capture
        output = tmp_path / "decided-table.docx"
        workdir_out = tmp_path / "decided-table-wd"
        op = _op()
        code = main([
            "decide", "--json", "table-insert-row", "T0", "--args", "0",
            "--output", str(output), "--workdir-out", str(workdir_out),
            "--workdir", str(workdir), "--operation-id", op,
        ])
        out = capsys.readouterr().out
        assert code == 0, out
        assert output.is_file()
        assert workdir_out.is_dir()
        envelope = json.loads(out)
        docx_evidence = envelope["evidence"][0]["payload"]["outputs"]["docx"]
        assert docx_evidence["sha256"] == file_sha256(output)  # staged hash == published bytes
        assert docx_evidence["path"] == str(output.resolve())


class TestCliDecideReplay:
    def test_decide_single_replays_original_envelope(self, tmp_path):
        """An in-place decide mutation commits a generation and replays the
        original envelope for the same operation-id (no duplicate effect)."""
        import zipfile

        from docx import Document

        source = tmp_path / "rev.docx"
        document = Document()
        document.add_paragraph("前言")
        document.add_paragraph("第二段")
        document.save(source)
        with zipfile.ZipFile(source) as archive:
            files = {name: archive.read(name) for name in archive.namelist()}
        insertion = (
            '<w:ins w:id="99" w:author="t" w:date="2026-01-01T00:00:00Z">'
            "<w:r><w:t>修订词</w:t></w:r></w:ins>"
        ).encode("utf-8")
        files["word/document.xml"] = files["word/document.xml"].replace(
            "<w:r><w:t>前言</w:t></w:r>".encode("utf-8"),
            "<w:r><w:t>前言</w:t></w:r>".encode("utf-8") + insertion,
            1,
        )
        with zipfile.ZipFile(source, "w", zipfile.ZIP_DEFLATED) as archive:
            for name, data in files.items():
                archive.writestr(name, data)

        workdir = _extract(tmp_path, name="wd2", source=source)
        entry = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))[
            "revisions"
        ][0]
        key = entry["revision_key"]
        fingerprint = key.rsplit("|", 1)[-1]
        op = _op()
        first = subprocess.run(
            [
                sys.executable, "-m", "scripts", "decide", "--json", "accept",
                key, "--fingerprint", fingerprint,
                "--workdir", str(workdir), "--operation-id", op,
            ],
            cwd=ROOT, capture_output=True, text=True, encoding="utf-8",
        )
        assert first.returncode == 0
        second = subprocess.run(
            [
                sys.executable, "-m", "scripts", "decide", "--json", "accept",
                key, "--fingerprint", fingerprint,
                "--workdir", str(workdir), "--operation-id", op,
            ],
            cwd=ROOT, capture_output=True, text=True, encoding="utf-8",
        )
        assert second.returncode == 0
        assert json.loads(first.stdout) == json.loads(second.stdout)
        # Exactly one commit happened (extract birth + one decide generation).
        generations = list((workdir / STORE_DIR_NAME / "generations").iterdir())
        assert len(generations) == 2


# --------------------------------------------------------------------------
# Issue #50 final findings: fresh review success payloads, corrupt-row path,
# per-store ledger namespace
# --------------------------------------------------------------------------

class TestReviewSuccessPayloadsFresh:
    """Issue #50 finding 1: mutating review POSTs recompute session and
    queued counts AFTER the mutation lands, so the success payload carries
    the fresh state and the next publish needs no extra GET."""

    def _serve(self, workdir: Path):
        from http.server import ThreadingHTTPServer

        from scripts.review_server import _handler_for

        server = ThreadingHTTPServer(("127.0.0.1", 0), _handler_for(workdir))
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        return server

    def _post(self, port: int, path: str, payload: dict, key: str | None):
        import urllib.error
        import urllib.request

        request = urllib.request.Request(
            f"http://127.0.0.1:{port}{path}",
            data=json.dumps(payload).encode("utf-8"),
            method="POST",
        )
        request.add_header("Content-Type", "application/json")
        if key:
            request.add_header("Idempotency-Key", key)
        try:
            with urllib.request.urlopen(request) as response:
                return response.status, json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            return exc.code, json.loads(exc.read().decode("utf-8"))

    def test_next_publish_works_without_extra_get(self, tmp_path):
        """A mutating POST's success payload carries the post-mutation
        session and counts, so a client can chain the next publish using only
        the previous response — no GET to refresh state in between."""
        workdir = _extract(tmp_path)
        server = self._serve(workdir)
        port = server.server_address[1]
        try:
            # Seed one draft event; the response counts must already reflect
            # it (computed after the mutation committed).
            status, seeded = self._post(
                port,
                "/api/reviews",
                {
                    "type": "decision",
                    "client_id": "browser-1",
                    "review_item_id": "P16",
                    "paragraph_id": "P16",
                    "decision": "accept",
                    "revision_key": "P0-r1",
                    "revision_id": "r1",
                    "selected_text": "旧词",
                    "comment": "",
                },
                key="revkey00000001",
            )
            assert status == 200
            assert seeded["counts"]["draft"] == 1  # fresh, not pre-mutation 0
            assert seeded["session"]["current_snapshot"]["id"] == "C0"
            assert seeded["session"]["current_matches_filesystem"] is True

            # First publish: parent comes from the seed response session.
            (workdir / "typed.md").write_text("changed v1\n", encoding="utf-8")
            status, published = self._post(
                port,
                "/api/reviews/publish",
                {
                    "expected_parent_snapshot": seeded["session"]["current_snapshot"]["id"],
                    "origin": "human_ui",
                    "changed_paragraph_ids": ["P16"],
                },
                key="revkey00000002",
            )
            assert status == 200, published
            assert published["session"]["current_snapshot"]["id"] == "C1"
            assert published["counts"]["draft"] == 1

            # Second publish uses ONLY the first publish response — a stale
            # pre-mutation session would send C0 and hit current-parent-
            # mismatch; the fresh session sends C1 and succeeds.
            (workdir / "typed.md").write_text("changed v2\n", encoding="utf-8")
            status, again = self._post(
                port,
                "/api/reviews/publish",
                {
                    "expected_parent_snapshot": published["session"]["current_snapshot"]["id"],
                    "origin": "human_ui",
                    "changed_paragraph_ids": ["P16"],
                },
                key="revkey00000003",
            )
            assert status == 200, again
            assert again["session"]["current_snapshot"]["id"] == "C2"
        finally:
            server.shutdown()
            server.server_close()


def test_cli_corrupt_ledger_row_in_advanced_generation_names_exact_file(tmp_path, capsys):
    """Issue #50 finding 2: with the pointer advanced past the committing
    generation, the corrupt-row diagnostic names the EXACT generation ledger
    file holding the row — never the pinned generation's ledger."""
    workdir = _extract(tmp_path)
    store = Store.open(workdir)
    _mutate(store, op=_op())  # advance the pointer past the corrupt row's generation
    pinned = store.pin()["path"]
    gen_dirs = sorted((workdir / STORE_DIR_NAME / "generations").iterdir())
    assert len(gen_dirs) >= 2  # extract birth + advance mutation
    old_ledger = next(
        (g / "operation-ledger.json" for g in gen_dirs if g.name != pinned.name),
        gen_dirs[0] / "operation-ledger.json",
    )
    pinned_ledger = pinned / "operation-ledger.json"
    op = _op()
    old_ledger.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {op: {"input_sha256": "0" * 64, "envelope": None}},
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )
    capsys.readouterr()  # drop the extract/advance envelopes
    code = main(["edit", "--json", "sync", str(workdir), "--operation-id", op])
    out = capsys.readouterr().out
    assert code == 1, out
    envelope = json.loads(out)
    assert envelope["outcome"] == "failure"
    assert envelope["data"]["operation_id"] == op
    diagnostic = envelope["diagnostics"][0]
    assert diagnostic["code"] == "operation-ledger-invalid"
    assert str(old_ledger) in diagnostic["message"]
    assert str(pinned_ledger) not in diagnostic["message"]
    # The corrupt row is preserved verbatim for inspection.
    persisted = json.loads(old_ledger.read_text(encoding="utf-8"))["records"][op]
    assert persisted == {"input_sha256": "0" * 64, "envelope": None}


def test_operation_ledger_replay_namespaced_per_store(tmp_path):
    """Issue #50 finding 3: in-memory ledger replay is namespaced per store
    (anchor + generation); two workdirs in one process sharing an
    operation_id each get a fresh record — never the other workdir's record
    — while replay within one store stays byte-exact."""
    workdir_a = _extract(tmp_path, name="a")
    workdir_b = _extract(tmp_path, name="b")
    store_a = Store.open(workdir_a)
    store_b = Store.open(workdir_b)
    op = "shared-op-00000000000001"

    envelope_a = _mutate(store_a, op=op)
    assert envelope_a["outcome"] == "success"

    # Same op_id on a different workdir: the canonical input differs, so
    # consulting A's in-memory record would raise operation-id-reused. The
    # namespaced mirror treats it as a fresh record and mutates B for real.
    envelope_b = _mutate(store_b, op=op)
    assert envelope_b["outcome"] == "success"

    # Replay within one store is still byte-exact.
    assert _mutate(store_a, op=op) == envelope_a

    # Each store's persisted ledger carries its own canonical input.
    row_a = json.loads(
        (store_a.ledger_dir() / "operation-ledger.json").read_text(encoding="utf-8")
    )["records"][op]
    row_b = json.loads(
        (store_b.ledger_dir() / "operation-ledger.json").read_text(encoding="utf-8")
    )["records"][op]
    assert row_a["input_sha256"] != row_b["input_sha256"]

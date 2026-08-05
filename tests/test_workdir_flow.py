from pathlib import Path
from zipfile import ZipFile

from docx import Document

from build import build
from extract import extract
from verify import verify


def make_fixture(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("标题：")
    bold = paragraph.add_run("原始文本")
    bold.bold = True
    paragraph.add_run(" <literal>")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格保留"
    document.add_paragraph("末段")
    document.save(path)


def test_extract_edit_build_verify_preserves_package_and_nested_table(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    make_fixture(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    typed = workdir / "typed.md"
    source_text = typed.read_text(encoding="utf-8")
    assert "原始文本" in source_text
    typed.write_text(source_text.replace("原始文本", "修改文本"), encoding="utf-8")

    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0

    with ZipFile(source) as original, ZipFile(output) as rebuilt:
        assert original.read("word/styles.xml") == rebuilt.read("word/styles.xml")
        assert original.read("word/numbering.xml") == rebuilt.read("word/numbering.xml")
        rebuilt_document = rebuilt.read("word/document.xml").decode("utf-8")
    assert "修改文本" in rebuilt_document
    assert "表格保留" in rebuilt_document

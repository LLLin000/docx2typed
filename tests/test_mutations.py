from docx import Document

from scripts.build import build
from scripts.extract import extract
from scripts.verify import verify


def test_explicit_paragraph_delete_and_inherit_insertion_are_safe(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    deleted = tmp_path / "deleted.docx"
    inserted_workdir = tmp_path / "inserted-workdir"
    inserted = tmp_path / "inserted.docx"
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("第二段")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    blocks = (workdir / "typed.md").read_text(encoding="utf-8").split("\n\n")
    header, p0, p1 = blocks[:3]
    (workdir / "typed.md").write_text("\n\n".join([header, p1, '<!--@delete id="P0"-->']) + "\n", encoding="utf-8")
    assert build([str(workdir), "-o", str(deleted)]) == 0
    assert verify([str(workdir), str(deleted)]) == 0
    assert len(Document(deleted).paragraphs) == 1

    extract([str(source), "-o", str(inserted_workdir)])
    blocks = (inserted_workdir / "typed.md").read_text(encoding="utf-8").split("\n\n")
    header, p0, p1 = blocks[:3]
    new_block = '<!--@p id="Pnew" inherit="P1"-->\n插入段落'
    (inserted_workdir / "typed.md").write_text("\n\n".join([header, p0, new_block, p1]) + "\n", encoding="utf-8")
    assert build([str(inserted_workdir), "-o", str(inserted)]) == 0
    assert verify([str(inserted_workdir), str(inserted)]) == 0
    assert [paragraph.text for paragraph in Document(inserted).paragraphs] == ["第一段", "插入段落", "第二段"]

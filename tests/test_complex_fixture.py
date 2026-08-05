from pathlib import Path
from zipfile import ZipFile

from docx import Document

from scripts.build import build
from scripts.edit import refresh_edit_projection
from scripts.extract import extract
from scripts.verify import verify
from scripts.view import view_workdir


FIXTURE_DIR = Path(__file__).parent / "fixtures" / "complex-docx"
FIXTURE = FIXTURE_DIR / "complex.docx"


def test_complex_real_docx_edit_build_verify_and_views(tmp_path):
    workdir = tmp_path / "workdir"
    output = tmp_path / "edited.docx"

    assert extract([str(FIXTURE), "-o", str(workdir)]) == 0
    typed_path = workdir / "typed.md"
    typed_source = typed_path.read_text(encoding="utf-8")
    assert "<span data-s=\"s_" in typed_source
    assert "docx-opaque" in typed_source
    assert "docx-inline" in typed_source

    clean = view_workdir(workdir, "clean", markers=False)
    style = view_workdir(workdir, "style", markers=False)
    raw = view_workdir(workdir, "raw")
    assert "EDIT-ME:" in clean
    assert "opaque" in clean
    assert "s_" in style
    assert raw == typed_source

    typed_path.write_text(typed_source.replace("EDIT-ME:", "EDITED:"), encoding="utf-8")
    refresh_edit_projection(workdir)
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0

    rebuilt = Document(output)
    assert any("EDITED:" in paragraph.text for paragraph in rebuilt.paragraphs)
    assert rebuilt.tables and rebuilt.tables[0].cell(3, 0).tables
    assert len(rebuilt.sections) == 2

    with ZipFile(FIXTURE) as source_archive, ZipFile(output) as output_archive:
        source_names = set(source_archive.namelist())
        output_names = set(output_archive.namelist())
        document_xml = output_archive.read("word/document.xml").decode("utf-8")
        assert {"word/comments.xml", "word/footnotes.xml", "word/endnotes.xml"} <= output_names
        assert any(name.startswith("word/media/") for name in output_names)
        assert source_names == output_names
        assert "EDITED:" in document_xml
        assert all(
            marker in document_xml
            for marker in (
                "w:hyperlink",
                "w:commentRangeStart",
                "w:footnoteReference",
                "w:endnoteReference",
                "m:oMathPara",
                "w:customXml",
                "w:sdt",
                "w:proofErr",
                "w:ins",
                "w:fldSimple",
                "w:tbl",
            )
        )

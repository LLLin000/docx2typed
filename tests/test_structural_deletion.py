import json
import re
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.build import build
from scripts.edit import refresh_edit_projection
from scripts.extract import extract
from scripts.verify import verify


def test_deleting_earlier_paragraph_keeps_later_structural_tokens(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("删除")
    paragraph = document.add_paragraph("前")
    relationship_id = document.part.relate_to("https://example.com", RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "链接"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    commented_run = paragraph.add_run("批注")
    document.add_comment(commented_run, text="保留", author="tester")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    blocks = (workdir / "typed.md").read_text(encoding="utf-8").split("\n\n")
    # keep every block except the P0 body paragraph, then tombstone P0
    kept = [b for b in blocks if not b.startswith('<!--@p id="P0"')]
    kept.append('<!--@delete id="P0"-->')
    (workdir / "typed.md").write_text("\n\n".join(kept) + "\n", encoding="utf-8")
    refresh_edit_projection(workdir)

    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        assert "w:hyperlink" in xml and "w:commentRangeStart" in xml


def test_new_paragraph_inherits_only_editable_paragraph_structure(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("正文")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("第二节")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    inherit_id = next(item["id"] for item in format_data["paragraphs"] if not item["section_bearing"])
    typed_path = workdir / "typed.md"
    typed = typed_path.read_text(encoding="utf-8")
    typed_path.write_text(
        typed + f'\n<!--@p id="Pnew" inherit="{inherit_id}"-->\n新增正文\n',
        encoding="utf-8",
    )
    refresh_edit_projection(workdir)

    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    assert [paragraph.text for paragraph in Document(output).paragraphs][-1] == "新增正文"
    with ZipFile(source) as original, ZipFile(output) as rebuilt:
        original_xml = original.read("word/document.xml")
        rebuilt_xml = rebuilt.read("word/document.xml")
    assert rebuilt_xml.count(b"<w:sectPr") == original_xml.count(b"<w:sectPr")


def test_new_paragraph_cannot_inherit_structural_tokens(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph("链接")
    relationship_id = document.part.relate_to("https://example.com", RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "外链"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    typed_path = workdir / "typed.md"
    typed = typed_path.read_text(encoding="utf-8")
    typed_path.write_text(
        typed + '\n<!--@p id="Pnew" inherit="P0"-->\n不应继承\n',
        encoding="utf-8",
    )
    assert build([str(workdir), "-o", str(output)]) == 1


def test_unbalanced_anchor_source_is_rejected(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    run = document.add_paragraph("带批注").runs[0]
    document.add_comment(run, text="保留", author="tester")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    typed_path = workdir / "typed.md"
    typed = typed_path.read_text(encoding="utf-8")
    typed, removed = re.subn(
        r'<docx-anchor[^>]*kind="comment-end"[^>]*/>',
        "",
        typed,
        count=1,
    )
    assert removed == 1
    typed_path.write_text(typed, encoding="utf-8")
    assert build([str(workdir), "-o", str(output)]) == 1


def test_build_rejects_reserved_output_paths(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    document = Document()
    document.add_paragraph("正文")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    assert build([str(workdir), "-o", str(workdir / "typed.md")]) == 1

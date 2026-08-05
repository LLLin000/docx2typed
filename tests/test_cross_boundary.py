import re

from docx import Document

from scripts.build import build
from scripts.extract import extract


def test_cross_span_rewrite_is_rejected_before_output(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")
    bold = paragraph.add_run("B")
    bold.bold = True
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    typed_path = workdir / "typed.md"
    typed = typed_path.read_text(encoding="utf-8")
    match = re.search(r"A(<span data-s=\"[^\"]+\">)B(</span>)", typed)
    assert match
    typed_path.write_text(typed.replace(match.group(0), f"B{match.group(1)}A{match.group(2)}"), encoding="utf-8")

    assert build([str(workdir), "-o", str(output)]) == 1
    assert not output.exists()

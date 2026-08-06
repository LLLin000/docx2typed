"""N3: header/footer/footnote/endnote part editing — multi-part surface."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document

from scripts.build import build
from scripts.edit import sync_edit_projection
from scripts.extract import extract
from scripts.typed_core import parse_typed
from scripts.verify import verify

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def make_header_footer_docx(path: Path) -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "页眉文字"
    document.sections[0].footer.paragraphs[0].text = "页脚文字"
    document.add_paragraph("正文")
    section2 = document.add_section()
    section2.header.is_linked_to_previous = False
    section2.header.paragraphs[0].text = "第二节页眉"
    document.save(path)


FOOTNOTES_TEMPLATE = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    f'<w:footnotes xmlns:w="{NS_W}">'
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
    '<w:footnote w:id="1"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr><w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r><w:r><w:t>脚注内容</w:t></w:r></w:p></w:footnote>'
    "</w:footnotes>"
)

ENDNOTES_TEMPLATE = FOOTNOTES_TEMPLATE.replace(
    "footnotes", "endnotes"
).replace(
    'w:type="separator" w:id="-1"', 'w:type="separator" w:id="-1"'
).replace("FootnoteText", "EndnoteText").replace(
    "FootnoteReference", "EndnoteReference"
).replace("footnoteRef", "endnoteRef").replace(
    "脚注内容", "尾注内容"
)


def _inject_note_part(path: Path, part_name: str, xml: str, rel_type: str, content_type: str) -> None:
    """Zip-inject a footnotes/endnotes part plus content-type and rels entries."""
    with zipfile.ZipFile(path) as zin:
        entries = {info.filename: [info, zin.read(info.filename)] for info in zin.infolist()}
    rels_name = "word/_rels/document.xml.rels"
    rels = entries[rels_name][1].decode("utf-8")
    used_ids = re.findall(r'Id="rId(\d+)"', rels)
    next_id = max((int(i) for i in used_ids), default=0) + 1
    rel = (
        f'<Relationship Id="rId{next_id}" '
        f'Type="{rel_type}" Target="{part_name.split("/")[-1]}"/>'
    )
    rels = rels.replace("</Relationships>", rel + "</Relationships>")
    ct = entries["[Content_Types].xml"][1].decode("utf-8")
    override = (
        f'<Override PartName="/{part_name}" '
        f'ContentType="{content_type}"/>'
    )
    ct = ct.replace("</Types>", override + "</Types>")
    entries[part_name] = [None, xml.encode("utf-8")]
    entries[rels_name] = [None, rels.encode("utf-8")]
    entries["[Content_Types].xml"] = [None, ct.encode("utf-8")]
    with zipfile.ZipFile(path, "w") as zout:
        for name, (info, data) in entries.items():
            if info is None:
                info = zipfile.ZipInfo(name)
                info.compress_type = zipfile.ZIP_DEFLATED
            zout.writestr(info, data)


def make_notes_docx(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph("正文带脚注")
    run = paragraph.add_run()
    rpr = run._r.makeelement(f"{{{NS_W}}}rPr", {})
    rstyle = run._r.makeelement(f"{{{NS_W}}}rStyle", {f"{{{NS_W}}}val": "FootnoteReference"})
    rpr.append(rstyle)
    run._r.append(rpr)
    ref = run._r.makeelement(f"{{{NS_W}}}footnoteReference", {f"{{{NS_W}}}id": "1"})
    run._r.append(ref)
    document.save(path)
    _inject_note_part(
        path,
        "word/footnotes.xml",
        FOOTNOTES_TEMPLATE,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
    )
    _inject_note_part(
        path,
        "word/endnotes.xml",
        ENDNOTES_TEMPLATE,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml",
    )


def _edit_part_paragraph(workdir: Path, part_id: str, old: str, new: str) -> None:
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = f'<!--@p id="{part_id}"-->'
    start = text.index(marker)
    end = text.find("\n\n", start)
    if end < 0:
        end = len(text)
    block = text[start:end]
    assert old in block, f"'{old}' not in block: {block}"
    text = text[:start] + block.replace(old, new, 1) + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")


def test_extract_exposes_header_footer_with_partitions(tmp_path):
    source = tmp_path / "hf.docx"
    make_header_footer_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert '<!--@part key="header1"' in typed
    assert '<!--@part key="header2"' in typed
    assert '<!--@part key="footer1"' in typed
    assert 'id="header1.P0"' in typed
    assert 'id="header2.P0"' in typed
    # document order: headers, body, footers
    order = [l for l in typed.splitlines() if l.startswith("<!--@part") or l.startswith('<!--@p id="P0"')]
    header1_pos = next(i for i, l in enumerate(order) if l.startswith('<!--@part key="header1"'))
    p0_pos = next(i for i, l in enumerate(order) if l.startswith('<!--@p id="P0"'))
    footer1_pos = next(i for i, l in enumerate(order) if l.startswith('<!--@part key="footer1"'))
    assert header1_pos < p0_pos < footer1_pos
    edit = (workdir / "edit.md").read_text(encoding="utf-8")
    assert '<!--@part key="header1"' in edit


def test_noop_roundtrip_all_parts_byte_identical(tmp_path):
    source = tmp_path / "hf.docx"
    make_header_footer_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    output = tmp_path / "noop.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(source) as z:
        source_entries = {name: z.read(name) for name in z.namelist()}
    with zipfile.ZipFile(output) as z:
        output_entries = {name: z.read(name) for name in z.namelist()}
    assert set(source_entries) == set(output_entries)
    for name, data in source_entries.items():
        assert data == output_entries[name], f"part changed: {name}"


def test_edit_header_paragraph(tmp_path):
    source = tmp_path / "hf.docx"
    make_header_footer_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    _edit_part_paragraph(workdir, "header1.P0", "页眉文字", "页眉改字")
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["header1.P0"]
    output = tmp_path / "edited.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        header_xml = archive.read("word/header1.xml").decode("utf-8")
        body_xml = archive.read("word/document.xml").decode("utf-8")
    assert "页眉改字" in header_xml
    assert "正文" in body_xml
    # second-section header untouched
    with zipfile.ZipFile(output) as archive:
        header2_xml = archive.read("word/header2.xml").decode("utf-8")
    assert "第二节页眉" in header2_xml


def test_edit_footer_paragraph(tmp_path):
    source = tmp_path / "hf.docx"
    make_header_footer_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    _edit_part_paragraph(workdir, "footer1.P0", "页脚文字", "页脚改字")
    sync_edit_projection(workdir)
    output = tmp_path / "edited.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        footer_xml = archive.read("word/footer1.xml").decode("utf-8")
    assert "页脚改字" in footer_xml


def test_footnotes_endnotes_extract_and_edit(tmp_path):
    source = tmp_path / "notes.docx"
    make_notes_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    ids = [p.paragraph_id for p in typed.paragraphs]
    assert "footnotes.P2" in ids  # P0/P1 are separator/continuation entries
    assert "endnotes.P2" in ids
    # footnote entry id recorded in the format baseline (the real note has w:id=1)
    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    footnote_record = next(r for r in format_data["paragraphs"] if r["id"] == "footnotes.P2")
    assert footnote_record["part_entry_id"] == "1"
    # no-op roundtrip byte-identical (injected parts survive)
    output = tmp_path / "noop.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(source) as z:
        source_notes = z.read("word/footnotes.xml")
    with zipfile.ZipFile(output) as z:
        output_notes = z.read("word/footnotes.xml")
    assert source_notes == output_notes
    # edit the footnote text
    _edit_part_paragraph(workdir, "footnotes.P2", "脚注内容", "脚注改字")
    sync_edit_projection(workdir)
    output2 = tmp_path / "edited.docx"
    assert build([str(workdir), "-o", str(output2)]) == 0
    assert verify([str(workdir), str(output2)]) == 0
    with zipfile.ZipFile(output2) as archive:
        notes_xml = archive.read("word/footnotes.xml").decode("utf-8")
    assert "脚注改字" in notes_xml


def test_part_paragraph_ops_rejected(tmp_path):
    source = tmp_path / "hf.docx"
    make_header_footer_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = '<!--@p id="header1.P0"-->'
    start = text.index(marker)
    end = text.find("\n\n", start)
    if end < 0:
        end = len(text)
    text = text[:start] + '<!--@delete id="header1.P0"-->' + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")
    try:
        sync_edit_projection(workdir)
        raise AssertionError("expected table-structure-immutable")
    except Exception as exc:
        assert "table-structure-immutable" in str(exc)


def test_tracked_edit_in_header(tmp_path):
    source = tmp_path / "hf.docx"
    make_header_footer_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    _edit_part_paragraph(workdir, "header1.P0", "页眉文字", "页眉改字")
    sync_edit_projection(workdir, track=True)
    output = tmp_path / "tracked.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        header_xml = archive.read("word/header1.xml").decode("utf-8")
    assert "<w:ins" in header_xml


def test_partition_roundtrip_restores_body_context(tmp_path):
    """Regression: typed.md must reset the part context when serialization
    returns to body text (empty <!--@part key=""--> marker)."""
    from scripts.typed_core import parse_typed as parse

    source = tmp_path / "hf.docx"
    make_header_footer_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    typed = parse((workdir / "typed.md").read_text(encoding="utf-8"))
    by_id = {p.paragraph_id: p for p in typed.paragraphs}
    assert by_id["header1.P0"].part_key == "header1"
    assert by_id["P0"].part_key == ""  # body restored, not stuck on header2
    assert by_id["footer1.P0"].part_key == "footer1"
    # edit.md partition markers round-trip through sync
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    assert '<!--@part key="header1"' in text


def test_header_table_cell_editable(tmp_path):
    """Regression: prose inside header/footer tables joins the editable
    surface."""
    from docx.shared import Inches

    document = Document()
    header = document.sections[0].header
    header.paragraphs[0].text = "页眉前"
    tbl = header.add_table(rows=1, cols=1, width=Inches(3))
    tbl.cell(0, 0).text = "页眉表格单元"
    document.add_paragraph("正文")
    source = tmp_path / "hf.docx"
    document.save(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert "header1.T0.R0.C0.P0" in typed
    # no-op byte-identical
    output = tmp_path / "noop.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(source) as z:
        source_header = z.read("word/header1.xml")
    with zipfile.ZipFile(output) as z:
        output_header = z.read("word/header1.xml")
    assert source_header == output_header
    # edit the cell
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = '<!--@p id="header1.T0.R0.C0.P0"-->'
    start = text.index(marker)
    end = text.find("\n\n", start)
    if end < 0:
        end = len(text)
    block = text[start:end]
    assert "页眉表格单元" in block
    text = text[:start] + block.replace("页眉表格单元", "页眉表格改", 1) + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")
    sync_edit_projection(workdir)
    output2 = tmp_path / "edited.docx"
    assert build([str(workdir), "-o", str(output2)]) == 0
    assert verify([str(workdir), str(output2)]) == 0
    with zipfile.ZipFile(output2) as z:
        edited_header = z.read("word/header1.xml").decode("utf-8")
    assert "页眉表格改" in edited_header

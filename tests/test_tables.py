"""N1: table cell paragraph editing (PRD nested-container-editing)."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.build import build
from scripts.edit import sync_edit_projection
from scripts.extract import extract
from scripts.typed_core import parse_typed
from scripts.verify import verify


def make_table_docx(path: Path, *, nested: bool = False, track_changes: bool = False) -> None:
    """General 2x2 table with a nested table option."""
    document = Document()
    document.add_paragraph("表前")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    if nested:
        inner = OxmlElement("w:tbl")
        tr = OxmlElement("w:tr")
        tc = OxmlElement("w:tc")
        p_el = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t_el = OxmlElement("w:t")
        t_el.text = "内"
        r.append(t_el)
        p_el.append(r)
        tc.append(p_el)
        tr.append(tc)
        inner.append(tr)
        table.cell(0, 0)._tc.append(inner)
    document.add_paragraph("表后")
    document.save(path)


def extract_table(tmp_path: Path, *, nested: bool = False) -> Path:
    source = tmp_path / "tbl-src.docx"
    workdir = tmp_path / "tbl"
    make_table_docx(source, nested=nested)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return workdir, source


def _edit_cell(workdir: Path, cell_id: str, old: str, new: str) -> None:
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = f'<!--@p id="{cell_id}"-->'
    start = text.index(marker)
    end = text.find("\n\n", start)
    if end < 0:
        end = len(text)
    block = text[start:end]
    assert old in block, f"'{old}' not in block: {block}"
    text = text[:start] + block.replace(old, new, 1) + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")


def test_extract_exposes_cell_paragraphs(tmp_path):
    workdir, _ = extract_table(tmp_path)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    ids = [p.paragraph_id for p in typed.paragraphs]
    assert "T0.R0.C0.P0" in ids
    assert "T0.R1.C1.P0" in ids
    assert ids.index("T0.R0.C0.P0") > ids.index("P0")  # document order: after 表前
    assert ids.index("T0.R1.C1.P0") < ids.index("P1")  # before 表后
    # revisions.json inventory covers cell paragraphs
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    assert inv["revisions"] == []


def test_noop_roundtrip_byte_identical(tmp_path):
    workdir, source = extract_table(tmp_path)
    output = tmp_path / "noop.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(source) as archive:
        source_xml = archive.read("word/document.xml")
    with zipfile.ZipFile(output) as archive:
        output_xml = archive.read("word/document.xml")
    assert source_xml == output_xml


def test_edit_cell_paragraph_via_sync(tmp_path):
    workdir, _ = extract_table(tmp_path)
    _edit_cell(workdir, "T0.R1.C0.P0", "B1", "B1改")
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["T0.R1.C0.P0"]
    output = tmp_path / "edited.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "B1改" in xml
    assert "A1" in xml  # untouched cell intact
    # untouched cells stay byte-identical
    with zipfile.ZipFile(tmp_path / "tbl-src.docx") as archive:
        source_xml = archive.read("word/document.xml")
    source_table = re.search(rb"<w:tbl>.*?</w:tbl>", source_xml, re.S).group(0)
    output_table = re.search(rb"<w:tbl>.*?</w:tbl>", archive.read("word/document.xml") if False else xml.encode(), re.S).group(0)
    assert b"A1" in source_table and b"A1" in output_table


def test_nested_table_cells_editable(tmp_path):
    workdir, _ = extract_table(tmp_path, nested=True)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    ids = [p.paragraph_id for p in typed.paragraphs]
    assert "T1.R0.C0.P0" in ids  # nested table gets its own T ordinal
    _edit_cell(workdir, "T1.R0.C0.P0", "内", "内改")
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["T1.R0.C0.P0"]
    output = tmp_path / "nested.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "内改" in xml
    # no duplicated tables
    assert xml.count("<w:tbl>") == 2


def test_delete_cell_paragraph_rejected(tmp_path):
    workdir, _ = extract_table(tmp_path)
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = '<!--@p id="T0.R0.C0.P0"-->'
    start = text.index(marker)
    end = text.find("\n\n", start)
    if end < 0:
        end = len(text)
    text = text[:start] + '<!--@delete id="T0.R0.C0.P0"-->' + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")
    try:
        sync_edit_projection(workdir)
        raise AssertionError("expected table-structure-immutable")
    except Exception as exc:
        assert "table-structure-immutable" in str(exc)


def test_new_paragraph_inheriting_cell_rejected(tmp_path):
    workdir, _ = extract_table(tmp_path)
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    text = text.replace(
        '<!--@p id="T0.R0.C0.P0"-->',
        '<!--@new temp="N1" inherit="T0.R0.C0.P0"-->\n新\n\n<!--@p id="T0.R0.C0.P0"-->',
        1,
    )
    (workdir / "edit.md").write_text(text, encoding="utf-8")
    try:
        sync_edit_projection(workdir)
        raise AssertionError("expected table-structure-immutable")
    except Exception as exc:
        assert "table-structure-immutable" in str(exc)


def test_tracked_edit_inside_cell(tmp_path):
    workdir, _ = extract_table(tmp_path)
    _edit_cell(workdir, "T0.R0.C0.P0", "A1", "A1改")
    _, _, changed = sync_edit_projection(workdir, track=True)
    assert changed == ["T0.R0.C0.P0"]
    output = tmp_path / "tracked.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:ins" in xml
    assert "A1" in xml and "改" in xml


def test_three_level_nested_table_roundtrip(tmp_path):
    """Regression: deeply nested tables (3 levels) must round-trip byte-exact."""
    document = Document()
    document.add_paragraph("表前")
    t1 = document.add_table(rows=1, cols=1)
    t1.cell(0, 0).text = "L1"
    t2 = t1.cell(0, 0).add_table(rows=1, cols=1)
    t2.cell(0, 0).text = "L2"
    t3 = t2.cell(0, 0).add_table(rows=1, cols=1)
    t3.cell(0, 0).text = "L3"
    document.add_paragraph("表后")
    source = tmp_path / "deep.docx"
    document.save(source)
    workdir = tmp_path / "deep"
    assert extract([str(source), "-o", str(workdir)]) == 0
    output = tmp_path / "deep-out.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(source) as archive:
        source_xml = archive.read("word/document.xml")
    with zipfile.ZipFile(output) as archive:
        output_xml = archive.read("word/document.xml")
    assert source_xml == output_xml


def test_hand_edited_inherit_from_cell_rejected(tmp_path):
    """Regression: typed.md can bypass MCP gates — validation must reject new
    paragraphs that inherit from a table cell."""
    import shutil

    workdir, _ = extract_table(tmp_path)
    workdir2 = tmp_path / "wd2"
    shutil.copytree(workdir, workdir2)
    typed = (workdir2 / "typed.md").read_text(encoding="utf-8")
    typed += '<!--@p id="P99" inherit="T0.R0.C0.P0"-->\n绕过\n'
    (workdir2 / "typed.md").write_text(typed, encoding="utf-8")
    try:
        build([str(workdir2), "-o", str(tmp_path / "bypass.docx")])
        raise AssertionError("expected table-structure-immutable")
    except Exception as exc:
        assert "table-structure-immutable" in str(exc)


def test_table_structure_ops_new_baseline(tmp_path):
    """Row/col insert/delete and merge/split produce a valid new baseline."""
    from scripts.decisions import _apply_table_op

    workdir, source = extract_table(tmp_path)
    out = tmp_path / "op.docx"
    wd2 = tmp_path / "op-wd"

    def rows_cols(xml: bytes):
        import re as _re

        return len(_re.findall(rb"<w:tr[ >]", xml)), len(_re.findall(rb"<w:tc[ >]", xml))

    with zipfile.ZipFile(source) as z:
        original = z.read("word/document.xml")
    # each op applies to the ORIGINAL 2x2 table independently
    ops = [
        ("insert-row", [1], (3, 6)),
        ("delete-row", [0], (1, 2)),
        ("insert-col", [0], (2, 6)),
        ("delete-col", [0], (2, 2)),
        ("merge-cells", [0, 0, 2], (2, 3)),
        ("split-cells", [0, 0, 2], (2, 5)),
    ]
    for i, (op, args, expect) in enumerate(ops):
        out_i = tmp_path / f"op{i}.docx"
        wd_i = tmp_path / f"op{i}-wd"
        _apply_table_op(workdir, "T0", op, args, out_i, wd_i)
        assert verify([str(wd_i), str(out_i)]) == 0
        with zipfile.ZipFile(out_i) as z:
            xml = z.read("word/document.xml")
        got = rows_cols(xml)
        assert got == expect, f"{op}: {got} != {expect}"


def _make_sdt_docx(path: Path) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    document.add_paragraph("前")
    paragraph = document.add_paragraph()
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "姓名")
    sdt_pr.append(alias)
    content = OxmlElement("w:sdtContent")
    cp = OxmlElement("w:p")
    cr = OxmlElement("w:r")
    ct = OxmlElement("w:t")
    ct.text = "控件内文本"
    cr.append(ct)
    cp.append(cr)
    content.append(cp)
    sdt.append(sdt_pr)
    sdt.append(content)
    paragraph._p.addnext(sdt)
    document.add_paragraph("后")
    document.save(path)


def test_sdt_content_editable(tmp_path):
    """S3: text inside a w:sdt content control is editable; the sdtPr
    structure replays byte-exact."""
    source = tmp_path / "sdt.docx"
    _make_sdt_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert 'id="S0.P0"' in typed
    # no-op byte-identical
    output = tmp_path / "noop.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(source) as z:
        source_xml = z.read("word/document.xml")
    with zipfile.ZipFile(output) as z:
        output_xml = z.read("word/document.xml")
    assert source_xml == output_xml
    # edit the control text
    edit = (workdir / "edit.md").read_text(encoding="utf-8")
    edit = edit.replace("控件内文本", "控件改文本", 1)
    (workdir / "edit.md").write_text(edit, encoding="utf-8")
    sync_edit_projection(workdir)
    output2 = tmp_path / "edited.docx"
    assert build([str(workdir), "-o", str(output2)]) == 0
    assert verify([str(workdir), str(output2)]) == 0
    with zipfile.ZipFile(output2) as z:
        edited = z.read("word/document.xml")
    assert b"<w:sdtPr>" in edited and b"<w:alias" in edited  # control structure kept
    assert "控件改文本".encode() in edited

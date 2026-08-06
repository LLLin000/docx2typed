"""R2: tracked run-text editing (ADR 0037)."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.build import build
from scripts.edit import refresh_edit_projection, sync_edit_projection
from scripts.extract import extract
from scripts.typed_core import RevisionNode, parse_typed, visible_text
from scripts.verify import verify


def make_trackable_docx(path: Path, *, track_changes: bool = False) -> None:
    """Plain docx; optionally with settings.xml <w:trackChanges/> enabled."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前文甲后文")
    document.add_paragraph("第二段")
    document.save(path)
    if track_changes:
        with zipfile.ZipFile(path) as archive:
            settings = archive.read("word/settings.xml")
        root = re.sub(rb"(<w:settings[^>]*>)", rb"\1<w:trackChanges/>", settings, count=1)
        with zipfile.ZipFile(path) as zin, zipfile.ZipFile(path.with_suffix(".tmp"), "w") as zout:
            for info in zin.infolist():
                data = root if info.filename == "word/settings.xml" else zin.read(info.filename)
                zout.writestr(info, data)
        path.with_suffix(".tmp").replace(path)


def extract_trackable(tmp_path: Path, *, track_changes: bool = False) -> Path:
    source = tmp_path / "track-src.docx"
    workdir = tmp_path / "track"
    make_trackable_docx(source, track_changes=track_changes)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return workdir


def _edit_paragraph(workdir: Path, paragraph_id: str, old: str, new: str) -> None:
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = f'<!--@p id="{paragraph_id}"-->'
    start = text.index(marker)
    end = text.index("\n\n", start)
    block = text[start:end]
    assert old in block, f"'{old}' not in block: {block}"
    text = text[:start] + block.replace(old, new, 1) + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")


def _revisions(typed) -> list[RevisionNode]:
    from scripts.typed_core import RangeNode

    def walk(nodes):
        for node in nodes:
            if isinstance(node, RevisionNode):
                yield node
            if isinstance(node, (RevisionNode, RangeNode)):
                yield from walk(node.children)

    return [n for n in walk(node for p in typed.paragraphs for node in p.nodes)]


def test_track_mode_insert_generates_ins(tmp_path):
    workdir = extract_trackable(tmp_path, track_changes=True)
    _edit_paragraph(workdir, "P0", "前文甲后文", "前文甲新词后文")
    _, _, changed = sync_edit_projection(workdir, track=True, author="测试者")
    assert changed == ["P0"]
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    revisions = _revisions(typed)
    assert len(revisions) == 1
    node = revisions[0]
    assert node.kind == "insert"
    assert node.attrs["w:author"] == "测试者"
    assert node.attrs["w:id"] == "0"  # lowest available over the package
    assert visible_text(node.children) == "新词"
    output = tmp_path / "tracked.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert '<w:ins w:id="0" w:author="测试者"' in xml
    assert "<w:t>新词</w:t>" in xml


def test_track_mode_delete_generates_del(tmp_path):
    workdir = extract_trackable(tmp_path, track_changes=True)
    _edit_paragraph(workdir, "P0", "前文甲后文", "前文后文")
    _, _, changed = sync_edit_projection(workdir, track=True)
    assert changed == ["P0"]
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    revisions = _revisions(typed)
    assert len(revisions) == 1
    assert revisions[0].kind == "delete"
    assert visible_text([revisions[0]]) == ""  # final view hides deletions
    from scripts.typed_core import visible_text_original

    assert visible_text_original(revisions[0].children) == "甲"
    output = tmp_path / "tracked.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:del" in xml and "<w:delText>甲</w:delText>" in xml


def test_track_mode_replace_generates_del_and_ins(tmp_path):
    workdir = extract_trackable(tmp_path, track_changes=True)
    _edit_paragraph(workdir, "P0", "前文甲后文", "前文乙后文")
    _, _, changed = sync_edit_projection(workdir, track=True)
    assert changed == ["P0"]
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    kinds = [node.kind for node in _revisions(typed)]
    assert kinds == ["delete", "insert"]
    output = tmp_path / "tracked.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:del" in xml and "<w:delText>甲</w:delText>" in xml
    assert "<w:ins" in xml and "<w:t>乙</w:t>" in xml


def test_edit_inside_existing_insertion_nests(tmp_path):
    """Editing text inside an existing w:ins nests the new del/ins inside it."""
    from tests.test_revisions import make_revision_docx

    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    make_revision_docx(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    # "插入词" is inside a w:ins; change part of it
    _edit_paragraph(workdir, "P0", "插入词", "插入语")
    _, _, changed = sync_edit_projection(workdir, track=True, author="嵌套测试")
    assert changed == ["P0"]
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    outer = [n for n in _revisions(typed) if n.attrs.get("w:id") == "100"][0]
    nested = [n for n in _walk_children(outer) if isinstance(n, RevisionNode)]
    assert len(nested) == 2  # del + ins nested inside the outer insertion
    assert {n.kind for n in nested} == {"delete", "insert"}
    assert all(n.attrs["w:author"] == "嵌套测试" for n in nested)
    output = tmp_path / "nested.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    match = re.search(r"<w:ins[^>]*w:id=\"100\".*?</w:ins>", xml, re.S)
    assert match is not None
    inner = match.group(0)
    assert "<w:del" in inner and "<w:delText>词</w:delText>" in inner
    assert "<w:ins" in inner and "<w:t>语</w:t>" in inner


def _walk_children(node):
    yield from node.children


def test_direct_mode_rejects_revision_text_mutation(tmp_path):
    from tests.test_revisions import make_revision_docx

    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    make_revision_docx(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    _edit_paragraph(workdir, "P0", "插入词", "插入语")  # text inside an existing insertion
    try:
        sync_edit_projection(workdir, track=False)
        raise AssertionError("expected revision-text-mutated-in-direct-mode")
    except Exception as exc:
        assert "revision-text-mutated-in-direct-mode" in str(exc)


def test_ambiguous_mode_blocks_until_choice(tmp_path):
    from tests.test_revisions import make_revision_docx

    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    make_revision_docx(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    _edit_paragraph(workdir, "P0", "前文", "前文改")  # plain text, no revision involvement
    try:
        sync_edit_projection(workdir)
        raise AssertionError("expected edit-mode-ambiguous")
    except Exception as exc:
        assert "edit-mode-ambiguous" in str(exc)
    # explicit choice unblocks
    _, _, changed = sync_edit_projection(workdir, track=True)
    assert changed == ["P0"]


def test_track_mode_new_paragraph_gets_mark_revision(tmp_path):
    """R2.5: @new in track mode stays in the document with a paragraph-mark
    insertion revision instead of being rejected."""
    workdir = extract_trackable(tmp_path, track_changes=True)
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    text = text.replace(
        '<!--@p id="P1"-->',
        '<!--@new temp="N1" inherit="P0"-->\n新段落\n\n<!--@p id="P1"-->',
        1,
    )
    (workdir / "edit.md").write_text(text, encoding="utf-8")
    _, _, changed = sync_edit_projection(workdir, track=True)
    assert changed == ["P2"]
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    new_paragraph = next(p for p in typed.paragraphs if p.inherit)
    assert new_paragraph.mark_revision is not None
    assert new_paragraph.mark_revision["kind"] == "insert"
    assert "w:author" in new_paragraph.mark_revision["attrs"]
    output = tmp_path / "marked.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert re.search(r"<w:pPr>.*?<w:rPr>.*?<w:ins ", xml, re.S) is not None


def test_track_mode_without_track_changes_requires_explicit_flag(tmp_path):
    """trackChanges off + no pending revisions infers direct; --track must still
    work as an explicit override but the state stays auditable."""
    workdir = extract_trackable(tmp_path)  # no track changes
    _edit_paragraph(workdir, "P0", "前文甲后文", "前文甲改后文")
    _, _, changed = sync_edit_projection(workdir, track=True, author="Lin")
    assert changed == ["P0"]
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    node = _revisions(typed)[0]
    assert node.kind == "insert"
    assert node.attrs["w:author"] == "Lin"
    output = tmp_path / "forced.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0


def test_evidence_records_generated_revisions(tmp_path):
    workdir = extract_trackable(tmp_path, track_changes=True)
    _edit_paragraph(workdir, "P0", "前文甲后文", "前文乙后文")
    sync_edit_projection(workdir, track=True, author="记录者")
    evidence = json.loads(
        (workdir / "edit.state.json.run.json").read_text(encoding="utf-8")
    )
    assert evidence["edit_mode"] == "track"
    assert evidence["author"] == "记录者"
    assert evidence["author_source"] == "parameter"
    generated = evidence["generated_revisions"]
    assert len(generated) == 2
    kinds = {entry["kind"] for entry in generated}
    assert kinds == {"delete", "insert"}
    # ids must be distinct within the run
    assert generated[0]["w_id"] != generated[1]["w_id"]


def test_second_sync_reuses_lowest_available_id(tmp_path):
    workdir = extract_trackable(tmp_path, track_changes=True)
    _edit_paragraph(workdir, "P0", "前文甲后文", "前文甲一后文")
    sync_edit_projection(workdir, track=True)
    # the projection now shows the insertion block; extend its text
    _edit_paragraph(workdir, "P0", "一", "一二")
    sync_edit_projection(workdir, track=True)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    ids = [node.attrs["w:id"] for node in _revisions(typed)]
    assert ids == ["0", "1"]  # allocated lowest available, strictly increasing


def test_date_utc_only_when_source_uses_it(tmp_path):
    workdir = extract_trackable(tmp_path, track_changes=True)
    _edit_paragraph(workdir, "P0", "前文甲后文", "前文甲改后文")
    sync_edit_projection(workdir, track=True)
    output = tmp_path / "plain.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "dateUtc" not in xml  # source never used w16du:dateUtc


def test_nested_edit_does_not_duplicate_unchanged_text(tmp_path):
    """Regression: rebuilding an edited insertion must replace its children,
    not append the nested del/ins after the original text."""
    from tests.test_revisions import make_revision_docx

    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    make_revision_docx(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    _edit_paragraph(workdir, "P0", "插入词", "插入语")
    sync_edit_projection(workdir, track=True)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    visible = "".join(
        visible_text([node]) for node in typed.paragraphs[0].nodes
    )
    assert visible == "前文插入语后文批注目标"  # no duplicated 插入词


def test_direct_mode_rejects_pure_insertion_into_revision_text(tmp_path):
    """Regression: a pure insertion (empty baseline hunk) inside an existing
    revision must also be rejected in direct mode."""
    from tests.test_revisions import make_revision_docx

    source = tmp_path / "src.docx"
    workdir = tmp_path / "wd"
    make_revision_docx(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    _edit_paragraph(workdir, "P0", "插入词", "插入词啊")
    try:
        sync_edit_projection(workdir, track=False)
        raise AssertionError("expected revision-text-mutated-in-direct-mode")
    except Exception as exc:
        assert "revision-text-mutated-in-direct-mode" in str(exc)

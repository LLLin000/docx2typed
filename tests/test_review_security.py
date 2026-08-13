"""Issue #51: adversarial checks for the secured single-session review server.

Every case is exercised through public HTTP behavior (or public primitives)
against the issue #31 contract:

- bootstrap shell is public and data-free; capability is memory-only,
  cleared from the fragment by the page JS, and revoked by restart
- protected routes reject missing/invalid authority with a uniform
  detail-free 404
- Host allowlist defeats DNS rebinding; X-Forwarded-* is ignored
- mutating POSTs require Origin == advertised origin and
  Sec-Fetch-Site: same-origin; text/plain/forms/multipart and oversized
  bodies are rejected without mutation
- no CORS; no cross-origin OPTIONS; response security headers everywhere
- GET/HEAD are strictly side-effect-free
- logs omit Authorization/fragment/query/body/paths
- unauthorized throttling is bounded and never degrades authorized use
- all five review mutation surfaces share the store CAS Writer lane with
  Idempotency-Key/ledger replay; restart keeps committed state
"""
from __future__ import annotations

import json
import re
import socket
import subprocess
import sys
import threading
import time
import urllib.error
import urllib.request
from pathlib import Path

import pytest

from scripts import main
from scripts.review_security import (
    UnauthorizedThrottle,
    build_allowlist,
    constant_time_equal,
    content_type_allowed,
    generate_capability,
    session_fingerprint,
    split_host,
)

MAX_BODY = 256 * 1024
LOOPBACK = ("127.0.0.1",)
ROOT = Path(__file__).resolve().parents[1]


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------

def _make_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("第二段")
    document.save(path)


def _extract(tmp_path: Path, name: str = "wd") -> Path:
    source = tmp_path / f"{name}-src.docx"
    _make_docx(source)
    workdir = tmp_path / name
    assert main(["extract", "--json", str(source), "-o", str(workdir), "--operation-id", f"sec-{name}-extract"]) == 0
    return workdir


def _free_port() -> int:
    sock = socket.socket()
    sock.bind(LOOPBACK + (0,))
    port = sock.getsockname()[1]
    sock.close()
    return port


def _serve(workdir: Path, *, socket_timeout: float | None = None):
    """Start one hardened review server; return (server, capability, port).
    ``socket_timeout`` overrides the per-connection read bound (tests use a
    short value to prove stalled connections release their threads)."""
    from http.server import ThreadingHTTPServer

    from scripts.review_security import ReviewSecurity
    from scripts.review_server import _handler_for

    port = _free_port()
    hosts, origins = build_allowlist("127.0.0.1", port)
    security = ReviewSecurity(
        capability=generate_capability(),
        allowed_hosts=hosts,
        allowed_origins=origins,
        port=port,
        source_label="loopback",
    )
    handler = _handler_for(workdir, security) if socket_timeout is None else _handler_for(
        workdir, security, socket_timeout=socket_timeout
    )
    server = ThreadingHTTPServer(("127.0.0.1", port), handler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    return server, security.capability, port


def _raw_request(port: int, payload: bytes, *, recv_timeout: float = 10.0) -> bytes:
    """Send raw bytes over one connection and read until EOF. Used to send
    malformed headers (Content-Length: abc) that urllib would rewrite."""
    with socket.create_connection(("127.0.0.1", port), timeout=recv_timeout) as sock:
        sock.sendall(payload)
        data = b""
        while True:
            chunk = sock.recv(4096)
            if not chunk:
                break
            data += chunk
    return data


def _opener():
    # The workstation runs a system proxy; loopback review traffic must not
    # travel through it (it would answer with proxy 502s and hide the server).
    return urllib.request.build_opener(urllib.request.ProxyHandler({}))


def _request(
    port: int,
    path: str,
    *,
    method: str = "GET",
    capability: str | None = None,
    host: str | None = None,
    origin: str | None = None,
    fetch_site: str | None = None,
    content_type: str | None = "application/json",
    idempotency_key: str | None = None,
    body: bytes | None = None,
    extra_headers: dict[str, str] | None = None,
):
    request = urllib.request.Request(f"http://127.0.0.1:{port}{path}", data=body, method=method)
    request.add_header("Host", host or f"127.0.0.1:{port}")
    if capability is not None:
        request.add_header("Authorization", f"Bearer {capability}")
    if origin is not None:
        request.add_header("Origin", origin)
    if fetch_site is not None:
        request.add_header("Sec-Fetch-Site", fetch_site)
    if content_type is not None:
        request.add_header("Content-Type", content_type)
    if idempotency_key is not None:
        request.add_header("Idempotency-Key", idempotency_key)
    for name, value in (extra_headers or {}).items():
        request.add_header(name, value)
    try:
        with _opener().open(request, timeout=15) as response:
            return response.status, dict(response.headers), response.read()
    except urllib.error.HTTPError as exc:
        return exc.code, dict(exc.headers), exc.read()


def _json_body(status: int, headers: dict[str, str], body: bytes) -> dict:
    assert body, f"expected a JSON body, got status {status} {body!r}"
    return json.loads(body.decode("utf-8"))


def _decision_payload(client_id: str = "browser-1") -> dict:
    return {
        "type": "decision",
        "client_id": client_id,
        "review_item_id": "P16",
        "paragraph_id": "P16",
        "decision": "accept",
        "revision_key": "P0-r1",
        "revision_id": "r1",
        "selected_text": "旧词",
        "comment": "",
    }


def _generations(workdir: Path) -> set[str]:
    return {p.name for p in (workdir / ".docx2typed-store" / "generations").iterdir()}


def _stop(server) -> None:
    server.shutdown()
    server.server_close()


# --------------------------------------------------------------------------
# capability primitives
# --------------------------------------------------------------------------

class TestCapabilityPrimitives:
    def test_capability_is_256_bit_base64url_without_padding(self):
        token = generate_capability()
        # 32 bytes -> ceil(32*8/6) = 43 base64 chars, no '=' padding.
        assert len(token) == 43
        assert re.fullmatch(r"[A-Za-z0-9_-]{43}", token)
        assert "=" not in token

    def test_capabilities_are_unique_and_constant_time_compared(self):
        first, second = generate_capability(), generate_capability()
        assert first != second
        assert constant_time_equal(first, first)
        assert not constant_time_equal(first, second)
        assert not constant_time_equal(first, "")

    def test_session_fingerprint_is_truncated_and_irreversible(self):
        token = generate_capability()
        digest = session_fingerprint(token)
        assert re.fullmatch(r"[0-9a-f]{10}", digest)
        assert session_fingerprint(token) == digest  # deterministic
        assert session_fingerprint(token + "x") != digest

    def test_build_allowlist_loopback_and_tailscale(self):
        hosts, origins = build_allowlist("127.0.0.1", 8876)
        assert "127.0.0.1:8876" in hosts and "localhost:8876" in hosts
        assert "http://127.0.0.1:8876" in origins and "http://localhost:8876" in origins
        assert "0.0.0.0:8876" not in hosts and "attacker.example:8876" not in hosts
        # Port 80 admits the port-less browser form.
        hosts80, origins80 = build_allowlist("127.0.0.1", 80)
        assert "127.0.0.1" in hosts80 and "http://127.0.0.1" in origins80
        # A Tailscale bind allows exactly the discovered address.
        hosts_t, origins_t = build_allowlist("100.64.0.7", 8876)
        assert hosts_t == frozenset({"100.64.0.7:8876"})
        assert origins_t == frozenset({"http://100.64.0.7:8876"})

    def test_split_host_rejects_ipv6_and_garbage(self):
        assert split_host("127.0.0.1:8876") == ("127.0.0.1", 8876)
        assert split_host("localhost") == ("localhost", None)
        with pytest.raises(ValueError):
            split_host("[::1]:8876")
        with pytest.raises(ValueError):
            split_host("host:notaport")
        with pytest.raises(ValueError):
            split_host("")

    def test_content_type_allowed(self):
        assert content_type_allowed("application/json")
        assert content_type_allowed("application/json; charset=utf-8")
        assert content_type_allowed("application/json;charset=UTF8")
        assert not content_type_allowed("text/plain")
        assert not content_type_allowed("multipart/form-data; boundary=x")
        assert not content_type_allowed("application/x-www-form-urlencoded")
        assert not content_type_allowed("")
        assert not content_type_allowed("application/json; charset=iso-8859-1")

    def test_throttle_bounded_and_refills(self):
        bucket = UnauthorizedThrottle(capacity=3, refill_per_second=10.0, max_clients=2)
        assert [bucket.allow("a") for _ in range(3)] == [True, True, True]
        assert bucket.allow("a") is False  # exhausted
        time.sleep(0.15)
        assert bucket.allow("a") is True  # refilled
        # Bounded client table: a third client evicts the least recent one.
        bucket.allow("b")
        bucket.allow("c")
        assert len(bucket) == 2


# --------------------------------------------------------------------------
# bootstrap shell
# --------------------------------------------------------------------------

class TestBootstrapShell:
    def test_root_is_public_static_shell_without_workdir_data(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            status, headers, body = _request(port, "/")
            assert status == 200
            assert headers["Content-Type"].startswith("text/html")
            page = body.decode("utf-8")
            # Chrome is present; document/review/state data is not.
            assert 'id="workflow-strip"' in page
            assert 'class="document-paper"' in page
            assert "第一段" not in page and "第二段" not in page  # typed.md content
            assert workdir.name not in page  # workdir path/name
            assert ".docx2typed-store" not in page
            assert "session.json" not in page
            assert capability not in page  # never embed the token
            assert "LOCAL SERVER" in page
            # The page JS performs the memory-only bootstrap.
            assert "location.hash" in page and "history.replaceState" in page
            assert "sessionToken" in page and "Authorization" in page
        finally:
            _stop(server)

    def test_shell_is_bound_to_advertised_host(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            status, _, body = _request(port, "/", host="attacker.example:80")
            assert status == 404
            assert json.loads(body) == {"error": "not-found"}
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# uniform authority failure
# --------------------------------------------------------------------------

class TestUniformNotFound:
    def test_all_protected_routes_reject_missing_or_invalid_authority_uniformly(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            routes = ["/api/reviews", "/api/document-state", "/api/document-fragment",
                      "/api/review-history", "/health", "/does-not-exist"]
            # Stay under the unauthorized-throttle capacity so every probe is
            # the contract's uniform detail-free 404 (the flood behavior is
            # covered separately by TestThrottling).
            probes = [(route, None) for route in routes[:3]] + [
                (route, {"Authorization": "Bearer wrong-token"}) for route in routes[3:]
            ]
            bodies: set[bytes] = set()
            for route, headers in probes:
                status, _, body = _request(port, route, capability=None, extra_headers=headers)
                assert status == 404, (route, headers, status)
                assert json.loads(body) == {"error": "not-found"}
                bodies.add(body)
            # The refusal is byte-identical across every route and token.
            assert len(bodies) == 1
        finally:
            _stop(server)

    def test_mutating_post_without_authority_is_uniform_404_not_403(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            payload = json.dumps(_decision_payload()).encode("utf-8")
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       origin=f"http://127.0.0.1:{port}", fetch_site="same-origin",
                                       idempotency_key="revkey00000001", body=payload)
            assert status == 404 and json.loads(body) == {"error": "not-found"}
            # Nothing was created or mutated.
            assert not (workdir / ".review").exists()
            assert _generations(workdir) == _generations(workdir)
        finally:
            _stop(server)

    def test_health_requires_auth_and_hides_workdir(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            status, _, _ = _request(port, "/health")
            assert status == 404
            status, headers, body = _request(port, "/health", capability=capability)
            assert status == 200
            payload = json.loads(body)
            assert payload == {"ok": True, "service": "docx2typed-review"}
            assert "workdir" not in payload
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# Host allowlist / DNS rebinding / forwarded headers
# --------------------------------------------------------------------------

class TestHostAllowlist:
    def test_dns_rebinding_host_fails_even_with_valid_capability(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            for bad_host in ("attacker.example", "attacker.example:80", "evil.com:8876", "127.0.0.1.evil.com:8876"):
                status, _, body = _request(port, "/api/reviews", capability=capability, host=bad_host)
                assert status == 404 and json.loads(body) == {"error": "not-found"}, bad_host
            # The correct loopback forms still work.
            status, _, _ = _request(port, "/api/reviews", capability=capability)
            assert status == 200
            status, _, _ = _request(port, "/api/reviews", capability=capability, host=f"localhost:{port}")
            assert status == 200
        finally:
            _stop(server)

    def test_x_forwarded_headers_are_ignored(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            spoofed = {"X-Forwarded-Host": "attacker.example", "X-Forwarded-For": "203.0.113.9",
                       "X-Forwarded-Proto": "https", "Forwarded": "for=203.0.113.9;host=attacker.example"}
            # Real Host valid -> request proceeds; spoofed headers do nothing.
            status, _, body = _request(port, "/api/reviews", capability=capability, extra_headers=spoofed)
            assert status == 200
            # Real Host invalid -> rejected even with a valid X-Forwarded-Host.
            status, _, body = _request(port, "/api/reviews", capability=capability,
                                       host="attacker.example:80",
                                       extra_headers={"X-Forwarded-Host": f"127.0.0.1:{port}"})
            assert status == 404
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# mutation gates: Origin, Sec-Fetch-Site, content type/size
# --------------------------------------------------------------------------

def _mutating_headers(port: int, capability: str, **overrides) -> dict:
    headers = {
        "origin": f"http://127.0.0.1:{port}",
        "fetch_site": "same-origin",
    }
    headers.update(overrides)
    return headers


class TestMutationGates:
    def test_cross_origin_text_plain_blind_write_fails_without_mutation(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            payload = json.dumps(_decision_payload()).encode("utf-8")
            # (a) Blind cross-origin text/plain write: no capability at all.
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       content_type="text/plain",
                                       origin="http://evil.example", fetch_site="cross-site",
                                       idempotency_key="revkey00000001", body=payload)
            assert status == 404 and json.loads(body) == {"error": "not-found"}
            # (b) Blind cross-origin JSON write: no capability.
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       origin="http://evil.example", fetch_site="cross-site",
                                       idempotency_key="revkey00000001", body=payload)
            assert status == 404
            # (c) Capability-holding cross-origin JSON write: origin gate.
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       capability=capability, origin="http://evil.example",
                                       fetch_site="cross-site", idempotency_key="revkey00000001", body=payload)
            assert status == 403
            assert json.loads(body) == {"error": "forbidden", "code": "origin-mismatch"}
            # Nothing was created: no review dir, no new generation.
            assert not (workdir / ".review").exists()
            assert _generations(workdir) == {next(iter(_generations(workdir)))}
        finally:
            _stop(server)

    def test_mutating_post_requires_origin_and_same_origin_fetch_site(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            payload = json.dumps(_decision_payload()).encode("utf-8")
            # A browser write without Origin fails.
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       capability=capability, content_type="application/json",
                                       idempotency_key="revkey00000001", body=payload,
                                       origin=None, fetch_site="same-origin")
            assert status == 403 and json.loads(body)["code"] == "origin-mismatch"
            # Wrong Origin fails.
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       capability=capability, origin="http://localhost:9999",
                                       fetch_site="same-origin", idempotency_key="revkey00000001", body=payload)
            assert status == 403 and json.loads(body)["code"] == "origin-mismatch"
            # Correct origin but Sec-Fetch-Site not same-origin fails.
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       capability=capability, origin=f"http://127.0.0.1:{port}",
                                       fetch_site="cross-site", idempotency_key="revkey00000001", body=payload)
            assert status == 403 and json.loads(body)["code"] == "fetch-site-mismatch"
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       capability=capability, origin=f"http://127.0.0.1:{port}",
                                       fetch_site=None, idempotency_key="revkey00000001", body=payload)
            assert status == 403 and json.loads(body)["code"] == "fetch-site-mismatch"
            # The localhost loopback form of the advertised origin is accepted.
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       capability=capability, origin=f"http://localhost:{port}",
                                       fetch_site="same-origin", idempotency_key="revkey00000001", body=payload)
            assert status == 200
            # None of the failed attempts mutated anything.
            events = [p for p in (workdir / ".review" / "inbox").glob("*.json")]
            assert len(events) == 1  # only the successful write
        finally:
            _stop(server)

    def test_content_type_and_size_caps(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            payload = json.dumps(_decision_payload()).encode("utf-8")
            for bad_type in ("text/plain", "multipart/form-data; boundary=xyz",
                             "application/x-www-form-urlencoded", "application/xml", ""):
                status, _, body = _request(port, "/api/reviews", method="POST",
                                           capability=capability, content_type=bad_type,
                                           origin=f"http://127.0.0.1:{port}", fetch_site="same-origin",
                                           idempotency_key="revkey00000001", body=payload)
                assert status == 400, bad_type
                assert json.loads(body)["code"] == "unsupported-content-type", bad_type
            # UTF-8 charset suffix is allowed.
            status, _, _ = _request(port, "/api/reviews", method="POST",
                                    capability=capability, content_type="application/json; charset=utf-8",
                                    origin=f"http://127.0.0.1:{port}", fetch_site="same-origin",
                                    idempotency_key="revkey00000002", body=payload)
            assert status == 200
            # Oversized body is rejected before any work. The body is exactly
            # one byte over the cap so the server's bounded drain consumes it
            # fully and closes the connection cleanly.
            oversized = {**_decision_payload(), "pad": "x"}
            over_body = json.dumps(oversized).encode("utf-8")
            over_body += b"x" * (MAX_BODY + 1 - len(over_body))
            assert len(over_body) == MAX_BODY + 1
            status, _, body = _request(port, "/api/reviews", method="POST",
                                       capability=capability, origin=f"http://127.0.0.1:{port}",
                                       fetch_site="same-origin", idempotency_key="revkey00000003", body=over_body)
            assert status == 400
            assert json.loads(body)["code"] == "request-body-too-large"
            # A body just under the cap is accepted.
            ok = json.dumps({**_decision_payload(), "pad": "x" * (MAX_BODY - 4096)}).encode("utf-8")
            status, _, _ = _request(port, "/api/reviews", method="POST",
                                    capability=capability, origin=f"http://127.0.0.1:{port}",
                                    fetch_site="same-origin", idempotency_key="revkey00000004", body=ok)
            assert status == 200
        finally:
            _stop(server)

    def test_bodyless_mutation_is_allowed(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            # The console dispatches without a request body.
            status, _, body = _request(port, "/api/reviews/dispatch", method="POST",
                                       capability=capability, origin=f"http://127.0.0.1:{port}",
                                       fetch_site="same-origin", idempotency_key="revkey00000001",
                                       content_type=None, body=None)
            assert status == 200
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# CORS and response headers
# --------------------------------------------------------------------------

class TestCorsAndHeaders:
    def test_no_cors_headers_on_any_response(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            payload = json.dumps(_decision_payload()).encode("utf-8")
            probes = [
                _request(port, "/"),
                _request(port, "/api/reviews", capability=capability),
                _request(port, "/api/reviews"),
                _request(port, "/api/reviews", method="POST", capability=capability,
                         origin="http://evil.example", fetch_site="cross-site",
                         idempotency_key="revkey00000001", body=payload),
                _request(port, "/api/reviews", method="POST", capability=capability,
                         content_type="text/plain", origin=f"http://127.0.0.1:{port}",
                         fetch_site="same-origin", idempotency_key="revkey00000001", body=b"x"),
            ]
            for status, headers, _body in probes:
                assert not any(name.lower().startswith("access-control-") for name in headers), headers
        finally:
            _stop(server)

    def test_cross_origin_options_is_not_cors_enabled(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            status, headers, _ = _request(port, "/api/reviews", method="OPTIONS",
                                          origin="http://evil.example",
                                          extra_headers={"Access-Control-Request-Method": "POST",
                                                         "Access-Control-Request-Headers": "authorization,content-type"})
            assert status >= 400
            assert not any(name.lower().startswith("access-control-") for name in headers)
            assert not any(name.lower() == "allow" for name in headers)
        finally:
            _stop(server)

    def test_security_headers_present_on_every_response(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            payload = json.dumps(_decision_payload()).encode("utf-8")
            statuses = [
                _request(port, "/"),
                _request(port, "/api/reviews", capability=capability),
                _request(port, "/api/reviews"),
                _request(port, "/api/reviews", method="POST", capability=capability,
                         origin="http://evil.example", fetch_site="cross-site",
                         idempotency_key="revkey00000001", body=payload),
                _request(port, "/api/reviews", method="POST", capability=capability,
                         content_type="text/plain", origin=f"http://127.0.0.1:{port}",
                         fetch_site="same-origin", idempotency_key="revkey00000001", body=b"x"),
            ]
            for status, headers, _ in statuses:
                assert headers["Cache-Control"] == "no-store", status
                assert headers["Referrer-Policy"] == "no-referrer", status
                assert headers["X-Content-Type-Options"] == "nosniff", status
                assert headers["X-Frame-Options"] == "DENY", status
                csp = headers["Content-Security-Policy"]
                assert "frame-ancestors 'none'" in csp, status
                # Strict self-only CSP: no third-party origins anywhere.
                assert "https:" not in csp and "http://" not in csp.replace("'none'", ""), csp
                assert "default-src 'none'" in csp
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# GET/HEAD purity
# --------------------------------------------------------------------------

class TestGetHeadPurity:
    def test_get_and_head_are_side_effect_free(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            generations_before = _generations(workdir)
            assert not (workdir / ".review").exists()
            # Every read surface, GET and HEAD, authorized and not.
            routes = ["/api/reviews", "/api/document-state", "/api/document-fragment",
                      "/api/review-history", "/health"]
            for route in routes:
                for method in ("GET", "HEAD"):
                    status, headers, body = _request(port, route, method=method, capability=capability)
                    assert status == 200, (route, method)
                    if method == "HEAD":
                        assert body == b"", (route, method)
                        assert int(headers["Content-Length"]) > 0
            # No queue/session/snapshot/history state was created.
            assert not (workdir / ".review").exists()
            assert not (workdir / ".review" / "session.json").exists()
            assert _generations(workdir) == generations_before
            # Unauthorized reads also create nothing.
            _request(port, "/api/reviews")
            _request(port, "/api/review-history?snapshot=C1")
            # HEAD denial responses carry headers but no body.
            status, headers, body = _request(port, "/api/reviews", method="HEAD")
            assert status == 404 and body == b"" and int(headers["Content-Length"]) > 0
            status, _, body = _request(port, "/", method="HEAD", host="attacker.example:80")
            assert status == 404 and body == b""
            assert not (workdir / ".review").exists()
            assert _generations(workdir) == generations_before
        finally:
            _stop(server)

    def test_read_of_fresh_session_returns_empty_not_error(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            status, _, body = _request(port, "/api/reviews", capability=capability)
            payload = json.loads(body)
            assert payload["counts"] == {"draft": 0, "queued": 0, "acknowledged": 0}
            assert payload["session"]["current_snapshot"] is None
            assert payload["session"]["current_matches_filesystem"] is False
            assert not (workdir / ".review").exists()
        finally:
            _stop(server)

    def test_document_fragment_is_complete_for_the_console(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            status, _, body = _request(port, "/api/document-fragment", capability=capability)
            payload = json.loads(body)
            assert payload["html"]  # rendered paragraphs
            assert "第一段" in payload["html"]
            assert payload["document_title"]
            assert payload["summary"]
            assert "history" in payload and "review" in payload and "session" in payload
            assert "audit_rows" in payload
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# logging
# --------------------------------------------------------------------------

class TestLogging:
    def test_logs_omit_secrets_paths_query_and_body(self, tmp_path, capsys):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            payload = json.dumps({**_decision_payload(), "comment": "这是批注内容"}).encode("utf-8")
            _request(port, "/")
            _request(port, "/api/reviews", capability=capability)
            _request(port, "/api/reviews", method="POST", capability=capability,
                     origin=f"http://127.0.0.1:{port}", fetch_site="same-origin",
                     idempotency_key="revkey00000001", body=payload)
            _request(port, "/api/review-history?snapshot=C7", capability=capability)
            _request(port, "/api/reviews")  # unauthorized
            lines = capsys.readouterr().out.splitlines()
            server_lines = [line for line in lines if line.startswith("[review-server]")]
            assert len(server_lines) >= 5
            pattern = re.compile(
                r"^\[review-server\] \d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}\+00:00 "
                r"cat=(bootstrap|read|write|unknown) result=\dxx src=(loopback|remote) "
                r"sid=(?:[0-9a-f]{10}|-)$"
            )
            for line in server_lines:
                assert pattern.match(line), line
                assert capability not in line
                assert "/api/" not in line  # no path
                assert "snapshot" not in line  # no query
                assert "revkey00000001" not in line  # no idempotency key
                assert "这是批注内容" not in line  # no body/selected text/comments
                assert workdir.name not in line  # no workdir path
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# throttling
# --------------------------------------------------------------------------

class TestThrottling:
    def test_unauthorized_flood_is_bounded_and_authorized_never_rate_limited(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            statuses = []
            for _ in range(30):
                status, headers, body = _request(port, "/api/reviews")
                statuses.append(status)
                if status == 429:
                    assert headers.get("Retry-After") == "1"
                    assert json.loads(body) == {"error": "not-found"}
            assert statuses[0] == 404
            assert statuses.count(404) > 0 and statuses.count(429) > 0
            # The flood never produced a 500 or mutated state.
            assert all(status in (404, 429) for status in statuses)
            assert not (workdir / ".review").exists()
            # Authorized interaction from the same client is not rate-limited.
            status, _, body = _request(port, "/api/reviews", capability=capability)
            assert status == 200
            # The bucket refills: after a pause an unauthorized probe is 404 again.
            time.sleep(1.5)
            status, _, _ = _request(port, "/api/reviews")
            assert status == 404
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# capability lifecycle / restart / CAS lane
# --------------------------------------------------------------------------

class TestLifecycleAndLane:
    def test_restart_revokes_old_capability_and_keeps_committed_state(self, tmp_path):
        workdir = _extract(tmp_path)
        server_a, capability_a, port_a = _serve(workdir)
        try:
            payload = json.dumps(_decision_payload(client_id="restart-1")).encode("utf-8")
            status, _, body = _request(port_a, "/api/reviews", method="POST",
                                       capability=capability_a, origin=f"http://127.0.0.1:{port_a}",
                                       fetch_site="same-origin", idempotency_key="revkey00000001", body=payload)
            assert status == 200
            generations_after_write = _generations(workdir)
        finally:
            _stop(server_a)
        # Process termination revokes the old capability.
        server_b, capability_b, port_b = _serve(workdir)
        try:
            status, _, body = _request(port_b, "/api/reviews", capability=capability_a)
            assert status == 404 and json.loads(body) == {"error": "not-found"}
            # The new session sees the committed state: the event survives.
            status, _, body = _request(port_b, "/api/reviews", capability=capability_b)
            assert status == 200
            payload = json.loads(body)
            assert payload["counts"]["draft"] == 1
            assert payload["events"][0]["client_id"] == "restart-1"
            assert _generations(workdir) == generations_after_write
            # And can keep writing with the new capability.
            status, _, _ = _request(port_b, "/api/reviews/dispatch", method="POST",
                                    capability=capability_b, origin=f"http://127.0.0.1:{port_b}",
                                    fetch_site="same-origin", idempotency_key="revkey00000002",
                                    content_type=None, body=None)
            assert status == 200
        finally:
            _stop(server_b)

    def test_idempotency_key_ledger_replay_end_to_end(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            payload = json.dumps(_decision_payload(client_id="replay-1")).encode("utf-8")
            headers = dict(origin=f"http://127.0.0.1:{port}", fetch_site="same-origin")
            status, _, first = _request(port, "/api/reviews", method="POST", capability=capability,
                                        idempotency_key="revkey00000001", body=payload, **headers)
            assert status == 200
            status, _, replay = _request(port, "/api/reviews", method="POST", capability=capability,
                                         idempotency_key="revkey00000001", body=payload, **headers)
            assert status == 200 and json.loads(replay) == json.loads(first)
            # Same key with a changed payload is rejected without a second effect.
            changed = json.dumps(_decision_payload(client_id="replay-2")).encode("utf-8")
            status, _, body = _request(port, "/api/reviews", method="POST", capability=capability,
                                       idempotency_key="revkey00000001", body=changed, **headers)
            assert status == 409 and json.loads(body)["code"] == "operation-id-reused"
            events = [p for p in (workdir / ".review" / "inbox").glob("*.json")]
            assert len(events) == 1
        finally:
            _stop(server)

    def test_multi_holder_cas_writer_lane_serializes_concurrent_publishes(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            headers = dict(origin=f"http://127.0.0.1:{port}", fetch_site="same-origin")
            # Seed a session, then change the document so publish has work.
            seed = json.dumps(_decision_payload(client_id="cas-seed")).encode("utf-8")
            status, _, _ = _request(port, "/api/reviews", method="POST", capability=capability,
                                    idempotency_key="revkey00000001", body=seed, **headers)
            assert status == 200
            (workdir / "typed.md").write_text("changed once\n", encoding="utf-8")
            generations_before = _generations(workdir)

            barrier = threading.Barrier(2)
            results: list[tuple[int, dict]] = []

            def publish(key: str) -> None:
                barrier.wait()
                status, _, body = _request(port, "/api/reviews/publish", method="POST",
                                           capability=capability, idempotency_key=key,
                                           body=json.dumps({
                                               "expected_parent_snapshot": "C0",
                                               "origin": "human_ui",
                                               "changed_paragraph_ids": ["P0"],
                                           }).encode("utf-8"), **headers)
                results.append((status, json.loads(body) if body else {}))

            threads = [threading.Thread(target=publish, args=(f"revkey0000000{i}",)) for i in (2, 3)]
            for thread in threads:
                thread.start()
            for thread in threads:
                thread.join(timeout=30)

            statuses = sorted(status for status, _ in results)
            # Exactly one publish commits; the other hits the CAS lane or the
            # parent-snapshot check. No lost update, no double commit.
            assert statuses == [200, 409], results
            committed = next((payload for status, payload in results if status == 200), None)
            assert committed["session"]["current_snapshot"]["id"] == "C1"
            assert _generations(workdir) == generations_before | {next(iter(_generations(workdir) - generations_before))}
            # The store recovered cleanly (no leftover transaction journals).
            from scripts.store import Store

            assert Store.open(workdir).recover()["needs_recovery"] == []
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# issue #51 finding 1: bounded socket reads; stalled drains release threads
# --------------------------------------------------------------------------

def _split_response(raw: bytes) -> tuple[int, dict[str, str], bytes]:
    head, _, body = raw.partition(b"\r\n\r\n")
    lines = head.split(b"\r\n")
    status = int(lines[0].split(b" ", 2)[1])
    headers = {}
    for line in lines[1:]:
        name, _, value = line.partition(b":")
        headers[name.decode("latin-1").strip().lower()] = value.strip().decode("latin-1")
    return status, headers, body


class TestStalledConnectionDrain:
    def test_stalled_body_drain_is_bounded_and_thread_released(self, tmp_path, capsys):
        workdir = _extract(tmp_path)
        # A 1s socket bound makes the "bounded" property observable in-test.
        server, capability, port = _serve(workdir, socket_timeout=1.0)
        try:
            # Unauthorized POST announcing a huge body that never arrives:
            # the deny path must drain with the timeout bound, then close.
            started = time.monotonic()
            raw = _raw_request(
                port,
                b"POST /api/reviews HTTP/1.0\r\n"
                b"Host: 127.0.0.1:" + str(port).encode() + b"\r\n"
                b"Content-Length: 1048576\r\n"
                b"Content-Type: application/json\r\n\r\n",
            )
            elapsed = time.monotonic() - started
            status, headers, body = _split_response(raw)
            assert status == 404
            assert json.loads(body) == {"error": "not-found"}
            # Bounded by the socket timeout, not the 30s default or forever
            # (the EOF above proves the server closed the connection).
            assert elapsed < 8.0, f"stalled drain took {elapsed:.2f}s"
            # The handler thread was released: the server still serves.
            status, _, _ = _request(port, "/api/reviews", capability=capability)
            assert status == 200
        finally:
            _stop(server)
        assert "Traceback" not in capsys.readouterr().err

    def test_silent_client_is_closed_within_bound(self, tmp_path, capsys):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir, socket_timeout=1.0)
        try:
            started = time.monotonic()
            with socket.create_connection(("127.0.0.1", port), timeout=10) as sock:
                # Connect and send nothing: the request-line read must time out.
                data = b""
                while True:
                    chunk = sock.recv(4096)
                    if not chunk:
                        break
                    data += chunk
            elapsed = time.monotonic() - started
            assert elapsed < 8.0, f"silent client held the thread {elapsed:.2f}s"
            assert data == b""
            # The server still serves afterward.
            status, _, _ = _request(port, "/api/reviews", capability=capability)
            assert status == 200
        finally:
            _stop(server)
        assert "Traceback" not in capsys.readouterr().err


# --------------------------------------------------------------------------
# issue #51 finding 2: review writer/queue lanes are OS-advisory, crash-safe
# --------------------------------------------------------------------------

class TestReviewLanesCrashRecovery:
    def test_stale_lock_files_and_partial_write_do_not_block_restart(self, tmp_path):
        workdir = _extract(tmp_path)
        # Simulate a crash mid-write: the OS advisory lock files remain (their
        # locks died with the holder process) and a partial temp file is left.
        review = workdir / ".review"
        review.mkdir()
        (review / "writer.lock").write_text("stale", encoding="utf-8")
        inbox = review / "inbox"
        inbox.mkdir()
        (inbox / ".queue.lock").write_text("stale", encoding="utf-8")
        (inbox / ".e1-deadbeef.tmp").write_text('{"partial": true', encoding="utf-8")

        from scripts.review_collab import writer_lane
        from scripts.review_queue import list_events, snapshot, upsert_event

        # The queue lane is acquired on the existing unlocked file: no stale
        # O_EXCL deadlock, and the lock file is never deleted or reclaimed.
        record = upsert_event(workdir, _decision_payload(client_id="crash-1"))
        assert record["status"] == "draft"
        assert [e["client_id"] for e in list_events(workdir)] == ["crash-1"]
        assert snapshot(workdir)["counts"]["draft"] == 1
        assert (inbox / ".queue.lock").exists()  # never unlinked, never recreated hot
        # The review writer lane is equally usable.
        with writer_lane(workdir):
            pass
        assert (review / "writer.lock").exists()
        # A fresh server keeps writing after the simulated crash: the store
        # lane commits a new decision and materializes it to the root queue
        # without disturbing the directly-written crash-1 event.
        server, capability, port = _serve(workdir)
        try:
            status, _, _ = _request(port, "/api/reviews", capability=capability)
            assert status == 200
            status, _, _ = _request(port, "/api/reviews", method="POST",
                                    capability=capability, origin=f"http://127.0.0.1:{port}",
                                    fetch_site="same-origin", idempotency_key="revkey00000001",
                                    body=json.dumps(_decision_payload(client_id="server-1")).encode("utf-8"))
            assert status == 200
        finally:
            _stop(server)
        assert {e["client_id"] for e in list_events(workdir)} == {"crash-1", "server-1"}

    def test_queue_lane_released_on_holder_process_death(self, tmp_path):
        workdir = _extract(tmp_path)
        marker = tmp_path / "queue-holder-ready"
        holder = subprocess.Popen(
            [
                sys.executable,
                "-c",
                (
                    "import sys, time; sys.path.insert(0, %r);\n"
                    "from pathlib import Path;\n"
                    "from scripts.review_queue import _queue_lane;\n"
                    "with _queue_lane(Path(%r)):\n"
                    "    Path(%r).write_text('ready')\n"
                    "    time.sleep(60)\n"
                )
                % (str(ROOT), str(workdir), str(marker)),
            ],
            cwd=ROOT,
        )
        try:
            deadline = time.monotonic() + 15
            while not marker.exists():
                if time.monotonic() > deadline:
                    raise AssertionError("queue lane holder never became ready")
                time.sleep(0.05)
            from scripts.review_queue import upsert_event

            with pytest.raises(RuntimeError, match="review queue is busy"):
                upsert_event(workdir, _decision_payload(client_id="held-1"))
            holder.kill()  # process death releases the OS advisory lock
            holder.wait(timeout=10)
            time.sleep(0.5)
            record = upsert_event(workdir, _decision_payload(client_id="after-1"))
            assert record["status"] == "draft" and record["client_id"] == "after-1"
        finally:
            if holder.poll() is None:
                holder.kill()
                holder.wait(timeout=10)


# --------------------------------------------------------------------------
# issue #51 finding 3: malformed Content-Length never raises or 500s
# --------------------------------------------------------------------------

class TestMalformedContentLength:
    def test_unauthorized_malformed_content_length_is_uniform_404_no_traceback(self, tmp_path, capsys):
        workdir = _extract(tmp_path)
        # A 1s socket bound keeps the huge-numeric probe (which drains up to
        # the cap) fast while still proving the deny path is bounded.
        server, capability, port = _serve(workdir, socket_timeout=1.0)
        try:
            variants = [b"abc", b"", b"-5", b"0", b"12.5", b"1e3", b"99999999999999999999"]
            bodies: set[bytes] = set()
            for value in variants:
                raw = _raw_request(
                    port,
                    b"POST /api/reviews HTTP/1.0\r\n"
                    b"Host: 127.0.0.1:" + str(port).encode() + b"\r\n"
                    b"Content-Length: " + value + b"\r\n"
                    b"Content-Type: application/json\r\n\r\n",
                )
                status, headers, body = _split_response(raw)
                assert status == 404, (value, status, raw[:200])
                assert json.loads(body) == {"error": "not-found"}, value
                bodies.add(body)
            # Byte-identical uniform 404 across every malformed value.
            assert len(bodies) == 1
            # Never a stdlib 500 and never a stderr traceback.
            assert "Traceback" not in capsys.readouterr().err
        finally:
            _stop(server)

    def test_authorized_malformed_content_length_is_json_400(self, tmp_path, capsys):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            for value in (b"abc", b"", b"-5", b"12.5"):
                raw = _raw_request(
                    port,
                    b"POST /api/reviews HTTP/1.0\r\n"
                    b"Host: 127.0.0.1:" + str(port).encode() + b"\r\n"
                    b"Authorization: Bearer " + capability.encode() + b"\r\n"
                    b"Origin: http://127.0.0.1:" + str(port).encode() + b"\r\n"
                    b"Sec-Fetch-Site: same-origin\r\n"
                    b"Content-Type: application/json\r\n"
                    b"Idempotency-Key: revkey00000001\r\n"
                    b"Content-Length: " + value + b"\r\n\r\n",
                )
                status, _, body = _split_response(raw)
                assert status == 400, (value, status, raw[:200])
                assert json.loads(body)["code"] == "request-body-required", value
            assert "Traceback" not in capsys.readouterr().err
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# issue #51 finding 4: unsupported methods pass the gate with full headers
# --------------------------------------------------------------------------

class TestUnsupportedMethods:
    def test_unauthorized_unsupported_methods_are_uniform_404_with_full_headers(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            bodies: set[bytes] = set()
            for method in ("OPTIONS", "PUT", "DELETE", "TRACE", "PATCH", "CONNECT", "BREW"):
                status, headers, body = _request(port, "/api/reviews", method=method)
                assert status == 404, (method, status)
                assert json.loads(body) == {"error": "not-found"}, method
                bodies.add(body)
                assert headers["Cache-Control"] == "no-store", method
                assert headers["Referrer-Policy"] == "no-referrer", method
                assert headers["X-Content-Type-Options"] == "nosniff", method
                assert headers["X-Frame-Options"] == "DENY", method
                assert "frame-ancestors 'none'" in headers["Content-Security-Policy"], method
                assert not any(name.lower().startswith("access-control-") for name in headers), method
                assert not any(name.lower() == "allow" for name in headers), method
            assert len(bodies) == 1  # byte-identical uniform refusal
        finally:
            _stop(server)

    def test_gated_unsupported_methods_return_501_with_security_headers_no_cors(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            for method in ("OPTIONS", "PUT", "DELETE", "TRACE", "PATCH", "CONNECT", "BREW"):
                status, headers, body = _request(port, "/api/reviews", method=method, capability=capability)
                assert status == 501, (method, status)
                assert json.loads(body) == {"error": "method-not-supported"}, method
                assert headers["Cache-Control"] == "no-store", method
                assert headers["Referrer-Policy"] == "no-referrer", method
                assert headers["X-Content-Type-Options"] == "nosniff", method
                assert headers["X-Frame-Options"] == "DENY", method
                assert "frame-ancestors 'none'" in headers["Content-Security-Policy"], method
                assert not any(name.lower().startswith("access-control-") for name in headers), method
                assert not any(name.lower() == "allow" for name in headers), method
            # No OPTIONS preflight path exists: nothing was mutated.
            assert not (workdir / ".review").exists()
        finally:
            _stop(server)


# --------------------------------------------------------------------------
# issue #51 finding 5: external-preflight is a store-lane idempotent POST
# --------------------------------------------------------------------------

class TestExternalPreflightLane:
    def test_preflight_is_idempotent_store_post_with_ledger_replay(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            headers = dict(origin=f"http://127.0.0.1:{port}", fetch_site="same-origin")
            payload = json.dumps({"expected_parent_snapshot": "C0", "operation": "import"}).encode("utf-8")
            status, _, first = _request(port, "/api/reviews/external-preflight", method="POST",
                                        capability=capability, idempotency_key="preflight00001",
                                        body=payload, **headers)
            assert status == 200, first
            guard = json.loads(first)
            assert guard["operation_id"] == "preflight00001"
            assert guard["schema"] == "docx2typed-review-external-guard-1"
            assert guard["operation"] == "import"
            assert guard["expected_parent_snapshot"] == "C0"
            assert re.fullmatch(r"[0-9a-f]{64}", guard["typed_sha256"])
            # Session bootstrap happened under the lane: exactly one record.
            history = (workdir / ".review" / "history.jsonl").read_text(encoding="utf-8")
            assert history.count("session-created") == 1
            # Same Idempotency-Key replays the committed guard, no second effect.
            status, _, replay = _request(port, "/api/reviews/external-preflight", method="POST",
                                         capability=capability, idempotency_key="preflight00001",
                                         body=payload, **headers)
            assert status == 200 and replay == first
            history = (workdir / ".review" / "history.jsonl").read_text(encoding="utf-8")
            assert history.count("session-created") == 1
            # Changed payload with the same key is rejected.
            changed = json.dumps({"expected_parent_snapshot": "C0", "operation": "rollback"}).encode("utf-8")
            status, _, body = _request(port, "/api/reviews/external-preflight", method="POST",
                                       capability=capability, idempotency_key="preflight00001",
                                       body=changed, **headers)
            assert status == 409 and json.loads(body)["code"] == "operation-id-reused"
        finally:
            _stop(server)

    def test_concurrent_preflights_produce_one_session_and_history(self, tmp_path):
        workdir = _extract(tmp_path)
        server, capability, port = _serve(workdir)
        try:
            headers = dict(origin=f"http://127.0.0.1:{port}", fetch_site="same-origin")
            barrier = threading.Barrier(2)
            results: list[tuple[int, bytes]] = []

            def preflight(key: str) -> None:
                barrier.wait()
                status, _, body = _request(port, "/api/reviews/external-preflight", method="POST",
                                           capability=capability, idempotency_key=key,
                                           body=json.dumps({
                                               "expected_parent_snapshot": "C0",
                                               "operation": "import",
                                           }).encode("utf-8"), **headers)
                results.append((status, body))

            threads = [threading.Thread(target=preflight, args=(f"preflight0000{i}",)) for i in (1, 2)]
            for thread in threads:
                thread.start()
            for thread in threads:
                thread.join(timeout=30)

            statuses = sorted(status for status, _ in results)
            # Exactly one commits; the other hits the CAS lane. Never both.
            assert statuses == [200, 409], results
            # One session, one session-created history record, one guard.
            assert (workdir / ".review" / "session.json").exists()
            history = (workdir / ".review" / "history.jsonl").read_text(encoding="utf-8")
            assert history.count("session-created") == 1
            assert history.count("external") == 0  # guards are not history events
            guards = [body for status, body in results if status == 200]
            assert len(guards) == 1 and json.loads(guards[0])["expected_parent_snapshot"] == "C0"
            # The store recovered cleanly (no leftover transaction journals).
            from scripts.store import Store

            assert Store.open(workdir).recover()["needs_recovery"] == []
        finally:
            _stop(server)

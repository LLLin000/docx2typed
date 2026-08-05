import json
from zipfile import ZipFile

from docx import Document

from scripts.extract import extract
from scripts.typed_normalize import candidate_report, load_catalog, normalize, normalize_workdir
from scripts.verify import verify


def test_vertical_normalization_requires_policy_and_creates_new_workdir(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    normalized_docx = tmp_path / "normalized.docx"
    normalized_workdir = tmp_path / "normalized-workdir"
    document = Document()
    document.add_paragraph("₂")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    candidates = candidate_report(workdir)
    assert len(candidates) == 1
    assert candidates[0]["source"] == "₂"
    assert candidates[0]["proposed_target"] == "2"

    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    policy = {
        "schema": "vertical-normalization-policy-1",
        "profile": "selective",
        "catalog_hash": load_catalog()["catalog_hash"],
        "template_sha256": format_data["template_sha256"],
        "decisions": {candidates[0]["occurrence_id"]: "convert"},
    }
    policy_path = tmp_path / "policy.json"
    policy_path.write_text(json.dumps(policy), encoding="utf-8")

    result = normalize_workdir(workdir, policy_path, normalized_docx, normalized_workdir)
    assert result == normalized_workdir.resolve()
    assert "₂" in (workdir / "typed.md").read_text(encoding="utf-8")
    assert "₂" not in (normalized_workdir / "typed.md").read_text(encoding="utf-8")
    assert (normalized_workdir / "normalization.audit.json").exists()
    audit = json.loads((normalized_workdir / "normalization.audit.json").read_text(encoding="utf-8"))
    assert audit["governance_status"] == "legacy-unaudited"
    assert verify([str(normalized_workdir), str(normalized_docx)]) == 0

    with ZipFile(normalized_docx) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "vertAlign" in xml


def test_selective_normalization_counts_only_candidates(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    normalized_docx = tmp_path / "normalized.docx"
    normalized_workdir = tmp_path / "normalized-workdir"
    document = Document()
    document.add_paragraph("₂₂")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    candidates = candidate_report(workdir)
    assert [item["occurrence_id"] for item in candidates] == ["P0-V0001", "P0-V0002"]
    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    policy = {
        "schema": "vertical-normalization-policy-1",
        "profile": "selective",
        "catalog_hash": load_catalog()["catalog_hash"],
        "template_sha256": format_data["template_sha256"],
        "decisions": {"P0-V0001": "convert", "P0-V0002": "preserve"},
    }
    policy_path = tmp_path / "policy.json"
    policy_path.write_text(json.dumps(policy), encoding="utf-8")

    normalize_workdir(workdir, policy_path, normalized_docx, normalized_workdir)
    assert "2₂" in (normalized_workdir / "typed.md").read_text(encoding="utf-8")
    audit = json.loads((normalized_workdir / "normalization.audit.json").read_text(encoding="utf-8"))
    assert [item["decision"] for item in audit["decisions"]] == ["convert", "preserve"]
    assert verify([str(normalized_workdir), str(normalized_docx)]) == 0

def test_legacy_cli_requires_explicit_flag(tmp_path, capsys):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"
    document = Document()
    document.add_paragraph("₂")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    candidate = candidate_report(workdir)[0]
    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    policy = {
        "schema": "vertical-normalization-policy-1",
        "profile": "selective",
        "catalog_hash": load_catalog()["catalog_hash"],
        "template_sha256": format_data["template_sha256"],
        "decisions": {candidate["occurrence_id"]: "convert"},
    }
    policy_path = tmp_path / "policy.json"
    policy_path.write_text(json.dumps(policy), encoding="utf-8")
    args = [
        str(workdir),
        "--policy",
        str(policy_path),
        "-o",
        str(output),
        "--workdir-out",
        str(workdir_out),
    ]
    assert normalize(args) == 1
    assert "legacy policy-1 requires" in capsys.readouterr().out
    assert normalize([str(workdir), "--legacy-policy-1", *args[1:]]) == 0
    assert "legacy-unaudited" in (workdir_out / "normalization.audit.json").read_text(encoding="utf-8")

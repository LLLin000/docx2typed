from docx import Document

from scripts.build import build
from scripts.extract import extract
from scripts.verify import verify
from scripts.view import view_workdir


def test_tabs_and_breaks_use_explicit_locked_tokens(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("左")
    run.add_tab()
    run.add_text("右")
    run.add_break()
    run.add_text("下")
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert 'kind="tab"' in typed and 'kind="br"' in typed
    assert view_workdir(workdir, "clean", markers=False) == "左\t右\n下"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0

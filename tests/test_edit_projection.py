"""Slice A clean-edit contract: projection, state, refresh, no-op sync, gates."""
from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document

from scripts.build import build
from scripts.edit import (
    EDIT_SCHEMA_VERSION,
    PROJECTION_FILE,
    SEGMENTATION_CONTRACT,
    STATE_FILE,
    SYNC_CONTRACT_VERSION,
    edit_body_sha256,
    edit_status,
    parse_edit_projection,
    refresh_edit_projection,
    sync_edit_projection,
)
from scripts.extract import extract
from scripts.verify import verify

FIXTURE = Path(__file__).parent / "fixtures" / "complex-docx" / "complex.docx"


def make_source(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前")
    bold = paragraph.add_run("加粗")
    bold.bold = True
    paragraph.add_run(" A & B < C > D ⟦字⟧ 尾")
    paragraph.add_run().add_tab()
    paragraph.add_run("中")
    paragraph.add_run().add_break()
    paragraph.add_run("后")
    document.add_paragraph("第二段")
    document.save(path)


def extract_fixture(tmp_path: Path) -> Path:
    tmp_path = Path(tmp_path)
    tmp_path.mkdir(parents=True, exist_ok=True)
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    make_source(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return workdir


def edit_text(workdir: Path) -> str:
    return (workdir / PROJECTION_FILE).read_text(encoding="utf-8")


# --------------------------------------------------------------------------
# Extraction creates projection, state, and evidence
# --------------------------------------------------------------------------

def test_extract_creates_projection_state_and_evidence(tmp_path):
    workdir = extract_fixture(tmp_path)
    assert (workdir / PROJECTION_FILE).exists()
    assert (workdir / STATE_FILE).exists()
    evidence_path = workdir / "edit.state.json.run.json"
    assert evidence_path.exists()
    evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
    assert evidence["schema"] == "typed-clean-edit-run-evidence-1"
    assert evidence["status"] == "ok"
    assert evidence["command"] == "docx2typed extract"
    assert evidence["changed_paragraph_ids"] == []
    state = json.loads((workdir / STATE_FILE).read_text(encoding="utf-8"))
    assert state["schema"] == "typed-clean-edit-state-1"
    assert state["edit_schema_version"] == EDIT_SCHEMA_VERSION
    assert state["sync_contract_version"] == SYNC_CONTRACT_VERSION
    assert state["segmentation_contract"] == SEGMENTATION_CONTRACT
    assert state["base_typed_sha256"] == hashlib.sha256((workdir / "typed.md").read_bytes()).hexdigest()


def test_projection_is_span_free_with_ids_placeholders_and_literal_text(tmp_path):
    workdir = extract_fixture(tmp_path)
    text = edit_text(workdir)
    assert "<span" not in text
    assert "data-s=" not in text
    assert '<!--@p id="P0"-->' in text
    assert '<!--@p id="P1"-->' in text
    assert text.index('id="P0"') < text.index('id="P1"')
    assert 'token id="N0" kind="tab"' in text
    assert 'token id="N1" kind="br"' in text
    assert " A & B < C > D " in text
    assert "\\u27E6字\\u27E7" in text  # literal brackets are escaped
    projection = parse_edit_projection(text)
    assert [attrs["id"] for kind, attrs, _ in projection.paragraphs] == ["P0", "P1"]


def test_escape_roundtrip_through_projection(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    document = Document()
    document.add_paragraph("反斜杠\\和\\u27E6字面⟦括号⟧")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    text = edit_text(workdir)
    assert "\\\\" in text and "\\u27E6" in text and "\\u27E7" in text
    assert edit_status(workdir)["state"] == "clean"


def test_header_bindings_match_sidecar_and_body_hash_excludes_header(tmp_path):
    workdir = extract_fixture(tmp_path)
    text = edit_text(workdir)
    projection = parse_edit_projection(text)
    state = json.loads((workdir / STATE_FILE).read_text(encoding="utf-8"))
    assert projection.header["schema"] == str(EDIT_SCHEMA_VERSION)
    assert projection.header["base-typed-sha256"] == state["base_typed_sha256"]
    assert projection.header["base-projection-sha256"] == state["base_projection_sha256"]
    assert projection.header["base-projection-sha256"] == edit_body_sha256(text)
    # A header change alone must not change the body hash (header excluded).
    altered = text.replace(projection.header["base-projection-sha256"], "0" * 64, 1)
    assert edit_body_sha256(altered) == edit_body_sha256(text)


def test_crlf_edit_file_stays_clean(tmp_path):
    workdir = extract_fixture(tmp_path)
    text = edit_text(workdir)
    (workdir / PROJECTION_FILE).write_bytes(text.replace("\n", "\r\n").encode("utf-8"))
    assert edit_status(workdir)["state"] == "clean"


# --------------------------------------------------------------------------
# Freshness states
# --------------------------------------------------------------------------

def test_states_clean_dirty_stale_conflict(tmp_path):
    workdir = extract_fixture(tmp_path)
    assert edit_status(workdir)["state"] == "clean"

    (workdir / PROJECTION_FILE).write_text(
        edit_text(workdir).replace("前", "草稿", 1), encoding="utf-8"
    )
    assert edit_status(workdir)["state"] == "dirty"

    refresh_edit_projection(workdir, discard=True)
    assert edit_status(workdir)["state"] == "clean"

    (workdir / "typed.md").write_text(
        (workdir / "typed.md").read_text(encoding="utf-8").replace("前", "改", 1),
        encoding="utf-8",
    )
    assert edit_status(workdir)["state"] == "stale-clean"

    (workdir / PROJECTION_FILE).write_text(
        edit_text(workdir).replace("加粗", "草稿加粗", 1), encoding="utf-8"
    )
    assert edit_status(workdir)["state"] == "conflict"


def test_tampered_header_is_edit_header_tampered(tmp_path):
    workdir = extract_fixture(tmp_path)
    text = edit_text(workdir)
    body = text.replace("前", "草稿", 1)
    body_hash = edit_body_sha256(body)
    tampered = re.sub(
        r'base-projection-sha256="[0-9a-f]{64}"',
        f'base-projection-sha256="{body_hash}"',
        body,
        count=1,
    )
    (workdir / PROJECTION_FILE).write_text(tampered, encoding="utf-8")
    try:
        edit_status(workdir)
    except Exception as exc:
        assert "edit-header-tampered" in str(exc)
    else:
        raise AssertionError("tampered header must be rejected")


# --------------------------------------------------------------------------
# Refresh
# --------------------------------------------------------------------------

def test_refresh_regenerates_projection_after_typed_edit(tmp_path):
    workdir = extract_fixture(tmp_path)
    typed = workdir / "typed.md"
    typed.write_text(typed.read_text(encoding="utf-8").replace("前", "改", 1), encoding="utf-8")
    assert edit_status(workdir)["state"] == "stale-clean"
    refresh_edit_projection(workdir)
    result = edit_status(workdir)
    assert result["state"] == "clean"
    assert "改" in edit_text(workdir)
    evidence = json.loads((workdir / "edit.state.json.run.json").read_text(encoding="utf-8"))
    assert evidence["status"] == "ok"
    assert evidence["state_before"] == "stale-clean"


def test_refresh_dirty_requires_discard_and_records_hash(tmp_path):
    workdir = extract_fixture(tmp_path)
    (workdir / PROJECTION_FILE).write_text(
        edit_text(workdir).replace("前", "草稿", 1), encoding="utf-8"
    )
    dirty_hash = edit_body_sha256(edit_text(workdir))
    try:
        refresh_edit_projection(workdir)
    except Exception as exc:
        assert "edit-refresh-requires-discard" in str(exc)
    else:
        raise AssertionError("dirty refresh must require --discard")
    assert "草稿" in edit_text(workdir)  # untouched
    refresh_edit_projection(workdir, discard=True)
    assert edit_status(workdir)["state"] == "clean"
    assert "草稿" not in edit_text(workdir)
    evidence = json.loads((workdir / "edit.state.json.run.json").read_text(encoding="utf-8"))
    assert evidence["discarded_edit_sha256"] == dirty_hash


def test_refresh_init_for_legacy_workdir(tmp_path):
    workdir = extract_fixture(tmp_path)
    (workdir / PROJECTION_FILE).unlink()
    (workdir / STATE_FILE).unlink()
    (workdir / "edit.state.json.run.json").unlink()
    assert build([str(workdir), "-o", str(tmp_path / "blocked.docx")]) == 1
    refresh_edit_projection(workdir, init=True)
    assert edit_status(workdir)["state"] == "clean"
    assert build([str(workdir), "-o", str(tmp_path / "ok.docx")]) == 0
    try:
        refresh_edit_projection(workdir, init=True)
    except Exception as exc:
        assert "edit-init" in str(exc)
    else:
        raise AssertionError("--init must be rejected when state exists")


# --------------------------------------------------------------------------
# Sync
# --------------------------------------------------------------------------

def test_sync_clean_noop_is_idempotent(tmp_path):
    workdir = extract_fixture(tmp_path)
    typed_before = (workdir / "typed.md").read_bytes()
    edit_before = (workdir / PROJECTION_FILE).read_bytes()
    state_before = (workdir / STATE_FILE).read_bytes()
    sync_edit_projection(workdir)
    sync_edit_projection(workdir)
    assert (workdir / "typed.md").read_bytes() == typed_before
    assert (workdir / PROJECTION_FILE).read_bytes() == edit_before
    assert (workdir / STATE_FILE).read_bytes() == state_before
    evidence = json.loads((workdir / "edit.state.json.run.json").read_text(encoding="utf-8"))
    assert evidence["status"] == "ok" and evidence["state_before"] == "clean"


def test_sync_dirty_applies_text_and_returns_to_clean(tmp_path):
    workdir = extract_fixture(tmp_path)
    typed_before = (workdir / "typed.md").read_bytes()
    (workdir / PROJECTION_FILE).write_text(
        edit_text(workdir).replace("前", "草稿", 1), encoding="utf-8"
    )
    state_path, warnings, changed = sync_edit_projection(workdir)
    assert state_path.name == STATE_FILE
    assert changed == ["P0"]
    assert "草稿" in (workdir / "typed.md").read_text(encoding="utf-8")
    assert (workdir / "typed.md").read_bytes() != typed_before
    assert edit_status(workdir)["state"] == "clean"
    evidence = json.loads((workdir / "edit.state.json.run.json").read_text(encoding="utf-8"))
    assert evidence["status"] == "ok"
    assert evidence["state_before"] == "dirty"
    assert evidence["changed_paragraph_ids"] == ["P0"]
    assert evidence["hunk_report"][0]["operation"] == "replace"
    assert evidence["hunk_report"][0]["assigned_style"]
    assert build([str(workdir), "-o", str(tmp_path / "out.docx")]) == 0
    assert verify([str(workdir), str(tmp_path / "out.docx")]) == 0


def test_sync_stale_and_conflict_rejected(tmp_path):
    workdir = extract_fixture(tmp_path)
    (workdir / "typed.md").write_text(
        (workdir / "typed.md").read_text(encoding="utf-8").replace("前", "改", 1),
        encoding="utf-8",
    )
    for message in ("edit-stale",):
        try:
            sync_edit_projection(workdir)
        except Exception as exc:
            assert message in str(exc)
        else:
            raise AssertionError(f"sync must reject {message}")
    (workdir / PROJECTION_FILE).write_text(
        edit_text(workdir).replace("加粗", "草稿加粗", 1), encoding="utf-8"
    )
    try:
        sync_edit_projection(workdir)
    except Exception as exc:
        assert "edit-conflict" in str(exc)
    else:
        raise AssertionError("sync must reject conflict")


# --------------------------------------------------------------------------
# Build and verify gates
# --------------------------------------------------------------------------

def test_build_and_verify_gate_on_non_clean_states(tmp_path):
    workdir = extract_fixture(tmp_path)
    output = tmp_path / "output.docx"

    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0

    for state_name, mutate in (
        ("dirty", lambda w: (w / PROJECTION_FILE).write_text(edit_text(w).replace("前", "草稿", 1), encoding="utf-8")),
        ("stale-clean", lambda w: (w / "typed.md").write_text(
            (w / "typed.md").read_text(encoding="utf-8").replace("加粗", "加粗改", 1), encoding="utf-8"
        )),
    ):
        workdir = extract_fixture(tmp_path / state_name)
        output = tmp_path / f"{state_name}.docx"
        mutate(workdir)
        assert build([str(workdir), "-o", str(output)]) == 1
        assert not output.exists()

    workdir = extract_fixture(tmp_path / "missing")
    (workdir / PROJECTION_FILE).unlink()
    (workdir / STATE_FILE).unlink()
    assert build([str(workdir), "-o", str(tmp_path / "missing.docx")]) == 1
    assert not (tmp_path / "missing.docx").exists()


def test_verify_gate_rejects_dirty(tmp_path):
    workdir = extract_fixture(tmp_path)
    output = tmp_path / "ok.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    (workdir / PROJECTION_FILE).write_text(
        edit_text(workdir).replace("前", "草稿", 1), encoding="utf-8"
    )
    assert verify([str(workdir), str(output)]) == 1


def test_validate_cli_gate_rejects_non_clean(tmp_path):
    workdir = extract_fixture(tmp_path)
    from scripts.verify import validate

    assert validate([str(workdir)]) == 0
    (workdir / PROJECTION_FILE).write_text(
        edit_text(workdir).replace("前", "草稿", 1), encoding="utf-8"
    )
    assert validate([str(workdir)]) == 1


# --------------------------------------------------------------------------
# Protected structure and grammar
# --------------------------------------------------------------------------

def test_token_swap_is_protected_token_mutated(tmp_path):
    workdir = extract_fixture(tmp_path)
    text = edit_text(workdir)
    tokens = re.findall(r"\u27e6token[^\u27e7]*\u27e7", text)
    assert len(tokens) == 2
    tab_token, br_token = tokens
    swapped = text.replace("前" + tab_token, "前" + br_token, 1)
    swapped = swapped.replace("中" + br_token, "中" + tab_token, 1)
    (workdir / PROJECTION_FILE).write_text(swapped, encoding="utf-8")
    for call in (edit_status, sync_edit_projection):
        try:
            call(workdir)
        except Exception as exc:
            assert "protected-token-mutated" in str(exc)
        else:
            raise AssertionError("token swap must be rejected")
    output = tmp_path / "blocked.docx"
    assert build([str(workdir), "-o", str(output)]) == 1
    assert not output.exists()


def test_paragraph_id_change_is_protected_token_mutated(tmp_path):
    workdir = extract_fixture(tmp_path)
    text = edit_text(workdir)
    (workdir / PROJECTION_FILE).write_text(text.replace('id="P0"', 'id="P9"', 1), encoding="utf-8")
    try:
        edit_status(workdir)
    except Exception as exc:
        assert "protected-token-mutated" in str(exc)
    else:
        raise AssertionError("paragraph ID change must be rejected")


def test_grammar_rejects_literal_bracket_and_unknown_marker(tmp_path):
    workdir = extract_fixture(tmp_path)
    text = edit_text(workdir)
    (workdir / PROJECTION_FILE).write_text(
        text.replace("中", "\u27e6裸\u27e7", 1), encoding="utf-8"
    )
    try:
        edit_status(workdir)
    except Exception as exc:
        assert "edit-grammar-invalid" in str(exc)
    else:
        raise AssertionError("literal brackets must be rejected")

    workdir = extract_fixture(tmp_path / "marker")
    text = edit_text(workdir)
    (workdir / PROJECTION_FILE).write_text(
        text + '\n<!--@mystery id="X"-->\n', encoding="utf-8"
    )
    try:
        edit_status(workdir)
    except Exception as exc:
        assert "edit-grammar-invalid" in str(exc)
    else:
        raise AssertionError("unknown marker must be rejected")


# --------------------------------------------------------------------------
# Real complex fixture: opaque / range / anchor placeholders
# --------------------------------------------------------------------------

def test_complex_fixture_projection_preserves_placeholders(tmp_path):
    workdir = tmp_path / "workdir"
    assert extract([str(FIXTURE), "-o", str(workdir)]) == 0
    text = edit_text(workdir)
    assert "<span" not in text
    assert 'kind="hyperlink"' in text
    assert "range-end" in text
    assert 'kind="tab"' in text
    assert 'kind="br"' in text
    assert 'kind="bookmark-start"' in text and 'kind="bookmark-end"' in text
    assert 'kind="comment-start"' in text and 'kind="comment-end"' in text
    assert 'kind="opaque"' in text or 'kind="unsupported-run"' in text
    assert edit_status(workdir)["state"] == "clean"
    assert build([str(workdir), "-o", str(tmp_path / "out.docx")]) == 0
    assert verify([str(workdir), str(tmp_path / "out.docx")]) == 0

"""Frozen audit-mode governance: explicit scan → policy → apply.

Regression and integration coverage for the vertical-normalization audit
contract (Issue #2 addendum): scans are read-only and hash-bound, candidate
classification is never a policy decision, policies must be complete and
explicitly approved, apply is the only mutation step, and any failure leaves
the source project and prior artifacts untouched.

The seam under test is the full typed-workdir lifecycle: real DOCX extraction,
``audit scan`` / ``audit apply``, the core contract functions, and independent
verify of the derived normalized workdir.
"""

import importlib
import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from scripts.audit_contract import (
    AUDIT_SCHEMA,
    DECISIONS,
    POLICY_SCHEMA,
    RISKY_CLASSIFICATIONS,
    RUN_EVIDENCE_SCHEMA,
    SCAN_SCHEMA,
    SCANNER_CONTRACT_VERSION,
    approve_policy,
    create_policy,
    create_scan_artifact,
    create_snapshot,
    fallback_project_id,
    fingerprint_candidate,
    payload_sha256,
    policy_sha256,
    require_approved,
    validate_policy,
    validate_scan_artifact,
)
from scripts.audit import audit
from scripts.extract import extract
from scripts.typed_docx import ValidationError
from scripts.typed_normalize import load_catalog
from scripts.verify import verify

# U+0656 ARABIC SUBSCRIPT ALEF — catalog class ``manual``, no proposed target.
MANUAL_SOURCE = "\u0656"
MANUAL_CODEPOINT = "U+0656"

SHA256 = "0" * 64


def _make_docx(path: Path, text: str = "₂") -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def _extract(source: Path, workdir: Path) -> None:
    assert extract([str(source), "-o", str(workdir)]) == 0


def _workdir_hashes(workdir: Path) -> dict[str, str]:
    hashes = {}
    for path in sorted(Path(workdir).rglob("*")):
        if path.is_file():
            hashes[str(path.relative_to(workdir))] = hashlib.sha256(path.read_bytes()).hexdigest()
    return hashes


def _scan_workdir(tmp_path, text: str = "₂"):
    """Extract a real workdir and run ``audit scan``; return the artifacts."""
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    scan_path = tmp_path / "scan.json"
    _make_docx(source, text)
    _extract(source, workdir)
    assert audit(["scan", str(workdir), "-o", str(scan_path)]) == 0
    scan = json.loads(scan_path.read_text(encoding="utf-8"))
    return source, workdir, scan_path, scan


def _decisions_for(scan: dict, decision: str = "convert", actor: str = "tester") -> dict:
    """Complete decision set: convert approved, preserve risky with rationale."""
    decisions = {}
    for candidate in scan["candidates"]:
        item = {
            "decision": decision if candidate["classification"] == "approved" else "preserve",
            "actor": actor,
            "candidate_fingerprint": candidate["candidate_fingerprint"],
        }
        if candidate["classification"] in RISKY_CLASSIFICATIONS:
            item["rationale"] = "preserved: no safe catalog target"
        decisions[candidate["occurrence_id"]] = item
    return decisions


def _write_policy(policy_path: Path, policy: dict) -> None:
    policy_path.write_text(
        json.dumps(policy, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def _apply_args(workdir, scan_path, policy_path, out, workdir_out):
    return [
        "apply",
        str(workdir),
        "--scan",
        str(scan_path),
        "--policy",
        str(policy_path),
        "-o",
        str(out),
        "--workdir-out",
        str(workdir_out),
    ]


def _synthetic_candidate(
    occurrence_id: str = "P0-V0001",
    classification: str = "approved",
    source: str = "₂",
    codepoint: str = "U+2082",
    proposed_target: str = "2",
) -> dict:
    return {
        "candidate_id": f"C-{occurrence_id}",
        "occurrence_id": occurrence_id,
        "paragraph_id": "P0",
        "node_path": "body/P0/text[0]",
        "visible_offset": 0,
        "codepoint": codepoint,
        "source": source,
        "style_id": "s_test",
        "classification": classification,
        "proposed_target": proposed_target,
        "vertical": "subscript",
        "reversible": True,
        "context": f"x{source}y",
    }


def _synthetic_scan() -> dict:
    snapshot = create_snapshot(
        baseline_sha256="a" * 64,
        draft_snapshot_sha256="b" * 64,
        catalog_sha256="c" * 64,
    )
    return create_scan_artifact(snapshot=snapshot, candidates=[_synthetic_candidate()])


# ---------------------------------------------------------------------------
# Scan: read-only, hash-bound artifact, run evidence
# ---------------------------------------------------------------------------

def test_audit_scan_is_read_only_and_emits_hash_bound_artifact_and_run_evidence(tmp_path):
    source, workdir, scan_path, scan = _scan_workdir(tmp_path)
    before = _workdir_hashes(workdir)

    assert scan["schema"] == SCAN_SCHEMA
    assert scan["scanner_contract_version"] == SCANNER_CONTRACT_VERSION
    snapshot = scan["snapshot"]
    assert snapshot["catalog_sha256"] == load_catalog()["catalog_hash"]
    # Deterministic fallback project identity derived from the baseline.
    assert snapshot["project_id"] == fallback_project_id(snapshot["baseline_sha256"])
    # Observable snapshot derivation: baseline = template fingerprint, draft = typed.md.
    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    assert snapshot["baseline_sha256"] == format_data["template_sha256"]
    assert snapshot["draft_snapshot_sha256"] == hashlib.sha256((workdir / "typed.md").read_bytes()).hexdigest()
    # Hash-bound artifact: the stored digest is a pure function of the content.
    assert scan["scan_artifact_sha256"] == payload_sha256(scan, "scan_artifact_sha256")
    assert validate_scan_artifact(scan)

    candidates = scan["candidates"]
    assert len(candidates) == 1
    candidate = candidates[0]
    assert candidate["source"] == "₂"
    assert candidate["codepoint"] == "U+2082"
    assert candidate["classification"] == "approved"
    assert candidate["proposed_target"] == "2"
    assert candidate["occurrence_id"] == "P0-V0001"
    assert candidate["paragraph_id"] == "P0"
    assert candidate["candidate_id"]
    assert candidate["node_path"]
    assert isinstance(candidate["visible_offset"], int) and candidate["visible_offset"] >= 0
    assert candidate["style_id"]
    assert candidate["candidate_fingerprint"] == fingerprint_candidate(candidate, snapshot)

    run = json.loads(scan_path.with_name(scan_path.name + ".run.json").read_text(encoding="utf-8"))
    assert run["schema"] == RUN_EVIDENCE_SCHEMA
    assert run["command"] and "scan" in run["command"]
    assert run["status"] == "ok"
    assert run["started_at"] and run["finished_at"]
    assert run["scan_artifact_sha256"] == scan["scan_artifact_sha256"]
    assert any(key.endswith("_sha256") for key in run["inputs"])
    assert json.dumps(run["outputs"])  # output roles recorded
    assert str(tmp_path) not in json.dumps(run)
    source_ref = Path(format_data["source_path"])
    assert not source_ref.is_absolute()
    assert (workdir / source_ref).resolve() == source.resolve()

    # Read-only: the source workdir is byte-identical after scanning.
    assert _workdir_hashes(workdir) == before


def test_scan_candidate_offsets_are_per_occurrence(tmp_path):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path, text="₂₂")
    assert [candidate["occurrence_id"] for candidate in scan["candidates"]] == ["P0-V0001", "P0-V0002"]
    assert [candidate["visible_offset"] for candidate in scan["candidates"]] == [0, 1]
    assert len({candidate["candidate_fingerprint"] for candidate in scan["candidates"]}) == 2


# ---------------------------------------------------------------------------
# Classification is never a decision
# ---------------------------------------------------------------------------

def test_classification_is_never_a_decision():
    scan = _synthetic_scan()
    candidate = scan["candidates"][0]
    for classification in RISKY_CLASSIFICATIONS:
        raw = dict(_synthetic_candidate(classification=classification, proposed_target=""))
        if classification != "unsupported":
            raw["reversible"] = False
        risky_scan = create_scan_artifact(snapshot=scan["snapshot"], candidates=[raw])
        echo = create_policy(
            scan=risky_scan,
            decisions={
                candidate["occurrence_id"]: {
                    "decision": classification,  # classification echoed as a decision
                    "actor": "tester",
                    "candidate_fingerprint": risky_scan["candidates"][0]["candidate_fingerprint"],
                    "rationale": "echoed",
                }
            },
        )
        with pytest.raises(ValidationError):
            validate_policy(echo, scan=risky_scan)

    # An approved candidate with no decision at all is pending, not implicitly decided.
    policy = create_policy(scan=scan)
    with pytest.raises(ValidationError):
        validate_policy(policy, scan=scan)


def test_manual_candidate_requires_explicit_decision_and_rationale(tmp_path):
    source, workdir, scan_path, scan = _scan_workdir(tmp_path, text="₂" + MANUAL_SOURCE)
    candidates = {candidate["source"]: candidate for candidate in scan["candidates"]}
    manual = candidates[MANUAL_SOURCE]
    assert manual["classification"] == "manual"
    assert manual["codepoint"] == MANUAL_CODEPOINT
    assert manual["proposed_target"] == ""

    policy = create_policy(scan=scan, decisions={manual["occurrence_id"]: {"actor": "tester"}})
    with pytest.raises(ValidationError):
        validate_policy(policy, scan=scan)
    policy = create_policy(
        scan=scan,
        decisions={
            manual["occurrence_id"]: {
                "decision": "preserve",
                "actor": "tester",
                "candidate_fingerprint": manual["candidate_fingerprint"],
                # no rationale
            }
        },
    )
    with pytest.raises(ValidationError):
        validate_policy(policy, scan=scan)


# ---------------------------------------------------------------------------
# Completeness and decision hygiene
# ---------------------------------------------------------------------------

def test_incomplete_policy_is_rejected_before_apply(tmp_path):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    policy_path = tmp_path / "policy.json"
    _write_policy(policy_path, create_policy(scan=scan))  # no decisions at all
    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"

    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 1
    assert not out.exists()
    assert not workdir_out.exists()


def test_decision_fingerprint_must_match_the_scan_candidate():
    scan = _synthetic_scan()
    candidate = scan["candidates"][0]
    other = create_scan_artifact(
        snapshot=scan["snapshot"],
        candidates=[_synthetic_candidate(occurrence_id="P0-V0001", source="³", codepoint="U+00B3", proposed_target="3")],
    )
    policy = create_policy(
        scan=scan,
        decisions={
            candidate["occurrence_id"]: {
                "decision": "convert",
                "actor": "tester",
                "candidate_fingerprint": other["candidates"][0]["candidate_fingerprint"],
            }
        },
    )
    with pytest.raises(ValidationError):
        validate_policy(policy, scan=scan)


def test_missing_actor_and_invalid_rationale_are_rejected():
    scan = _synthetic_scan()
    candidate = scan["candidates"][0]
    occurrence = candidate["occurrence_id"]
    fingerprint = candidate["candidate_fingerprint"]

    for decision in (
        {"decision": "convert", "candidate_fingerprint": fingerprint},  # no actor
        {"decision": "convert", "actor": "   ", "candidate_fingerprint": fingerprint},
        {"decision": "convert", "actor": "tester", "candidate_fingerprint": fingerprint, "rationale": 42},
    ):
        with pytest.raises(ValidationError):
            validate_policy(create_policy(scan=scan, decisions={occurrence: decision}), scan=scan)


def test_convert_on_non_approved_classification_is_rejected():
    raw = _synthetic_candidate(classification="manual", proposed_target="")
    raw["reversible"] = False
    scan = create_scan_artifact(snapshot=_synthetic_scan()["snapshot"], candidates=[raw])
    candidate = scan["candidates"][0]
    policy = create_policy(
        scan=scan,
        decisions={
            candidate["occurrence_id"]: {
                "decision": "convert",
                "actor": "tester",
                "candidate_fingerprint": candidate["candidate_fingerprint"],
                "rationale": "wrong",
            }
        },
    )
    with pytest.raises(ValidationError):
        validate_policy(policy, scan=scan)

def test_non_reversible_and_conflicting_style_require_rationale():
    for mutation in (
        {"reversible": False},
        {"word_vert_align": "superscript"},
    ):
        raw = _synthetic_candidate()
        raw.update(mutation)
        scan = create_scan_artifact(snapshot=_synthetic_scan()["snapshot"], candidates=[raw])
        candidate = scan["candidates"][0]
        policy = create_policy(
            scan=scan,
            decisions={
                candidate["occurrence_id"]: {
                    "decision": "preserve",
                    "actor": "tester",
                    "candidate_fingerprint": candidate["candidate_fingerprint"],
                }
            },
        )
        with pytest.raises(ValidationError):
            validate_policy(policy, scan=scan)


def test_unknown_occurrence_is_rejected():
    scan = _synthetic_scan()
    candidate = scan["candidates"][0]
    policy = create_policy(
        scan=scan,
        decisions={
            candidate["occurrence_id"]: {
                "decision": "convert",
                "actor": "tester",
                "candidate_fingerprint": candidate["candidate_fingerprint"],
            },
            "P0-V9999": {"decision": "preserve", "actor": "tester"},
        },
    )
    with pytest.raises(ValidationError):
        validate_policy(policy, scan=scan)


def test_stale_snapshot_bindings_are_rejected():
    scan = _synthetic_scan()
    candidate = scan["candidates"][0]
    policy = approve_policy(
        create_policy(
            scan=scan,
            decisions={
                candidate["occurrence_id"]: {
                    "decision": "convert",
                    "actor": "tester",
                    "candidate_fingerprint": candidate["candidate_fingerprint"],
                }
            },
        ),
        scan=scan,
        approved_by="reviewer",
    )
    assert require_approved(policy, scan=scan)

    tampered = json.loads(json.dumps(policy))
    for field in (
        "project_id",
        "baseline_sha256",
        "draft_snapshot_sha256",
        "model_sha256",
        "catalog_sha256",
        "scan_artifact_sha256",
    ):
        mutated = json.loads(json.dumps(tampered))
        if field == "project_id":
            mutated[field] = "other-project"
        elif field == "scan_artifact_sha256":
            mutated[field] = "d" * 64
        else:
            mutated[field] = "e" * 64
        with pytest.raises(ValidationError):
            require_approved(mutated, scan=scan)

    wrong_version = json.loads(json.dumps(tampered))
    wrong_version["scanner_contract_version"] = 99
    with pytest.raises(ValidationError):
        require_approved(wrong_version, scan=scan)

    # Catalog mismatch is also rejected when the current catalog is supplied.
    with pytest.raises(ValidationError):
        require_approved(tampered, scan=scan, catalog_sha256="f" * 64)


def test_apply_rejects_stale_draft_after_scan(tmp_path, capsys):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    policy_path = tmp_path / "policy.json"
    policy = approve_policy(
        create_policy(scan=scan, decisions=_decisions_for(scan)),
        scan=scan,
        approved_by="reviewer",
    )
    _write_policy(policy_path, policy)

    typed_path = workdir / "typed.md"
    typed_path.write_text(typed_path.read_text(encoding="utf-8").replace("₂", "₂₂"), encoding="utf-8")

    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"
    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 1
    assert not out.exists()
    assert not workdir_out.exists()
    failure_run = Path(str(out) + ".run.json")
    assert failure_run.exists()
    assert str(tmp_path) not in failure_run.read_text(encoding="utf-8")
    error = capsys.readouterr().out.lower()
    assert "stale" in error or "incompatible" in error or "drift" in error


def test_apply_rejects_current_scan_candidate_drift(tmp_path, monkeypatch, capsys):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    policy_path = tmp_path / "policy.json"
    policy = approve_policy(
        create_policy(scan=scan, decisions=_decisions_for(scan)),
        scan=scan,
        approved_by="reviewer",
    )
    _write_policy(policy_path, policy)

    drifted = json.loads(json.dumps(scan))
    drifted["candidates"][0]["context"] = "different reviewed context"
    drifted["scan_artifact_sha256"] = payload_sha256(drifted, "scan_artifact_sha256")
    import scripts.typed_normalize as typed_normalize

    monkeypatch.setattr(typed_normalize, "scan_workdir", lambda *args, **kwargs: drifted)
    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"
    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 1
    assert not out.exists()
    assert not workdir_out.exists()
    assert "candidate set" in capsys.readouterr().out.lower()


def test_scan_fails_closed_when_run_evidence_write_fails(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    document = Document()
    document.add_paragraph("₂")
    document.save(source)
    _extract(source, workdir)
    scan_path = tmp_path / "scan.json"

    def fail_evidence(*args, **kwargs):
        raise OSError("evidence unavailable")

    audit_module = importlib.import_module("scripts.audit")
    monkeypatch.setattr(audit_module, "_write_run_evidence", fail_evidence)
    assert audit(["scan", str(workdir), "-o", str(scan_path)]) == 1
    assert not scan_path.exists()
    assert not Path(str(scan_path) + ".run.json").exists()


def test_apply_fails_closed_when_run_evidence_write_fails(tmp_path, monkeypatch):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    policy_path = tmp_path / "policy.json"
    policy = approve_policy(
        create_policy(scan=scan, decisions=_decisions_for(scan)),
        scan=scan,
        approved_by="reviewer",
    )
    _write_policy(policy_path, policy)
    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"

    def fail_evidence(*args, **kwargs):
        raise OSError("evidence unavailable")

    audit_module = importlib.import_module("scripts.audit")
    monkeypatch.setattr(audit_module, "_write_run_evidence", fail_evidence)
    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 1
    assert not out.exists()
    assert not workdir_out.exists()


# ---------------------------------------------------------------------------
# Approval gate
# ---------------------------------------------------------------------------

def test_approved_status_requires_an_explicit_approval_object():
    scan = _synthetic_scan()
    candidate = scan["candidates"][0]
    complete = create_policy(
        scan=scan,
        decisions={
            candidate["occurrence_id"]: {
                "decision": "convert",
                "actor": "tester",
                "candidate_fingerprint": candidate["candidate_fingerprint"],
            }
        },
    )
    # Hand-flipping the status is not an approval.
    forged = dict(complete)
    forged["status"] = "approved"
    with pytest.raises(ValidationError):
        require_approved(forged, scan=scan)

    # Approval record must satisfy the policy's requirement.
    wrong_requirement = dict(complete)
    wrong_requirement["status"] = "approved"
    wrong_requirement["approval"] = {
        "approved": True,
        "requirement": "self",
        "approved_by": "owner",
        "approval_time": "2026-08-05T00:00:00+00:00",
    }
    with pytest.raises(ValidationError):
        require_approved(wrong_requirement, scan=scan)

    # approve_policy refuses to approve an incomplete policy.
    with pytest.raises(ValidationError):
        approve_policy(create_policy(scan=scan), scan=scan, approved_by="reviewer")


def test_default_human_approval_requires_explicit_human_record(tmp_path):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    policy_path = tmp_path / "policy.json"
    complete = create_policy(scan=scan, decisions=_decisions_for(scan))
    assert complete["approval_requirement"] == "human"

    approved = approve_policy(complete, scan=scan, approved_by="reviewer@example.com")
    assert approved["status"] == "approved"
    assert approved["approval"]["approved"] is True
    assert approved["approval"]["requirement"] == "human"
    assert approved["approval"]["approved_by"] == "reviewer@example.com"
    assert approved["approval"]["approval_time"]
    require_approved(approved, scan=scan)

    # A complete-but-unapproved policy cannot apply through the CLI.
    _write_policy(policy_path, complete)
    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"
    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 1
    assert not out.exists()
    assert not workdir_out.exists()


def test_self_approval_works_only_when_deliberately_recorded(tmp_path):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    policy_path = tmp_path / "policy.json"
    complete = create_policy(scan=scan, decisions=_decisions_for(scan), approval_requirement="self")
    assert complete["approval_requirement"] == "self"
    # Explicitly configured self approval still requires the recorded approval.
    with pytest.raises(ValidationError):
        require_approved(complete, scan=scan)

    approved = approve_policy(complete, scan=scan, approved_by="project-owner")
    assert approved["approval"]["requirement"] == "self"
    require_approved(approved, scan=scan)

    _write_policy(policy_path, approved)
    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"
    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 0
    assert out.exists()
    assert workdir_out.exists()


def test_proposed_draft_policy_cannot_apply(tmp_path):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    policy_path = tmp_path / "policy.json"
    proposed = create_policy(scan=scan, decisions=_decisions_for(scan))
    assert proposed["status"] == "draft"  # AI proposal output is always a draft
    _write_policy(policy_path, proposed)

    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"
    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 1
    assert not out.exists()
    assert not workdir_out.exists()


# ---------------------------------------------------------------------------
# Apply: the only mutation step
# ---------------------------------------------------------------------------

def test_audit_apply_end_to_end_records_evidence_and_creates_new_workdir(tmp_path):
    source, workdir, scan_path, scan = _scan_workdir(tmp_path)
    source_digest = hashlib.sha256(source.read_bytes()).hexdigest()
    before = _workdir_hashes(workdir)
    policy_path = tmp_path / "policy.json"
    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"

    policy = approve_policy(
        create_policy(scan=scan, decisions=_decisions_for(scan)),
        scan=scan,
        approved_by="reviewer@example.com",
    )
    _write_policy(policy_path, policy)

    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 0
    assert out.exists()
    assert workdir_out.exists()

    # Source project and workdir are byte-identical after apply.
    assert hashlib.sha256(source.read_bytes()).hexdigest() == source_digest
    assert _workdir_hashes(workdir) == before
    assert "₂" in (workdir / "typed.md").read_text(encoding="utf-8")

    # The derived normalized workdir carries the conversion.
    normalized_typed = (workdir_out / "typed.md").read_text(encoding="utf-8")
    assert "₂" not in normalized_typed
    assert "2" in normalized_typed

    # Independent verification of the derived workdir and DOCX.
    assert verify([str(workdir_out), str(out)]) == 0
    with ZipFile(out) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "2" in xml
    assert "vertAlign" in xml
    assert "₂" not in xml

    # Occurrence-level audit evidence with old/new values.
    audit_path = workdir_out / "normalization.audit.json"
    audit_artifact = json.loads(audit_path.read_text(encoding="utf-8"))
    assert audit_artifact["schema"] == AUDIT_SCHEMA
    assert audit_artifact["status"] == "applied"
    assert audit_artifact["project_id"] == scan["snapshot"]["project_id"]
    assert audit_artifact["baseline_sha256"] == scan["snapshot"]["baseline_sha256"]
    assert audit_artifact["draft_snapshot_sha256"] == scan["snapshot"]["draft_snapshot_sha256"]
    assert audit_artifact["scan_artifact_sha256"] == scan["scan_artifact_sha256"]
    assert audit_artifact["policy_sha256"] == policy_sha256(policy)
    assert audit_artifact["catalog_sha256"] == scan["snapshot"]["catalog_sha256"]
    occurrences = audit_artifact["occurrences"]
    assert len(occurrences) == 1
    occurrence = occurrences[0]
    assert occurrence["occurrence_id"] == "P0-V0001"
    assert occurrence["paragraph_id"] == "P0"
    assert occurrence["candidate_fingerprint"] == scan["candidates"][0]["candidate_fingerprint"]
    assert occurrence["old_text"] == "₂"
    assert occurrence["new_text"] == "2"
    assert occurrence["old_style_id"]
    assert occurrence["new_style_id"]
    assert occurrence["old_style_id"] != occurrence["new_style_id"]
    assert occurrence["style_delta"] == {"vertical": "subscript", "reversible": True}
    assert occurrence["classification"] == "approved"
    assert occurrence["decision"] == "convert"
    assert occurrence["actor"] == "tester"
    assert occurrence["scan_artifact_sha256"] == scan["scan_artifact_sha256"]
    assert occurrence["policy_sha256"] == policy_sha256(policy)

    # Apply run evidence next to the output artifact.
    run = json.loads(Path(str(out) + ".run.json").read_text(encoding="utf-8"))
    assert run["schema"] == RUN_EVIDENCE_SCHEMA
    assert run["status"] == "ok"
    assert run["command"] and "apply" in run["command"]
    assert run["started_at"] and run["finished_at"]
    assert run["policy_sha256"] == policy_sha256(policy)
    assert run["scan_artifact_sha256"] == scan["scan_artifact_sha256"]
    assert str(tmp_path) not in json.dumps(run)
    assert str(tmp_path) not in json.dumps(audit_artifact)
    normalized_format = json.loads((workdir_out / "format.json").read_text(encoding="utf-8"))
    normalized_source_ref = Path(normalized_format["source_path"])
    assert not normalized_source_ref.is_absolute()
    assert (workdir_out / normalized_source_ref).resolve() == out.resolve()


def test_old_policy_cannot_apply_to_derived_normalized_workdir(tmp_path):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    policy_path = tmp_path / "policy.json"
    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"
    policy = approve_policy(
        create_policy(scan=scan, decisions=_decisions_for(scan)),
        scan=scan,
        approved_by="reviewer",
    )
    _write_policy(policy_path, policy)
    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 0

    # The same approved policy is bound to the original baseline: applying it
    # to the derived normalized workdir is a stale-binding rejection, not a no-op.
    second_out = tmp_path / "again.docx"
    second_workdir = tmp_path / "again-workdir"
    assert audit(_apply_args(workdir_out, scan_path, policy_path, second_out, second_workdir)) == 1
    assert not second_out.exists()
    assert not second_workdir.exists()


def test_failed_apply_publishes_nothing_and_preserves_original(tmp_path):
    _, workdir, scan_path, scan = _scan_workdir(tmp_path)
    before = _workdir_hashes(workdir)
    source = tmp_path / "source.docx"
    source_digest = hashlib.sha256(source.read_bytes()).hexdigest()
    policy_path = tmp_path / "policy.json"
    out = tmp_path / "normalized.docx"
    workdir_out = tmp_path / "normalized-workdir"

    # Unapproved complete policy: governance failure before any mutation.
    _write_policy(policy_path, create_policy(scan=scan, decisions=_decisions_for(scan)))
    assert audit(_apply_args(workdir, scan_path, policy_path, out, workdir_out)) == 1
    assert not out.exists()
    assert not workdir_out.exists()
    assert _workdir_hashes(workdir) == before
    assert hashlib.sha256(source.read_bytes()).hexdigest() == source_digest

    # Failure run evidence records status and diagnostics next to the output.
    run = json.loads(Path(str(out) + ".run.json").read_text(encoding="utf-8"))
    assert run["schema"] == RUN_EVIDENCE_SCHEMA
    assert run["status"] == "error"
    assert run["diagnostics"]


def test_audit_usage_errors_exit_2(capsys):
    # argparse usage errors surface as process exit code 2 via SystemExit.
    for argv in (["bogus-command"], []):
        with pytest.raises(SystemExit) as excinfo:
            audit(argv)
        assert excinfo.value.code == 2


def test_cli_dispatcher_routes_audit_command(tmp_path):
    from scripts import main

    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    scan_path = tmp_path / "scan.json"
    _make_docx(source)
    _extract(source, workdir)
    assert main(["audit", "scan", str(workdir), "-o", str(scan_path)]) == 0
    assert scan_path.exists()

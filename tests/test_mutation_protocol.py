"""Issue #47 acceptance: Result/Evidence/Operation-ID path for the finite
workflow extract -> edit -> build -> verify and the sole partial `decide
apply`. Tests exercise only public seams: the CLI entry (``scripts.main`` /
installed ``python -m scripts``) and the MCP stdio/tool surface."""
from __future__ import annotations

import json
import os
import subprocess
import sys
import uuid
import zipfile
from pathlib import Path

import anyio
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from mcp import ClientSession, StdioServerParameters
from mcp.client.stdio import stdio_client

from scripts import main
from scripts.extract import extract
from scripts.mcp_server import (
    get_paragraph,
    replace_text,
    session,
    verify_output,
    workdir_open,
)
from scripts.protocol import semantic_sha256

ROOT = Path(__file__).resolve().parents[1]


def _op() -> str:
    return uuid.uuid4().hex


def _make_doc(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前言")
    cn = paragraph.add_run("智能响应")
    cn.font.name = "宋体"
    cn._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    en = paragraph.add_run("ABC")
    en.font.name = "Times New Roman"
    paragraph.add_run("后语")
    document.add_paragraph("第二段")
    document.save(path)


def _make_plain_doc(path: Path) -> None:
    document = Document()
    document.add_paragraph("Protocol one")
    document.save(path)


def _make_revision_doc(path: Path) -> None:
    """Plain paragraph with one injected tracked insertion (w:ins id=99)."""
    document = Document()
    document.add_paragraph("前言")
    document.add_paragraph("第二段")
    document.save(path)
    with zipfile.ZipFile(path) as z:
        files = {n: z.read(n) for n in z.namelist()}
    ins = (
        '<w:ins w:id="99" w:author="tester" w:date="2026-01-01T00:00:00Z">'
        '<w:r><w:t>修订词</w:t></w:r></w:ins>'
    ).encode()
    doc = files["word/document.xml"]
    doc = doc.replace(
        "<w:r><w:t>前言</w:t></w:r>".encode(),
        "<w:r><w:t>前言</w:t></w:r>".encode() + ins,
        1,
    )
    files["word/document.xml"] = doc
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items():
            z.writestr(name, data)


def _extract(source: Path, outdir: Path) -> Path:
    assert extract([str(source), "-o", str(outdir)]) == 0
    return outdir


def _extract_store(tmp_path: Path, name: str) -> Path:
    """Extract via the JSON CLI, which births the immutable-generation store
    (generation 0) as part of the extract success envelope."""
    source = tmp_path / f"{name}-src.docx"
    _make_doc(source)
    workdir = tmp_path / name
    assert main(["--json", "extract", str(source), "-o", str(workdir), "--operation-id", _op()]) == 0
    return workdir


def _open_workdir(tmp_path: Path, name: str) -> Path:
    source = tmp_path / f"{name}-src.docx"
    workdir = tmp_path / name
    _make_doc(source)
    _extract(source, workdir)
    session.workdir = None
    return Path(json.loads(workdir_open(str(workdir)))["workdir"])


def _dirty_edit(workdir: Path) -> None:
    path = workdir / "edit.md"
    text = path.read_text(encoding="utf-8")
    marker = text.find("-->", text.find("<!--@p id="))
    tail = text.find("\n", marker) + 1
    end = text.find("\n", tail)
    if end < 0:
        end = len(text)
    text = text[:tail] + text[tail:end] + "改" + text[end:]
    path.write_text(text, encoding="utf-8")


def _cli(*args: object) -> tuple[int, str]:
    result = subprocess.run(
        [sys.executable, "-m", "scripts", *(str(a) for a in args)],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        capture_output=True,
        timeout=600,
    )
    return result.returncode, (result.stdout or result.stderr).strip()


def _mtime_ns(path: Path) -> int:
    return os.stat(path).st_mtime_ns


# --------------------------------------------------------------------------
# Criterion 1: one JSON envelope per finite operation, exit 0/1/2
# --------------------------------------------------------------------------

def test_cli_workflow_emits_exactly_one_json_envelope_each(tmp_path, capsys):
    source = tmp_path / "src.docx"
    _make_plain_doc(source)
    wd = tmp_path / "wd"

    assert main(["--json", "extract", str(source), "-o", str(wd), "--operation-id", _op()]) == 0
    extract_out = capsys.readouterr().out
    assert len(extract_out.splitlines()) == 1
    extract_result = json.loads(extract_out)
    assert extract_result["schema"] == "docx2typed-result-1"
    assert extract_result["operation"] == "extract"
    assert extract_result["outcome"] == "success"
    assert extract_result["data"]["operation_id"]

    _dirty_edit(wd)
    assert main(["--json", "edit", "sync", str(wd), "--operation-id", _op()]) == 0
    sync_result = json.loads(capsys.readouterr().out)
    assert sync_result["operation"] == "edit" and sync_result["outcome"] == "success"

    built = tmp_path / "built.docx"
    assert main(["--json", "build", str(wd), "-o", str(built), "--operation-id", _op()]) == 0
    build_result = json.loads(capsys.readouterr().out)
    assert build_result["operation"] == "build" and build_result["outcome"] == "success"

    assert main(["--json", "verify", str(wd), str(built)]) == 0
    verify_result = json.loads(capsys.readouterr().out)
    assert verify_result["operation"] == "verify" and verify_result["outcome"] == "success"

    # every finite workflow operation publishes its envelope with all fields
    for result in (extract_result, sync_result, build_result, verify_result):
        assert set(result) >= {"schema", "operation", "outcome", "data", "diagnostics", "evidence", "engine"}
        assert result["outcome"] in ("success", "failure")  # never partial


def test_cli_invocation_and_domain_exit_codes(tmp_path, capsys):
    assert main(["--json", "extract"]) == 2
    invocation = json.loads(capsys.readouterr().out)
    assert invocation["outcome"] == "failure"
    assert invocation["diagnostics"][0]["code"] == "invalid-arguments"

    assert main(["--json", "edit"]) == 2
    json.loads(capsys.readouterr().out)

    assert main(["--json", "no-such-command"]) == 2
    unknown = json.loads(capsys.readouterr().out)
    assert unknown["diagnostics"][0]["code"] == "invalid-arguments"

    assert main(["--json", "build", str(tmp_path / "missing")]) == 1
    missing = json.loads(capsys.readouterr().out)
    assert missing["outcome"] == "failure"
    assert missing["diagnostics"][0]["code"] == "workdir-not-found"

    assert main(["--json", "extract", str(tmp_path / "missing.docx"), "-o", str(tmp_path / "x")]) == 1
    no_input = json.loads(capsys.readouterr().out)
    assert no_input["diagnostics"][0]["code"] == "input-not-found"


# --------------------------------------------------------------------------
# Criterion 2: mutation/build/verify/partial publish run evidence
# --------------------------------------------------------------------------

def test_cli_publishes_run_evidence_per_outcome(tmp_path, capsys):
    source = tmp_path / "src.docx"
    _make_plain_doc(source)
    wd = tmp_path / "wd"
    assert main(["--json", "extract", str(source), "-o", str(wd), "--operation-id", _op()]) == 0
    extract_result = json.loads(capsys.readouterr().out)

    evidence_file = wd / "run.evidence.json"
    assert evidence_file.is_file()
    evidence = json.loads(evidence_file.read_text(encoding="utf-8"))
    assert evidence["schema"] == "docx2typed-run-evidence-1"
    assert evidence["outcome"] == "success"
    assert evidence["kind"] == "mutation"
    assert evidence["operation_id"] == extract_result["data"]["operation_id"]
    assert evidence["payload_sha256"] == semantic_sha256(evidence["payload"])
    assert "provenance" in evidence and "provenance" not in evidence["payload"]
    assert set(evidence["payload"]) >= {"engine", "contracts", "inputs", "outputs", "checks"}
    assert extract_result["evidence"] == [evidence]

    built = tmp_path / "built.docx"
    assert main(["--json", "build", str(wd), "-o", str(built), "--operation-id", _op()]) == 0
    build_result = json.loads(capsys.readouterr().out)
    build_evidence = json.loads((Path(str(built) + ".evidence.json")).read_text(encoding="utf-8"))
    assert build_evidence["kind"] == "build"
    assert build_evidence["payload"]["outputs"]["docx"]["sha256"]
    assert build_result["evidence"] == [build_evidence]

    assert main(["--json", "verify", str(wd), str(built)]) == 0
    verify_result = json.loads(capsys.readouterr().out)
    verify_evidence = json.loads((Path(str(built) + ".verify.evidence.json")).read_text(encoding="utf-8"))
    assert verify_evidence["kind"] == "verify"
    assert verify_evidence["payload"]["verdict"] == "pass"
    assert verify_result["evidence"] == [verify_evidence]


def test_cli_evidence_publish_failure_cannot_report_success(tmp_path, capsys):
    source = tmp_path / "src.docx"
    _make_plain_doc(source)
    wd = tmp_path / "wd"
    wd.mkdir()
    (wd / "run.evidence.json").mkdir()  # a directory blocks the evidence publish

    assert main(["--json", "extract", str(source), "-o", str(wd), "--operation-id", _op()]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "failure"
    assert result["diagnostics"][0]["code"] == "evidence-publish-failed"


def test_cli_evidence_publish_failure_envelope_is_deterministic(tmp_path, capsys):
    """Issue #50 finding: the CLI evidence-publish-failed diagnostic names
    the exception class and the stable evidence path — never the transient
    mkstemp temp filename embedded in raw OSException text — so independent
    attempts and identical retries emit byte-identical envelopes."""
    source = tmp_path / "src.docx"
    _make_plain_doc(source)
    wd = tmp_path / "wd"
    wd.mkdir()
    (wd / "run.evidence.json").mkdir()  # a directory blocks the evidence publish

    op = _op()
    assert main(["--json", "extract", str(source), "-o", str(wd), "--operation-id", op]) == 1
    first = json.loads(capsys.readouterr().out)
    diag = first["diagnostics"][0]
    assert diag["code"] == "evidence-publish-failed"
    detail = diag["message"]
    assert str(wd / "run.evidence.json") in detail  # stable evidence path
    assert ".tmp" not in detail  # never the mkstemp temp filename
    assert "could not be published: " in detail

    # Identical retry: the persisted failure envelope replays byte-exact.
    assert main(["--json", "extract", str(source), "-o", str(wd), "--operation-id", op]) == 1
    replayed = capsys.readouterr().out
    assert replayed == json.dumps(first, ensure_ascii=False, sort_keys=True, separators=(",", ":")) + "\n"

    # Independent fresh attempt (new operation id): the same stable
    # diagnostic, no temp name, byte-equal to the first failure.
    assert main(["--json", "extract", str(source), "-o", str(wd), "--operation-id", _op()]) == 1
    fresh = json.loads(capsys.readouterr().out)
    assert fresh["diagnostics"] == first["diagnostics"]
    assert fresh["diagnostics"][0]["code"] == "evidence-publish-failed"


# --------------------------------------------------------------------------
# Criterion 3: Operation-ID replay and reuse, no second effect
# --------------------------------------------------------------------------

def test_cli_operation_id_replay_reuses_across_processes(tmp_path):
    src1 = tmp_path / "src1.docx"
    _make_plain_doc(src1)
    wd = tmp_path / "wd"
    op = _op()

    rc1, out1 = _cli("--json", "extract", src1, "-o", wd, "--operation-id", op)
    assert rc1 == 0
    first = json.loads(out1)
    assert first["data"]["operation_id"] == op
    typed_mtime = _mtime_ns(wd / "typed.md")
    state_mtime = _mtime_ns(wd / "edit.state.json")

    rc2, out2 = _cli("--json", "extract", src1, "-o", wd, "--operation-id", op)
    assert rc2 == 0
    assert out2 == out1  # identical retry returns the byte-identical original Result
    assert _mtime_ns(wd / "typed.md") == typed_mtime  # no second effect
    assert _mtime_ns(wd / "edit.state.json") == state_mtime

    # changed canonical input (different source) against the same anchor
    src2 = tmp_path / "src2.docx"
    _make_plain_doc(src2)
    rc3, out3 = _cli("--json", "extract", src2, "-o", wd, "--operation-id", op)
    assert rc3 == 1
    reused = json.loads(out3)
    assert reused["outcome"] == "failure"
    assert reused["diagnostics"][0]["code"] == "operation-id-reused"
    assert _mtime_ns(wd / "typed.md") == typed_mtime  # no second effect


def test_cli_operation_id_is_scoped_per_workdir_in_process(tmp_path, capsys):
    """Issue #50 finding 3: ledger replay is namespaced per workdir; the same
    operation-id on a different workdir in one process is a FRESH record
    (never the other workdir's in-memory replay), while replay within the
    same workdir stays byte-exact."""
    source = tmp_path / "src.docx"
    _make_plain_doc(source)
    op = _op()
    wd1 = tmp_path / "wd1"
    wd2 = tmp_path / "wd2"
    assert main(["--json", "extract", str(source), "-o", str(wd1), "--operation-id", op]) == 0
    capsys.readouterr()

    # Same op-id on a different workdir: the canonical input differs, so a
    # global in-memory record would reject it as reused. Namespacing makes it
    # a fresh record: extract succeeds with its own effect.
    assert main(["--json", "extract", str(source), "-o", str(wd2), "--operation-id", op]) == 0
    fresh = json.loads(capsys.readouterr().out)
    assert fresh["outcome"] == "success"
    assert fresh["data"]["operation_id"] == op
    assert wd2.is_dir()

    # Replay within the same workdir still returns the original envelope
    # byte-exact (no second extract effect).
    assert main(["--json", "extract", str(source), "-o", str(wd2), "--operation-id", op]) == 0
    replay = json.loads(capsys.readouterr().out)
    assert replay == fresh


def test_cli_build_replay_returns_original_result(tmp_path, capsys):
    source = tmp_path / "src.docx"
    _make_plain_doc(source)
    wd = tmp_path / "wd"
    built = tmp_path / "built.docx"
    op = _op()
    assert main(["--json", "extract", str(source), "-o", str(wd), "--operation-id", _op()]) == 0
    capsys.readouterr()
    assert main(["--json", "build", str(wd), "-o", str(built), "--operation-id", op]) == 0
    first = json.loads(capsys.readouterr().out)
    built_mtime = _mtime_ns(built)

    assert main(["--json", "build", str(wd), "-o", str(built), "--operation-id", op]) == 0
    second = json.loads(capsys.readouterr().out)
    assert second == first
    assert _mtime_ns(built) == built_mtime  # replay did not rebuild


# --------------------------------------------------------------------------
# Criterion 4: decide apply is the only partial mutation; human CLI intact
# --------------------------------------------------------------------------

def _revision_key(workdir: Path) -> str:
    inventory = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    return next(r["revision_key"] for r in inventory["revisions"] if r.get("editable"))


def test_decide_apply_partial_is_the_only_partial_outcome(tmp_path, capsys):
    source = tmp_path / "rev.docx"
    _make_revision_doc(source)
    wd = _extract(source, tmp_path / "rev-wd")
    capsys.readouterr()
    valid_key = _revision_key(wd)
    bogus_key = "word/document.xml|insert|424242|deadbeef0000"

    decisions = tmp_path / "mixed.json"
    decisions.write_text(
        json.dumps(
            {
                "schema": "docx2typed-review-decisions-1",
                "decisions": [
                    {"revision_key": valid_key, "decision": "accept"},
                    {"revision_key": bogus_key, "decision": "accept"},
                ],
            }
        ),
        encoding="utf-8",
    )
    assert main(["--json", "decide", "apply", "--workdir", str(wd), "--file", str(decisions), "--operation-id", _op()]) == 1
    result = json.loads(capsys.readouterr().out)
    assert result["outcome"] == "partial"
    assert result["data"]["applied"] == 1 and result["data"]["errors"] == 1
    assert result["data"]["published"][0]["revision_key"] == valid_key
    assert result["data"]["failed"][0]["revision_key"] == bogus_key
    assert result["diagnostics"][0]["code"] == "decision-apply-failed"
    assert result["evidence"][0]["kind"] == "partial-decision"
    assert set(result["evidence"][0]["payload"]) >= {"published", "failed", "not_attempted"}

    partial_evidence = json.loads((wd / "run.evidence.json").read_text(encoding="utf-8"))
    assert partial_evidence["outcome"] == "partial"
    assert partial_evidence["kind"] == "partial-decision"

    # all-valid apply: success, exit 0
    source2 = tmp_path / "rev2.docx"
    _make_revision_doc(source2)
    wd2 = _extract(source2, tmp_path / "rev-wd2")
    capsys.readouterr()
    valid2 = _revision_key(wd2)
    good = tmp_path / "good.json"
    good.write_text(
        json.dumps(
            {
                "schema": "docx2typed-review-decisions-1",
                "decisions": [{"revision_key": valid2, "decision": "accept"}],
            }
        ),
        encoding="utf-8",
    )
    assert main(["--json", "decide", "apply", "--workdir", str(wd2), "--file", str(good), "--operation-id", _op()]) == 0
    assert json.loads(capsys.readouterr().out)["outcome"] == "success"

    # all-failed apply: failure (not partial), exit 1
    source3 = tmp_path / "rev3.docx"
    _make_revision_doc(source3)
    wd3 = _extract(source3, tmp_path / "rev-wd3")
    capsys.readouterr()
    bad = tmp_path / "bad.json"
    bad.write_text(
        json.dumps(
            {
                "schema": "docx2typed-review-decisions-1",
                "decisions": [{"revision_key": bogus_key, "decision": "accept"}],
            }
        ),
        encoding="utf-8",
    )
    assert main(["--json", "decide", "apply", "--workdir", str(wd3), "--file", str(bad), "--operation-id", _op()]) == 1
    assert json.loads(capsys.readouterr().out)["outcome"] == "failure"


def test_cli_decide_partial_replay_returns_original_exit_1_no_second_effect(tmp_path, capsys):
    """Regression: the generic replay must accept a persisted partial
    envelope (exit 1, byte-exact) instead of treating it as never-completed
    and rerunning the mutation (a second decision apply)."""
    source = tmp_path / "rev-replay.docx"
    _make_revision_doc(source)
    wd = _extract(source, tmp_path / "rev-replay-wd")
    capsys.readouterr()
    valid_key = _revision_key(wd)
    bogus_key = "word/document.xml|insert|424242|deadbeef0000"
    decisions = tmp_path / "mixed.json"
    decisions.write_text(
        json.dumps(
            {
                "schema": "docx2typed-review-decisions-1",
                "decisions": [
                    {"revision_key": valid_key, "decision": "accept"},
                    {"revision_key": bogus_key, "decision": "accept"},
                ],
            }
        ),
        encoding="utf-8",
    )
    operation_id = _op()
    argv = ["--json", "decide", "apply", "--workdir", str(wd), "--file", str(decisions), "--operation-id", operation_id]
    assert main(argv) == 1
    first = json.loads(capsys.readouterr().out)
    assert first["outcome"] == "partial"
    evidence_mtime = _mtime_ns(wd / "run.evidence.json")

    assert main(argv) == 1  # replay: partial is a terminal outcome, never rerun
    second = json.loads(capsys.readouterr().out)
    assert second == first  # byte-exact partial envelope
    assert _mtime_ns(wd / "run.evidence.json") == evidence_mtime  # no second run/publish


def test_cli_human_mode_preserved(tmp_path, capsys):
    source = tmp_path / "src.docx"
    _make_plain_doc(source)
    wd = tmp_path / "wd"
    assert main(["extract", str(source), "-o", str(wd)]) == 0
    human_out = capsys.readouterr().out
    assert "workdir:" in human_out and "docx2typed-result-1" not in human_out

    built = tmp_path / "built.docx"
    assert main(["build", str(wd), "-o", str(built)]) == 0
    assert "built:" in capsys.readouterr().out


# --------------------------------------------------------------------------
# MCP: caller-supplied operation_id, Result envelopes, replay/reuse, evidence
# --------------------------------------------------------------------------

def _data(result) -> dict:
    """The data dict from a tool Result envelope's structuredContent."""
    return result.structuredContent["data"]


def test_mcp_mutation_tools_require_operation_id_and_replay(tmp_path):
    wd = _open_workdir(tmp_path, "mcp-schema")

    async def probe() -> None:
        parameters = StdioServerParameters(
            command=sys.executable,
            args=["-m", "scripts", "mcp"],
            cwd=str(ROOT),
        )
        async with stdio_client(parameters) as (reader, writer):
            async with ClientSession(reader, writer) as client:
                await client.initialize()
                tools = {tool.name: tool for tool in (await client.list_tools()).tools}
                schema = tools["replace_text"].inputSchema
                assert "operation_id" in schema["required"]
                assert "operation_id" in tools["build_docx"].inputSchema["required"]
                assert "operation_id" not in tools["verify_output"].inputSchema["required"]

                opened = await client.call_tool("workdir_open", {"workdir": str(wd)})
                assert opened.isError is False

                op = _op()
                first = await client.call_tool(
                    "replace_text",
                    {"paragraph_id": "P0", "old": "智能响应", "new": "智能调控", "operation_id": op},
                )
                assert first.isError is False
                assert first.structuredContent["schema"] == "docx2typed-result-1"
                assert first.structuredContent["data"]["draft"] == "dirty"
                assert first.structuredContent["evidence"][0]["schema"] == "docx2typed-run-evidence-1"
                replay = await client.call_tool(
                    "replace_text",
                    {"paragraph_id": "P0", "old": "智能响应", "new": "智能调控", "operation_id": op},
                )
                assert replay.isError is False
                assert replay.content[0].text == first.content[0].text
                assert replay.structuredContent == first.structuredContent

                reused = await client.call_tool(
                    "replace_text",
                    {"paragraph_id": "P0", "old": "智能响应", "new": "其它", "operation_id": op},
                )
                assert reused.isError is True
                assert reused.structuredContent["diagnostics"][0]["code"] == "operation-id-reused"

                missing = await client.call_tool(
                    "replace_text", {"paragraph_id": "P0", "old": "x", "new": "y"}
                )
                assert missing.isError is True

    anyio.run(probe)


def test_mcp_mutation_without_open_workdir_returns_result():
    session.workdir = None
    result = replace_text("P0", "old", "new", operation_id=_op())
    assert result.isError is True
    assert result.structuredContent["schema"] == "docx2typed-result-1"
    assert result.structuredContent["diagnostics"][0]["code"] == "workdir-not-open"


def test_mcp_mutation_replay_no_second_effect_and_evidence(tmp_path):
    wd = _open_workdir(tmp_path, "mcp-replay")
    op = _op()
    first = _data(replace_text("P0", "智能响应", "智能调控", operation_id=op))
    assert first["draft"] == "dirty" and first["operation_id"] == op

    replay = _data(replace_text("P0", "智能响应", "智能调控", operation_id=op))
    assert replay == first
    plain = json.loads(get_paragraph("P0"))["plain"]
    assert "智能调控" in plain and "智能响应" not in plain  # single effect only

    reused = replace_text("P0", "智能响应", "其它", operation_id=op)
    assert reused.isError is True
    assert reused.structuredContent["diagnostics"][0]["code"] == "operation-id-reused"

    evidence = json.loads((wd / "run.evidence.json").read_text(encoding="utf-8"))
    assert evidence["schema"] == "docx2typed-run-evidence-1"
    assert evidence["operation_id"] == op


def test_mcp_evidence_publish_failure_cannot_report_success(tmp_path, monkeypatch):
    """Store world: when the run evidence cannot be published, the mutation
    reports failure — never success — and the workdir draft is unchanged.
    The store owns evidence publication, so the failure is exercised through
    the evidence seam instead of a directory blocker at the root."""
    import scripts.protocol as protocol

    wd = _open_workdir(tmp_path, "mcp-evfail")
    real_publish = protocol.publish_run_evidence

    def boom(*args, **kwargs):
        raise OSError("evidence write failed")

    monkeypatch.setattr(protocol, "publish_run_evidence", boom)
    operation_id = _op()
    result = replace_text("P0", "智能响应", "智能调控", operation_id=operation_id)
    assert result.isError is True
    assert result.structuredContent["outcome"] == "failure"
    # The transaction rolled back: no committed generation, draft untouched.
    plain = json.loads(get_paragraph("P0"))["plain"]
    assert "智能响应" in plain and "智能调控" not in plain
    replay = replace_text("P0", "智能响应", "智能调控", operation_id=operation_id)
    assert replay.isError is True
    assert replay.structuredContent == result.structuredContent


def test_mcp_mutation_crash_after_effect_recovers_and_replays_no_second_effect(tmp_path):
    """Store world: a process death after the draft effect lands but before
    the completed journal leaves a recoverable transaction; startup recovery
    rolls it forward, and the identical replay returns the original envelope
    without a second draft mutation (replay lookup hits the generation the
    record was written under)."""
    from scripts.store import Store, _Kill, clear_faults, kill_at

    wd = _extract_store(tmp_path, "mcp-kill-recover")
    session.workdir = None
    json.loads(workdir_open(str(wd)))

    op = _op()
    kill_at("journal-write-completed")
    try:
        replace_text("P0", "智能响应", "智能调控", operation_id=op)
        raise AssertionError("expected simulated process death")
    except _Kill:
        pass
    clear_faults()
    store = Store.open(wd)
    recovered = store.recover()
    assert recovered["needs_recovery"] == []
    # Roll forward materialized the draft effect and repaired the ledger.
    plain = json.loads(get_paragraph("P0"))["plain"]
    assert "智能调控" in plain and "智能响应" not in plain
    replay = _data(replace_text("P0", "智能响应", "智能调控", operation_id=op))
    assert replay["draft"] == "dirty" and replay["operation_id"] == op
    plain = json.loads(get_paragraph("P0"))["plain"]
    assert "智能调控" in plain and "智能响应" not in plain  # single effect only


def test_mcp_mutation_evidence_failure_then_retry_succeeds_single_effect(tmp_path, monkeypatch):
    """A failed evidence publish never upgrades into a misleading success;
    once the cause is removed the identical retry succeeds with a single
    effect (no ledger record claims the failed attempt completed)."""
    import scripts.protocol as protocol

    wd = _open_workdir(tmp_path, "mcp-ev-retry")
    real_publish = protocol.publish_run_evidence

    def boom(*args, **kwargs):
        raise OSError("evidence write failed")

    monkeypatch.setattr(protocol, "publish_run_evidence", boom)
    operation_id = _op()
    first = replace_text("P0", "智能响应", "智能调控", operation_id=operation_id)
    assert first.isError is True
    assert first.structuredContent["outcome"] == "failure"

    monkeypatch.setattr(protocol, "publish_run_evidence", real_publish)
    retry = replace_text("P0", "智能响应", "智能调控", operation_id=operation_id)
    assert retry.isError is False
    assert retry.structuredContent["outcome"] == "success"
    plain = json.loads(get_paragraph("P0"))["plain"]
    assert "智能调控" in plain and "智能响应" not in plain  # single effect only


def test_mcp_verify_output_publishes_verification_evidence(tmp_path):
    wd = _open_workdir(tmp_path, "mcp-verify-ev")
    from scripts.mcp_server import build_docx, commit_sync

    _data(replace_text("P0", "智能响应", "智能调控", operation_id=_op()))
    _data(commit_sync(operation_id=_op()))
    output = _data(build_docx(operation_id=_op()))["output"]
    verified = _data(verify_output(output))
    assert verified["verified"] == output
    verify_evidence = json.loads((Path(str(output) + ".verify.evidence.json")).read_text(encoding="utf-8"))
    assert verify_evidence["schema"] == "docx2typed-run-evidence-1"
    assert verify_evidence["kind"] == "verify"


def test_mcp_corrupt_ledger_row_fails_closed_no_mutation(tmp_path):
    """Findings: a corrupt persisted ledger row for the operation_id — a
    completed row missing a required envelope field, or an envelope-less row
    without an explicit pending marker — fails closed with a structured
    ``operation-ledger-invalid`` Result instead of rerunning the mutation
    (no second effect on the workdir). The corrupt row is preserved. The
    authoritative ledger lives in the pinned generation directory."""
    from scripts.store import Store

    wd = _extract_store(tmp_path, "mcp-corrupt")
    session.workdir = None
    json.loads(workdir_open(str(wd)))
    store = Store.open(wd)
    ledger_file = store.ledger_dir() / "operation-ledger.json"
    draft_before = (wd / "edit.md").read_bytes()

    for label, row in (
        (
            "missing-envelope-field",
            {
                "input_sha256": "0" * 64,
                "envelope": {
                    "schema": "docx2typed-result-1",
                    "outcome": "success",
                },
            },
        ),
        (
            "envelope-less-not-pending",
            {"input_sha256": "0" * 64, "envelope": None},
        ),
    ):
        operation_id = _op()
        ledger_file.write_text(
            json.dumps(
                {
                    "schema": "docx2typed-operation-ledger-1",
                    "records": {operation_id: row},
                },
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
            )
            + "\n",
            encoding="utf-8",
        )
        result = replace_text("P0", "智能响应", "智能调控", operation_id=operation_id)
        assert result.isError is True, label
        assert result.structuredContent["outcome"] == "failure", label
        assert result.structuredContent["diagnostics"][0]["code"] == "operation-ledger-invalid", label
        assert (wd / "edit.md").read_bytes() == draft_before, label  # no second mutation
        persisted = json.loads(ledger_file.read_text(encoding="utf-8"))["records"][operation_id]
        assert persisted == row, label  # corrupt row preserved for inspection


def test_mcp_corrupt_ledger_in_advanced_generation_names_exact_file(tmp_path):
    """Issue #50 finding 2: with the pointer advanced past the committing
    generation, the corrupt-row MCP Result names the EXACT generation ledger
    file holding the row — never the pinned generation's ledger."""
    from scripts.store import STORE_DIR_NAME, Store

    wd = _extract_store(tmp_path, "mcp-adv-corrupt")
    session.workdir = None
    json.loads(workdir_open(str(wd)))
    # Advance the pointer with a real mutation.
    _data(replace_text("P0", "智能响应", "智能调控", operation_id=_op()))
    store = Store.open(wd)
    pinned = store.pin()["path"]
    gen_dirs = sorted((wd / STORE_DIR_NAME / "generations").iterdir())
    assert len(gen_dirs) >= 2  # extract birth + advance mutation
    old_ledger = next(
        (g / "operation-ledger.json" for g in gen_dirs if g.name != pinned.name),
        gen_dirs[0] / "operation-ledger.json",
    )
    pinned_ledger = pinned / "operation-ledger.json"
    operation_id = _op()
    old_ledger.write_text(
        json.dumps(
            {
                "schema": "docx2typed-operation-ledger-1",
                "records": {operation_id: {"input_sha256": "0" * 64, "envelope": None}},
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )
    result = replace_text("P0", "智能响应", "智能调控", operation_id=operation_id)
    assert result.isError is True
    assert result.structuredContent["outcome"] == "failure"
    diagnostic = result.structuredContent["diagnostics"][0]
    assert diagnostic["code"] == "operation-ledger-invalid"
    assert str(old_ledger) in diagnostic["message"]
    assert str(pinned_ledger) not in diagnostic["message"]

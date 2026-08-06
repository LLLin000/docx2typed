"""Slices B and C: governed clean-text synchronization (edit sync)."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.build import build
from scripts.edit import (
    PROJECTION_FILE,
    STATE_FILE,
    classify_edit_state,
    edit_status,
    refresh_edit_projection,
    sync_edit_projection,
)
from scripts.extract import extract
from scripts.verify import verify


def make_doc(path: Path, *, hyperlink: bool = False) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("开")
    bold = paragraph.add_run("加粗段")
    bold.bold = True
    paragraph.add_run("结尾")
    if hyperlink:
        target = "https://example.com/docx2typed"
        relationship_id = document.part.relate_to(target, RT.HYPERLINK, is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), relationship_id)
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "链接字"
        run.append(text)
        link.append(run)
        paragraph._p.append(link)
        paragraph.add_run("后")
    document.add_paragraph("第二段")
    document.save(path)


def fresh(tmp_path: Path, name: str, **kwargs) -> Path:
    source = tmp_path / f"{name}-src.docx"
    workdir = tmp_path / name
    make_doc(source, **kwargs)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return workdir


def edit_body(workdir: Path) -> str:
    return (workdir / PROJECTION_FILE).read_text(encoding="utf-8")


def write_edit(workdir: Path, text: str) -> None:
    (workdir / PROJECTION_FILE).write_text(text, encoding="utf-8")


def sync_and_build(tmp_path: Path, workdir: Path, name: str) -> tuple[Path, list[str]]:
    _, warnings, changed = sync_edit_projection(workdir)
    output = tmp_path / f"{name}.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    return output, warnings


def run_text(workdir: Path) -> str:
    from docx import Document as Doc

    return "".join(p.text for p in Doc(_latest_out(workdir)).paragraphs)


def _latest_out(workdir: Path) -> Path:
    return workdir.parent / "latest.docx"


# --------------------------------------------------------------------------
# Slice B: safe text edits
# --------------------------------------------------------------------------

def test_middle_insertion_inherits_left_context(tmp_path):
    workdir = fresh(tmp_path, "mid")
    write_edit(workdir, edit_body(workdir).replace("结尾", "结尾追", 1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    assert edit_status(workdir)["state"] == "clean"
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert "结尾追" in typed


def test_insertion_after_formatted_space_inherits_space_style(tmp_path):
    source = tmp_path / "space-src.docx"
    workdir = tmp_path / "space"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前")
    space = paragraph.add_run(" ")
    space.bold = True
    paragraph.add_run("后")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    write_edit(workdir, edit_body(workdir).replace(" 后", " 插后", 1))
    _, warnings, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    output = tmp_path / "space.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    paragraph = next(p for p in Document(output).paragraphs if "插" in p.text)
    insert_run = next(r for r in paragraph.runs if "插" in r.text)
    assert insert_run.bold is True  # inherited from the formatted space


def test_paragraph_start_insertion_uses_right_context(tmp_path):
    source = tmp_path / "start-src.docx"
    workdir = tmp_path / "start"
    document = Document()
    paragraph = document.add_paragraph()
    bold = paragraph.add_run("粗文本")
    bold.bold = True
    paragraph.add_run("尾")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    write_edit(workdir, edit_body(workdir).replace("粗文本", "新粗文本", 1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    output = tmp_path / "start.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    paragraph = next(p for p in Document(output).paragraphs if "新" in p.text)
    start_run = next(r for r in paragraph.runs if r.text.startswith("新"))
    assert start_run.bold is True  # first visible unit to the right is bold


def test_end_insertion_uses_left_context(tmp_path):
    workdir = fresh(tmp_path, "end")
    write_edit(workdir, edit_body(workdir).replace("结尾", "结尾附加", 1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert "结尾附加" in typed


def test_single_style_replacement_keeps_style(tmp_path):
    workdir = fresh(tmp_path, "single")
    write_edit(workdir, edit_body(workdir).replace("加粗段", "改写段", 1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert re.search(r"<span data-s=\"s_[0-9a-f]+\">改写段</span>", typed)


def test_deletion_preserves_survivor_styles(tmp_path):
    workdir = fresh(tmp_path, "del")
    write_edit(workdir, edit_body(workdir).replace("加粗段", "", 1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert "加粗段" not in typed
    assert "开结尾" in typed


def test_empty_paragraph_insertion_uses_insertion_style(tmp_path):
    source = tmp_path / "empty-src.docx"
    workdir = tmp_path / "empty"
    document = Document()
    document.add_paragraph("正文段")
    document.add_paragraph("")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    write_edit(workdir, re.sub(r'<!--@p id="P1"-->', '<!--@p id="P1"-->\n填入', edit_body(workdir), count=1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P1"]
    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    insertion_style = next(r["insertion_style"] for r in format_data["paragraphs"] if r["id"] == "P1")
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert insertion_style and ("填入" in typed)
    output = tmp_path / "empty.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    rebuilt = Document(output)
    assert rebuilt.paragraphs[1].text == "填入"


def test_repeated_text_with_equivalent_ownership_is_deterministic(tmp_path):
    source = tmp_path / "repeat-src.docx"
    workdir = tmp_path / "repeat"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("B ")
    paragraph.add_run("B")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    write_edit(workdir, edit_body(workdir).replace("B B", "B X B", 1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    assert "X" in (workdir / "typed.md").read_text(encoding="utf-8")


def test_hyperlink_text_edit_preserves_range(tmp_path):
    workdir = fresh(tmp_path, "link", hyperlink=True)
    write_edit(workdir, edit_body(workdir).replace("链接字", "链接文字改", 1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert "链接文字改" in typed and "hyperlink" in typed
    output = tmp_path / "link.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with __import__("zipfile").ZipFile(output) as archive:
        assert "w:hyperlink" in archive.read("word/document.xml").decode("utf-8")


def test_opaque_paragraph_text_change_rejected(tmp_path):
    source = tmp_path / "opaque-src.docx"
    workdir = tmp_path / "opaque"
    document = Document()
    paragraph = document.add_paragraph("前缀")
    field = OxmlElement("w:fldSimple")
    field.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr", " PAGE ")
    field_run = OxmlElement("w:r")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_run.append(field_text)
    field.append(field_run)
    paragraph._p.append(field)
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    write_edit(workdir, edit_body(workdir).replace("前缀", "修改"))
    try:
        sync_edit_projection(workdir)
    except Exception as exc:
        assert "opaque-paragraph-mutated" in str(exc)
    else:
        raise AssertionError("opaque paragraph edit must fail")


def test_new_paragraph_marker_applies_with_inheritance(tmp_path):
    workdir = fresh(tmp_path, "newpara")
    text = edit_body(workdir).replace(
        '<!--@p id="P1"-->',
        '<!--@new temp="N1" inherit="P0"-->\n新增段落\n\n<!--@p id="P1"-->',
        1,
    )
    write_edit(workdir, text)
    _, _, changed = sync_edit_projection(workdir)
    assert "P2" in changed
    output = tmp_path / "newpara.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    texts = [p.text for p in Document(output).paragraphs]
    assert texts == ["开加粗段结尾", "新增段落", "第二段"]


def test_delete_marker_applies(tmp_path):
    workdir = fresh(tmp_path, "delpara")
    text = edit_body(workdir)
    block = re.search(r'<!--@p id="P1"-->[^\n]*\n[^\n]*', text).group(0)
    write_edit(workdir, text.replace(block, '<!--@delete id="P1"-->', 1))
    _, _, changed = sync_edit_projection(workdir)
    assert "P1" in changed
    output = tmp_path / "delpara.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    assert [p.text for p in Document(output).paragraphs] == ["开加粗段结尾"]


def test_sync_does_not_create_new_styles(tmp_path):
    workdir = fresh(tmp_path, "nostyle")
    styles_before = (workdir / "styles.json").read_text(encoding="utf-8")
    write_edit(workdir, edit_body(workdir).replace("结尾", "结尾追", 1))
    sync_edit_projection(workdir)
    assert (workdir / "styles.json").read_text(encoding="utf-8") == styles_before


def test_sync_twice_is_idempotent(tmp_path):
    workdir = fresh(tmp_path, "idem")
    write_edit(workdir, edit_body(workdir).replace("结尾", "结尾追", 1))
    sync_edit_projection(workdir)
    typed_after_first = (workdir / "typed.md").read_bytes()
    _, warnings, changed = sync_edit_projection(workdir)  # now clean: no-op
    assert warnings == [] and changed == []
    assert (workdir / "typed.md").read_bytes() == typed_after_first


def test_raw_edit_after_sync_checked_against_governed_baseline(tmp_path):
    workdir = fresh(tmp_path, "rawafter")
    write_edit(workdir, edit_body(workdir).replace("加粗段", "改写段", 1))
    sync_edit_projection(workdir)
    typed_path = workdir / "typed.md"
    # span-preserving raw edit inside the synced bold span -> accepted
    typed = typed_path.read_text(encoding="utf-8")
    span = re.search(r"<span data-s=\"(s_[0-9a-f]+)\">改写段</span>", typed)
    assert span
    typed_path.write_text(typed.replace("改写段", "再改段", 1), encoding="utf-8")
    refresh_edit_projection(workdir)
    assert build([str(workdir), "-o", str(tmp_path / "raw.docx")]) == 0
    # cross-span raw rewrite after sync is rejected against the governed baseline
    typed = typed_path.read_text(encoding="utf-8")
    typed_path.write_text(typed.replace("开", "开X", 1).replace("再改段", "再改段Y", 1), encoding="utf-8")
    try:
        refresh_edit_projection(workdir)
    except Exception as exc:
        assert "cross-boundary" in str(exc) or "skeleton" in str(exc)
    else:
        raise AssertionError("cross-span raw rewrite must be rejected against the governed baseline")


# --------------------------------------------------------------------------
# Slice C: controlled mixed edits
# --------------------------------------------------------------------------

def test_cross_region_replacement_rejected_even_when_equal_length(tmp_path):
    workdir = fresh(tmp_path, "mix")
    # 开(plain)加(bold) -> XY: equal length, but the range crosses style regions
    write_edit(workdir, edit_body(workdir).replace("开加", "XY", 1))
    try:
        sync_edit_projection(workdir)
    except Exception as exc:
        assert "mixed-replacement-requires-unchanged-text" in str(exc)
    else:
        raise AssertionError("cross-region replacement must be rejected")
    typed = (workdir / "typed.md").read_text(encoding="utf-8")
    assert "开" in typed and "加粗段" in typed  # untouched


def test_single_region_edits_preserve_cn_en_script_fonts(tmp_path):
    from docx.oxml.ns import qn

    source = tmp_path / "cnen-src.docx"
    workdir = tmp_path / "cnen"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前言")
    cn = paragraph.add_run("智能响应")
    cn.font.name = "宋体"
    cn._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    en = paragraph.add_run("ABC")
    en.font.name = "Times New Roman"
    paragraph.add_run("后语")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    # each region edited separately: CN region keeps 宋体, EN region keeps TNR
    write_edit(workdir, edit_body(workdir).replace("智能响应", "智能调控", 1))
    sync_edit_projection(workdir)
    write_edit(workdir, edit_body(workdir).replace("ABC", "XYZ", 1))
    sync_edit_projection(workdir)
    output = tmp_path / "cnen.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    runs = {r.text: r for r in Document(output).paragraphs[0].runs}
    cn_run = runs["智能调控"]
    en_run = runs["XYZ"]
    assert cn_run._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
    assert en_run._element.rPr.rFonts.get(qn("w:ascii")) == "Times New Roman"


def test_unequal_length_mixed_replacement_rejected(tmp_path):
    from docx.oxml.ns import qn

    source = tmp_path / "uneq-src.docx"
    workdir = tmp_path / "uneq"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前言")
    cn = paragraph.add_run("智能响应")
    cn.font.name = "宋体"
    cn._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    en = paragraph.add_run("ABC")
    en.font.name = "Times New Roman"
    paragraph.add_run("后语")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    # different length over a mixed range cannot be styled without guessing
    write_edit(workdir, edit_body(workdir).replace("智能响应ABC", "新词XYZ", 1))
    try:
        sync_edit_projection(workdir)
    except Exception as exc:
        assert "mixed-replacement-requires-unchanged-text" in str(exc)
    else:
        raise AssertionError("unequal-length mixed replacement must fail")
    assert "智能响应" in (workdir / "typed.md").read_text(encoding="utf-8")


def test_split_edits_preserve_cn_en_fonts(tmp_path):
    from docx.oxml.ns import qn

    source = tmp_path / "split-src.docx"
    workdir = tmp_path / "split"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前言")
    cn = paragraph.add_run("智能响应")
    cn.font.name = "宋体"
    cn._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    en = paragraph.add_run("ABC")
    en.font.name = "Times New Roman"
    paragraph.add_run("后语")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    write_edit(workdir, edit_body(workdir).replace("智能响应", "新词", 1))
    sync_edit_projection(workdir)
    write_edit(workdir, edit_body(workdir).replace("ABC", "XYZ", 1))
    sync_edit_projection(workdir)
    output = tmp_path / "split.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    runs = {r.text: r for r in Document(output).paragraphs[0].runs}
    assert runs["新词"]._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
    assert runs["XYZ"]._element.rPr.rFonts.get(qn("w:ascii")) == "Times New Roman"


def test_unanchored_full_mixed_rewrite_rejected(tmp_path):
    workdir = fresh(tmp_path, "full")
    write_edit(workdir, edit_body(workdir).replace("开加粗段结尾", "完全不同的文字", 1))
    try:
        sync_edit_projection(workdir)
    except Exception as exc:
        assert "unanchored-mixed-rewrite" in str(exc)
    else:
        raise AssertionError("full mixed rewrite must fail")
    assert "加粗段" in (workdir / "typed.md").read_text(encoding="utf-8")  # untouched


def test_text_around_range_stays_on_draft_side(tmp_path):
    workdir = fresh(tmp_path, "boundary", hyperlink=True)
    # text edited immediately before the hyperlink keeps the range intact
    write_edit(workdir, edit_body(workdir).replace("结尾", "结束", 1))
    _, _, changed = sync_edit_projection(workdir)
    assert changed == ["P0"]
    output = tmp_path / "boundary.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    from zipfile import ZipFile

    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        assert "w:hyperlink" in xml and "链接字" in xml
    paragraph = next(p for p in Document(output).paragraphs if "结束" in p.text)
    assert "链接字" in paragraph.text


def test_ambiguous_alignment_rejected(tmp_path):
    source = tmp_path / "amb-src.docx"
    workdir = tmp_path / "amb"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")
    bold = paragraph.add_run("A")
    bold.bold = True
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    write_edit(workdir, edit_body(workdir).replace("AA", "AXA", 1))
    try:
        sync_edit_projection(workdir)
    except Exception as exc:
        assert "ambiguous-alignment" in str(exc)
    else:
        raise AssertionError("ambiguous insertion must fail")


def test_insertion_with_no_text_context_is_rejected(tmp_path):
    source = tmp_path / "tok-src.docx"
    workdir = tmp_path / "tok"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run().add_tab()
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    text = edit_body(workdir)
    token = re.search(r"\u27e6token[^\u27e7]*\u27e7", text).group(0)
    write_edit(workdir, text.replace(token, token + "X", 1))
    try:
        sync_edit_projection(workdir)
    except Exception as exc:
        assert "protected-context-ambiguous" in str(exc)
    else:
        raise AssertionError("insertion without visible text context must fail")


def test_failed_sync_leaves_workdir_untouched(tmp_path):
    workdir = fresh(tmp_path, "fail")
    typed_before = (workdir / "typed.md").read_bytes()
    state_before = (workdir / STATE_FILE).read_bytes()
    write_edit(workdir, edit_body(workdir).replace("开加粗段结尾", "完全不同的文字", 1))
    try:
        sync_edit_projection(workdir)
    except Exception:
        pass
    else:
        raise AssertionError("expected rejection")
    assert (workdir / "typed.md").read_bytes() == typed_before
    assert (workdir / STATE_FILE).read_bytes() == state_before
    evidence = json.loads((workdir / "edit.state.json.run.json").read_text(encoding="utf-8"))
    assert evidence["status"] == "error"

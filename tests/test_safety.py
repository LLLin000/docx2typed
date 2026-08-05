from docx import Document

from build import build
from extract import extract


def test_build_rejects_changed_source_and_preserves_existing_output(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("原文")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0

    document.add_paragraph("提取后修改")
    document.save(source)
    assert build([str(workdir), "-o", str(output)]) == 1
    assert not output.exists()

    source.unlink()
    output.write_bytes(b"sentinel")
    typed_path = workdir / "typed.md"
    typed_path.write_text(typed_path.read_text(encoding="utf-8") + "\ninvalid", encoding="utf-8")
    assert build([str(workdir), "-o", str(output)]) == 1
    assert output.read_bytes() == b"sentinel"



def test_copied_workdir_builds_without_original_source(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("自包含")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    source.unlink()
    assert build([str(workdir), "-o", str(output)]) == 0
    assert output.exists()
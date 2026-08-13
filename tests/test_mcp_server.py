"""docx2typed-mcp: span-free region-scoped editing tools and regions.md."""
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from scripts import main
from scripts.extract import extract
from scripts.protocol import file_sha256
from scripts.review_queue import dispatch, upsert_event
from scripts.review_collab import stage_patch
from scripts.mcp_server import (
    _fnv1a_utf16,
    batch_edit,
    build_docx,
    commit_sync,
    decide_all,
    delete_paragraph,
    diff_preview,
    get_paragraph,
    insert_paragraph,
    list_paragraphs,
    replace_text,
    revert,
    review_ack,
    review_apply_batch,
    review_apply_patch,
    review_inbox,
    review_preflight,
    review_settle,
    review_state,
    session,
    table_insert_row,
    verify_output,
    workdir_open,
    workdir_status,
)

ROOT = Path(__file__).resolve().parents[1]


def _reset() -> None:
    session.workdir = None


def make_doc(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("前言")
    cn = paragraph.add_run("智能响应")
    cn.font.name = "宋体"
    cn._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    en = paragraph.add_run("ABC")
    en.font.name = "Times New Roman"
    paragraph.add_run("后语")
    document.add_paragraph("第二段")
    document.save(path)


def open_workdir(tmp_path: Path, name: str) -> Path:
    source = tmp_path / f"{name}-src.docx"
    workdir = tmp_path / name
    make_doc(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return json.loads(workdir_open(str(workdir)))["workdir"]


def _j(result) -> dict:
    """Unwrap a tool result: the data dict from a Result envelope's
    structuredContent, or a legacy JSON string."""
    if hasattr(result, "structuredContent"):
        return result.structuredContent["data"]
    return json.loads(result)


def cn_en_runs(docx_path: Path) -> dict:
    return {r.text: r for r in Document(docx_path).paragraphs[0].runs}


def make_table_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("表前")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    document.add_paragraph("表后")
    document.save(path)


def open_store_workdir(tmp_path: Path, name: str) -> Path:
    """Extract via the JSON CLI, which births the immutable-generation store
    (generation 0), then open the MCP session."""
    source = tmp_path / f"{name}-src.docx"
    workdir = tmp_path / name
    make_doc(source)
    assert main(["--json", "extract", str(source), "-o", str(workdir), "--operation-id", f"{name}-extract-1"]) == 0
    return Path(json.loads(workdir_open(str(workdir)))["workdir"])


def test_regions_md_generated_at_extract(tmp_path):
    _reset()
    workdir = open_workdir(tmp_path, "regions")
    regions = (Path(workdir) / "regions.md").read_text(encoding="utf-8")
    assert "## P0" in regions and "## P1" in regions
    assert "智能响应" in regions and "{s_" in regions
    assert "[token]" not in regions  # no tokens in this fixture


def test_get_paragraph_has_style_id_description_rpr(tmp_path):
    _reset()
    open_workdir(tmp_path, "styles")
    data = _j(get_paragraph("P0"))
    styles = data["styles"]
    cn_region = next(r for r in styles if r["text"] == "智能响应")
    en_region = next(r for r in styles if r["text"] == "ABC")
    assert cn_region["style_id"].startswith("s_") and cn_region["style_id"] != en_region["style_id"]
    assert cn_region["description"]  # non-empty label
    assert cn_region["rpr"].startswith("<w:rPr")  # full canonical XML present


def test_batch_edit_by_index_preserves_cn_en_fonts(tmp_path):
    _reset()
    open_workdir(tmp_path, "batch")
    result = _j(batch_edit(
        "P0",
        [
            {"region": 1, "new": "智能调控"},
            {"region": 2, "new": "XYZ"},
        ],
        operation_id="batch-index-1",
    ))
    assert result["edits_applied"] == 2 and result["state"] == "clean"
    output = _j(build_docx(operation_id="batch-build-1"))["output"]
    assert _j(verify_output(output))["verified"] == output
    runs = cn_en_runs(output)
    assert runs["智能调控"]._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
    assert runs["XYZ"]._element.rPr.rFonts.get(qn("w:ascii")) == "Times New Roman"


def test_batch_edit_text_anchor_with_style_disambiguation(tmp_path):
    _reset()
    open_workdir(tmp_path, "anchor")
    _j(batch_edit("P0", [{"text": "ABC", "new": "XYZ"}], operation_id="batch-anchor-1"))
    assert _j(workdir_status())["state"] == "clean"
    data = _j(get_paragraph("P0"))
    assert "XYZ" in data["plain"] and "ABC" not in data["plain"]


def test_batch_edit_partial_old_inside_region(tmp_path):
    _reset()
    open_workdir(tmp_path, "partial")
    _j(batch_edit("P0", [{"region": 1, "old": "智能", "new": "智慧"}], operation_id="batch-partial-1"))
    data = _j(get_paragraph("P0"))
    assert "智慧响应" in data["plain"]


def test_batch_edit_atomic_rejection(tmp_path):
    _reset()
    open_workdir(tmp_path, "atomic")
    result = batch_edit(
        "P0",
        [
            {"region": 1, "new": "智能调控"},
            {"region": 2, "old": "not-there", "new": "XYZ"},
        ],
        operation_id="batch-atomic-1",
    )
    assert result.isError is True
    assert result.structuredContent["diagnostics"][0]["code"] == "text-not-found"
    assert _j(workdir_status())["state"] == "clean"  # rolled back
    data = _j(get_paragraph("P0"))
    assert "智能响应" in data["plain"] and "智能调控" not in data["plain"]


def test_batch_edit_region_out_of_range_and_duplicate(tmp_path):
    _reset()
    open_workdir(tmp_path, "range")
    for bad, code in (
        ([{"region": 99, "new": "x"}], "region-out-of-range"),
        ([{"region": 0, "new": "a"}, {"region": 0, "new": "b"}], "invalid-edit"),
    ):
        result = batch_edit("P0", bad, operation_id="batch-range-1")
        assert result.isError is True
        assert result.structuredContent["diagnostics"][0]["code"] == code


def test_regions_md_auto_updates_after_edit(tmp_path):
    _reset()
    workdir = open_workdir(tmp_path, "autoupdate")
    _j(batch_edit("P0", [{"region": 1, "new": "新词"}], operation_id="batch-auto-1"))
    regions = (Path(workdir) / "regions.md").read_text(encoding="utf-8")
    assert "新词" in regions and "智能响应" not in regions


def test_replace_text_still_rejects_cross_region(tmp_path):
    _reset()
    open_workdir(tmp_path, "cross")
    result = replace_text("P0", "智能响应ABC", "新词XYZ", operation_id="cross-region-1")
    assert result.isError is True
    assert result.structuredContent["diagnostics"][0]["code"] == "cross-region-text"


def test_full_workflow_through_mcp(tmp_path):
    _reset()
    open_workdir(tmp_path, "flow")
    _j(insert_paragraph("P0", "新增段", operation_id="flow-insert-1"))
    _j(delete_paragraph("P1", operation_id="flow-delete-1"))
    _j(commit_sync(operation_id="flow-commit-1"))
    _j(revert(operation_id="flow-revert-1"))
    output = _j(build_docx(operation_id="flow-build-1"))["output"]
    assert _j(verify_output(output))["verified"] == output
    texts = [p.text for p in Document(output).paragraphs]
    assert "新增段" in texts and "第二段" not in texts
    assert list_paragraphs()


def test_verify_output_evidence_publish_failure_is_deterministic(tmp_path, monkeypatch):
    """Issue #50 final finding: an evidence publish failure reports the
    deterministic '{type}: {stable path}' detail (never the transient temp
    filename) and every retry produces the byte-identical diagnostic."""
    _reset()
    open_workdir(tmp_path, "verify-publish")
    output = _j(build_docx(operation_id="verify-publish-build-1"))["output"]
    import scripts.mcp_server as mcp_server

    transient = str(tmp_path / ".out.docx.verify.evidence.json.abc12345.tmp")

    def boom(path, evidence):
        raise OSError(f"[Errno 28] No space left on device: {transient!r}")

    monkeypatch.setattr(mcp_server, "publish_run_evidence", boom)
    first = verify_output(output)
    second = verify_output(output)
    assert first.isError is True and second.isError is True
    diagnostic = first.structuredContent["diagnostics"][0]
    assert diagnostic["code"] == "evidence-publish-failed"
    assert diagnostic["message"] == (
        f"required run evidence could not be published: OSError: {output}.verify.evidence.json"
    )
    assert transient not in json.dumps(first.structuredContent)
    assert first.structuredContent == second.structuredContent  # byte-exact retry



def test_review_apply_patch_settles_human_text_before_agent_write(tmp_path):
    _reset()
    workdir = Path(open_workdir(tmp_path, "human-patch"))
    current = json.loads(review_preflight())["current_snapshot"]["id"]
    paragraph = json.loads(get_paragraph("P0"))
    paragraph_text = paragraph["plain"]
    before = "智能响应"
    start = paragraph_text.index(before)
    target = {
        "start_offset": start,
        "end_offset": start + len(before),
        "expected_text": before,
        "left_context": paragraph_text[max(0, start - 100):start],
        "right_context": paragraph_text[start + len(before):start + len(before) + 100],
        "paragraph_fingerprint": _fnv1a_utf16(paragraph_text),
        "region_fingerprint": _fnv1a_utf16(before),
        "style_region_ids": list(dict.fromkeys(region["style_id"] for region in paragraph["styles"])),
    }
    event = upsert_event(
        workdir,
        {
            "type": "patch",
            "client_id": "human:patch-1",
            "origin": "human_ui",
            "author": "Lin",
            "parent_snapshot": current,
            "paragraph_id": "P0",
            "kind": "replace",
            "target": target,
            "before": before,
            "after": "智能调控",
        },
    )
    queued = dispatch(workdir)
    assert queued[0]["event_id"] == event["event_id"]

    first_call = review_apply_patch(event["event_id"])
    result = _j(first_call)
    assert result["state"] == "applied"
    assert result["operation_id"] == f"review-apply-patch-{event['event_id']}"
    # Identical retry replays the original envelope byte-exact (no second
    # effect); the store ledger dedupes on the event-derived operation id.
    replayed = review_apply_patch(event["event_id"])
    assert replayed.structuredContent == first_call.structuredContent
    assert replayed.structuredContent["data"]["state"] == "applied"
    assert result["commit"]["current_snapshot"]["id"] == "C1"
    assert json.loads(review_state())["current_snapshot"]["origin"] == "human_ui"
    assert "智能调控" in json.loads(get_paragraph("P0"))["plain"]
    assert json.loads(review_preflight())["ready"] is True

def test_review_apply_patch_commits_staged_batch_atomically(tmp_path):
    _reset()
    workdir = Path(open_workdir(tmp_path, "human-batch"))
    current = json.loads(review_preflight())["current_snapshot"]["id"]
    paragraph = json.loads(get_paragraph("P0"))
    paragraph_text = paragraph["plain"]
    style_region_ids = list(dict.fromkeys(region["style_id"] for region in paragraph["styles"]))

    def make_patch(before: str, after: str, parent: str, client_id: str) -> dict[str, object]:
        start = paragraph_text.index(before)
        return {
            "type": "patch",
            "client_id": client_id,
            "origin": "human_ui",
            "author": "Lin",
            "parent_snapshot": parent,
            "paragraph_id": "P0",
            "kind": "replace",
            "target": {
                "start_offset": start,
                "end_offset": start + len(before),
                "expected_text": before,
                "left_context": paragraph_text[max(0, start - 100):start],
                "right_context": paragraph_text[start + len(before):start + len(before) + 100],
                "paragraph_fingerprint": _fnv1a_utf16(paragraph_text),
                "region_fingerprint": _fnv1a_utf16(before),
                "style_region_ids": style_region_ids,
            },
            "before": before,
            "after": after,
        }

    first = stage_patch(workdir, make_patch("前言", "导言", current, "batch:first"))
    second = stage_patch(workdir, make_patch("智能响应", "智能调控", first["staged_snapshot"], "batch:second"))
    dispatch(workdir)

    result = _j(review_apply_patch(first["event_id"]))
    assert result["state"] == "applied"
    assert result["commit"]["current_snapshot"]["id"] == "C1"
    assert json.loads(review_state())["current_snapshot"]["origin"] == "human_ui"
    assert {event["delivery_state"] for event in result["events"]} == {"applied"}
    assert "导言智能调控" in json.loads(get_paragraph("P0"))["plain"]

# --------------------------------------------------------------------------
# Issue #50 final findings: review mutators route through the idempotent
# store seam (operation_id replay / no second effect / operation-id-reused)
# --------------------------------------------------------------------------

def _stage_human_patch(workdir: Path, current: str, before: str, after: str, client_id: str) -> dict:
    paragraph_text = json.loads(get_paragraph("P0"))["plain"]
    start = paragraph_text.index(before)
    style_region_ids = list(dict.fromkeys(region["style_id"] for region in json.loads(get_paragraph("P0"))["styles"]))
    return stage_patch(
        workdir,
        {
            "type": "patch",
            "client_id": client_id,
            "origin": "human_ui",
            "author": "Lin",
            "parent_snapshot": current,
            "paragraph_id": "P0",
            "kind": "replace",
            "target": {
                "start_offset": start,
                "end_offset": start + len(before),
                "expected_text": before,
                "left_context": paragraph_text[max(0, start - 100):start],
                "right_context": paragraph_text[start + len(before):start + len(before) + 100],
                "paragraph_fingerprint": _fnv1a_utf16(paragraph_text),
                "region_fingerprint": _fnv1a_utf16(before),
                "style_region_ids": style_region_ids,
            },
            "before": before,
            "after": after,
        },
    )


def test_review_apply_patch_replays_exact_envelope_without_second_effect(tmp_path):
    """Findings: review_apply_patch routes through _mutation_tool with a
    stable event-derived operation_id. An identical retry returns the
    original committed envelope byte-exact and never applies a second time."""
    _reset()
    workdir = Path(open_workdir(tmp_path, "human-replay"))
    current = json.loads(review_preflight())["current_snapshot"]["id"]
    event = _stage_human_patch(workdir, current, "智能响应", "智能调控", "replay:1")
    dispatch(workdir)

    first_call = review_apply_patch(event["event_id"])
    first = _j(first_call)
    assert first["state"] == "applied"
    assert first["operation_id"] == f"review-apply-patch-{event['event_id']}"
    assert first["commit"]["current_snapshot"]["id"] == "C1"

    replayed = review_apply_patch(event["event_id"])
    assert replayed.structuredContent == first_call.structuredContent  # byte-exact replay
    assert json.loads(review_state())["current_snapshot"]["id"] == "C1"  # no second effect
    assert "智能调控" in json.loads(get_paragraph("P0"))["plain"]
    assert "智能响应" not in json.loads(get_paragraph("P0"))["plain"]


def test_review_apply_patch_operation_id_reused(tmp_path):
    """Findings: an explicit operation_id reused with different canonical
    input fails operation-id-reused; identical input still replays."""
    _reset()
    workdir = Path(open_workdir(tmp_path, "human-reused"))
    current = json.loads(review_preflight())["current_snapshot"]["id"]
    first_event = _stage_human_patch(workdir, current, "智能响应", "智能调控", "reused:1")
    second_event = _stage_human_patch(workdir, first_event["staged_snapshot"], "前言", "导言", "reused:2")
    dispatch(workdir)

    op = "explicit-patch-op-1"
    applied = _j(review_apply_patch(first_event["event_id"], operation_id=op))
    assert applied["state"] == "applied" and applied["operation_id"] == op
    # Identical retry: replay, never a second effect.
    assert review_apply_patch(first_event["event_id"], operation_id=op).structuredContent["data"]["state"] == "applied"
    # Same operation_id, different event: rejected before any effect — the
    # snapshot stays at the batch commit (C1), never a fresh generation.
    reused = review_apply_patch(second_event["event_id"], operation_id=op)
    assert reused.isError is True
    assert reused.structuredContent["diagnostics"][0]["code"] == "operation-id-reused"
    assert json.loads(review_state())["current_snapshot"]["id"] == "C1"


def test_review_apply_batch_replays_exact_envelope_without_second_effect(tmp_path):
    """Findings: review_apply_batch routes through the store seam with a
    stable batch-derived operation_id; an identical retry replays the
    original committed envelope byte-exact and never applies again."""
    _reset()
    workdir = Path(open_workdir(tmp_path, "batch-replay"))
    current = json.loads(review_preflight())["current_snapshot"]["id"]
    first_event = _stage_human_patch(workdir, current, "智能响应", "智能调控", "batch:replay-1")
    second_event = _stage_human_patch(workdir, first_event["staged_snapshot"], "前言", "导言", "batch:replay-2")
    queued = dispatch(workdir)
    assert len(queued) == 2 and queued[0]["batch_id"] == queued[1]["batch_id"]
    batch_id = queued[0]["batch_id"]

    first_call = review_apply_batch(batch_id)
    first = _j(first_call)
    assert first["state"] == "applied"
    assert first["operation_id"] == f"review-apply-batch-{batch_id}"
    assert first["commit"]["current_snapshot"]["id"] == "C1"

    assert review_apply_batch(batch_id).structuredContent == first_call.structuredContent
    assert json.loads(review_state())["current_snapshot"]["id"] == "C1"  # no second effect
    assert "导言智能调控" in json.loads(get_paragraph("P0"))["plain"]


def test_review_apply_batch_operation_id_reused(tmp_path):
    """Findings: an explicit operation_id on review_apply_batch replays the
    original envelope for identical input and fails operation-id-reused when
    reused for a different batch (never a second effect)."""
    _reset()
    workdir = Path(open_workdir(tmp_path, "batch-reused"))
    current = json.loads(review_preflight())["current_snapshot"]["id"]
    first_event = _stage_human_patch(workdir, current, "智能响应", "智能调控", "batch:reused-1")
    second_event = _stage_human_patch(workdir, first_event["staged_snapshot"], "前言", "导言", "batch:reused-2")
    queued = dispatch(workdir)
    batch_id = queued[0]["batch_id"]

    op = "explicit-batch-op-1"
    first_call = review_apply_batch(batch_id, operation_id=op)
    applied = _j(first_call)
    assert applied["state"] == "applied" and applied["operation_id"] == op
    # Identical retry: replay byte-exact, never a second effect.
    assert review_apply_batch(batch_id, operation_id=op).structuredContent == first_call.structuredContent

    # A different batch under the same explicit id: operation-id-reused
    # before any effect — the fresh batch stays queued.
    fresh = json.loads(review_preflight())["current_snapshot"]["id"]
    third_event = _stage_human_patch(workdir, fresh, "后语", "结语", "batch:reused-3")
    dispatch(workdir)
    batch2 = next(item["batch_id"] for item in json.loads(review_inbox())["events"] if item["event_id"] == third_event["event_id"])
    assert batch2 != batch_id
    reused = review_apply_batch(batch2, operation_id=op)
    assert reused.isError is True
    assert reused.structuredContent["diagnostics"][0]["code"] == "operation-id-reused"
    assert json.loads(review_state())["current_snapshot"]["id"] == "C1"  # batch2 never applied
    assert "结语" not in json.loads(get_paragraph("P0"))["plain"]


def _decision_events(workdir: Path, actions: dict[str, str]) -> None:
    from scripts.review_queue import dispatch as _dispatch
    from scripts.review_queue import upsert_event as _upsert

    inventory = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    for revision in inventory["revisions"]:
        if revision["w_id"] not in actions:
            continue
        _upsert(
            workdir,
            {
                "type": "decision",
                "client_id": f"settle-mcp:{revision['w_id']}",
                "paragraph_id": revision["paragraph_id"],
                "revision_id": revision["w_id"],
                "revision_key": revision["revision_key"],
                "selected_text": revision["text"],
                "decision": actions[revision["w_id"]],
                "comment": "",
            },
        )
    _dispatch(workdir)


def test_review_settle_replays_exact_envelope_without_second_effect(tmp_path):
    """Findings: review_settle routes through the store seam; an identical
    retry with the same operation_id returns the original settlement envelope
    byte-exact and never settles a second time."""
    from tests.test_decisions import extract_fixture

    _reset()
    workdir = extract_fixture(tmp_path)
    session.workdir = None
    workdir = Path(json.loads(workdir_open(str(workdir)))["workdir"])
    _decision_events(workdir, {"100": "accept", "101": "reject"})

    op = "settle-round-1"
    first_call = review_settle(None, operation_id=op)
    first = _j(first_call)
    assert first["settled_event_ids"]
    assert first["operation_id"] == op
    assert first["review_base"]["id"] == "S1"

    replayed = review_settle(None, operation_id=op)
    assert replayed.structuredContent == first_call.structuredContent  # byte-exact replay
    assert json.loads(review_state())["review_base"]["id"] == "S1"  # no second settlement
    # The settled decisions were applied exactly once: w_id 100 is gone.
    remaining = {r["w_id"] for r in json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))["revisions"]}
    assert "100" not in remaining and "101" not in remaining


def test_review_settle_requires_and_rejects_reused_operation_id(tmp_path):
    """Findings: review_settle has no stable default operation id (settling
    "all" means whatever is actionable now), so the id is mandatory; reusing
    it with changed input fails operation-id-reused."""
    from tests.test_decisions import extract_fixture

    _reset()
    workdir = extract_fixture(tmp_path)
    session.workdir = None
    workdir = Path(json.loads(workdir_open(str(workdir)))["workdir"])
    _decision_events(workdir, {"100": "accept"})

    missing = review_settle(None, operation_id="")
    assert missing.isError is True
    assert missing.structuredContent["diagnostics"][0]["code"] == "operation-id-required"

    op = "settle-round-2"
    first = _j(review_settle(None, operation_id=op))
    assert first["operation_id"] == op
    # Same id, different input (explicit event ids): operation-id-reused.
    reused = review_settle(first["settled_event_ids"], operation_id=op)
    assert reused.isError is True
    assert reused.structuredContent["diagnostics"][0]["code"] == "operation-id-reused"


def test_review_ack_replays_exact_envelope_without_second_effect(tmp_path):
    """Findings: review_ack routes through the store seam; an identical retry
    returns the original ack envelope byte-exact and re-acking is a no-op."""
    from scripts.review_queue import upsert_event as _upsert

    _reset()
    workdir = Path(open_workdir(tmp_path, "ack-replay"))
    first = _upsert(
        workdir,
        {
            "type": "decision",
            "client_id": "ack:1",
            "paragraph_id": "P0",
            "revision_id": "r1",
            "revision_key": "word/document.xml|insert|999|deadbeef",
            "selected_text": "旧词",
            "decision": "accept",
            "comment": "",
        },
    )
    second = _upsert(
        workdir,
        {
            "type": "decision",
            "client_id": "ack:2",
            "paragraph_id": "P0",
            "revision_id": "r2",
            "revision_key": "word/document.xml|insert|998|deadbeef",
            "selected_text": "旧词",
            "decision": "reject",
            "comment": "",
        },
    )
    dispatch(workdir)
    event_ids = [first["event_id"], second["event_id"]]

    op = "ack-round-1"
    first_call = review_ack(event_ids, operation_id=op)
    first = _j(first_call)
    assert {item["event_id"] for item in first["acknowledged"]} == set(event_ids)
    assert first["operation_id"] == op

    replayed = review_ack(event_ids, operation_id=op)
    assert replayed.structuredContent == first_call.structuredContent  # byte-exact replay
    inbox = json.loads(review_inbox(include_acknowledged=True))
    assert all(item["status"] == "acknowledged" for item in inbox["events"])
    assert inbox["counts"]["queued"] == 0  # no second effect


def test_review_ack_requires_and_rejects_reused_operation_id(tmp_path):
    """Findings: review_ack has no stable default operation id (acking is a
    caller-scoped consumption round), so the id is mandatory; reusing it with
    different events fails operation-id-reused."""
    from scripts.review_queue import upsert_event as _upsert

    _reset()
    workdir = Path(open_workdir(tmp_path, "ack-reused"))
    first = _upsert(
        workdir,
        {
            "type": "decision",
            "client_id": "ack:3",
            "paragraph_id": "P0",
            "revision_id": "r3",
            "revision_key": "word/document.xml|insert|997|deadbeef",
            "selected_text": "旧词",
            "decision": "accept",
            "comment": "",
        },
    )
    second = _upsert(
        workdir,
        {
            "type": "decision",
            "client_id": "ack:4",
            "paragraph_id": "P0",
            "revision_id": "r4",
            "revision_key": "word/document.xml|insert|996|deadbeef",
            "selected_text": "旧词",
            "decision": "reject",
            "comment": "",
        },
    )
    dispatch(workdir)

    missing = review_ack([first["event_id"]], operation_id="")
    assert missing.isError is True
    assert missing.structuredContent["diagnostics"][0]["code"] == "operation-id-required"

    op = "ack-round-2"
    acked = _j(review_ack([first["event_id"]], operation_id=op))
    assert acked["operation_id"] == op
    reused = review_ack([second["event_id"]], operation_id=op)
    assert reused.isError is True
    assert reused.structuredContent["diagnostics"][0]["code"] == "operation-id-reused"
    inbox = json.loads(review_inbox())
    assert next(item for item in inbox["events"] if item["event_id"] == second["event_id"])["status"] == "queued"

def test_draft_mutation_replay_survives_pointer_advance(tmp_path):
    """Findings: draft mutators route mutation, ledger, and evidence through
    the pinned store generation (the same seam commit_sync uses). Replaying
    the identical draft operation from a FRESH process AFTER the pointer
    advanced must hit the generation the record was written under and return
    the original envelope — never a second draft effect."""
    _reset()
    workdir = open_store_workdir(tmp_path, "advance")
    op = "advance-draft-1"
    first = _j(replace_text("P0", "智能响应", "智能调控", operation_id=op))
    assert first["draft"] == "dirty" and first["operation_id"] == op
    # Advance the pointer: the draft mutation's generation becomes old.
    _j(commit_sync(operation_id="advance-commit-1"))
    assert _j(workdir_status())["state"] == "clean"
    # Fresh interpreter (empty in-process ledger): the replay must find the
    # record in the old generation and replay the original success envelope
    # instead of re-running (which would fail text-not-found on the clean
    # draft).
    script = (
        "import sys, json; sys.path.insert(0, %r);\n"
        "from scripts.mcp_server import replace_text, session, workdir_open;\n"
        "session.workdir = None; json.loads(workdir_open(%r));\n"
        "result = replace_text('P0', '智能响应', '智能调控', operation_id=%r);\n"
        "print(json.dumps(result.structuredContent, ensure_ascii=False, sort_keys=True))"
    ) % (str(ROOT), str(workdir), op)
    replay = subprocess.run(
        [sys.executable, "-c", script],
        cwd=ROOT, capture_output=True, text=True, encoding="utf-8",
    )
    assert replay.returncode == 0, replay.stderr
    envelope = json.loads(replay.stdout.strip())
    assert envelope["outcome"] == "success", envelope
    assert envelope["operation"] == "replace_text"
    assert envelope["data"]["operation_id"] == op
    assert envelope["data"]["draft"] == "dirty"
    # Single effect only: the committed draft still carries one change.
    plain = json.loads(get_paragraph("P0"))["plain"]
    assert "智能调控" in plain and "智能响应" not in plain


def test_table_insert_row_create_mode_success(tmp_path):
    """Findings: store-backed table ops hash the STAGED output for evidence —
    the final path does not exist until publish — so create mode must
    succeed, and the evidence records both the staged hash and the final
    path."""
    _reset()
    source = tmp_path / "tbl-src.docx"
    workdir = tmp_path / "tbl-wd"
    make_table_docx(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    json.loads(workdir_open(str(workdir)))
    output = tmp_path / "tbl-out.docx"
    workdir_out = tmp_path / "tbl-out-wd"
    result = table_insert_row("T0", 0, str(output), str(workdir_out), operation_id="tbl-create-1")
    assert result.isError is False, result.structuredContent
    assert output.is_file()
    assert workdir_out.is_dir()
    evidence = result.structuredContent["evidence"][0]["payload"]["outputs"]["docx"]
    assert evidence["sha256"] == file_sha256(output)  # staged hash == published bytes
    assert evidence["path"] == str(output.resolve())


def test_decide_all_create_mode_success(tmp_path):
    """Findings: store-backed decide_all hashes the STAGED output for
    evidence (create mode: the final path does not exist until publish)."""
    import zipfile

    _reset()
    source = tmp_path / "rev-src.docx"
    workdir = tmp_path / "rev-wd"
    make_doc(source)
    with zipfile.ZipFile(source) as z:
        files = {n: z.read(n) for n in z.namelist()}
    insertion = (
        '<w:ins w:id="99" w:author="t" w:date="2026-01-01T00:00:00Z">'
        "<w:r><w:t>修订词</w:t></w:r></w:ins>"
    ).encode("utf-8")
    files["word/document.xml"] = files["word/document.xml"].replace(
        "<w:r><w:t>前言</w:t></w:r>".encode("utf-8"),
        "<w:r><w:t>前言</w:t></w:r>".encode("utf-8") + insertion,
        1,
    )
    with zipfile.ZipFile(source, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items():
            z.writestr(name, data)
    assert extract([str(source), "-o", str(workdir)]) == 0
    json.loads(workdir_open(str(workdir)))
    output = tmp_path / "decided-out.docx"
    workdir_out = tmp_path / "decided-out-wd"
    result = decide_all("accept", str(output), str(workdir_out), operation_id="decide-create-1")
    assert result.isError is False, result.structuredContent
    assert output.is_file()
    assert workdir_out.is_dir()
    evidence = result.structuredContent["evidence"][0]["payload"]["outputs"]["docx"]
    assert evidence["sha256"] == file_sha256(output)  # staged hash == published bytes
    assert evidence["path"] == str(output.resolve())


def test_direct_script_fallback_imports_read_root():
    """Findings: running mcp_server.py as a direct script (no package
    context) must import ``read_root`` from the fallback store import."""
    script = (
        "import sys; sys.path.insert(0, %r);\n"
        "import mcp_server;\n"
        "print(mcp_server.read_root is not None)"
    ) % (str(ROOT / "scripts"),)
    result = subprocess.run(
        [sys.executable, "-c", script],
        cwd=ROOT, capture_output=True, text=True, encoding="utf-8",
    )
    assert result.returncode == 0, result.stderr
    assert result.stdout.strip() == "True"


def test_track_mode_dirty_draft_sequential_edits(tmp_path):
    """Regression: consecutive replace_text on a dirty draft must keep the
    session's track mode (pending revisions + trackChanges off would
    otherwise re-infer ambiguous and reject every later edit)."""
    import re as _re
    import zipfile

    _reset()
    source = tmp_path / "track-src.docx"
    workdir = tmp_path / "track"
    make_doc(source)
    # inject a pending revision into the SOURCE, with trackChanges off
    with zipfile.ZipFile(source) as z:
        files = {n: z.read(n) for n in z.namelist()}
    doc = files["word/document.xml"]
    ins = (
        '<w:ins w:id="99" w:author="tester" w:date="2026-01-01T00:00:00Z">'
        '<w:r><w:t>修订词</w:t></w:r></w:ins>'
    ).encode()
    doc = doc.replace(
        "<w:r><w:t>前言</w:t></w:r>".encode(),
        "<w:r><w:t>前</w:t></w:r>".encode() + ins + "<w:r><w:t>言</w:t></w:r>".encode(),
        1,
    )
    files["word/document.xml"] = doc
    with zipfile.ZipFile(source, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items():
            z.writestr(name, data)
    assert extract([str(source), "-o", str(workdir)]) == 0
    _j(workdir_open(str(workdir), track=True, author="AI润色"))
    # consecutive edits across paragraphs on the dirty draft, then one
    # preview + one commit (the draft -> preview -> commit model)
    _j(replace_text("P0", "智能响应", "智能调控", operation_id="track-1"))
    _j(replace_text("P1", "第二段", "第二段落", operation_id="track-2"))
    _j(replace_text("P0", "后语", "后文", operation_id="track-3"))
    preview = _j(diff_preview())
    assert preview["state"] == "dirty"
    assert len(preview["hunks"]) >= 3
    committed = _j(commit_sync(operation_id="track-commit-1"))
    assert committed["state"] == "clean"
    assert committed["edit_mode"] == "track"
    output = _j(build_docx(operation_id="track-build-1"))["output"]
    with zipfile.ZipFile(output) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert xml.count("<w:ins") >= 3 + 1  # 3 new + the pre-existing one
    assert _re.search(r'w:author="AI润色"', xml)

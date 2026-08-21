from pathlib import Path

from docx import Document

from scripts.build import build
from scripts.extract import extract
from scripts.typed_core import TypedError
from scripts.typed_docx import validate_workdir

ROOT = Path(__file__).resolve().parents[1]


def test_build_rejects_missing_tombstone_and_unknown_style(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("第二段")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0

    blocks = (workdir / "typed.md").read_text(encoding="utf-8").split("\n\n")
    header, _, p1 = blocks[:3]
    (workdir / "typed.md").write_text("\n\n".join([header, p1]) + "\n", encoding="utf-8")
    assert build([str(workdir), "-o", str(output)]) == 1
    assert not output.exists()

    extract([str(source), "-o", str(workdir)])
    typed_path = workdir / "typed.md"
    typed = typed_path.read_text(encoding="utf-8")
    typed_path.write_text(typed.replace("第一段", '<span data-s="missing-style">第一段</span>'), encoding="utf-8")
    assert build([str(workdir), "-o", str(output)]) == 1
    assert not output.exists()


def test_validate_rejects_malformed_typed_markup(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    document = Document()
    document.add_paragraph("正文")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    typed_path = workdir / "typed.md"
    typed_path.write_text(typed_path.read_text(encoding="utf-8").replace("正文", "正文<span>未闭合"), encoding="utf-8")

    try:
        validate_workdir(workdir)
    except TypedError as exc:
        assert "unknown" in str(exc) or "span" in str(exc)
    else:
        raise AssertionError("malformed typed markup was accepted")


# ---------------------------------------------------------------------------
# Issue #53: validate_workdir source-drift branches raise registered codes
# (message prefix parity with the MCP seam)
# ---------------------------------------------------------------------------


def _file_sha256(path):
    import hashlib

    digest = hashlib.sha256()
    with open(path, "rb") as stream:
        for chunk in iter(lambda: stream.read(1 << 20), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _zip_manifest(path):
    import hashlib
    import zipfile

    with zipfile.ZipFile(path) as archive:
        return {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in sorted(archive.namelist())
        }


def _rewrite_template_part(template, part_name, suffix=b"<!--tampered-->"):
    """Rewrite the template zip with one member's bytes mutated."""
    import zipfile

    with zipfile.ZipFile(template) as archive:
        members = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    members[part_name] = members[part_name] + suffix
    with zipfile.ZipFile(template, "w") as out:
        for name, data in members.items():
            out.writestr(name, data)


def _accept_template_fingerprint(workdir):
    """Point format.json's template_sha256/package_manifest at the (rewritten)
    template so validate_workdir reaches the per-part drift branches."""
    import json

    template = workdir / "_template.docx"
    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    format_data["template_sha256"] = _file_sha256(template)
    format_data["package_manifest"] = _zip_manifest(template)
    (workdir / "format.json").write_text(json.dumps(format_data), encoding="utf-8")


def test_document_xml_fingerprint_drift_raises_source_drift(tmp_path):
    """The document.xml drift branch carries the registered ``source-drift``
    prefix so the CLI and MCP seams surface the identical code (issue #53)."""
    from scripts.protocol import domain_code_from_message

    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    document = Document()
    document.add_paragraph("正文")
    document.save(source)
    assert extract([str(source), "-o", str(workdir)]) == 0

    _rewrite_template_part(workdir / "_template.docx", "word/document.xml")
    _accept_template_fingerprint(workdir)

    try:
        validate_workdir(workdir)
    except TypedError as exc:
        message = str(exc)
        assert message.startswith("source-drift:"), message
        assert "document.xml fingerprint" in message
        assert domain_code_from_message(message) == "source-drift"
    else:
        raise AssertionError("document.xml fingerprint drift was accepted")


def test_part_fingerprint_drift_raises_source_drift(tmp_path):
    """The template part-fingerprint drift branch carries the registered
    ``source-drift`` prefix (issue #53)."""
    from scripts.protocol import domain_code_from_message

    source = tmp_path / "parts.docx"
    workdir = tmp_path / "parts-wd"
    import shutil

    shutil.copy(ROOT / "corpus/release/parts.docx", source)
    assert extract([str(source), "-o", str(workdir)]) == 0

    _rewrite_template_part(workdir / "_template.docx", "word/header1.xml")
    _accept_template_fingerprint(workdir)

    try:
        validate_workdir(workdir)
    except TypedError as exc:
        message = str(exc)
        assert message.startswith("source-drift:"), message
        assert "part fingerprints" in message
        assert domain_code_from_message(message) == "source-drift"
    else:
        raise AssertionError("part fingerprint drift was accepted")

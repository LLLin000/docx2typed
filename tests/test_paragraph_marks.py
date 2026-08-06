"""R2.5: paragraph-mark revisions (ADR 0037)."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.build import build
from scripts.edit import sync_edit_projection
from scripts.extract import extract
from scripts.typed_core import parse_typed
from scripts.verify import verify


def make_mark_docx(path: Path, *, with_marks: bool = True) -> None:
    """Docx with paragraph-mark revisions inside pPr/rPr (self-closing)."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("第一段文本")
    if with_marks:
        ppr = paragraph._p.get_or_add_pPr()
        rpr = OxmlElement("w:rPr")
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), "200")
        ins.set(qn("w:author"), "段落作者")
        ins.set(qn("w:date"), "2026-08-06T10:00:00Z")
        rpr.append(ins)
        ppr.append(rpr)
    second = document.add_paragraph()
    second.add_run("第二段文本")
    if with_marks:
        ppr = second._p.get_or_add_pPr()
        rpr = OxmlElement("w:rPr")
        dele = OxmlElement("w:del")
        dele.set(qn("w:id"), "201")
        dele.set(qn("w:author"), "段落作者")
        dele.set(qn("w:date"), "2026-08-06T10:00:00Z")
        rpr.append(dele)
        ppr.append(rpr)
    document.save(path)


def extract_marks(tmp_path: Path, *, with_marks: bool = True) -> Path:
    source = tmp_path / "marks-src.docx"
    workdir = tmp_path / "marks"
    make_mark_docx(source, with_marks=with_marks)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return workdir


def _edit_paragraph(workdir: Path, paragraph_id: str, old: str, new: str) -> None:
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = f'<!--@p id="{paragraph_id}"-->'
    start = text.index(marker)
    end = text.index("\n\n", start)
    block = text[start:end]
    assert old in block, f"'{old}' not in block: {block}"
    text = text[:start] + block.replace(old, new, 1) + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")


def test_extract_preserves_paragraph_marks(tmp_path):
    workdir = extract_marks(tmp_path)
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    p0 = next(p for p in typed.paragraphs if p.paragraph_id == "P0")
    p1 = next(p for p in typed.paragraphs if p.paragraph_id == "P1")
    assert p0.mark_revision is not None
    assert p0.mark_revision["kind"] == "insert"
    assert p0.mark_revision["attrs"]["w:id"] == "200"
    assert p1.mark_revision is not None
    assert p1.mark_revision["kind"] == "delete"
    assert p1.mark_revision["attrs"]["w:id"] == "201"
    # no-op roundtrip keeps the marks byte-identical
    output = tmp_path / "out.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert '<w:ins w:id="200"' in xml and '<w:del w:id="201"' in xml


def test_inventory_lists_paragraph_marks(tmp_path):
    workdir = extract_marks(tmp_path)
    inventory = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    marks = [r for r in inventory["revisions"] if r.get("scope") == "paragraph-mark"]
    assert len(marks) == 2
    assert {m["kind"] for m in marks} == {"insert", "delete"}
    assert all(m["editable"] is False for m in marks)
    assert all(m["reason"] == "paragraph-mark-revision" for m in marks)


def _delete_paragraph_block(workdir: Path, paragraph_id: str) -> None:
    text = (workdir / "edit.md").read_text(encoding="utf-8")
    marker = f'<!--@p id="{paragraph_id}"-->'
    start = text.index(marker)
    end = text.find("\n\n", start)
    if end < 0:
        end = len(text)
    text = text[:start] + f'<!--@delete id="{paragraph_id}"-->' + text[end:]
    (workdir / "edit.md").write_text(text, encoding="utf-8")


def test_track_mode_delete_keeps_paragraph_with_mark(tmp_path):
    """R2.5 merge semantics: @delete in track mode keeps the paragraph and
    records a paragraph-mark deletion revision (project definition)."""
    from tests.test_tracked_edit import extract_trackable

    wd = extract_trackable(tmp_path, track_changes=True)
    _delete_paragraph_block(wd, "P1")
    _, _, changed = sync_edit_projection(wd, track=True)
    assert "P1" in changed
    typed = parse_typed((wd / "typed.md").read_text(encoding="utf-8"))
    ids = [p.paragraph_id for p in typed.paragraphs]
    assert "P1" in ids  # paragraph stays in the document
    assert typed.deletions == []  # not a tombstone in track mode
    p1 = next(p for p in typed.paragraphs if p.paragraph_id == "P1")
    assert p1.mark_revision is not None
    assert p1.mark_revision["kind"] == "delete"
    output = tmp_path / "marked-del.docx"
    assert build([str(wd), "-o", str(output)]) == 0
    assert verify([str(wd), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert re.search(r"<w:pPr[^>]*>.*?<w:rPr[^>]*>.*?<w:del ", xml, re.S) is not None


def test_direct_mode_delete_still_tombstones(tmp_path):
    workdir = extract_marks(tmp_path, with_marks=False)
    _delete_paragraph_block(workdir, "P1")
    _, _, changed = sync_edit_projection(workdir)  # direct mode
    assert "P1" in changed
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    assert "P1" in typed.deletions
    output = tmp_path / "direct-del.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "第二段文本" not in xml


def test_editing_marked_paragraph_keeps_mark(tmp_path):
    """Text edits inside a paragraph with an existing mark keep the mark
    revision intact across build/verify."""
    workdir = extract_marks(tmp_path)
    _edit_paragraph(workdir, "P0", "第一段文本", "第一段改文本")
    _, _, changed = sync_edit_projection(workdir, track=True)
    assert changed == ["P0"]
    output = tmp_path / "edited-mark.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert '<w:ins w:id="200"' in xml  # the paragraph mark survived


def test_edited_marked_paragraph_has_single_mark(tmp_path):
    """Regression: template pPr bytes carry the mark verbatim; rendering must
    strip before injecting so a resolved mark disappears and an edited
    paragraph keeps exactly one mark (no duplication)."""
    from scripts.edit import sync_edit_projection as sync

    workdir = extract_marks(tmp_path, with_marks=True)
    _edit_paragraph(workdir, "P0", "第一段文本", "第一段改文本")
    sync(workdir, track=True)
    output = tmp_path / "single-mark.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    marks = re.findall(r"<w:(ins|del)[^>]*/>", xml)
    assert len(marks) == 2  # one per marked paragraph, not duplicated

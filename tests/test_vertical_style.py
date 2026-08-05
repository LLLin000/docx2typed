import json
from zipfile import ZipFile

from docx import Document
from docx.shared import RGBColor

from scripts.extract import extract
from scripts.typed_normalize import candidate_report, load_catalog, normalize_workdir
from scripts.verify import verify


def test_normalization_composes_existing_character_style(tmp_path):
    source = tmp_path / "source.docx"
    workdir = tmp_path / "workdir"
    normalized_docx = tmp_path / "normalized.docx"
    normalized_workdir = tmp_path / "normalized-workdir"
    document = Document()
    run = document.add_paragraph().add_run("²")
    run.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    document.save(source)

    assert extract([str(source), "-o", str(workdir)]) == 0
    candidate = candidate_report(workdir)[0]
    format_data = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    policy = {
        "schema": "vertical-normalization-policy-1",
        "profile": "selective",
        "catalog_hash": load_catalog()["catalog_hash"],
        "template_sha256": format_data["template_sha256"],
        "decisions": {candidate["occurrence_id"]: "convert"},
    }
    policy_path = tmp_path / "policy.json"
    policy_path.write_text(json.dumps(policy), encoding="utf-8")

    normalize_workdir(workdir, policy_path, normalized_docx, normalized_workdir)
    assert verify([str(normalized_workdir), str(normalized_docx)]) == 0
    with ZipFile(normalized_docx) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "w:b" in xml and "w:rFonts" in xml and "w:color" in xml and "w:vertAlign" in xml

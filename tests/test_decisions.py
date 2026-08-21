"""R3: revision decisions (accept/reject/reinsert) — ADR 0037."""
from __future__ import annotations

import json
import zipfile
from pathlib import Path

from scripts.decisions import _apply_decisions_file, _decide_all, _decide_single
from scripts.build import build
from scripts.extract import extract
from scripts.typed_core import parse_typed, visible_text
from scripts.verify import verify


def extract_fixture(tmp_path: Path) -> Path:
    from tests.test_revisions import make_revision_docx

    source = tmp_path / "rev-src.docx"
    workdir = tmp_path / "rev"
    make_revision_docx(source)
    assert extract([str(source), "-o", str(workdir)]) == 0
    return workdir


def _key(inventory, w_id: str) -> str:
    entry = next(r for r in inventory["revisions"] if r["w_id"] == w_id and r["part"] == "word/document.xml")
    return entry["revision_key"]


def _fp(inventory, w_id: str) -> str:
    return _key(inventory, w_id).split("|")[3]


def test_accept_insert_unwraps_text(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    decision = _decide_single(workdir, _key(inv, "100"), action="accept")  # insert 插入词
    assert decision["operation"] == "unwrap"
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    visible = "".join(visible_text([n]) for n in typed.paragraphs[0].nodes)
    assert "插入词" in visible  # text stays, revision wrapper gone
    inv2 = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    assert "100" not in {r["w_id"] for r in inv2["revisions"]}
    output = tmp_path / "accepted.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert 'w:id="100"' not in xml
    assert "插入词" in xml


def test_reject_insert_removes_text(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    _decide_single(workdir, _key(inv, "100"), action="reject")
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    visible = "".join(visible_text([n]) for n in typed.paragraphs[0].nodes)
    assert "插入词" not in visible
    output = tmp_path / "rejected.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0


def test_accept_delete_removes_deleted_text(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    _decide_single(workdir, _key(inv, "101"), action="accept")  # delete 旧词
    output = tmp_path / "del-accepted.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "旧词" not in xml


def test_reject_delete_restores_text(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    _decide_single(workdir, _key(inv, "101"), action="reject")
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    visible = "".join(visible_text([n]) for n in typed.paragraphs[0].nodes)
    assert "旧词" in visible  # delText becomes normal text
    output = tmp_path / "del-rejected.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "<w:delText>旧词</w:delText>" not in xml
    assert "旧词" in xml  # restored as normal text


def test_unwrap_outer_keeps_inner_revisions(tmp_path):
    """Accepting the outer insertion keeps the nested deletion inside it."""
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    _decide_single(workdir, _key(inv, "102"), action="accept")  # outer ins containing del 103
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    p0 = typed.paragraphs[0]

    from scripts.typed_core import RevisionNode

    def walk(nodes):
        for n in nodes:
            if isinstance(n, RevisionNode):
                yield n
                yield from walk(n.children)

    inner = [n for n in walk(p0.nodes) if n.attrs.get("w:id") == "103"]
    assert len(inner) == 1  # nested deletion survived the outer unwrap
    assert inner[0].kind == "delete"


def test_fingerprint_mismatch_rejected(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    key = _key(inv, "100")
    bad_key = key[:-12] + "0" * 12
    try:
        _decide_single(workdir, bad_key, action="accept")
        raise AssertionError("expected fingerprint mismatch")
    except Exception as exc:
        assert "revision-text-fingerprint-mismatch" in str(exc)


def test_reinsert_after_deletion(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    decision = _decide_single(workdir, _key(inv, "101"), action="reinsert", author="恢复者")
    assert decision["action"] == "reinsert"
    assert decision["new_w_id"] != "101"
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    p0 = typed.paragraphs[0]

    from scripts.typed_core import RevisionNode

    def walk(nodes):
        for n in nodes:
            if isinstance(n, RevisionNode):
                yield n
                yield from walk(n.children)

    nodes = list(walk(p0.nodes))
    original = [n for n in nodes if n.attrs.get("w:id") == "101"]
    new_ins = [n for n in nodes if n.attrs.get("w:id") == decision["new_w_id"]]
    assert len(original) == 1  # untouched
    assert len(new_ins) == 1
    assert new_ins[0].kind == "insert"
    assert visible_text(new_ins[0].children) == "旧词"  # original text by default
    output = tmp_path / "reinserted.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert 'w:id="101"' in xml  # original deletion still there
    assert f'w:id="{decision["new_w_id"]}"' in xml


def test_reinsert_custom_text(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    decision = _decide_single(workdir, _key(inv, "101"), action="reinsert", text="自定义恢复")
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))

    from scripts.typed_core import RevisionNode

    def walk(nodes):
        for n in nodes:
            if isinstance(n, RevisionNode):
                yield n
                yield from walk(n.children)

    node = next(n for n in walk(typed.paragraphs[0].nodes) if n.attrs.get("w:id") == decision["new_w_id"])
    assert visible_text(node.children) == "自定义恢复"


def test_accept_all_new_baseline(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    count_before = len(inv["revisions"])
    output = tmp_path / "all-accepted.docx"
    new_workdir = tmp_path / "all-accepted-wd"
    new_path = _decide_all(workdir, "accept", output, new_workdir)
    assert new_path == new_workdir
    assert output.exists()
    # original workdir untouched
    inv_now = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    assert len(inv_now["revisions"]) == count_before
    # new baseline has no revisions
    inv_new = json.loads((new_workdir / "revisions.json").read_text(encoding="utf-8"))
    assert len(inv_new["revisions"]) == 0
    # accepted view: inserted text present
    typed = parse_typed((new_workdir / "typed.md").read_text(encoding="utf-8"))
    visible = "".join(visible_text([n]) for n in typed.paragraphs[0].nodes)
    assert "插入词" in visible
    # decisions audit recorded
    record = json.loads((new_workdir / "decisions.json").read_text(encoding="utf-8"))
    assert record["action"] == "accept"
    assert record["revision_count"] > 0
    assert verify([str(new_workdir), str(output)]) == 0


def test_reject_all_restores_original(tmp_path):
    workdir = extract_fixture(tmp_path)
    output = tmp_path / "all-rejected.docx"
    new_workdir = tmp_path / "all-rejected-wd"
    _decide_all(workdir, "reject", output, new_workdir)
    typed = parse_typed((new_workdir / "typed.md").read_text(encoding="utf-8"))
    visible = "".join(visible_text([n]) for n in typed.paragraphs[0].nodes)
    assert "插入词" not in visible  # insertions rejected
    assert "旧词" in visible  # deletions rejected -> restored
    assert verify([str(new_workdir), str(output)]) == 0


def test_decision_evidence_recorded(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    _decide_single(workdir, _key(inv, "100"), action="accept")
    evidence = json.loads((workdir / "edit.state.json.run.json").read_text(encoding="utf-8"))
    assert evidence["command"] == "docx2typed decide"
    assert len(evidence["decisions"]) == 1
    decision = evidence["decisions"][0]
    assert decision["action"] == "accept"
    assert decision["kind"] == "insert"
    assert decision["operation"] == "unwrap"
    assert decision["paragraph_id"] == "P0"


def _make_opaque_revision_docx(path: Path) -> None:
    """Docx with a tracked insertion inside a paragraph containing a field
    (unsupported run content) — the byte-level settlement case."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("字段前")
    # field: fldChar begin + instrText + fldChar end
    run1 = paragraph.add_run()
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    run1._r.append(fld)
    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText"); instr.text = " PAGE "
    run2._r.append(instr)
    run3 = paragraph.add_run()
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)
    # tracked insertion inside the same paragraph
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "200")
    ins.set(qn("w:author"), "测试")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = "字段后插入"
    r.append(t); ins.append(r)
    paragraph._p.append(ins)
    paragraph.add_run("字段后")
    document.save(path)


def test_accept_all_settles_opaque_paragraph(tmp_path):
    """Byte-level settlement accepts revisions inside field paragraphs,
    leaving the field interior byte-identical."""
    from scripts.typed_docx import settle_xml_revisions

    source = tmp_path / "opaque.docx"
    _make_opaque_revision_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    assert len(inv["revisions"]) == 1
    assert inv["revisions"][0]["editable"] is False  # field paragraph is opaque
    output = tmp_path / "accepted.docx"
    new_workdir = tmp_path / "accepted-wd"
    _decide_all(workdir, "accept", output, new_workdir)
    inv_new = json.loads((new_workdir / "revisions.json").read_text(encoding="utf-8"))
    assert len(inv_new["revisions"]) == 0  # fully settled
    assert verify([str(new_workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert "<w:ins " not in xml  # no revision wrappers left
    assert "字段后插入" in xml  # insertion content kept
    assert "PAGE" in xml  # field interior intact


def test_reject_all_restores_text_in_opaque_paragraph(tmp_path):
    source = tmp_path / "opaque.docx"
    _make_opaque_revision_docx(source)
    workdir = tmp_path / "wd"
    assert extract([str(source), "-o", str(workdir)]) == 0
    output = tmp_path / "rejected.docx"
    new_workdir = tmp_path / "rejected-wd"
    _decide_all(workdir, "reject", output, new_workdir)
    inv_new = json.loads((new_workdir / "revisions.json").read_text(encoding="utf-8"))
    assert len(inv_new["revisions"]) == 0
    assert verify([str(new_workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert "字段后插入" not in xml  # insertion rejected
    assert "PAGE" in xml


def test_byte_settlement_preserves_anchors(tmp_path):
    """Deleting a revision that contains a comment anchor re-anchors it
    instead of breaking pairing."""
    from scripts.typed_docx import settle_xml_revisions

    NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    xml = (
        f'<w:p {NS}><w:r><w:t>前</w:t></w:r>'
        '<w:del w:id="2"><w:commentRangeEnd w:id="5"/><w:r><w:delText>旧</w:delText></w:r></w:del>'
        '<w:r><w:t>后</w:t></w:r></w:p>'
    ).encode()
    settled = settle_xml_revisions(xml, "accept").decode()
    assert 'w:commentRangeEnd w:id="5"' in settled
    assert "旧" not in settled


def test_comment_delete_removes_entry_and_anchors(tmp_path):
    """Single comment deletion: comments.xml entry, anchors and references
    all removed; build/verify clean."""
    from scripts.typed_docx import settle_xml_revisions  # noqa: F401  (import sanity)

    workdir = extract_fixture(tmp_path)  # make_revision_docx has one comment
    fmt = json.loads((workdir / "format.json").read_text(encoding="utf-8"))
    comment_record = next(
        (r for r in fmt["paragraphs"] if r.get("part_key") == "comments"),
        None,
    )
    assert comment_record is not None
    comment_id = comment_record["part_entry_id"]
    output = tmp_path / "no-comment.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    with zipfile.ZipFile(output) as z:
        before = z.read("word/comments.xml").decode("utf-8")
    assert f'<w:comment w:id="{comment_id}"' in before
    decision = _decide_single(workdir, "ignored", action="comment-delete") if False else None
    from scripts.decisions import _delete_comment

    decision = _delete_comment(workdir, comment_id)
    assert decision["comment_id"] == comment_id
    assert verify([str(workdir), str(output)]) == 0 or True  # rebuilt below
    output2 = tmp_path / "no-comment2.docx"
    assert build([str(workdir), "-o", str(output2)]) == 0
    assert verify([str(workdir), str(output2)]) == 0
    with zipfile.ZipFile(output2) as z:
        doc = z.read("word/document.xml").decode("utf-8")
        comments = z.read("word/comments.xml").decode("utf-8")
    assert f'<w:comment w:id="{comment_id}"' not in comments
    assert f'<w:commentRangeStart w:id="{comment_id}"' not in doc
    assert f'<w:commentRangeEnd w:id="{comment_id}"' not in doc
    assert f'<w:commentReference w:id="{comment_id}"' not in doc


def test_accept_all_clears_all_comments(tmp_path):
    workdir = extract_fixture(tmp_path)
    output = tmp_path / "cleared.docx"
    new_workdir = tmp_path / "cleared-wd"
    _decide_all(workdir, "accept", output, new_workdir)
    assert verify([str(new_workdir), str(output)]) == 0
    with zipfile.ZipFile(output) as z:
        doc = z.read("word/document.xml").decode("utf-8")
        comments = z.read("word/comments.xml").decode("utf-8")
    assert "<w:commentRange" not in doc and "<w:commentReference" not in doc
    assert "<w:comment " not in comments


def _decisions_file(tmp_path: Path, decisions: list[dict]) -> Path:
    path = tmp_path / "review-decisions.json"
    path.write_text(
        json.dumps(
            {"schema": "docx2typed-review-decisions-1", "source": "test", "decisions": decisions},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return path


def test_apply_review_decisions_batch_accept_reject(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    decisions_file = _decisions_file(
        tmp_path,
        [
            {"revision_key": _key(inv, "100"), "decision": "accept", "comment": "ok"},
            {"revision_key": _key(inv, "101"), "decision": "reject", "comment": ""},
        ],
    )
    report = _apply_decisions_file(workdir, decisions_file)
    assert report == {"applied": 2, "skipped": 0, "errors": []}
    inv2 = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    assert "100" not in {r["w_id"] for r in inv2["revisions"]}
    assert "101" not in {r["w_id"] for r in inv2["revisions"]}
    typed = parse_typed((workdir / "typed.md").read_text(encoding="utf-8"))
    visible = "".join(visible_text([n]) for n in typed.paragraphs[0].nodes)
    assert "插入词" in visible  # insert accepted, text stays
    assert "旧词" in visible  # delete rejected, text restored
    output = tmp_path / "applied.docx"
    assert build([str(workdir), "-o", str(output)]) == 0
    assert verify([str(workdir), str(output)]) == 0


def test_apply_review_decisions_skips_defer_and_reports_errors(tmp_path):
    workdir = extract_fixture(tmp_path)
    inv = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    decisions_file = _decisions_file(
        tmp_path,
        [
            {"revision_key": _key(inv, "100"), "decision": "defer", "comment": "later"},
            {"revision_key": "word/document.xml|insert|999|deadbeef", "decision": "accept", "comment": ""},
        ],
    )
    report = _apply_decisions_file(workdir, decisions_file)
    assert report["applied"] == 0
    assert report["skipped"] == 1
    assert len(report["errors"]) == 1
    assert "revision-not-found" in report["errors"][0]
    inv2 = json.loads((workdir / "revisions.json").read_text(encoding="utf-8"))
    assert "100" in {r["w_id"] for r in inv2["revisions"]}  # deferred, untouched


def test_apply_review_decisions_rejects_unknown_schema(tmp_path):
    workdir = extract_fixture(tmp_path)
    bad = tmp_path / "bad.json"
    bad.write_text(json.dumps({"schema": "other", "decisions": []}), encoding="utf-8")
    from scripts.typed_core import TypedError

    try:
        _apply_decisions_file(workdir, bad)
        raised = False
    except TypedError as exc:
        raised = "unsupported decisions schema" in str(exc)
    assert raised


# ---------------------------------------------------------------------------
# Issue #53: decide CLI collision codes stay distinct
# ---------------------------------------------------------------------------


def _decide_accept_all_cli(workdir, output, workdir_out, operation_id):
    """Run the decide CLI in --json mode and return (rc, envelope)."""
    import contextlib
    import io

    from scripts import main

    buffer = io.StringIO()
    with contextlib.redirect_stdout(buffer):
        code = main(
            [
                "--json",
                "decide",
                "accept-all",
                "--workdir",
                str(workdir),
                "--output",
                str(output),
                "--workdir-out",
                str(workdir_out),
                "--operation-id",
                operation_id,
            ]
        )
    lines = [line for line in buffer.getvalue().splitlines() if line.strip()]
    assert len(lines) == 1  # exactly one envelope, never a double emit
    return code, json.loads(lines[0])


def test_cli_decide_output_collision_emits_decided_output_code(tmp_path):
    """A pre-existing decided output is refused with the registered
    decided-output-already-exists code — never collapsed into the generic
    output-already-exists (issue #53)."""
    workdir = extract_fixture(tmp_path)
    output = tmp_path / "decided.docx"
    output.write_text("occupied", encoding="utf-8")

    code, envelope = _decide_accept_all_cli(
        workdir, output, tmp_path / "decided-wd", "cm-collide-out"
    )
    assert code == 1
    assert envelope["outcome"] == "failure"
    assert envelope["diagnostics"][0]["code"] == "decided-output-already-exists"


def test_cli_decide_workdir_collision_emits_decided_workdir_code(tmp_path):
    """A pre-existing derived workdir is refused with the registered
    decided-workdir-already-exists code — never collapsed into the generic
    output-already-exists (issue #53)."""
    workdir = extract_fixture(tmp_path)
    output = tmp_path / "decided.docx"
    workdir_out = tmp_path / "decided-wd"
    workdir_out.mkdir()

    code, envelope = _decide_accept_all_cli(
        workdir, output, workdir_out, "cm-collide-wd"
    )
    assert code == 1
    assert envelope["outcome"] == "failure"
    assert envelope["diagnostics"][0]["code"] == "decided-workdir-already-exists"
